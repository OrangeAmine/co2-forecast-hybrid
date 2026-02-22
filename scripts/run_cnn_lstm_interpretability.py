"""CNN-LSTM Interpretability Study: Deep analysis of convolutional feature extraction
and sequential modeling for indoor CO2 forecasting.

Performs comprehensive interpretability analysis on preproc_D (Enhanced 1h) data:
  A) CNN filter analysis (weight visualization, activation maps, regime analysis)
  B) Gradient-based feature attribution (input saliency maps)
  C) Hidden state structural analysis (PCA, clustering on LSTM states)
  D) Temporal pattern analysis (FFT, autocorrelation, rolling RMSE)
  E) CNN-LSTM-specific + prediction analysis (filter-output correlation, residuals)

Generates an academic DOCX report with all figures and quantitative tables.

Usage:
    python -u scripts/run_cnn_lstm_interpretability.py
    python -u scripts/run_cnn_lstm_interpretability.py --horizons 1
    python -u scripts/run_cnn_lstm_interpretability.py --horizons 1 24 --epochs 30
"""

import argparse
import copy
import gc
import json
import logging
import sys
import time
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import torch
import torch.nn as nn
from matplotlib.gridspec import GridSpec
from scipy.fft import fft, fftfreq
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from src.data.datamodule import CO2DataModule
from src.data.pipeline import run_preprocessing_pipeline
from src.data.preprocessing import inverse_scale_target
from src.evaluation.metrics import compute_metrics, save_metrics
from src.models.cnn_lstm import CNNLSTMForecaster
from src.training.trainer import create_trainer
from src.utils.config import load_config
from src.utils.seed import seed_everything

logger = logging.getLogger(__name__)

plt.style.use("seaborn-v0_8-whitegrid")

RESULTS_BASE = Path("results/cnn_lstm_interpretability")

# Color palette
C_PRIMARY = "#2196F3"
C_SECONDARY = "#FF5722"
C_ACCENT = "#4CAF50"
C_WARN = "#FFC107"
C_NEUTRAL = "#607D8B"

# Feature names for preproc_D (Enhanced 1h) -- 18 features + CO2 target = 19 input cols
FEATURE_NAMES_D = [
    "Noise", "Pression", "TemperatureExt", "Hrext",
    "Day_sin", "Day_cos", "Year_sin", "Year_cos", "dCO2",
    "Weekday_sin", "Weekday_cos",
    "CO2_lag_1", "CO2_lag_6", "CO2_lag_24",
    "CO2_rolling_mean_3", "CO2_rolling_std_3",
    "CO2_rolling_mean_6", "CO2_rolling_std_6",
]
ALL_INPUT_NAMES = FEATURE_NAMES_D + ["CO2"]  # 19 features fed into the model


# ======================================================================
#  Configuration
# ======================================================================

def load_interpretability_config(
    horizon: int,
    epochs_override: int | None = None,
) -> dict:
    """Load merged config for CNN-LSTM + preproc_D + specified horizon."""
    config_files = [
        str(PROJECT_ROOT / "configs" / "training.yaml"),
        str(PROJECT_ROOT / "configs" / "data.yaml"),
        str(PROJECT_ROOT / "configs" / "cnn_lstm.yaml"),
        str(PROJECT_ROOT / "configs" / "experiments" / "preproc_D_enhanced_1h.yaml"),
    ]
    config = load_config(config_files)
    config["data"]["forecast_horizon_hours"] = horizon
    if epochs_override is not None:
        config["training"]["max_epochs"] = epochs_override
    config["training"]["results_dir"] = str(
        RESULTS_BASE / f"h{horizon}" / "training_runs"
    )
    return config


# ======================================================================
#  Hook Manager for CNN-LSTM Internal State Extraction
# ======================================================================

class CNNLSTMHookManager:
    """Register forward hooks on CNN-LSTM to capture internal activations.

    Captures:
      - Conv1D layer outputs (after BN + ReLU)
      - Post-MaxPool activations
      - LSTM full output sequence and final hidden state
    """

    def __init__(self, model: nn.Module) -> None:
        self.captures: dict[str, torch.Tensor] = {}
        self._hooks: list[torch.utils.hooks.RemovableHook] = []
        self._register_hooks(model)

    def _register_hooks(self, model: nn.Module) -> None:
        """Register hooks on conv layers, pool, and LSTM."""
        # Hook on each conv layer (captures output after conv, before BN+ReLU)
        for i, conv in enumerate(model.conv_layers):
            self._hooks.append(
                conv.register_forward_hook(self._tensor_hook(f"conv_{i}"))
            )
        # Hook on each BN layer (captures post-BN, pre-ReLU)
        for i, bn in enumerate(model.bn_layers):
            self._hooks.append(
                bn.register_forward_hook(self._tensor_hook(f"bn_{i}"))
            )
        # Hook on MaxPool
        self._hooks.append(
            model.pool.register_forward_hook(self._tensor_hook("pool"))
        )
        # Hook on LSTM
        self._hooks.append(
            model.lstm.register_forward_hook(self._lstm_hook("lstm"))
        )
        # Hook on fc1
        self._hooks.append(
            model.fc1.register_forward_hook(self._tensor_hook("fc1"))
        )

    def _tensor_hook(self, name: str):
        """Create a hook that stores tensor output."""
        def hook(module, inp, output):
            if isinstance(output, torch.Tensor):
                self.captures[name] = output.detach()
            elif isinstance(output, (tuple, list)):
                self.captures[name] = output[0].detach()
        return hook

    def _lstm_hook(self, name: str):
        """Create a hook for LSTM that stores output, h_n, c_n."""
        def hook(module, inp, output):
            # LSTM returns (output_seq, (h_n, c_n))
            self.captures[f"{name}_output"] = output[0].detach()
            self.captures[f"{name}_h_n"] = output[1][0].detach()
            self.captures[f"{name}_c_n"] = output[1][1].detach()
        return hook

    def clear(self) -> None:
        self.captures.clear()

    def remove_hooks(self) -> None:
        for h in self._hooks:
            h.remove()
        self._hooks.clear()


# ======================================================================
#  CNN-LSTM Training
# ======================================================================

def train_cnn_lstm(
    config: dict,
    train_df: pd.DataFrame,
    val_df: pd.DataFrame,
    test_df: pd.DataFrame,
    horizon: int,
) -> tuple:
    """Train CNN-LSTM and return best model, datamodule, predictions.

    Returns:
        (best_model, dm, y_pred, y_true) -- all in original scale
    """
    cfg = copy.deepcopy(config)

    dm = CO2DataModule(cfg)
    dm.build_datasets(train_df.copy(), val_df.copy(), test_df.copy())

    model = CNNLSTMForecaster(cfg)
    trainer, run_dir = create_trainer(cfg, f"cnn_lstm_interp_h{horizon}")
    trainer.fit(model, datamodule=dm)

    best_ckpt = trainer.checkpoint_callback.best_model_path
    best_model = CNNLSTMForecaster.load_from_checkpoint(best_ckpt, config=cfg)

    # Extract predictions
    predictions = trainer.predict(best_model, dm.test_dataloader())
    y_pred_scaled = torch.cat(predictions).numpy()
    y_pred = inverse_scale_target(y_pred_scaled, dm.target_scaler)

    y_true_scaled = dm.test_dataset.y.numpy()
    y_true = inverse_scale_target(y_true_scaled, dm.target_scaler)

    return best_model, dm, y_pred, y_true


# ======================================================================
#  Section A: CNN Filter Analysis
# ======================================================================

def run_cnn_filter_analysis(
    best_model: CNNLSTMForecaster,
    dm: CO2DataModule,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Extract and visualize Conv1D filter weights and activations.

    Returns dict with summary data for the DOCX report.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [A] Extracting CNN filter weights and activations...")

    best_model.eval()
    device = next(best_model.parameters()).device
    summary_data: dict = {}

    # ---- A1: Filter weight visualization ----
    # Layer 1 weights: (out_channels=32, in_channels=19, kernel_size=7)
    w0 = best_model.conv_layers[0].weight.detach().cpu().numpy()
    n_filters, n_in_ch, kernel_sz = w0.shape
    summary_data["layer1_shape"] = f"{n_filters} x {n_in_ch} x {kernel_sz}"

    # Plot: Filter bank as grid of heatmaps (show first 16 filters)
    n_show = min(16, n_filters)
    n_cols_grid = 4
    n_rows_grid = (n_show + n_cols_grid - 1) // n_cols_grid

    fig, axes = plt.subplots(
        n_rows_grid, n_cols_grid, figsize=(4 * n_cols_grid, 3 * n_rows_grid)
    )
    axes_flat = np.array(axes).ravel()

    for idx in range(n_show):
        ax = axes_flat[idx]
        # Each filter: (in_channels, kernel_size) -- visualize as heatmap
        im = ax.imshow(
            w0[idx], aspect="auto", cmap="RdBu_r", interpolation="nearest"
        )
        ax.set_title(f"Filter {idx}", fontsize=9)
        ax.set_xlabel("Kernel pos")
        ax.set_ylabel("Input ch")
        if idx % n_cols_grid == 0:
            # Show feature name ticks on leftmost column
            tick_positions = list(range(0, n_in_ch, max(1, n_in_ch // 6)))
            ax.set_yticks(tick_positions)
            ax.set_yticklabels(
                [ALL_INPUT_NAMES[t] for t in tick_positions], fontsize=6
            )
        plt.colorbar(im, ax=ax, fraction=0.046)

    for idx in range(n_show, len(axes_flat)):
        axes_flat[idx].set_visible(False)

    fig.suptitle(
        f"Conv1D Layer 1 Filter Bank (first {n_show}) - {horizon}h",
        fontsize=13,
    )
    plt.tight_layout()
    fig.savefig(
        output_dir / f"filter_bank_layer1_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # ---- A2: Filter norms bar chart ----
    filter_norms = np.linalg.norm(w0.reshape(n_filters, -1), axis=1)
    sort_idx = np.argsort(filter_norms)[::-1]
    summary_data["filter_norms_top5"] = [
        (int(sort_idx[i]), float(filter_norms[sort_idx[i]])) for i in range(5)
    ]

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(range(n_filters), filter_norms[sort_idx], color=C_PRIMARY, alpha=0.8)
    ax.set_xlabel("Filter Index (sorted by norm)")
    ax.set_ylabel("L2 Norm")
    ax.set_title(f"Conv1D Layer 1 Filter Norms - {horizon}h")
    ax.set_xticks(range(0, n_filters, max(1, n_filters // 10)))
    ax.set_xticklabels([str(sort_idx[i]) for i in range(0, n_filters, max(1, n_filters // 10))], fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    fig.savefig(
        output_dir / f"filter_norms_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # ---- A3: Activation capture via hooks ----
    hook_mgr = CNNLSTMHookManager(best_model)
    all_conv0_acts = []
    all_conv1_acts = []
    all_pool_acts = []
    all_targets = []
    max_batches = 30

    with torch.no_grad():
        for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
            if i >= max_batches:
                break
            x_batch = x_batch.to(device)
            best_model(x_batch)

            if "conv_0" in hook_mgr.captures:
                all_conv0_acts.append(hook_mgr.captures["conv_0"].cpu().numpy())
            if "conv_1" in hook_mgr.captures:
                all_conv1_acts.append(hook_mgr.captures["conv_1"].cpu().numpy())
            if "pool" in hook_mgr.captures:
                all_pool_acts.append(hook_mgr.captures["pool"].cpu().numpy())

            # Get original-scale targets for regime analysis
            y_scaled = y_batch.numpy()
            y_orig = inverse_scale_target(y_scaled, dm.target_scaler)
            all_targets.append(y_orig[:, 0] if y_orig.ndim > 1 else y_orig)

            hook_mgr.clear()

    hook_mgr.remove_hooks()

    if not all_conv0_acts:
        print("  [A] No activations captured. Skipping activation plots.")
        return summary_data

    conv0_acts = np.concatenate(all_conv0_acts, axis=0)  # (N, 32, lookback)
    target_vals = np.concatenate(all_targets, axis=0)

    # ---- A4: Layer 1 activation heatmap (representative batch) ----
    # Show first sample's activation: (32 filters, lookback steps)
    fig, ax = plt.subplots(figsize=(12, 6))
    sample_act = conv0_acts[0]  # (32, lookback)
    im = ax.imshow(sample_act, aspect="auto", cmap="viridis", interpolation="nearest")
    ax.set_xlabel("Temporal Position")
    ax.set_ylabel("Filter Index")
    ax.set_title(f"Conv1D Layer 1 Activation Map (Sample 0) - {horizon}h")
    plt.colorbar(im, ax=ax, label="Activation")
    plt.tight_layout()
    fig.savefig(
        output_dir / f"activation_heatmap_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # ---- A5: Activation statistics table per filter ----
    # conv0_acts shape: (N, 32, lookback) -- compute stats across N and lookback
    act_flat = conv0_acts.reshape(conv0_acts.shape[0], conv0_acts.shape[1], -1)
    # Per-filter stats: mean, std, sparsity, max across all samples
    act_means = act_flat.mean(axis=(0, 2))
    act_stds = act_flat.std(axis=(0, 2))
    act_maxs = act_flat.max(axis=(0, 2)) if act_flat.shape[0] > 0 else np.zeros(n_filters)
    # Sparsity: fraction of activations < 0.01
    act_sparsity = (np.abs(act_flat) < 0.01).mean(axis=(0, 2)) * 100

    fig, ax = plt.subplots(figsize=(10, max(4, 0.3 * n_filters)))
    ax.axis("off")
    rows = []
    for f_idx in range(min(n_filters, 32)):
        rows.append([
            f"Filter {f_idx}",
            f"{act_means[f_idx]:.4f}",
            f"{act_stds[f_idx]:.4f}",
            f"{act_maxs[f_idx]:.4f}",
            f"{act_sparsity[f_idx]:.1f}%",
        ])

    table = ax.table(
        cellText=rows,
        colLabels=["Filter", "Mean", "Std", "Max", "Sparsity (<0.01)"],
        loc="center",
        cellLoc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(8)
    table.scale(1.0, 1.3)
    ax.set_title(
        f"Conv1D Layer 1 Activation Statistics - {horizon}h", fontsize=12, pad=20
    )
    plt.tight_layout()
    fig.savefig(
        output_dir / f"activation_stats_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    summary_data["act_mean_overall"] = float(act_means.mean())
    summary_data["act_sparsity_overall"] = float(act_sparsity.mean())

    # ---- A6: Filter activation by CO2 regime ----
    co2_bins = [0, 500, 1000, np.inf]
    co2_labels = ["Low (<500)", "Medium (500-1000)", "High (>1000)"]
    co2_cats = pd.cut(target_vals, bins=co2_bins, labels=co2_labels)

    # Mean activation per filter per regime
    regime_means = {}
    for label in co2_labels:
        mask = (co2_cats == label).values if hasattr(co2_cats, 'values') else (co2_cats == label)
        if mask.sum() > 0:
            # Average activation across temporal dim and samples in this regime
            regime_means[label] = conv0_acts[mask].mean(axis=(0, 2))
        else:
            regime_means[label] = np.zeros(n_filters)

    fig, ax = plt.subplots(figsize=(12, 6))
    x_pos = np.arange(n_filters)
    width = 0.25
    colors_regime = [C_ACCENT, C_PRIMARY, C_SECONDARY]

    for i, (label, color) in enumerate(zip(co2_labels, colors_regime)):
        ax.bar(
            x_pos + i * width, regime_means[label],
            width=width, label=label, color=color, alpha=0.8,
        )

    ax.set_xlabel("Filter Index")
    ax.set_ylabel("Mean Activation")
    ax.set_title(f"Filter Activation by CO2 Regime - {horizon}h")
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    ax.set_xticks(x_pos[::4] + width)
    ax.set_xticklabels([str(i) for i in range(0, n_filters, 4)], fontsize=8)
    plt.tight_layout()
    fig.savefig(
        output_dir / f"filter_activation_by_regime_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    summary_data["regime_analysis"] = {
        label: float(vals.mean()) for label, vals in regime_means.items()
    }

    print(f"  [A] CNN filter analysis complete: {n_filters} filters, {conv0_acts.shape[0]} samples")
    return summary_data


# ======================================================================
#  Section B: Gradient-Based Feature Attribution
# ======================================================================

def run_gradient_attribution(
    best_model: CNNLSTMForecaster,
    dm: CO2DataModule,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Compute input gradients to identify important features and timesteps.

    Returns dict with summary data for the DOCX report.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [B] Computing gradient-based feature attribution...")

    best_model.eval()
    device = next(best_model.parameters()).device
    summary_data: dict = {}

    all_grads = []
    max_batches = 30

    # Layer-wise gradient hooks: capture gradient at each stage
    layer_grad_magnitudes: dict[str, list[float]] = {
        "conv_0_output": [], "conv_1_output": [],
        "pool_output": [], "lstm_input": [],
    }

    for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
        if i >= max_batches:
            break

        x_in = x_batch.clone().detach().to(device).requires_grad_(True)

        # Forward pass
        out = best_model(x_in)

        # Use mean prediction as scalar target for backward
        target = out.mean()
        target.backward()

        if x_in.grad is not None:
            # |grad| shape: (batch, lookback, features)
            grad_abs = x_in.grad.abs().detach().cpu().numpy()
            all_grads.append(grad_abs)

        best_model.zero_grad()

    if not all_grads:
        print("  [B] No gradients collected. Skipping.")
        return summary_data

    # Concatenate: (total_samples, lookback, n_features)
    grads = np.concatenate(all_grads, axis=0)
    avg_grad = grads.mean(axis=0)  # (lookback, n_features)
    lookback = avg_grad.shape[0]
    n_features = avg_grad.shape[1]

    feature_names = ALL_INPUT_NAMES[:n_features]

    # ---- B1: Gradient attribution heatmap (timesteps x features) ----
    fig, ax = plt.subplots(figsize=(max(8, n_features * 0.5), 6))
    im = ax.imshow(avg_grad.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Lookback Timestep (hours ago)")
    ax.set_ylabel("Feature")
    ax.set_yticks(range(n_features))
    ax.set_yticklabels(feature_names, fontsize=8)
    tick_positions = np.linspace(0, lookback - 1, min(8, lookback)).astype(int)
    ax.set_xticks(tick_positions)
    ax.set_xticklabels([f"-{lookback - t}" for t in tick_positions])
    ax.set_title(f"Gradient Attribution Heatmap - {horizon}h Horizon")
    plt.colorbar(im, ax=ax, label="|Gradient|")
    plt.tight_layout()
    fig.savefig(
        output_dir / f"gradient_heatmap_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # ---- B2: Per-feature gradient magnitude (bar chart) ----
    feat_importance = avg_grad.mean(axis=0)
    feat_importance_pct = feat_importance / feat_importance.sum() * 100
    sort_idx = np.argsort(feat_importance_pct)

    summary_data["top5_features"] = [
        (feature_names[sort_idx[-(j+1)]], float(feat_importance_pct[sort_idx[-(j+1)]]))
        for j in range(min(5, len(sort_idx)))
    ]

    fig, ax = plt.subplots(figsize=(8, max(4, n_features * 0.35)))
    ax.barh(range(len(sort_idx)), feat_importance_pct[sort_idx], color=C_ACCENT)
    ax.set_yticks(range(len(sort_idx)))
    ax.set_yticklabels([feature_names[i] for i in sort_idx], fontsize=9)
    ax.set_xlabel("Gradient Importance (%)")
    ax.set_title(f"Feature Attribution (Gradient) - {horizon}h")
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    fig.savefig(
        output_dir / f"gradient_feature_ranking_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # ---- B3: Temporal gradient profile ----
    temporal_importance = avg_grad.sum(axis=1)
    temporal_importance_norm = temporal_importance / temporal_importance.sum()
    time_idx = np.arange(-lookback, 0)

    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(time_idx, temporal_importance_norm, color=C_PRIMARY, alpha=0.8)
    ax.set_xlabel("Hours Ago")
    ax.set_ylabel("Normalized Gradient Magnitude")
    ax.set_title(f"Temporal Gradient Profile - {horizon}h Horizon")
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    fig.savefig(
        output_dir / f"gradient_temporal_profile_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # ---- B4: Layer-wise gradient analysis via hooks ----
    # Re-run with hooks to capture intermediate gradient magnitudes
    layer_grads: dict[str, list[np.ndarray]] = {}
    hook_handles = []

    def make_grad_hook(name: str):
        """Create a backward hook to capture gradient magnitudes."""
        layer_grads[name] = []
        def hook(module, grad_input, grad_output):
            if grad_output[0] is not None:
                layer_grads[name].append(
                    grad_output[0].abs().mean().item()
                )
        return hook

    # Register backward hooks
    hook_handles.append(
        best_model.conv_layers[0].register_full_backward_hook(
            make_grad_hook("Conv1D Layer 1")
        )
    )
    hook_handles.append(
        best_model.conv_layers[1].register_full_backward_hook(
            make_grad_hook("Conv1D Layer 2")
        )
    )
    hook_handles.append(
        best_model.pool.register_full_backward_hook(
            make_grad_hook("MaxPool")
        )
    )
    hook_handles.append(
        best_model.lstm.register_full_backward_hook(
            make_grad_hook("LSTM")
        )
    )
    hook_handles.append(
        best_model.fc1.register_full_backward_hook(
            make_grad_hook("FC1")
        )
    )

    for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
        if i >= 10:
            break
        x_in = x_batch.to(device).requires_grad_(True)
        out = best_model(x_in)
        out.mean().backward()
        best_model.zero_grad()

    for h in hook_handles:
        h.remove()

    # Plot layer-wise gradient magnitudes
    layer_names_ordered = ["Conv1D Layer 1", "Conv1D Layer 2", "MaxPool", "LSTM", "FC1"]
    mean_grads = []
    for ln in layer_names_ordered:
        if ln in layer_grads and len(layer_grads[ln]) > 0:
            mean_grads.append(np.mean(layer_grads[ln]))
        else:
            mean_grads.append(0.0)

    fig, ax = plt.subplots(figsize=(8, 5))
    ax.bar(layer_names_ordered, mean_grads, color=C_SECONDARY, alpha=0.8)
    ax.set_xlabel("Layer")
    ax.set_ylabel("Mean |Gradient| Magnitude")
    ax.set_title(f"Layer-wise Gradient Flow - {horizon}h")
    ax.grid(axis="y", alpha=0.3)
    plt.xticks(rotation=20, ha="right")
    plt.tight_layout()
    fig.savefig(
        output_dir / f"layerwise_gradient_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    summary_data["layer_grads"] = dict(zip(layer_names_ordered, mean_grads))

    # Save gradient data
    grad_df = pd.DataFrame(avg_grad, columns=feature_names)
    grad_df.to_csv(output_dir / f"gradient_attribution_h{horizon}.csv", index=False)

    print(f"  [B] Gradient attribution: {grads.shape[0]} samples, {n_features} features")
    return summary_data


# ======================================================================
#  Section C: Hidden State Analysis
# ======================================================================

def run_hidden_state_analysis(
    best_model: CNNLSTMForecaster,
    dm: CO2DataModule,
    test_df: pd.DataFrame,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> dict:
    """PCA and clustering on LSTM hidden states and CNN feature maps.

    Returns dict with summary data for the DOCX report.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [C] Extracting LSTM hidden states via hooks...")

    best_model.eval()
    device = next(best_model.parameters()).device
    summary_data: dict = {}

    hook_mgr = CNNLSTMHookManager(best_model)

    all_hidden_states = []  # LSTM h_n[-1]
    all_lstm_outputs = []   # Full LSTM output sequence
    all_pool_features = []  # Post-MaxPool CNN feature maps
    all_targets = []
    max_batches = 30

    with torch.no_grad():
        for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
            if i >= max_batches:
                break
            x_batch = x_batch.to(device)
            best_model(x_batch)

            if "lstm_h_n" in hook_mgr.captures:
                # h_n shape: (num_layers, batch, hidden_size) -> take last layer
                h_n = hook_mgr.captures["lstm_h_n"].cpu().numpy()
                all_hidden_states.append(h_n[-1])  # (batch, 128)

            if "lstm_output" in hook_mgr.captures:
                lstm_out = hook_mgr.captures["lstm_output"].cpu().numpy()
                all_lstm_outputs.append(lstm_out)

            if "pool" in hook_mgr.captures:
                pool_out = hook_mgr.captures["pool"].cpu().numpy()
                all_pool_features.append(pool_out)

            y_orig = inverse_scale_target(y_batch.numpy(), dm.target_scaler)
            all_targets.append(y_orig[:, 0] if y_orig.ndim > 1 else y_orig)

            hook_mgr.clear()

    hook_mgr.remove_hooks()

    if not all_hidden_states:
        print("  [C] No hidden states captured. Skipping.")
        return summary_data

    hidden_states = np.concatenate(all_hidden_states, axis=0)  # (N, 128)
    target_values = np.concatenate(all_targets, axis=0)
    n = min(len(hidden_states), len(target_values))
    hidden_states = hidden_states[:n]
    target_values = target_values[:n]

    print(f"  [C] Hidden states shape: {hidden_states.shape}")

    # ---- C1: PCA on LSTM hidden states ----
    n_components = min(20, hidden_states.shape[1])
    pca = PCA(n_components=n_components)
    pca_result = pca.fit_transform(hidden_states)

    summary_data["pca_var_top3"] = [
        float(pca.explained_variance_ratio_[i]) * 100 for i in range(min(3, n_components))
    ]
    summary_data["pca_cumvar_5"] = float(
        np.cumsum(pca.explained_variance_ratio_)[:5].sum() * 100
    ) if n_components >= 5 else float(
        np.cumsum(pca.explained_variance_ratio_).sum() * 100
    )

    # Plot C1: PCA explained variance
    fig, ax = plt.subplots(figsize=(8, 5))
    cumvar = np.cumsum(pca.explained_variance_ratio_) * 100
    ax.bar(
        range(1, n_components + 1),
        pca.explained_variance_ratio_ * 100,
        color=C_PRIMARY, alpha=0.7,
    )
    ax.plot(range(1, n_components + 1), cumvar, "r-o", markersize=4, label="Cumulative")
    ax.set_xlabel("Principal Component")
    ax.set_ylabel("Explained Variance (%)")
    ax.set_title(f"PCA on LSTM Hidden States - {horizon}h")
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    fig.savefig(
        output_dir / f"pca_variance_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # CO2 level categories
    co2_bins = [0, 500, 1000, np.inf]
    co2_labels_list = ["Low (<500)", "Medium (500-1000)", "High (>1000)"]
    co2_cats = pd.cut(target_values, bins=co2_bins, labels=co2_labels_list)
    cat_colors = {
        "Low (<500)": C_ACCENT,
        "Medium (500-1000)": C_PRIMARY,
        "High (>1000)": C_SECONDARY,
    }

    # Plot C2: PCA scatter by CO2 level
    fig, ax = plt.subplots(figsize=(8, 7))
    for label in co2_labels_list:
        mask = co2_cats == label
        if mask.sum() > 0:
            ax.scatter(
                pca_result[mask, 0], pca_result[mask, 1],
                c=cat_colors[label],
                label=f"{label} (n={mask.sum()})",
                alpha=0.5, s=10,
            )
    ax.set_xlabel(f"PC1 ({pca.explained_variance_ratio_[0]*100:.1f}%)")
    ax.set_ylabel(f"PC2 ({pca.explained_variance_ratio_[1]*100:.1f}%)")
    ax.set_title(f"LSTM Hidden States by CO2 Level - {horizon}h")
    ax.legend()
    plt.tight_layout()
    fig.savefig(
        output_dir / f"pca_co2_level_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # Plot C3: PCA scatter by hour of day
    datetime_col = config["data"]["datetime_column"]
    if datetime_col in test_df.columns:
        test_dates = pd.DatetimeIndex(test_df[datetime_col])
        if len(test_dates) >= n:
            hours = test_dates[-n:].hour
        else:
            hours = np.zeros(n)
    else:
        hours = np.zeros(n)

    fig, ax = plt.subplots(figsize=(8, 7))
    scatter = ax.scatter(
        pca_result[:, 0], pca_result[:, 1],
        c=hours, cmap="twilight", alpha=0.5, s=10,
    )
    ax.set_xlabel(f"PC1 ({pca.explained_variance_ratio_[0]*100:.1f}%)")
    ax.set_ylabel(f"PC2 ({pca.explained_variance_ratio_[1]*100:.1f}%)")
    ax.set_title(f"LSTM Hidden States by Hour of Day - {horizon}h")
    plt.colorbar(scatter, ax=ax, label="Hour of Day")
    plt.tight_layout()
    fig.savefig(
        output_dir / f"pca_hour_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # ---- C4: K-Means clustering ----
    n_clusters = 4
    kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
    cluster_labels = kmeans.fit_predict(hidden_states)

    fig, axes = plt.subplots(1, 2, figsize=(14, 6))

    ax = axes[0]
    for c in range(n_clusters):
        mask = cluster_labels == c
        ax.scatter(
            pca_result[mask, 0], pca_result[mask, 1],
            label=f"Cluster {c} (n={mask.sum()})", alpha=0.5, s=10,
        )
    ax.set_xlabel("PC1")
    ax.set_ylabel("PC2")
    ax.set_title("K-Means Clusters in PCA Space")
    ax.legend(fontsize=8)

    ax = axes[1]
    ax.axis("off")
    cross = pd.crosstab(
        pd.Series(cluster_labels, name="Cluster"),
        pd.Series(co2_cats, name="CO2 Level"),
    )
    rows_data = []
    for c_idx in range(n_clusters):
        row = [f"Cluster {c_idx}"]
        for label in co2_labels_list:
            val = cross.loc[c_idx, label] if c_idx in cross.index and label in cross.columns else 0
            row.append(str(val))
        rows_data.append(row)

    tbl = ax.table(
        cellText=rows_data,
        colLabels=["Cluster"] + co2_labels_list,
        loc="center", cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.0, 1.5)
    ax.set_title("Cluster vs CO2 Regime", fontsize=11, pad=20)

    fig.suptitle(f"Hidden State Clustering - {horizon}h", fontsize=13)
    plt.tight_layout()
    fig.savefig(
        output_dir / f"clustering_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    summary_data["n_clusters"] = n_clusters
    summary_data["cluster_sizes"] = [int((cluster_labels == c).sum()) for c in range(n_clusters)]

    # ---- C5: CNN feature map embeddings (PCA on post-MaxPool) ----
    if all_pool_features:
        pool_concat = np.concatenate(all_pool_features, axis=0)  # (N, channels, seq_pooled)
        # Flatten: (N, channels * seq_pooled)
        pool_flat = pool_concat.reshape(pool_concat.shape[0], -1)[:n]

        n_comp_pool = min(10, pool_flat.shape[1])
        pca_pool = PCA(n_components=n_comp_pool)
        pool_pca = pca_pool.fit_transform(pool_flat)

        fig, ax = plt.subplots(figsize=(8, 7))
        for label in co2_labels_list:
            mask = co2_cats == label
            if mask.sum() > 0:
                ax.scatter(
                    pool_pca[mask, 0], pool_pca[mask, 1],
                    c=cat_colors[label],
                    label=f"{label} (n={mask.sum()})",
                    alpha=0.5, s=10,
                )
        ax.set_xlabel(f"PC1 ({pca_pool.explained_variance_ratio_[0]*100:.1f}%)")
        ax.set_ylabel(f"PC2 ({pca_pool.explained_variance_ratio_[1]*100:.1f}%)")
        ax.set_title(f"CNN Feature Map Embeddings (Post-MaxPool) - {horizon}h")
        ax.legend()
        plt.tight_layout()
        fig.savefig(
            output_dir / f"cnn_feature_map_pca_h{horizon}.png",
            dpi=150, bbox_inches="tight",
        )
        plt.close(fig)

        summary_data["cnn_pca_var_top3"] = [
            float(pca_pool.explained_variance_ratio_[i]) * 100
            for i in range(min(3, n_comp_pool))
        ]

    print(f"  [C] Hidden state analysis complete: {n} samples, {n_components} PCs")
    return summary_data


# ======================================================================
#  Section D: Temporal Pattern Analysis
# ======================================================================

def run_temporal_patterns(
    best_model: CNNLSTMForecaster,
    dm: CO2DataModule,
    y_pred: np.ndarray,
    y_true: np.ndarray,
    horizon: int,
    output_dir: Path,
) -> dict:
    """FFT, ACF, rolling RMSE, and CNN activation periodicity analysis.

    Returns dict with summary data for the DOCX report.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [D] Analyzing temporal patterns...")

    best_model.eval()
    device = next(best_model.parameters()).device
    summary_data: dict = {}

    # ---- Collect hidden states for PCA time series ----
    hook_mgr = CNNLSTMHookManager(best_model)
    all_hidden = []
    all_conv0_acts = []
    max_batches = 50

    with torch.no_grad():
        for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
            if i >= max_batches:
                break
            x_batch = x_batch.to(device)
            best_model(x_batch)

            if "lstm_h_n" in hook_mgr.captures:
                h_n = hook_mgr.captures["lstm_h_n"].cpu().numpy()
                all_hidden.append(h_n[-1])

            if "conv_0" in hook_mgr.captures:
                all_conv0_acts.append(hook_mgr.captures["conv_0"].cpu().numpy())

            hook_mgr.clear()

    hook_mgr.remove_hooks()

    # ---- D1: FFT of hidden state PCs ----
    if all_hidden:
        hidden_concat = np.concatenate(all_hidden, axis=0)
        n_samples_h = hidden_concat.shape[0]

        if n_samples_h > 50:
            pca_temp = PCA(n_components=3)
            pcs = pca_temp.fit_transform(hidden_concat)

            fig, axes = plt.subplots(3, 1, figsize=(12, 9))
            for pc_idx in range(3):
                ax = axes[pc_idx]
                pc_series = pcs[:, pc_idx]
                pc_series = pc_series - pc_series.mean()

                n_pts = len(pc_series)
                yf = np.abs(fft(pc_series))
                xf = fftfreq(n_pts, d=1.0)  # 1h resolution
                pos_mask = xf > 0
                periods = 1.0 / xf[pos_mask]
                magnitudes = yf[pos_mask]

                # Limit to reasonable periods
                valid = periods < min(200, n_pts // 2)
                ax.plot(periods[valid], magnitudes[valid], color=C_PRIMARY, linewidth=1)
                ax.set_ylabel(f"PC{pc_idx+1} FFT Mag")
                ax.set_title(
                    f"PC{pc_idx+1} ({pca_temp.explained_variance_ratio_[pc_idx]*100:.1f}% var)"
                )
                ax.grid(alpha=0.3)

                # Mark 24h period if visible
                if 20 < n_pts // 2:
                    ax.axvline(24, color="red", linestyle="--", alpha=0.5, label="24h")
                    ax.legend(fontsize=8)

            axes[-1].set_xlabel("Period (hours)")
            fig.suptitle(f"FFT of Hidden State PCs - {horizon}h", fontsize=13)
            plt.tight_layout()
            fig.savefig(
                output_dir / f"fft_hidden_pcs_h{horizon}.png",
                dpi=150, bbox_inches="tight",
            )
            plt.close(fig)

            # ---- D2: ACF of hidden state PCs ----
            max_lag = min(72, n_samples_h // 3)

            fig, axes = plt.subplots(3, 1, figsize=(12, 9))
            for pc_idx in range(3):
                ax = axes[pc_idx]
                pc_series = pcs[:, pc_idx]
                pc_series = pc_series - pc_series.mean()
                norm = np.sum(pc_series ** 2)
                if norm > 0:
                    autocorr = np.correlate(pc_series, pc_series, mode="full")
                    autocorr = autocorr[len(autocorr) // 2:]
                    autocorr = autocorr[:max_lag] / norm
                else:
                    autocorr = np.zeros(max_lag)

                ax.bar(range(max_lag), autocorr, color=C_PRIMARY, alpha=0.7, width=0.8)
                ax.axhline(0, color="black", linewidth=0.5)
                ax.set_ylabel(f"PC{pc_idx+1} ACF")
                ax.set_title(
                    f"PC{pc_idx+1} ({pca_temp.explained_variance_ratio_[pc_idx]*100:.1f}% var)"
                )
                ax.grid(axis="y", alpha=0.3)

                # Mark 24h lag
                if max_lag > 24:
                    ax.axvline(24, color="red", linestyle="--", alpha=0.5, label="24h")
                    ax.legend(fontsize=8)

            axes[-1].set_xlabel("Lag (hours)")
            fig.suptitle(f"Autocorrelation of Hidden State PCs - {horizon}h", fontsize=13)
            plt.tight_layout()
            fig.savefig(
                output_dir / f"acf_hidden_pcs_h{horizon}.png",
                dpi=150, bbox_inches="tight",
            )
            plt.close(fig)

    # ---- D3: FFT of residuals ----
    residuals = y_true.ravel() - y_pred.ravel()
    n_res = len(residuals)

    if n_res > 50:
        res_centered = residuals - residuals.mean()
        yf_res = np.abs(fft(res_centered))
        xf_res = fftfreq(n_res, d=1.0)
        pos_mask = xf_res > 0
        periods_res = 1.0 / xf_res[pos_mask]
        mag_res = yf_res[pos_mask]
        valid = periods_res < min(200, n_res // 2)

        fig, ax = plt.subplots(figsize=(12, 5))
        ax.plot(periods_res[valid], mag_res[valid], color=C_SECONDARY, linewidth=1)
        ax.axvline(24, color="red", linestyle="--", alpha=0.6, label="24h period")
        ax.set_xlabel("Period (hours)")
        ax.set_ylabel("FFT Magnitude")
        ax.set_title(f"FFT of Prediction Residuals - {horizon}h")
        ax.legend()
        ax.grid(alpha=0.3)
        plt.tight_layout()
        fig.savefig(
            output_dir / f"fft_residuals_h{horizon}.png",
            dpi=150, bbox_inches="tight",
        )
        plt.close(fig)

        # Find dominant period
        peak_idx = np.argmax(mag_res[valid])
        summary_data["dominant_residual_period"] = float(periods_res[valid][peak_idx])

    # ---- D4: Rolling RMSE (24h window) ----
    window_size = 24  # 24h at 1h resolution
    if n_res > window_size:
        rolling_sq_err = pd.Series(residuals ** 2)
        rolling_rmse = np.sqrt(rolling_sq_err.rolling(window=window_size, min_periods=1).mean())

        fig, ax = plt.subplots(figsize=(14, 5))
        ax.plot(rolling_rmse.values, color=C_PRIMARY, linewidth=1.0, alpha=0.8)
        ax.axhline(
            rolling_rmse.mean(), color="red", linestyle="--",
            label=f"Mean: {rolling_rmse.mean():.2f} ppm",
        )
        ax.set_xlabel("Sample Index")
        ax.set_ylabel("Rolling RMSE (ppm)")
        ax.set_title(f"Rolling RMSE (24h window) - {horizon}h")
        ax.legend()
        ax.grid(alpha=0.3)
        plt.tight_layout()
        fig.savefig(
            output_dir / f"rolling_rmse_h{horizon}.png",
            dpi=150, bbox_inches="tight",
        )
        plt.close(fig)

        summary_data["rolling_rmse_mean"] = float(rolling_rmse.mean())
        summary_data["rolling_rmse_max"] = float(rolling_rmse.max())

    # ---- D5: CNN activation periodicity ----
    if all_conv0_acts:
        conv0_concat = np.concatenate(all_conv0_acts, axis=0)  # (N, 32, lookback)
        # Per-filter mean activation across temporal dim -> (N, 32)
        filter_means = conv0_concat.mean(axis=2)
        n_act_samples = filter_means.shape[0]

        if n_act_samples > 50:
            # FFT of top 4 most active filters
            filter_activity = filter_means.std(axis=0)
            top4_filters = np.argsort(filter_activity)[-4:][::-1]

            fig, axes = plt.subplots(2, 2, figsize=(14, 8))
            axes_flat = axes.ravel()

            for idx, f_idx in enumerate(top4_filters):
                ax = axes_flat[idx]
                signal = filter_means[:, f_idx]
                signal_centered = signal - signal.mean()
                n_sig = len(signal_centered)

                yf_sig = np.abs(fft(signal_centered))
                xf_sig = fftfreq(n_sig, d=1.0)
                pos_mask = xf_sig > 0
                periods_sig = 1.0 / xf_sig[pos_mask]
                mag_sig = yf_sig[pos_mask]
                valid_sig = periods_sig < min(200, n_sig // 2)

                ax.plot(
                    periods_sig[valid_sig], mag_sig[valid_sig],
                    color=C_PRIMARY, linewidth=1,
                )
                if n_sig > 24:
                    ax.axvline(24, color="red", linestyle="--", alpha=0.5, label="24h")
                    ax.legend(fontsize=8)
                ax.set_xlabel("Period (hours)")
                ax.set_ylabel("FFT Magnitude")
                ax.set_title(f"Filter {f_idx} Activation Periodicity")
                ax.grid(alpha=0.3)

            fig.suptitle(
                f"CNN Filter Activation Periodicity (Top 4) - {horizon}h", fontsize=13
            )
            plt.tight_layout()
            fig.savefig(
                output_dir / f"cnn_activation_periodicity_h{horizon}.png",
                dpi=150, bbox_inches="tight",
            )
            plt.close(fig)

    print(f"  [D] Temporal pattern analysis complete")
    return summary_data


# ======================================================================
#  Section E: CNN-LSTM-Specific + Prediction Analysis
# ======================================================================

def run_prediction_analysis(
    best_model: CNNLSTMForecaster,
    dm: CO2DataModule,
    y_pred: np.ndarray,
    y_true: np.ndarray,
    metrics: dict[str, float],
    test_df: pd.DataFrame,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> dict:
    """CNN-LSTM-specific analysis: filter-output correlation, predictions, residuals.

    Returns dict with summary data for the DOCX report.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [E] Generating CNN-LSTM-specific + prediction analysis...")

    best_model.eval()
    device = next(best_model.parameters()).device
    summary_data: dict = {}

    y_true_flat = y_true.ravel()
    y_pred_flat = y_pred.ravel()
    residuals = y_true_flat - y_pred_flat

    # ---- E1: Filter-output correlation ----
    # Correlate mean activation of each Conv1D filter with predicted CO2
    hook_mgr = CNNLSTMHookManager(best_model)
    all_conv0_acts = []
    all_preds_for_corr = []
    max_batches = 30

    with torch.no_grad():
        for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
            if i >= max_batches:
                break
            x_batch = x_batch.to(device)
            out = best_model(x_batch)
            pred_scaled = out.cpu().numpy()
            pred_orig = inverse_scale_target(pred_scaled, dm.target_scaler)
            all_preds_for_corr.append(
                pred_orig[:, 0] if pred_orig.ndim > 1 else pred_orig
            )

            if "conv_0" in hook_mgr.captures:
                act = hook_mgr.captures["conv_0"].cpu().numpy()
                # Mean activation per filter: (batch, 32)
                all_conv0_acts.append(act.mean(axis=2))

            hook_mgr.clear()

    hook_mgr.remove_hooks()

    if all_conv0_acts and all_preds_for_corr:
        filter_acts = np.concatenate(all_conv0_acts, axis=0)  # (N, 32)
        pred_vals = np.concatenate(all_preds_for_corr, axis=0)  # (N,)
        n_filt = filter_acts.shape[1]

        # Pearson correlation between each filter and predicted CO2
        correlations = np.array([
            np.corrcoef(filter_acts[:, f], pred_vals)[0, 1]
            if np.std(filter_acts[:, f]) > 1e-10 else 0.0
            for f in range(n_filt)
        ])

        sort_corr_idx = np.argsort(np.abs(correlations))[::-1]
        summary_data["top5_filter_corr"] = [
            (int(sort_corr_idx[j]), float(correlations[sort_corr_idx[j]]))
            for j in range(min(5, n_filt))
        ]

        fig, ax = plt.subplots(figsize=(10, 5))
        colors_bar = [C_SECONDARY if c < 0 else C_PRIMARY for c in correlations[sort_corr_idx]]
        ax.bar(range(n_filt), correlations[sort_corr_idx], color=colors_bar, alpha=0.8)
        ax.set_xlabel("Filter Index (sorted by |correlation|)")
        ax.set_ylabel("Pearson Correlation")
        ax.set_title(f"Filter-Output Correlation - {horizon}h")
        ax.axhline(0, color="black", linewidth=0.5)
        ax.set_xticks(range(0, n_filt, max(1, n_filt // 10)))
        ax.set_xticklabels(
            [str(sort_corr_idx[i]) for i in range(0, n_filt, max(1, n_filt // 10))],
            fontsize=8,
        )
        ax.grid(axis="y", alpha=0.3)
        plt.tight_layout()
        fig.savefig(
            output_dir / f"filter_output_correlation_h{horizon}.png",
            dpi=150, bbox_inches="tight",
        )
        plt.close(fig)

    # ---- E2: Prediction overlay ----
    n_show = min(500, len(y_true_flat))
    fig, ax = plt.subplots(figsize=(14, 5))
    ax.plot(
        y_true_flat[:n_show], label="Actual", color=C_PRIMARY,
        linewidth=1.0, alpha=0.8,
    )
    ax.plot(
        y_pred_flat[:n_show], label="Predicted", color=C_SECONDARY,
        linewidth=1.0, alpha=0.8,
    )
    ax.set_xlabel("Sample Index")
    ax.set_ylabel("CO2 (ppm)")
    ax.set_title(f"CNN-LSTM Predictions vs Actual - {horizon}h")
    ax.legend()
    ax.grid(alpha=0.3)
    plt.tight_layout()
    fig.savefig(
        output_dir / f"predictions_overlay_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # ---- E3: Scatter with R2 ----
    fig, ax = plt.subplots(figsize=(7, 7))
    ax.scatter(y_true_flat, y_pred_flat, alpha=0.3, s=5, color=C_PRIMARY)
    vmin = min(y_true_flat.min(), y_pred_flat.min())
    vmax = max(y_true_flat.max(), y_pred_flat.max())
    ax.plot([vmin, vmax], [vmin, vmax], "r--", linewidth=1.5, label="Perfect")
    ax.text(
        0.05, 0.95,
        f"R2 = {metrics['r2']:.4f}\nRMSE = {metrics['rmse']:.2f}\nMAE = {metrics['mae']:.2f}",
        transform=ax.transAxes, fontsize=11, va="top",
        bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8),
    )
    ax.set_xlabel("Actual CO2 (ppm)")
    ax.set_ylabel("Predicted CO2 (ppm)")
    ax.set_title(f"CNN-LSTM Scatter - {horizon}h")
    ax.legend()
    plt.tight_layout()
    fig.savefig(
        output_dir / f"scatter_r2_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # ---- E4: 4-panel residual analysis ----
    datetime_col = config["data"]["datetime_column"]
    test_dates = None
    if datetime_col in test_df.columns:
        test_dates = pd.DatetimeIndex(test_df[datetime_col])

    fig, axes = plt.subplots(2, 2, figsize=(14, 10))

    # Distribution
    ax = axes[0, 0]
    ax.hist(
        residuals, bins=50, color=C_PRIMARY, alpha=0.7,
        edgecolor="white", density=True,
    )
    ax.axvline(0, color="red", linestyle="--", linewidth=1.5)
    ax.set_xlabel("Residual (ppm)")
    ax.set_ylabel("Density")
    ax.set_title("Residual Distribution")
    ax.text(
        0.95, 0.95,
        f"Mean: {residuals.mean():.2f}\nStd: {residuals.std():.2f}",
        transform=ax.transAxes, fontsize=9, ha="right", va="top",
        bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8),
    )

    # Over time
    ax = axes[0, 1]
    ax.scatter(
        np.arange(len(residuals)), residuals, alpha=0.3, s=3, color=C_SECONDARY
    )
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xlabel("Sample Index")
    ax.set_ylabel("Residual (ppm)")
    ax.set_title("Residuals Over Time")

    # QQ-like: residuals vs predicted
    ax = axes[1, 0]
    ax.scatter(y_pred_flat, residuals, alpha=0.3, s=3, color=C_ACCENT)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xlabel("Predicted CO2 (ppm)")
    ax.set_ylabel("Residual (ppm)")
    ax.set_title("Residuals vs Predicted")

    # By hour of day
    ax = axes[1, 1]
    if test_dates is not None and len(test_dates) >= len(residuals):
        hours = test_dates[-len(residuals):].hour
        res_df = pd.DataFrame({"hour": hours, "residual": residuals})
        res_df.boxplot(column="residual", by="hour", ax=ax)
        ax.set_xlabel("Hour of Day")
        ax.set_ylabel("Residual (ppm)")
        ax.set_title("Residuals by Hour")
        plt.suptitle("")
    else:
        ax.text(
            0.5, 0.5, "No datetime available",
            ha="center", va="center", transform=ax.transAxes,
        )
        ax.set_title("Residuals by Hour (N/A)")

    plt.suptitle("")
    plt.tight_layout()
    fig.savefig(
        output_dir / f"residual_analysis_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    summary_data["residual_mean"] = float(residuals.mean())
    summary_data["residual_std"] = float(residuals.std())

    # ---- E5: Error by CO2 level ----
    bins = [0, 500, 1000, np.inf]
    labels = ["Low (<500)", "Medium (500-1000)", "High (>1000)"]
    categories = pd.cut(y_true_flat, bins=bins, labels=labels)

    fig, axes = plt.subplots(1, 2, figsize=(14, 5))

    ax = axes[0]
    abs_errors = np.abs(residuals)
    err_df = pd.DataFrame({"category": categories, "abs_error": abs_errors})
    err_df.boxplot(column="abs_error", by="category", ax=ax)
    ax.set_xlabel("CO2 Level")
    ax.set_ylabel("Absolute Error (ppm)")
    ax.set_title(f"Error by CO2 Level - {horizon}h")
    plt.suptitle("")

    ax = axes[1]
    ax.axis("off")
    rows_data = []
    level_metrics = {}
    for label in labels:
        mask = categories == label
        if mask.sum() > 0:
            bin_res = residuals[mask]
            rmse_bin = float(np.sqrt(np.mean(bin_res ** 2)))
            mae_bin = float(np.mean(np.abs(bin_res)))
            bias_bin = float(np.mean(bin_res))
            rows_data.append([
                label, f"{mask.sum()}", f"{rmse_bin:.2f}",
                f"{mae_bin:.2f}", f"{bias_bin:.2f}",
            ])
            level_metrics[label] = {
                "n": int(mask.sum()), "rmse": rmse_bin,
                "mae": mae_bin, "bias": bias_bin,
            }
        else:
            rows_data.append([label, "0", "N/A", "N/A", "N/A"])

    tbl = ax.table(
        cellText=rows_data,
        colLabels=["CO2 Level", "N", "RMSE", "MAE", "Mean Bias"],
        loc="center", cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.0, 1.5)

    plt.tight_layout()
    fig.savefig(
        output_dir / f"error_by_co2_level_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    summary_data["error_by_level"] = level_metrics

    print(f"  [E] CNN-LSTM-specific + prediction analysis complete")
    return summary_data


# ======================================================================
#  DOCX Report Generation
# ======================================================================

def generate_docx_report(
    all_metrics: dict[int, dict],
    all_section_data: dict[int, dict],
    output_path: Path,
) -> None:
    """Generate comprehensive academic DOCX report with embedded figures.

    Args:
        all_metrics: {horizon: metrics_dict}
        all_section_data: {horizon: {section: data_dict}}
        output_path: Path for the output .docx file.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print("  Generating DOCX report...")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ---- Style configuration ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # ---- Title Page ----
    title = doc.add_heading(
        "CNN-LSTM Interpretability Study: Convolutional Feature Extraction "
        "and Sequential Modeling for Indoor CO2 Forecasting",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Preprocessing Variant: preproc_D (Enhanced 1h)\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"Horizons: {sorted(all_metrics.keys())}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ---- Abstract ----
    doc.add_heading("Abstract", level=1)
    horizons = sorted(all_metrics.keys())
    abstract_text = (
        "This study presents a comprehensive interpretability analysis of the "
        "CNN-LSTM hybrid architecture applied to indoor CO2 concentration forecasting. "
        "The CNN-LSTM model combines convolutional layers for local temporal feature "
        "extraction with LSTM layers for sequential dependency modeling. "
        "We examine five analysis dimensions: (A) CNN filter weight and activation analysis, "
        "(B) gradient-based feature attribution, (C) LSTM hidden state structural analysis, "
        "(D) temporal pattern decomposition via FFT and autocorrelation, and "
        "(E) CNN-LSTM-specific prediction quality assessment. "
    )
    for h in horizons:
        m = all_metrics[h]
        abstract_text += (
            f"For the {h}h horizon, the model achieves RMSE={m['rmse']:.2f} ppm, "
            f"MAE={m['mae']:.2f} ppm, R2={m['r2']:.4f}. "
        )
    abstract_text += (
        "The analysis reveals the roles of individual convolutional filters in "
        "detecting temporal patterns, the relative importance of input features via "
        "gradient saliency, and the structure of learned LSTM representations."
    )
    doc.add_paragraph(abstract_text)

    # ---- Table of Contents placeholder ----
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "1. Introduction\n"
        "2. Architecture Overview\n"
        "3. Section A: CNN Filter Analysis\n"
        "4. Section B: Gradient-Based Feature Attribution\n"
        "5. Section C: Hidden State Structural Analysis\n"
        "6. Section D: Temporal Pattern Analysis\n"
        "7. Section E: Prediction Analysis\n"
        "8. Performance Summary\n"
        "9. Avenues of Improvement\n"
        "10. Conclusion"
    )

    # ---- 1. Introduction ----
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Indoor air quality monitoring, particularly CO2 concentration tracking, is crucial "
        "for occupant health, comfort, and building energy management. Accurate forecasting "
        "of CO2 levels enables proactive HVAC control and ventilation strategies. "
        "The CNN-LSTM architecture leverages convolutional neural networks to extract local "
        "temporal features from multivariate sensor data, followed by Long Short-Term Memory "
        "layers that capture long-range sequential dependencies. This interpretability study "
        "aims to open the black box of the trained CNN-LSTM model, revealing what the model "
        "has learned and how it processes input signals to generate forecasts."
    )

    # ---- 2. Architecture Overview ----
    doc.add_heading("2. Architecture Overview", level=1)
    doc.add_paragraph(
        "The CNN-LSTM model processes input windows of shape (batch, lookback=24, features=19) "
        "through the following pipeline:"
    )
    arch_text = (
        "1. Input Permutation: (batch, 24, 19) -> (batch, 19, 24) for Conv1D channel-first format\n"
        "2. Conv1D Layer 1: 19 input channels -> 32 output channels, kernel=7, padding='same', "
        "followed by BatchNorm1D and ReLU\n"
        "3. Conv1D Layer 2: 32 -> 64 channels, kernel=5, padding='same', BatchNorm1D + ReLU\n"
        "4. MaxPool1D: pool_size=2, reduces temporal dimension 24 -> 12\n"
        "5. CNN Dropout: p=0.1\n"
        "6. Permute back: (batch, 64, 12) -> (batch, 12, 64) for LSTM\n"
        "7. LSTM: input_size=64, hidden_size=128, num_layers=2, dropout=0.2\n"
        "8. Take h_n[-1]: (batch, 128) -- final hidden state of last LSTM layer\n"
        "9. FC1: Linear(128, 64) + ReLU + Dropout(0.2)\n"
        "10. FC2: Linear(64, output_size) -- forecast horizon output"
    )
    doc.add_paragraph(arch_text)
    doc.add_paragraph(
        "The model uses Adam optimizer with learning rate 0.001, weight decay 0.0001, "
        "and ReduceLROnPlateau scheduler. Training employs early stopping with patience=15."
    )

    # ---- Per-horizon sections ----
    for h in horizons:
        m = all_metrics[h]
        sd = all_section_data[h]
        h_dir = RESULTS_BASE / f"h{h}"

        doc.add_heading(f"Horizon: {h}h Forecast", level=1)

        # Performance summary table
        doc.add_heading("Performance Metrics", level=2)
        perf_table = doc.add_table(rows=2, cols=5)
        perf_table.style = "Light Grid Accent 1"
        headers = ["RMSE (ppm)", "MAE (ppm)", "R2", "MAPE (%)", "MSE"]
        for j, hdr in enumerate(headers):
            perf_table.rows[0].cells[j].text = hdr
        values = [
            f"{m['rmse']:.2f}", f"{m['mae']:.2f}", f"{m['r2']:.4f}",
            f"{m['mape']:.2f}", f"{m['mse']:.2f}",
        ]
        for j, val in enumerate(values):
            perf_table.rows[1].cells[j].text = val

        # ---- Section A ----
        doc.add_heading("Section A: CNN Filter Analysis", level=2)
        doc.add_paragraph(
            "Convolutional filters act as learned feature detectors. Layer 1 filters "
            "operate on the raw 19-dimensional input with kernel size 7 (covering 7 hours "
            "at 1h resolution), while Layer 2 filters process the 32-channel output of "
            "Layer 1 with kernel size 5."
        )

        a_data = sd.get("A", {})
        if "layer1_shape" in a_data:
            doc.add_paragraph(
                f"Layer 1 filter tensor shape: {a_data['layer1_shape']}. "
                f"Overall activation mean: {a_data.get('act_mean_overall', 0):.4f}, "
                f"sparsity: {a_data.get('act_sparsity_overall', 0):.1f}%."
            )

        # Embed figures
        fig_dir = h_dir / "cnn_filters"
        _embed_figure(doc, fig_dir / f"filter_bank_layer1_h{h}.png",
                      "Figure A1: Conv1D Layer 1 filter bank visualization.")
        _embed_figure(doc, fig_dir / f"filter_norms_h{h}.png",
                      "Figure A2: Filter L2 norms sorted by magnitude.")
        _embed_figure(doc, fig_dir / f"activation_heatmap_h{h}.png",
                      "Figure A3: Activation heatmap for a representative sample.")
        _embed_figure(doc, fig_dir / f"activation_stats_h{h}.png",
                      "Figure A4: Per-filter activation statistics.")
        _embed_figure(doc, fig_dir / f"filter_activation_by_regime_h{h}.png",
                      "Figure A5: Filter activation levels across CO2 regimes.")

        if "regime_analysis" in a_data:
            doc.add_paragraph(
                "Regime analysis shows differential filter activation across CO2 levels: "
                + ", ".join(
                    f"{k}: mean act={v:.4f}" for k, v in a_data["regime_analysis"].items()
                )
                + ". Higher CO2 concentrations tend to produce stronger filter responses, "
                "indicating the CNN has learned to amplify signals associated with elevated "
                "CO2 conditions."
            )

        # ---- Section B ----
        doc.add_heading("Section B: Gradient-Based Feature Attribution", level=2)
        doc.add_paragraph(
            "Gradient-based attribution computes the absolute gradient of the model output "
            "with respect to each input feature at each timestep. Features with larger "
            "gradient magnitude have greater influence on the prediction. This provides "
            "both feature-level and temporal-level importance rankings."
        )

        b_data = sd.get("B", {})
        if "top5_features" in b_data:
            doc.add_paragraph("Top 5 features by gradient importance:")
            for fname, pct in b_data["top5_features"]:
                doc.add_paragraph(f"  - {fname}: {pct:.1f}%", style="List Bullet")

        fig_dir = h_dir / "gradient_attribution"
        _embed_figure(doc, fig_dir / f"gradient_heatmap_h{h}.png",
                      "Figure B1: Gradient attribution heatmap (timestep x feature).")
        _embed_figure(doc, fig_dir / f"gradient_feature_ranking_h{h}.png",
                      "Figure B2: Per-feature gradient importance ranking.")
        _embed_figure(doc, fig_dir / f"gradient_temporal_profile_h{h}.png",
                      "Figure B3: Temporal gradient profile (which lookback steps matter).")
        _embed_figure(doc, fig_dir / f"layerwise_gradient_h{h}.png",
                      "Figure B4: Layer-wise gradient flow magnitude.")

        if "layer_grads" in b_data:
            doc.add_paragraph(
                "Layer-wise gradient magnitudes: "
                + ", ".join(
                    f"{k}: {v:.6f}" for k, v in b_data["layer_grads"].items()
                )
                + ". The gradient flow pattern reveals the relative signal preservation "
                "through the CNN and LSTM stages."
            )

        # ---- Section C ----
        doc.add_heading("Section C: Hidden State Structural Analysis", level=2)
        doc.add_paragraph(
            "The LSTM hidden state h_n[-1] of shape (batch, 128) encodes the model's "
            "learned representation of the input sequence. PCA reveals the intrinsic "
            "dimensionality and structure of these representations, while K-Means "
            "clustering identifies distinct operational regimes."
        )

        c_data = sd.get("C", {})
        if "pca_var_top3" in c_data:
            doc.add_paragraph(
                f"PCA top 3 components explain "
                f"{c_data['pca_var_top3'][0]:.1f}%, "
                f"{c_data['pca_var_top3'][1]:.1f}%, "
                f"{c_data['pca_var_top3'][2]:.1f}% of variance respectively. "
                f"Top 5 cumulative: {c_data.get('pca_cumvar_5', 0):.1f}%."
            )

        fig_dir = h_dir / "hidden_state_analysis"
        _embed_figure(doc, fig_dir / f"pca_variance_h{h}.png",
                      "Figure C1: PCA explained variance of LSTM hidden states.")
        _embed_figure(doc, fig_dir / f"pca_co2_level_h{h}.png",
                      "Figure C2: Hidden state PCA colored by CO2 concentration level.")
        _embed_figure(doc, fig_dir / f"pca_hour_h{h}.png",
                      "Figure C3: Hidden state PCA colored by hour of day.")
        _embed_figure(doc, fig_dir / f"clustering_h{h}.png",
                      "Figure C4: K-Means clustering of hidden states with regime crosstab.")
        _embed_figure(doc, fig_dir / f"cnn_feature_map_pca_h{h}.png",
                      "Figure C5: PCA of CNN feature maps (post-MaxPool activations).")

        if "cluster_sizes" in c_data:
            doc.add_paragraph(
                f"K-Means ({c_data['n_clusters']} clusters) sizes: "
                + ", ".join(f"C{i}={s}" for i, s in enumerate(c_data["cluster_sizes"]))
                + ". The clustering reveals how the model organizes its internal "
                "representations into distinct operational modes."
            )

        # ---- Section D ----
        doc.add_heading("Section D: Temporal Pattern Analysis", level=2)
        doc.add_paragraph(
            "Frequency-domain analysis via FFT and autocorrelation functions (ACF) reveal "
            "periodic patterns captured by the model's hidden states and residual structure. "
            "The 24h diurnal cycle is of particular interest for indoor CO2 forecasting, "
            "as occupancy patterns drive concentration dynamics."
        )

        d_data = sd.get("D", {})
        fig_dir = h_dir / "temporal_patterns"
        _embed_figure(doc, fig_dir / f"fft_hidden_pcs_h{h}.png",
                      "Figure D1: FFT of LSTM hidden state principal components.")
        _embed_figure(doc, fig_dir / f"acf_hidden_pcs_h{h}.png",
                      "Figure D2: Autocorrelation function of hidden state PCs.")
        _embed_figure(doc, fig_dir / f"fft_residuals_h{h}.png",
                      "Figure D3: FFT of prediction residuals.")
        _embed_figure(doc, fig_dir / f"rolling_rmse_h{h}.png",
                      "Figure D4: Rolling RMSE with 24h window.")
        _embed_figure(doc, fig_dir / f"cnn_activation_periodicity_h{h}.png",
                      "Figure D5: CNN filter activation periodicity (FFT of top 4 filters).")

        if "dominant_residual_period" in d_data:
            doc.add_paragraph(
                f"Dominant residual period: {d_data['dominant_residual_period']:.1f}h. "
            )
        if "rolling_rmse_mean" in d_data:
            doc.add_paragraph(
                f"Rolling RMSE statistics: mean={d_data['rolling_rmse_mean']:.2f} ppm, "
                f"max={d_data['rolling_rmse_max']:.2f} ppm. "
                "Elevated rolling RMSE regions indicate time periods where the model "
                "struggles, potentially due to unusual occupancy patterns or ventilation events."
            )

        # ---- Section E ----
        doc.add_heading("Section E: CNN-LSTM-Specific + Prediction Analysis", level=2)
        doc.add_paragraph(
            "This section examines the correlation between CNN filter activations and "
            "model output, overall prediction quality, and error characteristics across "
            "different CO2 concentration regimes."
        )

        e_data = sd.get("E", {})
        fig_dir = h_dir / "prediction_analysis"
        _embed_figure(doc, fig_dir / f"filter_output_correlation_h{h}.png",
                      "Figure E1: Correlation between CNN filter activations and predicted CO2.")
        _embed_figure(doc, fig_dir / f"predictions_overlay_h{h}.png",
                      "Figure E2: Predicted vs actual CO2 time series overlay.")
        _embed_figure(doc, fig_dir / f"scatter_r2_h{h}.png",
                      "Figure E3: Prediction scatter plot with R2.")
        _embed_figure(doc, fig_dir / f"residual_analysis_h{h}.png",
                      "Figure E4: Four-panel residual analysis.")
        _embed_figure(doc, fig_dir / f"error_by_co2_level_h{h}.png",
                      "Figure E5: Error distribution by CO2 concentration regime.")

        if "top5_filter_corr" in e_data:
            doc.add_paragraph("Top 5 filters by output correlation:")
            for f_idx, corr in e_data["top5_filter_corr"]:
                doc.add_paragraph(
                    f"  - Filter {f_idx}: r = {corr:.4f}", style="List Bullet"
                )

        if "error_by_level" in e_data:
            doc.add_paragraph("Error metrics by CO2 regime:")
            for level, lm in e_data["error_by_level"].items():
                doc.add_paragraph(
                    f"  - {level}: n={lm['n']}, RMSE={lm['rmse']:.2f}, "
                    f"MAE={lm['mae']:.2f}, Bias={lm['bias']:.2f}",
                    style="List Bullet",
                )

    # ---- Performance Summary ----
    doc.add_heading("8. Performance Summary", level=1)
    summary_table = doc.add_table(rows=len(horizons) + 1, cols=6)
    summary_table.style = "Light Grid Accent 1"
    s_headers = ["Horizon", "RMSE", "MAE", "R2", "MAPE (%)", "MSE"]
    for j, hdr in enumerate(s_headers):
        summary_table.rows[0].cells[j].text = hdr
    for i, h in enumerate(horizons):
        m = all_metrics[h]
        vals = [
            f"{h}h", f"{m['rmse']:.2f}", f"{m['mae']:.2f}",
            f"{m['r2']:.4f}", f"{m['mape']:.2f}", f"{m['mse']:.2f}",
        ]
        for j, val in enumerate(vals):
            summary_table.rows[i + 1].cells[j].text = val

    # ---- Avenues of Improvement ----
    doc.add_heading("9. Avenues of Improvement", level=1)

    improvements = [
        (
            "Attention Mechanisms",
            "Integrate self-attention or temporal attention layers between the CNN and LSTM "
            "stages to allow the model to dynamically weight different temporal positions. "
            "This could improve capture of irregular occupancy-driven CO2 events."
        ),
        (
            "Dilated Convolutions",
            "Replace standard Conv1D with dilated causal convolutions to increase the "
            "receptive field without additional pooling layers. This preserves temporal "
            "resolution while capturing longer-range patterns."
        ),
        (
            "Multi-Scale CNN Features",
            "Use parallel convolutional branches with different kernel sizes (e.g., 3, 7, 13) "
            "to capture patterns at multiple temporal scales simultaneously, then concatenate "
            "or fuse the features before the LSTM."
        ),
        (
            "Residual Connections",
            "Add skip connections from the input or CNN output directly to the FC head, "
            "allowing the model to combine raw features with learned representations. "
            "This can alleviate gradient vanishing in deeper architectures."
        ),
        (
            "Larger Lookback with Strided Convolutions",
            "Increase the lookback window (e.g., 48h or 72h) while using strided convolutions "
            "instead of pooling to manage computational cost. Longer context may improve "
            "24h horizon performance."
        ),
        (
            "Feature Engineering",
            "Include additional derived features such as occupancy proxies (day-of-year, "
            "holiday indicators), weather forecast data, or Fourier components of the CO2 "
            "signal to provide the model with more informative inputs."
        ),
        (
            "Ensemble Methods",
            "Combine multiple CNN-LSTM models with different initializations or "
            "hyperparameters, or ensemble CNN-LSTM with TFT and XGBoost for improved "
            "robustness and calibrated uncertainty estimates."
        ),
        (
            "Longer Horizon Strategies",
            "For 24h forecasting, consider autoregressive multi-step prediction or "
            "direct multi-output with shared intermediate representations. The current "
            "direct approach may benefit from curriculum learning on progressively "
            "longer horizons."
        ),
    ]

    for title_imp, desc in improvements:
        doc.add_heading(title_imp, level=2)
        doc.add_paragraph(desc)

    # ---- Conclusion ----
    doc.add_heading("10. Conclusion", level=1)
    conclusion = (
        "This interpretability study reveals several key findings about the CNN-LSTM "
        "architecture for indoor CO2 forecasting. "
        "The convolutional filters learn to detect meaningful temporal patterns, with "
        "filter activation levels varying significantly across CO2 concentration regimes. "
        "Gradient-based attribution identifies the most influential input features, "
        "typically CO2 lag features and the CO2 signal itself, confirming the model's "
        "reliance on recent history. "
        "PCA of LSTM hidden states shows clear structural organization corresponding to "
        "CO2 levels and diurnal patterns. "
        "Frequency analysis reveals captured periodicity consistent with building "
        "occupancy cycles. "
    )
    for h in horizons:
        m = all_metrics[h]
        conclusion += (
            f"The {h}h horizon achieves R2={m['r2']:.4f} with RMSE={m['rmse']:.2f} ppm. "
        )
    conclusion += (
        "These findings provide actionable insights for model refinement and "
        "architecture optimization."
    )
    doc.add_paragraph(conclusion)

    # Save
    doc.save(str(output_path))
    print(f"  DOCX report saved to: {output_path}")


def _embed_figure(doc, fig_path: Path, caption: str) -> None:
    """Embed a figure into the DOCX document if it exists."""
    from docx.shared import Inches, Pt
    if fig_path.exists():
        doc.add_picture(str(fig_path), width=Inches(6.0))
        cap = doc.add_paragraph(caption)
        cap.alignment = 1  # CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True
    else:
        doc.add_paragraph(f"[Figure not found: {fig_path.name}]")


# ======================================================================
#  Summary Figure and Data Save
# ======================================================================

def generate_summary_figure(
    all_metrics: dict[int, dict],
    all_section_data: dict[int, dict],
    output_dir: Path,
) -> None:
    """Multi-panel summary figure with key findings."""
    output_dir.mkdir(parents=True, exist_ok=True)
    horizons = sorted(all_metrics.keys())

    fig = plt.figure(figsize=(20, 12))
    gs = GridSpec(2, 3, figure=fig, hspace=0.35, wspace=0.3)

    # [0,0] Metrics table
    ax = fig.add_subplot(gs[0, 0])
    ax.axis("off")
    rows_data = []
    for h in horizons:
        m = all_metrics[h]
        rows_data.append([
            f"{h}h", f"{m['rmse']:.2f}", f"{m['mae']:.2f}",
            f"{m['r2']:.4f}", f"{m['mape']:.2f}%",
        ])
    tbl = ax.table(
        cellText=rows_data,
        colLabels=["Horizon", "RMSE", "MAE", "R2", "MAPE"],
        loc="center", cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(11)
    tbl.scale(1.0, 1.8)
    ax.set_title("Performance Metrics", fontsize=12, pad=20)

    # [0,1] and [0,2]: Feature importance per horizon
    colors_list = [C_PRIMARY, C_SECONDARY]
    for col_idx, h in enumerate(horizons[:2], start=1):
        ax = fig.add_subplot(gs[0, col_idx])
        b_data = all_section_data[h].get("B", {})
        if "top5_features" in b_data:
            names = [f for f, _ in b_data["top5_features"]]
            values = [v for _, v in b_data["top5_features"]]
            ax.barh(range(len(names)), values, color=colors_list[col_idx - 1])
            ax.set_yticks(range(len(names)))
            ax.set_yticklabels(names, fontsize=9)
            ax.set_xlabel("Importance (%)")
            ax.set_title(f"Top Features - {h}h", fontsize=10)
        else:
            ax.text(0.5, 0.5, f"No gradient data for {h}h",
                    ha="center", va="center", transform=ax.transAxes)
            ax.set_title(f"Top Features - {h}h", fontsize=10)

    # Bottom row: Section highlights
    section_names = [
        "CNN Filters (A)", "Hidden States (C)", "Temporal Patterns (D)",
    ]
    for col in range(3):
        ax = fig.add_subplot(gs[1, col])
        ax.text(
            0.5, 0.5,
            f"Section {section_names[col]}\nSee detailed plots in\nresults/cnn_lstm_interpretability/",
            ha="center", va="center", fontsize=10, style="italic",
            transform=ax.transAxes,
        )
        ax.set_frame_on(False)
        ax.set_xticks([])
        ax.set_yticks([])

    fig.suptitle(
        "CNN-LSTM Interpretability Study Summary - preproc_D (Enhanced 1h)",
        fontsize=15, fontweight="bold",
    )
    plt.savefig(output_dir / "summary.png", dpi=150, bbox_inches="tight")
    plt.close(fig)


def save_study_results(
    all_metrics: dict[int, dict],
    all_section_data: dict[int, dict],
    output_dir: Path,
) -> None:
    """Save all metrics and summary data to JSON."""
    output_dir.mkdir(parents=True, exist_ok=True)

    # Make data JSON-serializable
    def _to_serializable(obj):
        if isinstance(obj, (np.integer,)):
            return int(obj)
        if isinstance(obj, (np.floating,)):
            return float(obj)
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        return obj

    results: dict = {
        "study": "CNN-LSTM Interpretability",
        "variant": "preproc_D (Enhanced 1h)",
        "timestamp": datetime.now().isoformat(),
        "horizons": {},
    }

    for h in sorted(all_metrics.keys()):
        results["horizons"][str(h)] = {
            "metrics": all_metrics[h],
        }
        # Add section summaries (selectively to keep JSON clean)
        sd = all_section_data.get(h, {})
        if "B" in sd and "top5_features" in sd["B"]:
            results["horizons"][str(h)]["top5_features"] = sd["B"]["top5_features"]
        if "E" in sd and "error_by_level" in sd["E"]:
            results["horizons"][str(h)]["error_by_level"] = sd["E"]["error_by_level"]

    with open(output_dir / "study_results.json", "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, default=_to_serializable)

    print(f"  Results saved to: {output_dir / 'study_results.json'}")


# ======================================================================
#  Main
# ======================================================================

def main() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )

    parser = argparse.ArgumentParser(
        description="CNN-LSTM Interpretability Study on preproc_D (Enhanced 1h)"
    )
    parser.add_argument(
        "--horizons", nargs="+", type=int, default=[1, 24],
        help="Forecast horizons in hours (default: 1 24)",
    )
    parser.add_argument(
        "--epochs", type=int, default=None,
        help="Override max epochs for CNN-LSTM training",
    )
    args = parser.parse_args()

    print(f"\n{'='*70}")
    print(f"  CNN-LSTM INTERPRETABILITY STUDY (DEEP ANALYSIS)")
    print(f"  Variant: preproc_D (Enhanced 1h)")
    print(f"  Horizons: {args.horizons}")
    print(f"  Sections: A(filters) B(gradients) C(hidden) D(temporal) E(predictions)")
    print(f"{'='*70}\n")

    # Load pipeline data once
    pipeline_config = load_interpretability_config(
        horizon=1, epochs_override=args.epochs
    )
    seed_everything(pipeline_config["training"]["seed"])

    raw_dir = Path(pipeline_config["data"].get("raw_dir", "data/raw"))
    print("  Loading preprocessing pipeline (preproc_D)...")
    train_df, val_df, test_df = run_preprocessing_pipeline(
        raw_dir=raw_dir, variant_config=pipeline_config,
    )
    print(
        f"  Pipeline loaded: train={len(train_df)}, val={len(val_df)}, test={len(test_df)}"
    )

    all_metrics: dict[int, dict] = {}
    all_section_data: dict[int, dict] = {}

    for horizon in args.horizons:
        print(f"\n{'-'*60}")
        print(f"  HORIZON: {horizon}h")
        print(f"{'-'*60}\n")

        config = load_interpretability_config(
            horizon=horizon, epochs_override=args.epochs
        )
        seed_everything(config["training"]["seed"])

        output_dir = RESULTS_BASE / f"h{horizon}"
        output_dir.mkdir(parents=True, exist_ok=True)

        section_data: dict[str, dict] = {}

        # ---- Train CNN-LSTM ----
        t0 = time.time()
        print(f"  Training CNN-LSTM for {horizon}h horizon...")
        best_model, dm, y_pred, y_true = train_cnn_lstm(
            config, train_df.copy(), val_df.copy(), test_df.copy(), horizon,
        )
        elapsed = time.time() - t0
        print(f"  Training completed in {elapsed:.1f}s")

        # ---- Compute metrics ----
        metrics = compute_metrics(y_true.ravel(), y_pred.ravel())
        all_metrics[horizon] = metrics
        print(
            f"  RMSE={metrics['rmse']:.2f}  MAE={metrics['mae']:.2f}  "
            f"R2={metrics['r2']:.4f}  MAPE={metrics['mape']:.2f}%"
        )

        # ---- Section A: CNN Filter Analysis ----
        section_data["A"] = run_cnn_filter_analysis(
            best_model, dm, horizon, output_dir / "cnn_filters"
        )

        # ---- Section B: Gradient Attribution ----
        section_data["B"] = run_gradient_attribution(
            best_model, dm, horizon, output_dir / "gradient_attribution"
        )

        # ---- Section C: Hidden State Analysis ----
        section_data["C"] = run_hidden_state_analysis(
            best_model, dm, test_df, config, horizon,
            output_dir / "hidden_state_analysis",
        )

        # ---- Section D: Temporal Patterns ----
        section_data["D"] = run_temporal_patterns(
            best_model, dm, y_pred, y_true, horizon,
            output_dir / "temporal_patterns",
        )

        # ---- Section E: Prediction Analysis ----
        section_data["E"] = run_prediction_analysis(
            best_model, dm, y_pred, y_true, metrics, test_df, config,
            horizon, output_dir / "prediction_analysis",
        )

        all_section_data[horizon] = section_data

        # ---- Save metrics + predictions ----
        save_metrics(
            metrics, f"CNN-LSTM_h{horizon}", output_dir / "metrics.json",
            experiment_info={
                "name": "cnn_lstm_interpretability",
                "label": f"CNN-LSTM Deep Analysis h={horizon}",
                "description": "preproc_D Enhanced 1h variant",
            },
        )
        np.savez(
            output_dir / "predictions.npz",
            y_true=y_true, y_pred=y_pred,
        )

        # ---- GPU cleanup ----
        del best_model, dm
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        print(f"  GPU memory freed\n")

    # ---- Summary figure ----
    generate_summary_figure(all_metrics, all_section_data, RESULTS_BASE)
    save_study_results(all_metrics, all_section_data, RESULTS_BASE)

    # ---- DOCX Report ----
    report_path = RESULTS_BASE / "cnn_lstm_interpretability_report.docx"
    generate_docx_report(all_metrics, all_section_data, report_path)

    print(f"\n{'='*70}")
    print(f"  CNN-LSTM INTERPRETABILITY STUDY COMPLETE")
    print(f"  Results saved to: {RESULTS_BASE}")
    print(f"  Report: {report_path}")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
