"""CatBoost Interpretability Study: Ordered boosting and oblivious tree analysis.

Performs comprehensive interpretability analysis on preproc_D (Enhanced 1h) data:
  A) Built-in feature importance (PredictionValuesChange, feature interactions)
  B) SHAP analysis (CatBoost native SHAP)
  C) Partial dependence analysis
  D) Prediction quality analysis (overlay, scatter, residuals, error by CO2 level)

Generates a DOCX academic report with all figures and quantitative analysis.

Usage:
    python -u scripts/run_catboost_interpretability.py
    python -u scripts/run_catboost_interpretability.py --horizons 1
"""

import argparse
import gc
import json
import logging
import sys
import time
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from src.data.datamodule import CO2DataModule
from src.data.pipeline import run_preprocessing_pipeline
from src.data.preprocessing import inverse_scale_target
from src.evaluation.metrics import compute_metrics, save_metrics
from src.models.catboost_model import CatBoostForecaster
from src.utils.config import load_config
from src.utils.seed import seed_everything

logger = logging.getLogger(__name__)
plt.style.use("seaborn-v0_8-whitegrid")
RESULTS_BASE = Path("results/catboost_interpretability")

C_PRIMARY = "#2196F3"
C_SECONDARY = "#FF5722"
C_ACCENT = "#4CAF50"
C_WARN = "#FFC107"
C_NEUTRAL = "#607D8B"


def load_interpretability_config(horizon: int) -> dict:
    config_files = [
        str(PROJECT_ROOT / "configs" / "training.yaml"),
        str(PROJECT_ROOT / "configs" / "data.yaml"),
        str(PROJECT_ROOT / "configs" / "catboost.yaml"),
        str(PROJECT_ROOT / "configs" / "experiments" / "preproc_D_enhanced_1h.yaml"),
    ]
    config = load_config(config_files)
    config["data"]["forecast_horizon_hours"] = horizon
    config["training"]["results_dir"] = str(RESULTS_BASE / f"h{horizon}" / "training_runs")
    return config


def get_feature_names(config: dict) -> list[str]:
    return config["data"]["feature_columns"] + [config["data"]["target_column"]]


def get_flat_feature_names(feature_names: list[str], lookback: int) -> list[str]:
    return [f"{feat}_lag{lag}" for lag in range(lookback) for feat in feature_names]


# ======================================================================
#  Training
# ======================================================================

def train_catboost(config, train_df, val_df, test_df, horizon):
    dm = CO2DataModule(config)
    dm.build_datasets(train_df, val_df, test_df)

    X_train = dm.train_dataset.X.numpy()
    y_train = dm.train_dataset.y.numpy()
    X_val = dm.val_dataset.X.numpy()
    y_val = dm.val_dataset.y.numpy()
    X_test = dm.test_dataset.X.numpy()
    y_test_scaled = dm.test_dataset.y.numpy()

    model = CatBoostForecaster(config)
    model.fit(X_train, y_train, X_val, y_val)

    y_pred_scaled = model.predict(X_test)
    y_pred = inverse_scale_target(y_pred_scaled, dm.target_scaler)
    y_true = inverse_scale_target(y_test_scaled, dm.target_scaler)

    metrics = compute_metrics(y_true.ravel(), y_pred.ravel())

    return model, dm, {
        "y_true": y_true, "y_pred": y_pred, "metrics": metrics,
        "X_test": X_test, "X_train": X_train,
    }


# ======================================================================
#  Section A: Built-in Feature Importance
# ======================================================================

def run_builtin_importance(model, config, horizon, output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [A] Extracting built-in feature importance...")

    feature_names = get_feature_names(config)
    n_features = len(feature_names)
    lookback = config["data"]["lookback_hours"] * config["data"]["samples_per_hour"]
    n_steps = len(model.model_.estimators_)

    # PredictionValuesChange importance
    all_importances = []
    for est in model.model_.estimators_:
        imp = est.get_feature_importance(type="PredictionValuesChange")
        all_importances.append(imp)

    all_importances = np.array(all_importances)
    avg_importance = all_importances.mean(axis=0)
    imp_2d = avg_importance.reshape(lookback, n_features)

    # Plot A1: Importance heatmap
    fig, ax = plt.subplots(figsize=(max(10, n_features * 0.5), 6))
    im = ax.imshow(imp_2d.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Lookback Step (hours ago)")
    ax.set_ylabel("Feature")
    ax.set_yticks(range(n_features))
    ax.set_yticklabels(feature_names, fontsize=8)
    tick_pos = np.linspace(0, lookback - 1, min(8, lookback)).astype(int)
    ax.set_xticks(tick_pos)
    ax.set_xticklabels([f"-{lookback - t}" for t in tick_pos])
    ax.set_title(f"CatBoost PredictionValuesChange Importance - {horizon}h")
    plt.colorbar(im, ax=ax, label="Importance")
    plt.tight_layout()
    plt.savefig(output_dir / f"importance_heatmap_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot A2: Per-feature importance
    feat_imp = imp_2d.sum(axis=0)
    feat_imp_pct = feat_imp / feat_imp.sum() * 100 if feat_imp.sum() > 0 else feat_imp
    sort_idx = np.argsort(feat_imp_pct)

    fig, ax = plt.subplots(figsize=(8, max(4, n_features * 0.35)))
    ax.barh(range(len(sort_idx)), feat_imp_pct[sort_idx], color=C_ACCENT)
    ax.set_yticks(range(len(sort_idx)))
    ax.set_yticklabels([feature_names[i] for i in sort_idx], fontsize=9)
    ax.set_xlabel("Importance (%)")
    ax.set_title(f"CatBoost Feature Importance - {horizon}h")
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"feature_importance_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot A3: Per-lag importance
    lag_imp = imp_2d.sum(axis=1)
    lag_imp_pct = lag_imp / lag_imp.sum() * 100 if lag_imp.sum() > 0 else lag_imp

    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(range(-lookback, 0), lag_imp_pct, color=C_PRIMARY, alpha=0.8)
    ax.set_xlabel("Hours Ago")
    ax.set_ylabel("Importance (%)")
    ax.set_title(f"CatBoost Temporal Importance - {horizon}h")
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"temporal_importance_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot A4: Feature interactions
    try:
        sub_model = model.model_.estimators_[0]
        interactions = sub_model.get_feature_importance(type="Interaction")
        if interactions is not None and len(interactions) > 0:
            interactions = np.array(interactions)
            # interactions shape: (n_pairs, 3) = [feat_i, feat_j, strength]
            n_show = min(20, len(interactions))
            top_inter = interactions[np.argsort(-interactions[:, 2])[:n_show]]

            flat_names = get_flat_feature_names(feature_names, lookback)

            fig, ax = plt.subplots(figsize=(10, max(4, n_show * 0.3)))
            labels = []
            for row in top_inter:
                fi, fj = int(row[0]), int(row[1])
                ni = feature_names[fi % n_features] if fi < len(flat_names) else f"f{fi}"
                nj = feature_names[fj % n_features] if fj < len(flat_names) else f"f{fj}"
                labels.append(f"{ni} x {nj}")
            ax.barh(range(n_show), top_inter[:, 2], color=C_WARN)
            ax.set_yticks(range(n_show))
            ax.set_yticklabels(labels, fontsize=8)
            ax.set_xlabel("Interaction Strength")
            ax.set_title(f"Top Feature Interactions - {horizon}h")
            ax.grid(axis="x", alpha=0.3)
            plt.tight_layout()
            plt.savefig(output_dir / f"interactions_h{horizon}.png", dpi=150, bbox_inches="tight")
            plt.close(fig)
    except Exception as e:
        print(f"    Feature interaction analysis failed: {e}")

    imp_df = pd.DataFrame(imp_2d, columns=feature_names)
    imp_df.to_csv(output_dir / f"importance_h{horizon}.csv", index=False)

    ranking = [(feature_names[i], float(feat_imp_pct[i])) for i in np.argsort(feat_imp_pct)[::-1]]
    print(f"  [A] Top 5: {', '.join(f'{n}({v:.1f}%)' for n, v in ranking[:5])}")

    return {"imp_2d": imp_2d, "feat_imp_pct": feat_imp_pct, "ranking": ranking}


# ======================================================================
#  Section B: SHAP Analysis (CatBoost native)
# ======================================================================

def run_shap_analysis(model, X_test, config, horizon, output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [B] Computing CatBoost native SHAP values...")

    feature_names = get_feature_names(config)
    n_features = len(feature_names)
    lookback = config["data"]["lookback_hours"] * config["data"]["samples_per_hour"]

    X_test_flat = X_test.reshape(X_test.shape[0], -1).astype(np.float64)
    n_shap = min(500, X_test_flat.shape[0])
    X_shap = X_test_flat[:n_shap]

    try:
        from catboost import Pool
        sub_model = model.model_.estimators_[0]
        pool = Pool(data=X_shap)
        shap_raw = sub_model.get_feature_importance(type="ShapValues", data=pool)
        # Shape: (n_samples, n_features+1), last col is base value
        shap_values = shap_raw[:, :-1]
        base_value = shap_raw[0, -1]
    except Exception as e:
        print(f"  [B] CatBoost SHAP failed: {e}. Trying shap library...")
        try:
            import shap
            sub_model = model.model_.estimators_[0]
            explainer = shap.TreeExplainer(sub_model)
            shap_values = explainer.shap_values(X_shap)
            base_value = explainer.expected_value
            if isinstance(base_value, np.ndarray):
                base_value = float(base_value[0])
        except Exception as e2:
            print(f"  [B] SHAP analysis failed completely: {e2}")
            return {}

    shap_3d = shap_values.reshape(n_shap, lookback, n_features)
    mean_abs_shap_2d = np.abs(shap_3d).mean(axis=0)

    # Plot B1: SHAP feature importance
    feat_shap = mean_abs_shap_2d.sum(axis=0)
    feat_shap_pct = feat_shap / feat_shap.sum() * 100 if feat_shap.sum() > 0 else feat_shap
    sort_idx = np.argsort(feat_shap_pct)

    fig, ax = plt.subplots(figsize=(8, max(4, n_features * 0.35)))
    ax.barh(range(len(sort_idx)), feat_shap_pct[sort_idx], color=C_SECONDARY)
    ax.set_yticks(range(len(sort_idx)))
    ax.set_yticklabels([feature_names[i] for i in sort_idx], fontsize=9)
    ax.set_xlabel("Mean |SHAP| (%)")
    ax.set_title(f"CatBoost SHAP Feature Importance - {horizon}h")
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"shap_feature_importance_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot B2: SHAP temporal heatmap
    fig, ax = plt.subplots(figsize=(max(10, n_features * 0.5), 6))
    im = ax.imshow(mean_abs_shap_2d.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Lookback Step (hours ago)")
    ax.set_ylabel("Feature")
    ax.set_yticks(range(n_features))
    ax.set_yticklabels(feature_names, fontsize=8)
    tick_pos = np.linspace(0, lookback - 1, min(8, lookback)).astype(int)
    ax.set_xticks(tick_pos)
    ax.set_xticklabels([f"-{lookback - t}" for t in tick_pos])
    ax.set_title(f"CatBoost SHAP Temporal Heatmap - {horizon}h")
    plt.colorbar(im, ax=ax, label="Mean |SHAP|")
    plt.tight_layout()
    plt.savefig(output_dir / f"shap_temporal_heatmap_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot B3: SHAP temporal profile
    lag_shap = mean_abs_shap_2d.sum(axis=1)
    lag_shap_pct = lag_shap / lag_shap.sum() * 100 if lag_shap.sum() > 0 else lag_shap

    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(range(-lookback, 0), lag_shap_pct, color=C_PRIMARY, alpha=0.8)
    ax.set_xlabel("Hours Ago")
    ax.set_ylabel("Mean |SHAP| (%)")
    ax.set_title(f"CatBoost SHAP Temporal Profile - {horizon}h")
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"shap_temporal_profile_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot B4: Waterfall for 3 samples
    flat_names = get_flat_feature_names(feature_names, lookback)
    sample_indices = [n_shap // 2, 0, n_shap - 1]
    sample_labels = ["Median sample", "First sample", "Last sample"]

    for si, (idx, label) in enumerate(zip(sample_indices, sample_labels)):
        sv = shap_values[idx]
        top_k = 15
        abs_sv = np.abs(sv)
        top_feat_idx = np.argsort(abs_sv)[-top_k:][::-1]

        fig, ax = plt.subplots(figsize=(10, 6))
        vals = sv[top_feat_idx]
        names = [flat_names[i] if i < len(flat_names) else f"f{i}" for i in top_feat_idx]
        colors = [C_SECONDARY if v > 0 else C_PRIMARY for v in vals]
        ax.barh(range(len(vals)), vals, color=colors)
        ax.set_yticks(range(len(vals)))
        ax.set_yticklabels(names, fontsize=8)
        ax.set_xlabel("SHAP Value")
        ax.set_title(f"CatBoost SHAP: {label} - {horizon}h")
        ax.axvline(0, color="black", linewidth=0.8)
        ax.grid(axis="x", alpha=0.3)
        plt.tight_layout()
        plt.savefig(output_dir / f"shap_waterfall_{si}_h{horizon}.png", dpi=150, bbox_inches="tight")
        plt.close(fig)

    shap_df = pd.DataFrame(mean_abs_shap_2d, columns=feature_names)
    shap_df.to_csv(output_dir / f"shap_importance_h{horizon}.csv", index=False)

    shap_ranking = [(feature_names[i], float(feat_shap_pct[i])) for i in np.argsort(feat_shap_pct)[::-1]]
    print(f"  [B] Top 5 SHAP: {', '.join(f'{n}({v:.1f}%)' for n, v in shap_ranking[:5])}")

    return {"feat_shap_pct": feat_shap_pct, "shap_ranking": shap_ranking}


# ======================================================================
#  Section C: Partial Dependence
# ======================================================================

def run_partial_dependence(model, X_test, config, horizon, output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [C] Computing partial dependence plots...")

    from sklearn.inspection import partial_dependence

    feature_names = get_feature_names(config)
    n_features = len(feature_names)
    lookback = config["data"]["lookback_hours"] * config["data"]["samples_per_hour"]

    X_test_flat = X_test.reshape(X_test.shape[0], -1)
    n_pdp = min(300, X_test_flat.shape[0])
    X_pdp = X_test_flat[:n_pdp]

    sub_model = model.model_.estimators_[0]
    recent_start = (lookback - 1) * n_features
    recent_imp = sub_model.get_feature_importance(type="PredictionValuesChange")
    recent_imp_slice = recent_imp[recent_start:recent_start + n_features]
    top6_feat = np.argsort(recent_imp_slice)[-6:][::-1]
    top6_flat = [recent_start + fi for fi in top6_feat]

    fig, axes = plt.subplots(2, 3, figsize=(15, 9))
    axes_flat = axes.ravel()

    for idx, (feat_flat_idx, feat_orig_idx) in enumerate(zip(top6_flat, top6_feat)):
        ax = axes_flat[idx]
        try:
            pdp_results = partial_dependence(
                sub_model, X_pdp, features=[feat_flat_idx],
                kind="average", grid_resolution=50,
            )
            ax.plot(pdp_results["grid_values"][0], pdp_results["average"][0],
                    color=C_PRIMARY, linewidth=2)
            ax.set_xlabel(f"{feature_names[feat_orig_idx]}", fontsize=9)
            ax.set_ylabel("Partial Dependence")
            ax.set_title(feature_names[feat_orig_idx], fontsize=10)
            ax.grid(alpha=0.3)
        except Exception as e:
            ax.text(0.5, 0.5, f"Error: {str(e)[:30]}", ha="center", va="center",
                    transform=ax.transAxes, fontsize=8)

    fig.suptitle(f"CatBoost Partial Dependence (Most Recent Lag) - {horizon}h", fontsize=13)
    plt.tight_layout()
    plt.savefig(output_dir / f"pdp_top6_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    print(f"  [C] Partial dependence analysis complete")


# ======================================================================
#  Section D: Prediction Quality
# ======================================================================

def run_prediction_analysis(results, test_df, config, horizon, output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [D] Generating prediction analysis plots...")

    y_true = results["y_true"].ravel()
    y_pred = results["y_pred"].ravel()
    residuals = y_true - y_pred
    metrics = results["metrics"]
    n_show = min(500, len(y_true))

    # Predictions overlay
    fig, ax = plt.subplots(figsize=(14, 5))
    ax.plot(y_true[:n_show], label="Actual", color=C_PRIMARY, linewidth=1.0, alpha=0.8)
    ax.plot(y_pred[:n_show], label="Predicted", color=C_SECONDARY, linewidth=1.0, alpha=0.8)
    ax.set_xlabel("Sample Index")
    ax.set_ylabel("CO2 (ppm)")
    ax.set_title(f"CatBoost Predictions vs Actual - {horizon}h")
    ax.legend()
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"predictions_overlay_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Scatter
    fig, ax = plt.subplots(figsize=(7, 7))
    ax.scatter(y_true, y_pred, alpha=0.3, s=5, color=C_PRIMARY)
    vmin, vmax = min(y_true.min(), y_pred.min()), max(y_true.max(), y_pred.max())
    ax.plot([vmin, vmax], [vmin, vmax], "r--", linewidth=1.5, label="Perfect")
    ax.text(0.05, 0.95,
            f"R2 = {metrics['r2']:.4f}\nRMSE = {metrics['rmse']:.2f}\nMAE = {metrics['mae']:.2f}",
            transform=ax.transAxes, fontsize=11, va="top",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8))
    ax.set_xlabel("Actual CO2 (ppm)")
    ax.set_ylabel("Predicted CO2 (ppm)")
    ax.set_title(f"CatBoost Scatter - {horizon}h")
    ax.legend()
    plt.tight_layout()
    plt.savefig(output_dir / f"scatter_r2_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # 4-panel residual analysis
    datetime_col = config["data"].get("datetime_column", "datetime")
    test_dates = pd.DatetimeIndex(test_df[datetime_col]) if datetime_col in test_df.columns else None

    fig, axes = plt.subplots(2, 2, figsize=(14, 10))

    ax = axes[0, 0]
    ax.hist(residuals, bins=50, color=C_PRIMARY, alpha=0.7, edgecolor="white", density=True)
    ax.axvline(0, color="red", linestyle="--", linewidth=1.5)
    ax.set_xlabel("Residual (ppm)")
    ax.set_ylabel("Density")
    ax.set_title("Residual Distribution")

    ax = axes[0, 1]
    ax.scatter(np.arange(len(residuals)), residuals, alpha=0.3, s=3, color=C_SECONDARY)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xlabel("Sample Index")
    ax.set_ylabel("Residual (ppm)")
    ax.set_title("Residuals Over Time")

    ax = axes[1, 0]
    if test_dates is not None and len(test_dates) >= len(residuals):
        hours = test_dates[-len(residuals):].hour
        pd.DataFrame({"hour": hours, "residual": residuals}).boxplot(column="residual", by="hour", ax=ax)
        ax.set_title("Residuals by Hour")
        plt.suptitle("")
    else:
        ax.set_title("Residuals by Hour (N/A)")

    ax = axes[1, 1]
    if test_dates is not None and len(test_dates) >= len(residuals):
        dow = test_dates[-len(residuals):].dayofweek
        pd.DataFrame({"day": dow, "residual": residuals}).boxplot(column="residual", by="day", ax=ax)
        ax.set_xticklabels(["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"])
        ax.set_title("Residuals by Day")
        plt.suptitle("")
    else:
        ax.set_title("Residuals by Day (N/A)")

    plt.suptitle("")
    plt.tight_layout()
    plt.savefig(output_dir / f"residual_analysis_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Error by CO2 level
    bins = [0, 500, 1000, np.inf]
    labels_co2 = ["Low (<500)", "Medium (500-1000)", "High (>1000)"]
    categories = pd.cut(y_true, bins=bins, labels=labels_co2)

    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    ax = axes[0]
    pd.DataFrame({"category": categories, "abs_error": np.abs(residuals)}).boxplot(
        column="abs_error", by="category", ax=ax)
    ax.set_title(f"Error by CO2 Level - {horizon}h")
    plt.suptitle("")

    ax = axes[1]
    ax.axis("off")
    rows_data = []
    for label in labels_co2:
        mask = categories == label
        if mask.sum() > 0:
            bin_res = residuals[mask]
            rows_data.append([label, f"{mask.sum()}", f"{np.sqrt(np.mean(bin_res**2)):.2f}",
                              f"{np.mean(np.abs(bin_res)):.2f}", f"{np.mean(bin_res):.2f}"])
        else:
            rows_data.append([label, "0", "N/A", "N/A", "N/A"])
    tbl = ax.table(cellText=rows_data, colLabels=["CO2 Level", "N", "RMSE", "MAE", "Bias"],
                   loc="center", cellLoc="center")
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.0, 1.5)
    plt.tight_layout()
    plt.savefig(output_dir / f"error_by_co2_level_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Per-step RMSE
    y_true_2d, y_pred_2d = results["y_true"], results["y_pred"]
    if y_true_2d.ndim == 2 and y_true_2d.shape[1] > 1:
        step_rmse = np.sqrt(np.mean((y_true_2d - y_pred_2d)**2, axis=0))
        fig, ax = plt.subplots(figsize=(10, 5))
        ax.plot(range(1, len(step_rmse)+1), step_rmse, marker="o", color=C_PRIMARY, markersize=4)
        ax.set_xlabel("Forecast Step")
        ax.set_ylabel("RMSE (ppm)")
        ax.set_title(f"CatBoost Per-Step RMSE - {horizon}h")
        ax.grid(alpha=0.3)
        plt.tight_layout()
        plt.savefig(output_dir / f"perstep_rmse_h{horizon}.png", dpi=150, bbox_inches="tight")
        plt.close(fig)

    print(f"  [D] Prediction analysis complete")


# ======================================================================
#  DOCX Report
# ======================================================================

def generate_docx_report(all_results, all_builtin, all_shap, horizons, output_dir):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print("  Generating DOCX report...")
    doc = Document()

    title = doc.add_heading(
        "CatBoost Interpretability Study: Ordered Boosting and Oblivious Tree Analysis "
        "for Indoor CO2 Forecasting", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Preprocessing Variant: preproc_D (Enhanced 1h)\n"
        f"Horizons: {', '.join(str(h) + 'h' for h in horizons)}\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.italic = True
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    for item in ["1. Executive Summary", "2. Built-in Feature Importance",
                  "3. SHAP Analysis", "4. Partial Dependence",
                  "5. Prediction Quality", "6. Avenues of Improvement"]:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    def add_figure(fig_path, caption, width=6.0):
        if fig_path.exists():
            doc.add_picture(str(fig_path), width=Inches(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.font.size = Pt(9)
            r.font.italic = True
        else:
            doc.add_paragraph(f"[Figure not found: {fig_path.name}]")

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This report presents a comprehensive interpretability study of a CatBoost "
        "gradient-boosted tree ensemble for indoor CO2 forecasting. CatBoost features "
        "ordered boosting (reducing target leakage) and oblivious decision trees (symmetric "
        "splits at each depth level). The study covers built-in importance, native SHAP values, "
        "feature interactions, partial dependence, and prediction quality analysis.")

    doc.add_heading("Performance Summary", level=2)
    table = doc.add_table(rows=1 + len(horizons), cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Horizon", "RMSE", "MAE", "R2", "MAPE"]):
        table.rows[0].cells[i].text = h
    for ri, h in enumerate(horizons):
        m = all_results[h]["metrics"]
        table.rows[ri+1].cells[0].text = f"{h}h"
        table.rows[ri+1].cells[1].text = f"{m['rmse']:.2f}"
        table.rows[ri+1].cells[2].text = f"{m['mae']:.2f}"
        table.rows[ri+1].cells[3].text = f"{m['r2']:.4f}"
        table.rows[ri+1].cells[4].text = f"{m['mape']:.2f}"
    doc.add_page_break()

    # 2. Built-in importance
    doc.add_heading("2. Built-in Feature Importance", level=1)
    doc.add_paragraph(
        "CatBoost provides PredictionValuesChange importance, which measures how much "
        "each feature changes the model's predictions on average. Additionally, CatBoost "
        "detects pairwise feature interactions, revealing which feature combinations "
        "synergistically influence predictions.")
    for h in horizons:
        doc.add_heading(f"Horizon: {h}h", level=2)
        d = output_dir / f"h{h}" / "builtin_importance"
        add_figure(d / f"importance_heatmap_h{h}.png", f"PredictionValuesChange importance heatmap at {h}h.")
        add_figure(d / f"feature_importance_h{h}.png", f"Per-feature importance (aggregated across lags) at {h}h.")
        add_figure(d / f"temporal_importance_h{h}.png", f"Temporal importance profile at {h}h.")
        add_figure(d / f"interactions_h{h}.png", f"Top pairwise feature interactions detected by CatBoost.")
    doc.add_page_break()

    # 3. SHAP
    doc.add_heading("3. SHAP Analysis", level=1)
    doc.add_paragraph(
        "CatBoost provides native SHAP value computation, offering exact Shapley value "
        "decomposition for each prediction. The SHAP values are reshaped from the flattened "
        "feature space to reveal temporal and feature-level attribution patterns.")
    for h in horizons:
        doc.add_heading(f"Horizon: {h}h", level=2)
        d = output_dir / f"h{h}" / "shap_analysis"
        add_figure(d / f"shap_feature_importance_h{h}.png", f"SHAP feature importance at {h}h.")
        add_figure(d / f"shap_temporal_heatmap_h{h}.png", f"SHAP temporal heatmap at {h}h.")
        add_figure(d / f"shap_temporal_profile_h{h}.png", f"SHAP temporal profile at {h}h.")
        for si in range(3):
            add_figure(d / f"shap_waterfall_{si}_h{h}.png", f"SHAP waterfall plot for sample {si}.")
    doc.add_page_break()

    # 4. PDP
    doc.add_heading("4. Partial Dependence Analysis", level=1)
    for h in horizons:
        d = output_dir / f"h{h}" / "partial_dependence"
        add_figure(d / f"pdp_top6_h{h}.png", f"1D partial dependence for top features at {h}h.")
    doc.add_page_break()

    # 5. Predictions
    doc.add_heading("5. Prediction Quality Analysis", level=1)
    for h in horizons:
        doc.add_heading(f"Horizon: {h}h", level=2)
        d = output_dir / f"h{h}" / "prediction_analysis"
        add_figure(d / f"predictions_overlay_h{h}.png", f"Predictions overlay at {h}h.")
        add_figure(d / f"scatter_r2_h{h}.png", f"Scatter plot at {h}h.")
        add_figure(d / f"residual_analysis_h{h}.png", f"Residual analysis at {h}h.")
        add_figure(d / f"error_by_co2_level_h{h}.png", f"Error by CO2 level at {h}h.")
        add_figure(d / f"perstep_rmse_h{h}.png", f"Per-step RMSE at {h}h.")
    doc.add_page_break()

    # 6. Avenues
    doc.add_heading("6. Avenues of Improvement", level=1)
    for title_t, desc in [
        ("Ordered Boosting Tuning",
         "CatBoost's ordered boosting reduces overfitting but adds computational cost. "
         "Investigating different permutation counts and their impact on generalization "
         "could optimize the bias-variance tradeoff."),
        ("Symmetric Tree Depth Optimization",
         "Oblivious trees with depth 6 create 64-leaf trees. Shallower trees (depth 4-5) "
         "may reduce overfitting while deeper trees (7-8) could capture more complex patterns."),
        ("Feature Interaction Engineering",
         "The detected feature interactions suggest explicit engineering of interaction "
         "terms could improve model efficiency."),
        ("Temporal Feature Reduction",
         "SHAP analysis shows most importance concentrated in recent lags. Reducing the "
         "lookback window or using exponentially-weighted features could improve efficiency."),
        ("Comparison with LightGBM",
         "A third gradient boosting variant (LightGBM) with leaf-wise growth could provide "
         "a useful comparison point for understanding the impact of tree structure choices."),
    ]:
        doc.add_heading(title_t, level=2)
        doc.add_paragraph(desc)

    report_path = output_dir / "catboost_interpretability_report.docx"
    doc.save(str(report_path))
    print(f"  Report saved to: {report_path}")
    return report_path


# ======================================================================
#  Main
# ======================================================================

def main():
    logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
    parser = argparse.ArgumentParser(description="CatBoost Interpretability Study")
    parser.add_argument("--horizons", nargs="+", type=int, default=[1, 24])
    parser.add_argument("--epochs", type=int, default=None)
    args = parser.parse_args()

    print(f"\n{'='*70}")
    print(f"  CATBOOST INTERPRETABILITY STUDY (DEEP ANALYSIS)")
    print(f"  Variant: preproc_D (Enhanced 1h)")
    print(f"  Horizons: {args.horizons}")
    print(f"{'='*70}\n")

    pipeline_config = load_interpretability_config(horizon=1)
    seed_everything(pipeline_config["training"]["seed"])

    raw_dir = Path(pipeline_config["data"].get("raw_dir", "data/raw"))
    print("  Loading preprocessing pipeline (preproc_D)...")
    train_df, val_df, test_df = run_preprocessing_pipeline(raw_dir=raw_dir, variant_config=pipeline_config)
    print(f"  Pipeline loaded: train={len(train_df)}, val={len(val_df)}, test={len(test_df)}")

    all_results, all_builtin, all_shap = {}, {}, {}

    for horizon in args.horizons:
        print(f"\n{'-'*60}")
        print(f"  HORIZON: {horizon}h")
        print(f"{'-'*60}\n")

        config = load_interpretability_config(horizon=horizon)
        seed_everything(config["training"]["seed"])
        output_dir = RESULTS_BASE / f"h{horizon}"
        output_dir.mkdir(parents=True, exist_ok=True)

        t0 = time.time()
        model, dm, results = train_catboost(config, train_df.copy(), val_df.copy(), test_df.copy(), horizon)
        print(f"  Training completed in {time.time()-t0:.1f}s")

        m = results["metrics"]
        all_results[horizon] = results
        print(f"  RMSE={m['rmse']:.2f}  MAE={m['mae']:.2f}  R2={m['r2']:.4f}  MAPE={m['mape']:.2f}%")

        all_builtin[horizon] = run_builtin_importance(model, config, horizon, output_dir / "builtin_importance")
        all_shap[horizon] = run_shap_analysis(model, results["X_test"], config, horizon, output_dir / "shap_analysis")
        run_partial_dependence(model, results["X_test"], config, horizon, output_dir / "partial_dependence")
        run_prediction_analysis(results, test_df, config, horizon, output_dir / "prediction_analysis")

        save_metrics(m, f"CatBoost_h{horizon}", output_dir / "metrics.json",
                     experiment_info={"name": "catboost_interpretability"})
        np.savez(output_dir / "predictions.npz", y_true=results["y_true"], y_pred=results["y_pred"])

        del model, dm
        gc.collect()
        print(f"  Memory freed\n")

    report_path = generate_docx_report(all_results, all_builtin, all_shap, args.horizons, RESULTS_BASE)

    print(f"\n{'='*70}")
    print(f"  CATBOOST INTERPRETABILITY STUDY COMPLETE")
    print(f"  Results: {RESULTS_BASE}")
    print(f"  Report: {report_path}")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
