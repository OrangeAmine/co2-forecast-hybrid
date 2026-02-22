"""HMM-LSTM Interpretability Study: Regime-Aware Sequential Modeling Analysis.

Performs comprehensive interpretability analysis on the HMM-LSTM hybrid model
using preproc_D (Enhanced 1h) data:
  A) HMM Regime Analysis (unique to this model: transition matrix, emission
     distributions, state-colored time series, duration histograms, entropy)
  B) Gradient-based feature attribution (22-dim input: 19 original + 3 HMM)
  C) Hidden state structural analysis (PCA, clustering, HMM state crosstab)
  D) Temporal pattern analysis (FFT, autocorrelation, rolling RMSE)
  E) HMM-LSTM-specific analysis (regime-conditioned error, HMM channel
     permutation importance, prediction quality)

Generates a DOCX academic report with all figures and quantitative analysis.

Usage:
    python -u scripts/run_hmm_lstm_interpretability.py
    python -u scripts/run_hmm_lstm_interpretability.py --horizons 1
    python -u scripts/run_hmm_lstm_interpretability.py --horizons 1 24 --epochs 30
"""

import argparse
import copy
import gc
import json
import logging
import sys
import time
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import Ellipse
import numpy as np
import pandas as pd
import torch
import torch.nn as nn
from scipy.fft import fft, fftfreq
from scipy.stats import norm as scipy_norm
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import pytorch_lightning as pl
from pytorch_lightning.callbacks import (
    EarlyStopping,
    LearningRateMonitor,
    ModelCheckpoint,
    TQDMProgressBar,
)
from pytorch_lightning.loggers import TensorBoardLogger

from src.data.datamodule import CO2DataModule
from src.data.pipeline import run_preprocessing_pipeline
from src.data.preprocessing import inverse_scale_target
from src.evaluation.metrics import compute_metrics, save_metrics
from src.models.hmm_lstm import HMMRegimeDetector, HMMLSTMForecaster
from src.utils.config import load_config
from src.utils.seed import seed_everything

logger = logging.getLogger(__name__)

plt.style.use("seaborn-v0_8-whitegrid")

RESULTS_BASE = Path("results/hmm_lstm_interpretability")

# Color palette
C_PRIMARY = "#2196F3"
C_SECONDARY = "#FF5722"
C_ACCENT = "#4CAF50"
C_WARN = "#FFC107"
C_NEUTRAL = "#607D8B"

# HMM state colors
STATE_COLORS = ["#2196F3", "#FF5722", "#4CAF50"]
STATE_LABELS_DEFAULT = ["State 0", "State 1", "State 2"]


# ======================================================================
#  Configuration
# ======================================================================

def load_interpretability_config(
    horizon: int,
    epochs_override: int | None = None,
) -> dict:
    """Load merged config for HMM-LSTM + preproc_D + specified horizon."""
    config_files = [
        str(PROJECT_ROOT / "configs" / "training.yaml"),
        str(PROJECT_ROOT / "configs" / "data.yaml"),
        str(PROJECT_ROOT / "configs" / "hmm_lstm.yaml"),
        str(PROJECT_ROOT / "configs" / "experiments" / "preproc_D_enhanced_1h.yaml"),
    ]
    config = load_config(config_files)
    config["data"]["forecast_horizon_hours"] = horizon
    if epochs_override is not None:
        config["training"]["max_epochs"] = epochs_override
    config["training"]["results_dir"] = str(
        RESULTS_BASE / f"h{horizon}" / "training_runs"
    )
    return config


# ======================================================================
#  HMM-LSTM Training and Prediction
# ======================================================================

def train_hmm_lstm(
    config: dict,
    train_df: pd.DataFrame,
    val_df: pd.DataFrame,
    test_df: pd.DataFrame,
    horizon: int,
) -> tuple:
    """Train HMM-LSTM and return best model, datamodule, detector, predictions.

    Returns:
        (best_model, dm, detector, y_true, y_pred, train_df, val_df, test_df)
        where train/val/test_df are HMM-augmented.
    """
    cfg = copy.deepcopy(config)

    # Make copies to avoid mutating originals
    tr_df = train_df.copy()
    va_df = val_df.copy()
    te_df = test_df.copy()

    # Stage 1: Fit HMM on UNSCALED training data
    hmm_cfg = cfg["model"]
    detector = HMMRegimeDetector(
        n_states=hmm_cfg.get("hmm_n_states", 3),
        covariance_type=hmm_cfg.get("hmm_covariance_type", "full"),
        n_iter=hmm_cfg.get("hmm_n_iter", 100),
        hmm_features=hmm_cfg.get("hmm_features", ["CO2", "Noise", "TemperatureExt"]),
    )
    detector.fit(tr_df)

    # Augment all splits with HMM posterior probabilities
    for df in [tr_df, va_df, te_df]:
        probs = detector.predict_proba(df)
        for k in range(detector.n_states):
            df[f"hmm_state_{k}"] = probs[:, k]

    # Update config to include HMM features
    hmm_state_cols = [f"hmm_state_{k}" for k in range(detector.n_states)]
    cfg["data"]["feature_columns"] = cfg["data"]["feature_columns"] + hmm_state_cols

    # Build DataModule with augmented data
    dm = CO2DataModule.from_dataframes(tr_df, va_df, te_df, cfg)

    # Compute input_size: features + target (CO2) already counted in the data
    # The from_dataframes path builds sequences with all feature_columns + target
    # So input_size = len(feature_columns) + 1 (target)
    n_features = len(cfg["data"]["feature_columns"]) + 1  # +1 for target column
    # HMM states are already in feature_columns now, so we must NOT double-add
    # The HMMLSTMForecaster constructor adds hmm_n_states by default,
    # but since we already include them in feature_columns, pass input_size directly.
    model = HMMLSTMForecaster(cfg, input_size=n_features, output_size=None)

    # Training
    training_cfg = cfg["training"]
    results_dir = Path(training_cfg["results_dir"])
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    run_dir = results_dir / f"hmm_lstm_interp_h{horizon}_{timestamp}"

    trainer = pl.Trainer(
        max_epochs=training_cfg["max_epochs"],
        accelerator=training_cfg.get("accelerator", "auto"),
        devices=training_cfg.get("devices", 1),
        precision=training_cfg.get("precision", 32),
        callbacks=[
            EarlyStopping(
                monitor="val_loss",
                patience=training_cfg["patience"],
                mode="min",
                verbose=True,
            ),
            ModelCheckpoint(
                dirpath=run_dir / "checkpoints",
                filename="best-{epoch}-{val_loss:.4f}",
                monitor="val_loss",
                mode="min",
                save_top_k=1,
            ),
            LearningRateMonitor(logging_interval="epoch"),
            TQDMProgressBar(refresh_rate=20),
        ],
        gradient_clip_val=training_cfg.get("gradient_clip_val", 1.0),
        log_every_n_steps=training_cfg.get("log_every_n_steps", 10),
        enable_progress_bar=True,
        logger=TensorBoardLogger(save_dir=str(run_dir), name="tb_logs"),
        deterministic=True,
    )

    trainer.fit(model, datamodule=dm)

    best_ckpt = trainer.checkpoint_callback.best_model_path  # type: ignore[union-attr]
    best_model = HMMLSTMForecaster.load_from_checkpoint(
        best_ckpt, config=cfg, input_size=n_features,
    )

    # Predictions
    best_model.eval()
    predictions = trainer.predict(best_model, dm.test_dataloader())
    y_pred_scaled = torch.cat(predictions).numpy()
    y_pred = inverse_scale_target(y_pred_scaled, dm.target_scaler)

    y_true_scaled = dm.test_dataset.y.numpy()
    y_true = inverse_scale_target(y_true_scaled, dm.target_scaler)

    return best_model, dm, detector, y_true, y_pred, tr_df, va_df, te_df


# ======================================================================
#  Section A: HMM Regime Analysis
# ======================================================================

def label_states_by_co2(detector: HMMRegimeDetector) -> list:
    """Create human-readable labels for HMM states based on CO2 means."""
    means = detector.hmm.means_
    # CO2 is typically the first HMM feature
    co2_idx = detector.hmm_features.index("CO2") if "CO2" in detector.hmm_features else 0
    co2_means = means[:, co2_idx]
    # Sort states by CO2 mean to assign labels
    order = np.argsort(co2_means)
    labels = [""] * len(order)
    descriptors = ["Low CO2 (baseline)", "Medium CO2 (transition)", "High CO2 (occupied)"]
    if len(order) > 3:
        descriptors = [f"State {i}" for i in range(len(order))]
    for rank, state_idx in enumerate(order):
        if rank < len(descriptors):
            labels[state_idx] = descriptors[rank]
        else:
            labels[state_idx] = f"State {state_idx}"
    return labels


def run_hmm_regime_analysis(
    detector: HMMRegimeDetector,
    test_df: pd.DataFrame,
    train_df: pd.DataFrame,
    y_true: np.ndarray,
    y_pred: np.ndarray,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Comprehensive HMM regime analysis - the unique interpretability section.

    Returns:
        Dictionary with analysis results for DOCX report.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [A] HMM Regime Analysis...")

    n_states = detector.n_states
    hmm = detector.hmm
    hmm_features = detector.hmm_features

    # Extract HMM parameters
    means = hmm.means_          # (n_states, n_features)
    covars = hmm.covars_        # (n_states, n_features, n_features) for "full"
    transmat = hmm.transmat_    # (n_states, n_states)

    # State labels
    state_labels = label_states_by_co2(detector)

    # Viterbi states on test data
    test_states = detector.predict_states(test_df)
    test_probs = detector.predict_proba(test_df)

    results = {
        "state_labels": state_labels,
        "means": means.tolist(),
        "transmat": transmat.tolist(),
    }

    # ---- A1: Transition Matrix Heatmap ----
    fig, ax = plt.subplots(figsize=(7, 6))
    im = ax.imshow(transmat, cmap="YlOrRd", vmin=0, vmax=1)
    for i in range(n_states):
        for j in range(n_states):
            color = "white" if transmat[i, j] > 0.5 else "black"
            ax.text(j, i, f"{transmat[i, j]:.3f}", ha="center", va="center",
                    fontsize=12, color=color, fontweight="bold")
    ax.set_xticks(range(n_states))
    ax.set_yticks(range(n_states))
    ax.set_xticklabels([f"To S{i}" for i in range(n_states)], fontsize=10)
    ax.set_yticklabels([f"From S{i}" for i in range(n_states)], fontsize=10)
    ax.set_title(f"HMM Transition Matrix - {horizon}h", fontsize=13)
    plt.colorbar(im, ax=ax, label="Transition Probability")
    # Add state labels as annotation
    for i in range(n_states):
        ax.annotate(state_labels[i], xy=(n_states + 0.3, i), fontsize=8,
                    va="center", ha="left", annotation_clip=False)
    plt.tight_layout()
    plt.savefig(output_dir / f"A1_transition_matrix_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- A2: State Emission Distributions ----
    fig, axes = plt.subplots(1, len(hmm_features), figsize=(6 * len(hmm_features), 5))
    if len(hmm_features) == 1:
        axes = [axes]

    for f_idx, feat_name in enumerate(hmm_features):
        ax = axes[f_idx]
        # Get range from training data
        feat_vals = train_df[feat_name].values
        x_range = np.linspace(
            feat_vals.min() - 0.1 * (feat_vals.max() - feat_vals.min()),
            feat_vals.max() + 0.1 * (feat_vals.max() - feat_vals.min()),
            300,
        )

        for s in range(n_states):
            mu = means[s, f_idx]
            sigma = np.sqrt(covars[s, f_idx, f_idx])
            pdf = scipy_norm.pdf(x_range, mu, sigma)
            ax.plot(x_range, pdf, color=STATE_COLORS[s % len(STATE_COLORS)],
                    linewidth=2, label=f"S{s}: {state_labels[s]}")
            ax.fill_between(x_range, pdf, alpha=0.15,
                            color=STATE_COLORS[s % len(STATE_COLORS)])
            ax.axvline(mu, color=STATE_COLORS[s % len(STATE_COLORS)],
                       linestyle="--", alpha=0.5, linewidth=1)

        ax.set_xlabel(feat_name, fontsize=11)
        ax.set_ylabel("Density", fontsize=11)
        ax.set_title(f"{feat_name} Emission Distribution", fontsize=11)
        ax.legend(fontsize=8)
        ax.grid(alpha=0.3)

    fig.suptitle(f"HMM State Emission Distributions - {horizon}h", fontsize=13)
    plt.tight_layout()
    plt.savefig(output_dir / f"A2_emission_distributions_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- A3: State Means Table ----
    fig, ax = plt.subplots(figsize=(10, 3))
    ax.axis("off")
    rows_data = []
    for s in range(n_states):
        row = [f"S{s}", state_labels[s]]
        for f_idx in range(len(hmm_features)):
            row.append(f"{means[s, f_idx]:.1f}")
        rows_data.append(row)

    col_labels = ["State", "Label"] + hmm_features
    tbl = ax.table(
        cellText=rows_data,
        colLabels=col_labels,
        loc="center",
        cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(11)
    tbl.scale(1.2, 1.8)
    ax.set_title(f"HMM State Means - {horizon}h", fontsize=13, pad=25)
    plt.tight_layout()
    plt.savefig(output_dir / f"A3_state_means_table_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- A4: CO2 vs Noise scatter colored by Viterbi state + 95% confidence ellipses ----
    co2_idx = hmm_features.index("CO2") if "CO2" in hmm_features else 0
    noise_idx = hmm_features.index("Noise") if "Noise" in hmm_features else 1

    fig, ax = plt.subplots(figsize=(9, 7))
    test_co2 = test_df["CO2"].values
    test_noise = test_df["Noise"].values

    for s in range(n_states):
        mask = test_states == s
        if mask.sum() == 0:
            continue
        ax.scatter(test_co2[mask], test_noise[mask], c=STATE_COLORS[s % len(STATE_COLORS)],
                   alpha=0.3, s=8, label=f"S{s}: {state_labels[s]} (n={mask.sum()})")

        # 95% confidence ellipse from HMM parameters
        mu_co2 = means[s, co2_idx]
        mu_noise = means[s, noise_idx]
        cov_2d = np.array([
            [covars[s, co2_idx, co2_idx], covars[s, co2_idx, noise_idx]],
            [covars[s, noise_idx, co2_idx], covars[s, noise_idx, noise_idx]],
        ])
        eigenvalues, eigenvectors = np.linalg.eigh(cov_2d)
        # chi2(2) 95% quantile = 5.991
        scale = np.sqrt(5.991)
        angle = np.degrees(np.arctan2(eigenvectors[1, 1], eigenvectors[0, 1]))
        width = 2 * scale * np.sqrt(eigenvalues[1])
        height = 2 * scale * np.sqrt(eigenvalues[0])
        ellipse = Ellipse(xy=(mu_co2, mu_noise), width=width, height=height,
                          angle=angle, edgecolor=STATE_COLORS[s % len(STATE_COLORS)],
                          facecolor="none", linewidth=2, linestyle="--")
        ax.add_patch(ellipse)

    ax.set_xlabel("CO2 (ppm)", fontsize=11)
    ax.set_ylabel("Noise (dB)", fontsize=11)
    ax.set_title(f"CO2 vs Noise by HMM State (95% Confidence Ellipses) - {horizon}h",
                 fontsize=12)
    ax.legend(fontsize=9)
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"A4_co2_noise_scatter_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- A5: CO2 time series with background color-coded by Viterbi state ----
    fig, ax = plt.subplots(figsize=(16, 5))
    n_test = len(test_co2)
    x_idx = np.arange(n_test)

    # Background color bands for each state
    for s in range(n_states):
        mask = test_states == s
        if mask.sum() == 0:
            continue
        # Find contiguous segments
        diff = np.diff(mask.astype(int))
        starts = np.where(diff == 1)[0] + 1
        ends = np.where(diff == -1)[0] + 1
        if mask[0]:
            starts = np.insert(starts, 0, 0)
        if mask[-1]:
            ends = np.append(ends, n_test)
        for s_start, s_end in zip(starts, ends):
            ax.axvspan(s_start, s_end, alpha=0.15,
                       color=STATE_COLORS[s % len(STATE_COLORS)])

    ax.plot(x_idx, test_co2, color="black", linewidth=0.5, alpha=0.8)
    ax.set_xlabel("Test Sample Index", fontsize=11)
    ax.set_ylabel("CO2 (ppm)", fontsize=11)
    ax.set_title(f"CO2 Time Series with HMM State Background - {horizon}h", fontsize=12)

    # Legend
    patches = [mpatches.Patch(color=STATE_COLORS[s % len(STATE_COLORS)], alpha=0.3,
                              label=f"S{s}: {state_labels[s]}") for s in range(n_states)]
    ax.legend(handles=patches, fontsize=9, loc="upper right")
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"A5_co2_state_timeline_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- A6: Zoomed 1-week view of state sequence ----
    # 1 week at 1h resolution = 168 samples
    week_len = min(168, n_test)
    fig, axes = plt.subplots(2, 1, figsize=(16, 8), gridspec_kw={"height_ratios": [3, 1]})

    ax = axes[0]
    ax.plot(x_idx[:week_len], test_co2[:week_len], color="black", linewidth=1.0)
    for s in range(n_states):
        mask_wk = test_states[:week_len] == s
        diff = np.diff(mask_wk.astype(int))
        starts = np.where(diff == 1)[0] + 1
        ends = np.where(diff == -1)[0] + 1
        if mask_wk[0]:
            starts = np.insert(starts, 0, 0)
        if mask_wk[-1]:
            ends = np.append(ends, week_len)
        for s_start, s_end in zip(starts, ends):
            ax.axvspan(s_start, s_end, alpha=0.2,
                       color=STATE_COLORS[s % len(STATE_COLORS)])
    ax.set_ylabel("CO2 (ppm)", fontsize=11)
    ax.set_title(f"1-Week Zoomed View with HMM States - {horizon}h", fontsize=12)
    ax.legend(handles=patches, fontsize=8, loc="upper right")
    ax.grid(alpha=0.3)

    # State sequence bar
    ax2 = axes[1]
    for t in range(week_len):
        ax2.barh(0, 1, left=t, color=STATE_COLORS[test_states[t] % len(STATE_COLORS)],
                 edgecolor="none")
    ax2.set_xlabel("Sample Index", fontsize=11)
    ax2.set_yticks([])
    ax2.set_xlim(0, week_len)
    ax2.set_title("Viterbi State Sequence", fontsize=10)

    plt.tight_layout()
    plt.savefig(output_dir / f"A6_zoomed_week_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- A7: State Duration Histogram ----
    fig, ax = plt.subplots(figsize=(10, 5))
    for s in range(n_states):
        durations = []
        mask = test_states == s
        diff = np.diff(mask.astype(int))
        starts = np.where(diff == 1)[0] + 1
        ends = np.where(diff == -1)[0] + 1
        if mask[0]:
            starts = np.insert(starts, 0, 0)
        if mask[-1]:
            ends = np.append(ends, len(mask))
        for s_start, s_end in zip(starts, ends):
            durations.append(s_end - s_start)
        if durations:
            ax.hist(durations, bins=30, alpha=0.5,
                    color=STATE_COLORS[s % len(STATE_COLORS)],
                    label=f"S{s}: {state_labels[s]} (median={np.median(durations):.0f}h)",
                    edgecolor="white")
            results[f"state_{s}_median_duration"] = float(np.median(durations))
            results[f"state_{s}_mean_duration"] = float(np.mean(durations))

    ax.set_xlabel("Duration (hours)", fontsize=11)
    ax.set_ylabel("Count", fontsize=11)
    ax.set_title(f"HMM State Duration Distribution - {horizon}h", fontsize=12)
    ax.legend(fontsize=9)
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"A7_state_duration_hist_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- A8: Posterior Probability Time Series ----
    fig, axes = plt.subplots(2, 1, figsize=(16, 8), gridspec_kw={"height_ratios": [2, 1]})

    # Posterior probabilities
    ax = axes[0]
    for s in range(n_states):
        ax.plot(test_probs[:, s], color=STATE_COLORS[s % len(STATE_COLORS)],
                linewidth=0.7, alpha=0.8, label=f"S{s}: {state_labels[s]}")
    ax.set_ylabel("Posterior Probability", fontsize=11)
    ax.set_title(f"HMM Posterior Probabilities - {horizon}h", fontsize=12)
    ax.legend(fontsize=9, loc="upper right")
    ax.set_ylim(-0.05, 1.05)
    ax.grid(alpha=0.3)

    # CO2 overlay
    ax2 = axes[1]
    ax2.plot(test_co2, color="black", linewidth=0.5, alpha=0.8)
    ax2.set_xlabel("Test Sample Index", fontsize=11)
    ax2.set_ylabel("CO2 (ppm)", fontsize=11)
    ax2.set_title("CO2 Reference", fontsize=10)
    ax2.grid(alpha=0.3)

    plt.tight_layout()
    plt.savefig(output_dir / f"A8_posterior_probabilities_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- A9: Posterior Entropy + Correlation with Prediction Error ----
    # Entropy: H = -sum(p * log(p))
    eps = 1e-10
    entropy = -np.sum(test_probs * np.log(test_probs + eps), axis=1)

    # Align prediction error with test samples
    # y_true and y_pred are per-window; each corresponds to a test sample
    n_preds = min(len(y_true), len(y_pred))
    # Take first-step prediction error for simplicity
    if y_true.ndim == 2:
        pred_error = np.abs(y_true[:, 0] - y_pred[:, 0])
    else:
        pred_error = np.abs(y_true.ravel()[:n_preds] - y_pred.ravel()[:n_preds])

    # Align entropy with prediction windows
    # Test windows start from lookback onwards, so entropy needs alignment
    lookback = 24  # 24h lookback at 1h resolution
    n_align = min(len(entropy) - lookback, n_preds)
    entropy_aligned = entropy[lookback:lookback + n_align]
    error_aligned = pred_error[:n_align]

    fig, axes = plt.subplots(2, 1, figsize=(14, 8))

    ax = axes[0]
    ax.plot(entropy, color=C_PRIMARY, linewidth=0.5, alpha=0.8)
    ax.set_ylabel("Posterior Entropy (nats)", fontsize=11)
    ax.set_title(f"HMM Posterior Entropy Over Time - {horizon}h", fontsize=12)
    ax.grid(alpha=0.3)

    ax = axes[1]
    if n_align > 10:
        corr = np.corrcoef(entropy_aligned, error_aligned)[0, 1]
        ax.scatter(entropy_aligned, error_aligned, alpha=0.2, s=5, color=C_SECONDARY)
        ax.set_xlabel("Posterior Entropy", fontsize=11)
        ax.set_ylabel("|Prediction Error| (ppm)", fontsize=11)
        ax.set_title(f"Entropy vs Prediction Error (r={corr:.3f}) - {horizon}h", fontsize=12)
        ax.grid(alpha=0.3)
        results["entropy_error_correlation"] = float(corr)
    else:
        ax.text(0.5, 0.5, "Not enough aligned data", ha="center", va="center",
                transform=ax.transAxes)

    plt.tight_layout()
    plt.savefig(output_dir / f"A9_entropy_error_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # State distribution stats
    for s in range(n_states):
        pct = (test_states == s).sum() / len(test_states) * 100
        results[f"state_{s}_pct"] = float(pct)
        print(f"    S{s} ({state_labels[s]}): {pct:.1f}% of test samples")

    print(f"  [A] HMM Regime Analysis complete: {n_states} states, "
          f"{len(test_states)} test samples")
    return results


# ======================================================================
#  Section B: Gradient-Based Feature Attribution
# ======================================================================

def run_gradient_attribution(
    best_model: HMMLSTMForecaster,
    dm: CO2DataModule,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> None:
    """Gradient saliency on the 22-dim HMM-LSTM input."""
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [B] Computing gradient-based feature attribution...")

    best_model.eval()
    device = next(best_model.parameters()).device

    all_grads = []
    max_batches = 50

    test_dl = dm.test_dataloader()

    for i, (x, y) in enumerate(test_dl):
        if i >= max_batches:
            break
        x_in = x.clone().detach().to(device).requires_grad_(True)
        out = best_model(x_in)
        # Use mean of predictions as scalar target
        target = out.mean()
        target.backward()

        if x_in.grad is not None:
            grad_abs = x_in.grad.abs().detach().cpu().numpy()
            all_grads.append(grad_abs)

        best_model.zero_grad()

    if not all_grads:
        print("  [B] No gradients collected. Skipping.")
        return

    # Mean across all samples: (lookback, n_features)
    grads = np.concatenate(all_grads, axis=0)
    avg_grad = grads.mean(axis=0)

    # Feature names: feature_columns + target (CO2)
    feature_names = config["data"]["feature_columns"] + [config["data"]["target_column"]]
    n_input = avg_grad.shape[1]
    if len(feature_names) != n_input:
        feature_names = [f"feat_{i}" for i in range(n_input)]

    # Identify HMM channels
    hmm_mask = np.array(["hmm_state" in f for f in feature_names])
    sensor_mask = ~hmm_mask

    # ---- Plot B1: Full gradient heatmap (lookback x features) ----
    fig, ax = plt.subplots(figsize=(max(10, n_input * 0.5), 7))
    im = ax.imshow(avg_grad.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Lookback Timestep (hours ago)", fontsize=11)
    ax.set_ylabel("Feature", fontsize=11)
    ax.set_yticks(range(n_input))
    ax.set_yticklabels(feature_names, fontsize=7)
    lookback = avg_grad.shape[0]
    tick_positions = np.linspace(0, lookback - 1, min(8, lookback)).astype(int)
    ax.set_xticks(tick_positions)
    ax.set_xticklabels([f"-{lookback - t}" for t in tick_positions])
    ax.set_title(f"Gradient Attribution Heatmap (22 features) - {horizon}h", fontsize=12)
    plt.colorbar(im, ax=ax, label="|Gradient|")
    plt.tight_layout()
    plt.savefig(output_dir / f"B1_gradient_heatmap_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- Plot B2: Per-feature bar chart highlighting HMM channels ----
    feat_importance = avg_grad.mean(axis=0)
    feat_importance_pct = feat_importance / feat_importance.sum() * 100
    sort_idx = np.argsort(feat_importance_pct)

    fig, ax = plt.subplots(figsize=(9, max(5, n_input * 0.35)))
    colors = [C_SECONDARY if hmm_mask[i] else C_PRIMARY for i in sort_idx]
    ax.barh(range(len(sort_idx)), feat_importance_pct[sort_idx], color=colors)
    ax.set_yticks(range(len(sort_idx)))
    ax.set_yticklabels([feature_names[i] for i in sort_idx], fontsize=8)
    ax.set_xlabel("Gradient Importance (%)", fontsize=11)
    ax.set_title(f"Feature Attribution - HMM channels highlighted - {horizon}h", fontsize=12)
    # Legend
    sensor_patch = mpatches.Patch(color=C_PRIMARY, label="Original/Sensor")
    hmm_patch = mpatches.Patch(color=C_SECONDARY, label="HMM State Posterior")
    ax.legend(handles=[sensor_patch, hmm_patch], fontsize=9)
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"B2_feature_ranking_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- Plot B3: Temporal gradient profile split by sensor vs HMM ----
    sensor_temporal = avg_grad[:, sensor_mask].sum(axis=1)
    hmm_temporal = avg_grad[:, hmm_mask].sum(axis=1)
    total_temporal = sensor_temporal + hmm_temporal
    total_temporal = np.where(total_temporal == 0, 1, total_temporal)

    time_idx = np.arange(-lookback, 0)

    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(time_idx, sensor_temporal / total_temporal, color=C_PRIMARY, alpha=0.8,
           label="Sensor Features")
    ax.bar(time_idx, hmm_temporal / total_temporal, bottom=sensor_temporal / total_temporal,
           color=C_SECONDARY, alpha=0.8, label="HMM State Posteriors")
    ax.set_xlabel("Hours Ago", fontsize=11)
    ax.set_ylabel("Gradient Share", fontsize=11)
    ax.set_title(f"Temporal Gradient Profile: Sensor vs HMM - {horizon}h", fontsize=12)
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"B3_temporal_gradient_split_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Save gradient data
    grad_df = pd.DataFrame(avg_grad, columns=feature_names)
    grad_df.to_csv(output_dir / f"gradient_attribution_h{horizon}.csv", index=False)

    # Print summary
    hmm_total_pct = feat_importance_pct[hmm_mask].sum()
    print(f"  [B] Gradient attribution: {grads.shape[0]} samples, "
          f"{n_input} features, HMM channels={hmm_total_pct:.1f}% total")


# ======================================================================
#  Section C: Hidden State Analysis
# ======================================================================

def run_hidden_state_analysis(
    best_model: HMMLSTMForecaster,
    dm: CO2DataModule,
    detector: HMMRegimeDetector,
    test_df: pd.DataFrame,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> None:
    """PCA on LSTM h_n[-1] across test samples, with HMM state coloring."""
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [C] Extracting LSTM hidden states...")

    best_model.eval()
    device = next(best_model.parameters()).device

    all_hidden = []
    all_targets = []
    max_batches = 80

    # Hook to capture h_n from LSTM
    hidden_captures = []

    def lstm_hook(module, inp, output):
        # output = (output_seq, (h_n, c_n))
        h_n = output[1][0].detach().cpu()  # (num_layers, batch, hidden)
        hidden_captures.append(h_n[-1])  # last layer

    hook_handle = best_model.lstm.register_forward_hook(lstm_hook)

    test_dl = dm.test_dataloader()

    with torch.no_grad():
        for i, (x, y) in enumerate(test_dl):
            if i >= max_batches:
                break
            x_dev = x.to(device)
            best_model(x_dev)
            all_targets.append(y.numpy())

    hook_handle.remove()

    if not hidden_captures:
        print("  [C] No hidden states captured. Skipping.")
        return

    hidden_states = torch.cat(hidden_captures, dim=0).numpy()
    target_values = np.concatenate(all_targets, axis=0)
    if target_values.ndim == 2:
        target_first_step = target_values[:, 0]
    else:
        target_first_step = target_values.ravel()

    n = min(len(hidden_states), len(target_first_step))
    hidden_states = hidden_states[:n]
    target_first_step = target_first_step[:n]

    print(f"  [C] Hidden states shape: {hidden_states.shape}")

    # PCA
    n_components = min(20, hidden_states.shape[1])
    pca = PCA(n_components=n_components)
    pca_result = pca.fit_transform(hidden_states)

    # ---- C1: PCA Explained Variance ----
    fig, ax = plt.subplots(figsize=(8, 5))
    cumvar = np.cumsum(pca.explained_variance_ratio_) * 100
    ax.bar(range(1, n_components + 1), pca.explained_variance_ratio_ * 100,
           color=C_PRIMARY, alpha=0.7)
    ax.plot(range(1, n_components + 1), cumvar, "r-o", markersize=4, label="Cumulative")
    ax.set_xlabel("Principal Component", fontsize=11)
    ax.set_ylabel("Explained Variance (%)", fontsize=11)
    ax.set_title(f"PCA on LSTM Hidden States - {horizon}h", fontsize=12)
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"C1_pca_variance_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Inverse-scale target for CO2 level coloring
    target_co2 = inverse_scale_target(target_first_step, dm.target_scaler)

    co2_bins = [0, 500, 1000, np.inf]
    co2_labels = ["Low (<500)", "Medium (500-1000)", "High (>1000)"]
    co2_cats = pd.cut(target_co2, bins=co2_bins, labels=co2_labels)
    cat_colors = {
        "Low (<500)": C_ACCENT,
        "Medium (500-1000)": C_PRIMARY,
        "High (>1000)": C_SECONDARY,
    }

    # ---- C2: PCA scatter by CO2 level ----
    fig, ax = plt.subplots(figsize=(8, 7))
    for label in co2_labels:
        mask = co2_cats == label
        if mask.sum() > 0:
            ax.scatter(pca_result[mask, 0], pca_result[mask, 1],
                       c=cat_colors[label],
                       label=f"{label} (n={mask.sum()})",
                       alpha=0.5, s=10)
    ax.set_xlabel(f"PC1 ({pca.explained_variance_ratio_[0]*100:.1f}%)", fontsize=11)
    ax.set_ylabel(f"PC2 ({pca.explained_variance_ratio_[1]*100:.1f}%)", fontsize=11)
    ax.set_title(f"LSTM Hidden States by CO2 Level - {horizon}h", fontsize=12)
    ax.legend(fontsize=9)
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"C2_pca_co2_level_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- C3: PCA scatter by HMM regime (unique to HMM-LSTM!) ----
    # Get HMM states for the test windows
    test_viterbi = detector.predict_states(test_df)
    lookback = config["data"]["lookback_hours"] * config["data"]["samples_per_hour"]
    # Each window starts at index i, uses i:i+lookback, the state at end of lookback
    # Approximate: use state at lookback position for each window
    hmm_states_for_windows = test_viterbi[lookback:lookback + n]

    fig, ax = plt.subplots(figsize=(8, 7))
    state_labels_local = label_states_by_co2(detector)
    for s in range(detector.n_states):
        mask = hmm_states_for_windows == s
        if mask.sum() > 0:
            ax.scatter(pca_result[mask, 0], pca_result[mask, 1],
                       c=STATE_COLORS[s % len(STATE_COLORS)],
                       label=f"S{s}: {state_labels_local[s]} (n={mask.sum()})",
                       alpha=0.5, s=10)
    ax.set_xlabel(f"PC1 ({pca.explained_variance_ratio_[0]*100:.1f}%)", fontsize=11)
    ax.set_ylabel(f"PC2 ({pca.explained_variance_ratio_[1]*100:.1f}%)", fontsize=11)
    ax.set_title(f"LSTM Hidden States by HMM Regime - {horizon}h", fontsize=12)
    ax.legend(fontsize=9)
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"C3_pca_hmm_regime_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- C4: K-Means clustering vs HMM state crosstab ----
    n_clusters = detector.n_states
    kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
    cluster_labels = kmeans.fit_predict(hidden_states)

    fig, axes = plt.subplots(1, 2, figsize=(14, 6))

    # Scatter colored by cluster
    ax = axes[0]
    for c in range(n_clusters):
        mask = cluster_labels == c
        ax.scatter(pca_result[mask, 0], pca_result[mask, 1],
                   label=f"Cluster {c} (n={mask.sum()})", alpha=0.5, s=10)
    ax.set_xlabel("PC1", fontsize=11)
    ax.set_ylabel("PC2", fontsize=11)
    ax.set_title("K-Means Clusters in PCA Space", fontsize=11)
    ax.legend(fontsize=8)
    ax.grid(alpha=0.3)

    # Crosstab: K-Means vs HMM state
    ax = axes[1]
    ax.axis("off")
    cross = pd.crosstab(
        pd.Series(cluster_labels[:len(hmm_states_for_windows)], name="Cluster"),
        pd.Series(hmm_states_for_windows[:len(cluster_labels)], name="HMM State"),
    )
    rows_data = []
    hmm_col_labels = [f"HMM S{s}" for s in range(n_clusters)]
    for c_idx in range(n_clusters):
        row = [f"Cluster {c_idx}"]
        for s_idx in range(n_clusters):
            val = cross.loc[c_idx, s_idx] if c_idx in cross.index and s_idx in cross.columns else 0
            row.append(str(val))
        rows_data.append(row)

    tbl = ax.table(
        cellText=rows_data,
        colLabels=["Cluster"] + hmm_col_labels,
        loc="center",
        cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.0, 1.5)
    ax.set_title("K-Means Cluster vs HMM State Crosstab", fontsize=11, pad=20)

    fig.suptitle(f"Hidden State Clustering vs HMM Regimes - {horizon}h", fontsize=13)
    plt.tight_layout()
    plt.savefig(output_dir / f"C4_clustering_crosstab_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    print(f"  [C] Hidden state analysis complete: {n} samples, {n_components} PCs")


# ======================================================================
#  Section D: Temporal Pattern Analysis
# ======================================================================

def run_temporal_patterns(
    best_model: HMMLSTMForecaster,
    dm: CO2DataModule,
    detector: HMMRegimeDetector,
    test_df: pd.DataFrame,
    y_true: np.ndarray,
    y_pred: np.ndarray,
    horizon: int,
    output_dir: Path,
) -> None:
    """FFT, ACF of hidden state PCs, FFT of residuals, rolling RMSE,
    and HMM state periodicity."""
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [D] Analyzing temporal patterns...")

    best_model.eval()
    device = next(best_model.parameters()).device

    # Collect hidden states
    hidden_captures = []

    def lstm_hook(module, inp, output):
        h_n = output[1][0].detach().cpu()
        hidden_captures.append(h_n[-1])

    hook_handle = best_model.lstm.register_forward_hook(lstm_hook)
    test_dl = dm.test_dataloader()
    max_batches = 100

    with torch.no_grad():
        for i, (x, y) in enumerate(test_dl):
            if i >= max_batches:
                break
            best_model(x.to(device))

    hook_handle.remove()

    if hidden_captures:
        hidden_concat = torch.cat(hidden_captures, dim=0).numpy()
    else:
        hidden_concat = np.zeros((1, 128))

    n_samples_h = hidden_concat.shape[0]

    # ---- D1: FFT of hidden state PCs ----
    if n_samples_h > 50:
        pca_temp = PCA(n_components=3)
        pcs = pca_temp.fit_transform(hidden_concat)

        fig, axes = plt.subplots(3, 1, figsize=(12, 10))
        for pc_idx in range(3):
            ax = axes[pc_idx]
            pc_series = pcs[:, pc_idx]
            pc_series = pc_series - pc_series.mean()
            n_pts = len(pc_series)
            yf = np.abs(fft(pc_series))
            xf = fftfreq(n_pts, d=1.0)
            pos_mask = xf > 0
            periods = 1.0 / xf[pos_mask]
            magnitudes = yf[pos_mask]
            valid = periods < min(200, n_pts // 2)
            ax.plot(periods[valid], magnitudes[valid],
                    color=C_PRIMARY, linewidth=1.0)
            ax.set_ylabel(f"PC{pc_idx+1} FFT Mag", fontsize=10)
            ax.set_title(f"PC{pc_idx+1} ({pca_temp.explained_variance_ratio_[pc_idx]*100:.1f}% var)")
            ax.grid(alpha=0.3)

        axes[-1].set_xlabel("Period (hours)", fontsize=11)
        fig.suptitle(f"FFT of Hidden State Principal Components - {horizon}h", fontsize=13)
        plt.tight_layout()
        plt.savefig(output_dir / f"D1_fft_hidden_pcs_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    # ---- D2: ACF of hidden state PCs ----
    if n_samples_h > 50:
        max_lag = min(72, n_samples_h // 3)

        fig, axes = plt.subplots(3, 1, figsize=(12, 9))
        for pc_idx in range(3):
            ax = axes[pc_idx]
            pc_series = pcs[:, pc_idx]
            pc_series = pc_series - pc_series.mean()
            norm_val = np.sum(pc_series ** 2)
            if norm_val > 0:
                autocorr = np.correlate(pc_series, pc_series, mode="full")
                autocorr = autocorr[len(autocorr) // 2:]
                autocorr = autocorr[:max_lag] / norm_val
            else:
                autocorr = np.zeros(max_lag)

            ax.bar(range(max_lag), autocorr, color=C_PRIMARY, alpha=0.7, width=0.8)
            ax.axhline(0, color="black", linewidth=0.5)
            ax.set_ylabel(f"PC{pc_idx+1} ACF", fontsize=10)
            ax.set_title(f"PC{pc_idx+1} ({pca_temp.explained_variance_ratio_[pc_idx]*100:.1f}% var)")
            ax.grid(axis="y", alpha=0.3)

        axes[-1].set_xlabel("Lag (hours)", fontsize=11)
        fig.suptitle(f"Autocorrelation of Hidden State PCs - {horizon}h", fontsize=13)
        plt.tight_layout()
        plt.savefig(output_dir / f"D2_acf_hidden_pcs_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    # ---- D3: FFT of residuals ----
    if y_true.ndim == 2:
        residuals = y_true[:, 0] - y_pred[:, 0]
    else:
        residuals = y_true.ravel() - y_pred.ravel()

    n_res = len(residuals)
    if n_res > 20:
        fig, ax = plt.subplots(figsize=(12, 5))
        res_centered = residuals - residuals.mean()
        yf = np.abs(fft(res_centered))
        xf = fftfreq(n_res, d=1.0)
        pos_mask = xf > 0
        periods = 1.0 / xf[pos_mask]
        magnitudes = yf[pos_mask]
        valid = periods < min(200, n_res // 2)
        ax.plot(periods[valid], magnitudes[valid], color=C_SECONDARY, linewidth=1.0)
        ax.set_xlabel("Period (hours)", fontsize=11)
        ax.set_ylabel("FFT Magnitude", fontsize=11)
        ax.set_title(f"FFT of Prediction Residuals - {horizon}h", fontsize=12)
        ax.grid(alpha=0.3)

        # Mark key periodicities
        if len(periods[valid]) > 0:
            peak_idx = np.argmax(magnitudes[valid])
            peak_period = periods[valid][peak_idx]
            ax.axvline(peak_period, color="red", linestyle="--", alpha=0.6,
                       label=f"Peak at {peak_period:.1f}h")
            if 20 < peak_period or peak_period < 30:
                ax.axvline(24, color="green", linestyle=":", alpha=0.6, label="24h reference")
            ax.legend(fontsize=9)

        plt.tight_layout()
        plt.savefig(output_dir / f"D3_fft_residuals_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    # ---- D4: Rolling RMSE ----
    if n_res > 50:
        window_size = min(24, n_res // 5)
        rolling_sq_err = pd.Series(residuals ** 2).rolling(window=window_size, min_periods=1).mean()
        rolling_rmse = np.sqrt(rolling_sq_err.values)

        fig, ax = plt.subplots(figsize=(14, 5))
        ax.plot(rolling_rmse, color=C_PRIMARY, linewidth=1.0, alpha=0.8)
        ax.axhline(np.sqrt(np.mean(residuals ** 2)), color="red", linestyle="--",
                   linewidth=1.5, label=f"Overall RMSE = {np.sqrt(np.mean(residuals**2)):.2f}")
        ax.set_xlabel("Test Sample Index", fontsize=11)
        ax.set_ylabel(f"Rolling RMSE ({window_size}h window)", fontsize=11)
        ax.set_title(f"Rolling RMSE Over Test Set - {horizon}h", fontsize=12)
        ax.legend(fontsize=10)
        ax.grid(alpha=0.3)
        plt.tight_layout()
        plt.savefig(output_dir / f"D4_rolling_rmse_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    # ---- D5: HMM State Periodicity (ACF of posterior probabilities) ----
    test_probs = detector.predict_proba(test_df)
    n_test = test_probs.shape[0]
    if n_test > 50:
        max_lag_hmm = min(168, n_test // 3)  # up to 1 week

        fig, axes = plt.subplots(detector.n_states, 1,
                                 figsize=(12, 3 * detector.n_states))
        if detector.n_states == 1:
            axes = [axes]

        state_labels_local = label_states_by_co2(detector)
        for s in range(detector.n_states):
            ax = axes[s]
            prob_series = test_probs[:, s]
            prob_series = prob_series - prob_series.mean()
            norm_val = np.sum(prob_series ** 2)
            if norm_val > 0:
                autocorr = np.correlate(prob_series, prob_series, mode="full")
                autocorr = autocorr[len(autocorr) // 2:]
                autocorr = autocorr[:max_lag_hmm] / norm_val
            else:
                autocorr = np.zeros(max_lag_hmm)

            ax.bar(range(max_lag_hmm), autocorr,
                   color=STATE_COLORS[s % len(STATE_COLORS)], alpha=0.7, width=0.8)
            ax.axhline(0, color="black", linewidth=0.5)
            ax.set_ylabel("ACF", fontsize=10)
            ax.set_title(f"S{s}: {state_labels_local[s]}", fontsize=10)
            ax.grid(axis="y", alpha=0.3)

        axes[-1].set_xlabel("Lag (hours)", fontsize=11)
        fig.suptitle(f"Autocorrelation of HMM Posterior Probabilities - {horizon}h",
                     fontsize=13)
        plt.tight_layout()
        plt.savefig(output_dir / f"D5_acf_hmm_posteriors_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    print(f"  [D] Temporal pattern analysis complete")


# ======================================================================
#  Section E: HMM-LSTM-Specific Analysis
# ======================================================================

def run_hmm_lstm_specific_analysis(
    best_model: HMMLSTMForecaster,
    dm: CO2DataModule,
    detector: HMMRegimeDetector,
    test_df: pd.DataFrame,
    config: dict,
    y_true: np.ndarray,
    y_pred: np.ndarray,
    metrics: dict,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Regime-conditioned error, HMM channel permutation importance,
    prediction quality plots."""
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [E] HMM-LSTM-Specific Analysis...")

    results = {}

    if y_true.ndim == 2:
        y_true_flat = y_true[:, 0]
        y_pred_flat = y_pred[:, 0]
    else:
        y_true_flat = y_true.ravel()
        y_pred_flat = y_pred.ravel()

    residuals = y_true_flat - y_pred_flat
    n_preds = len(y_true_flat)

    # Get HMM states aligned with prediction windows
    test_viterbi = detector.predict_states(test_df)
    lookback = config["data"]["lookback_hours"] * config["data"]["samples_per_hour"]
    # States at the end of each lookback window
    hmm_states_aligned = test_viterbi[lookback:lookback + n_preds]
    n_align = min(len(hmm_states_aligned), n_preds)
    hmm_states_aligned = hmm_states_aligned[:n_align]
    y_true_aligned = y_true_flat[:n_align]
    y_pred_aligned = y_pred_flat[:n_align]
    residuals_aligned = residuals[:n_align]

    state_labels_local = label_states_by_co2(detector)

    # ---- E1: Regime-conditioned RMSE ----
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))

    # Bar chart of RMSE per state
    ax = axes[0]
    state_rmses = []
    state_counts = []
    for s in range(detector.n_states):
        mask = hmm_states_aligned == s
        if mask.sum() > 0:
            rmse_s = np.sqrt(np.mean(residuals_aligned[mask] ** 2))
            state_rmses.append(rmse_s)
            state_counts.append(mask.sum())
            results[f"rmse_state_{s}"] = float(rmse_s)
            results[f"count_state_{s}"] = int(mask.sum())
        else:
            state_rmses.append(0)
            state_counts.append(0)

    bars = ax.bar(range(detector.n_states), state_rmses,
                  color=[STATE_COLORS[s % len(STATE_COLORS)] for s in range(detector.n_states)])
    ax.set_xticks(range(detector.n_states))
    ax.set_xticklabels([f"S{s}\n({state_counts[s]})" for s in range(detector.n_states)],
                       fontsize=10)
    ax.set_ylabel("RMSE (ppm)", fontsize=11)
    ax.set_title(f"RMSE by HMM State - {horizon}h", fontsize=12)
    # Add value labels
    for bar_obj, val in zip(bars, state_rmses):
        ax.text(bar_obj.get_x() + bar_obj.get_width() / 2, bar_obj.get_height() + 0.5,
                f"{val:.1f}", ha="center", va="bottom", fontsize=10)
    ax.grid(axis="y", alpha=0.3)

    # Table
    ax = axes[1]
    ax.axis("off")
    rows_data = []
    for s in range(detector.n_states):
        mask = hmm_states_aligned == s
        if mask.sum() > 0:
            res_s = residuals_aligned[mask]
            rows_data.append([
                f"S{s}: {state_labels_local[s]}",
                str(mask.sum()),
                f"{np.sqrt(np.mean(res_s**2)):.2f}",
                f"{np.mean(np.abs(res_s)):.2f}",
                f"{np.mean(res_s):.2f}",
            ])
        else:
            rows_data.append([f"S{s}: {state_labels_local[s]}", "0", "N/A", "N/A", "N/A"])

    tbl = ax.table(
        cellText=rows_data,
        colLabels=["State", "N", "RMSE", "MAE", "Mean Bias"],
        loc="center",
        cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.0, 1.5)
    ax.set_title("Error Metrics by HMM Regime", fontsize=11, pad=20)

    fig.suptitle(f"Regime-Conditioned Error Analysis - {horizon}h", fontsize=13)
    plt.tight_layout()
    plt.savefig(output_dir / f"E1_regime_conditioned_error_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- E2: HMM Channel Permutation Importance ----
    print("    Computing HMM channel permutation importance...")
    best_model.eval()
    device = next(best_model.parameters()).device

    test_dl = dm.test_dataloader()

    # Baseline MSE
    baseline_preds = []
    baseline_targets = []
    with torch.no_grad():
        for x, y in test_dl:
            pred = best_model(x.to(device)).cpu()
            baseline_preds.append(pred)
            baseline_targets.append(y)

    baseline_pred_cat = torch.cat(baseline_preds).numpy()
    baseline_target_cat = torch.cat(baseline_targets).numpy()
    baseline_mse = np.mean((baseline_pred_cat - baseline_target_cat) ** 2)

    # Feature names
    feature_names = config["data"]["feature_columns"] + [config["data"]["target_column"]]
    n_input = len(feature_names)

    # Identify HMM channel indices
    hmm_indices = [i for i, f in enumerate(feature_names) if "hmm_state" in f]

    # Permutation test: shuffle each HMM channel individually, and all together
    rng = np.random.default_rng(42)
    n_permutations = 5
    perm_results = {}

    for channel_idx in hmm_indices:
        channel_name = feature_names[channel_idx]
        mse_increases = []

        for perm_i in range(n_permutations):
            perm_preds = []
            with torch.no_grad():
                for x, y in test_dl:
                    x_perm = x.clone()
                    # Shuffle this channel across the batch
                    perm_order = rng.permutation(x_perm.shape[0])
                    x_perm[:, :, channel_idx] = x_perm[perm_order, :, channel_idx]
                    pred = best_model(x_perm.to(device)).cpu()
                    perm_preds.append(pred)
            perm_pred_cat = torch.cat(perm_preds).numpy()
            perm_mse = np.mean((perm_pred_cat - baseline_target_cat) ** 2)
            mse_increases.append(perm_mse - baseline_mse)

        mean_increase = np.mean(mse_increases)
        perm_results[channel_name] = mean_increase

    # Shuffle ALL HMM channels simultaneously
    all_hmm_mse_increases = []
    for perm_i in range(n_permutations):
        perm_preds = []
        with torch.no_grad():
            for x, y in test_dl:
                x_perm = x.clone()
                for ch_idx in hmm_indices:
                    perm_order = rng.permutation(x_perm.shape[0])
                    x_perm[:, :, ch_idx] = x_perm[perm_order, :, ch_idx]
                pred = best_model(x_perm.to(device)).cpu()
                perm_preds.append(pred)
        perm_pred_cat = torch.cat(perm_preds).numpy()
        perm_mse = np.mean((perm_pred_cat - baseline_target_cat) ** 2)
        all_hmm_mse_increases.append(perm_mse - baseline_mse)

    perm_results["All HMM Channels"] = np.mean(all_hmm_mse_increases)

    # Plot
    fig, ax = plt.subplots(figsize=(8, max(4, len(perm_results) * 0.6)))
    names = list(perm_results.keys())
    values = [perm_results[n] for n in names]
    sort_idx = np.argsort(values)
    colors = [C_SECONDARY if "All" in names[i] else C_PRIMARY for i in sort_idx]
    ax.barh(range(len(sort_idx)), [values[i] for i in sort_idx], color=colors)
    ax.set_yticks(range(len(sort_idx)))
    ax.set_yticklabels([names[i] for i in sort_idx], fontsize=10)
    ax.set_xlabel("MSE Increase (Permutation Importance)", fontsize=11)
    ax.set_title(f"HMM Channel Permutation Importance - {horizon}h", fontsize=12)
    ax.axvline(0, color="black", linewidth=0.5)
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"E2_hmm_permutation_importance_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    results["hmm_permutation_importance"] = perm_results
    print(f"    HMM channel permutation importance: "
          f"All channels MSE increase = {perm_results['All HMM Channels']:.4f}")

    # ---- E3: Prediction overlay ----
    n_show = min(500, n_preds)
    # Inverse-scale for original CO2 ppm
    y_true_ppm = inverse_scale_target(y_true_flat, dm.target_scaler)
    y_pred_ppm = inverse_scale_target(y_pred_flat, dm.target_scaler)

    fig, ax = plt.subplots(figsize=(14, 5))
    ax.plot(y_true_ppm[:n_show], label="Actual", color=C_PRIMARY, linewidth=1.0, alpha=0.8)
    ax.plot(y_pred_ppm[:n_show], label="Predicted", color=C_SECONDARY, linewidth=1.0, alpha=0.8)
    ax.set_xlabel("Sample Index", fontsize=11)
    ax.set_ylabel("CO2 (ppm)", fontsize=11)
    ax.set_title(f"HMM-LSTM Predictions vs Actual - {horizon}h", fontsize=12)
    ax.legend(fontsize=10)
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"E3_predictions_overlay_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- E4: Scatter with R2 ----
    fig, ax = plt.subplots(figsize=(7, 7))
    ax.scatter(y_true_ppm, y_pred_ppm, alpha=0.3, s=5, color=C_PRIMARY)
    vmin = min(y_true_ppm.min(), y_pred_ppm.min())
    vmax = max(y_true_ppm.max(), y_pred_ppm.max())
    ax.plot([vmin, vmax], [vmin, vmax], "r--", linewidth=1.5, label="Perfect")
    ax.text(
        0.05, 0.95,
        f"R2 = {metrics['r2']:.4f}\nRMSE = {metrics['rmse']:.2f}\nMAE = {metrics['mae']:.2f}",
        transform=ax.transAxes, fontsize=11, va="top",
        bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8),
    )
    ax.set_xlabel("Actual CO2 (ppm)", fontsize=11)
    ax.set_ylabel("Predicted CO2 (ppm)", fontsize=11)
    ax.set_title(f"HMM-LSTM Scatter - {horizon}h", fontsize=12)
    ax.legend(fontsize=10)
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"E4_scatter_r2_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- E5: 4-panel residual analysis ----
    residuals_ppm = y_true_ppm - y_pred_ppm

    datetime_col = config["data"]["datetime_column"]
    test_dates = None
    if datetime_col in test_df.columns:
        test_dates = pd.DatetimeIndex(test_df[datetime_col])

    fig, axes = plt.subplots(2, 2, figsize=(14, 10))

    # Distribution
    ax = axes[0, 0]
    ax.hist(residuals_ppm, bins=50, color=C_PRIMARY, alpha=0.7, edgecolor="white", density=True)
    ax.axvline(0, color="red", linestyle="--", linewidth=1.5)
    ax.set_xlabel("Residual (ppm)", fontsize=10)
    ax.set_ylabel("Density", fontsize=10)
    ax.set_title("Residual Distribution")
    ax.text(0.95, 0.95,
            f"Mean: {residuals_ppm.mean():.2f}\nStd: {residuals_ppm.std():.2f}",
            transform=ax.transAxes, fontsize=9, ha="right", va="top",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8))

    # Over time
    ax = axes[0, 1]
    ax.scatter(np.arange(n_preds), residuals_ppm, alpha=0.3, s=3, color=C_SECONDARY)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xlabel("Sample Index", fontsize=10)
    ax.set_ylabel("Residual (ppm)", fontsize=10)
    ax.set_title("Residuals Over Time")

    # By hour of day
    ax = axes[1, 0]
    if test_dates is not None and len(test_dates) >= lookback + n_preds:
        hours = test_dates[lookback:lookback + n_preds].hour
        res_df = pd.DataFrame({"hour": hours.values, "residual": residuals_ppm})
        res_df.boxplot(column="residual", by="hour", ax=ax)
        ax.set_xlabel("Hour of Day", fontsize=10)
        ax.set_ylabel("Residual (ppm)", fontsize=10)
        ax.set_title("Residuals by Hour")
        plt.suptitle("")
    else:
        ax.text(0.5, 0.5, "No datetime available", ha="center", va="center",
                transform=ax.transAxes)
        ax.set_title("Residuals by Hour (N/A)")

    # By HMM state
    ax = axes[1, 1]
    if n_align > 0:
        res_state_df = pd.DataFrame({
            "state": [f"S{s}" for s in hmm_states_aligned],
            "residual": residuals_ppm[:n_align],
        })
        res_state_df.boxplot(column="residual", by="state", ax=ax)
        ax.set_xlabel("HMM State", fontsize=10)
        ax.set_ylabel("Residual (ppm)", fontsize=10)
        ax.set_title("Residuals by HMM State")
        plt.suptitle("")
    else:
        ax.text(0.5, 0.5, "Not enough aligned data", ha="center", va="center",
                transform=ax.transAxes)
        ax.set_title("Residuals by HMM State (N/A)")

    plt.suptitle("")
    fig.suptitle(f"HMM-LSTM Residual Analysis - {horizon}h", fontsize=13)
    plt.tight_layout()
    plt.savefig(output_dir / f"E5_residual_analysis_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    # ---- E6: Error by CO2 level ----
    bins = [0, 500, 1000, np.inf]
    labels = ["Low (<500)", "Medium (500-1000)", "High (>1000)"]
    categories = pd.cut(y_true_ppm, bins=bins, labels=labels)

    fig, axes = plt.subplots(1, 2, figsize=(14, 5))

    ax = axes[0]
    abs_errors = np.abs(residuals_ppm)
    err_df = pd.DataFrame({"category": categories, "abs_error": abs_errors})
    err_df.boxplot(column="abs_error", by="category", ax=ax)
    ax.set_xlabel("CO2 Level", fontsize=10)
    ax.set_ylabel("Absolute Error (ppm)", fontsize=10)
    ax.set_title(f"Error by CO2 Level - {horizon}h")
    plt.suptitle("")

    ax = axes[1]
    ax.axis("off")
    rows_data = []
    for label in labels:
        mask = categories == label
        if mask.sum() > 0:
            bin_res = residuals_ppm[mask]
            rows_data.append([
                label,
                f"{mask.sum()}",
                f"{np.sqrt(np.mean(bin_res**2)):.2f}",
                f"{np.mean(np.abs(bin_res)):.2f}",
                f"{np.mean(bin_res):.2f}",
            ])
        else:
            rows_data.append([label, "0", "N/A", "N/A", "N/A"])

    tbl = ax.table(
        cellText=rows_data,
        colLabels=["CO2 Level", "N", "RMSE", "MAE", "Mean Bias"],
        loc="center",
        cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.0, 1.5)

    plt.tight_layout()
    plt.savefig(output_dir / f"E6_error_by_co2_level_h{horizon}.png",
                dpi=150, bbox_inches="tight")
    plt.close(fig)

    print(f"  [E] HMM-LSTM-specific analysis complete")
    return results


# ======================================================================
#  DOCX Report Generation
# ======================================================================

def generate_docx_report(
    all_metrics: dict,
    all_hmm_results: dict,
    all_specific_results: dict,
    horizons: list,
    output_dir: Path,
) -> None:
    """Generate comprehensive DOCX academic report."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    output_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ---- Title ----
    title = doc.add_heading(
        "HMM-LSTM Interpretability Study: Regime-Aware Sequential Modeling "
        "for Indoor CO2 Forecasting",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"Variant: preproc_D (Enhanced 1h) | "
        f"Horizons: {', '.join(str(h) + 'h' for h in horizons)}"
    )

    # ---- Abstract ----
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This report presents a comprehensive interpretability study of the HMM-LSTM "
        "hybrid model for indoor CO2 concentration forecasting. The HMM-LSTM model "
        "employs a two-stage architecture: a Gaussian Hidden Markov Model (HMM) first "
        "identifies latent environmental regimes from CO2, Noise, and external temperature "
        "signals, then an LSTM network uses these regime posterior probabilities alongside "
        "the original sensor features for multi-step prediction. We analyze the model across "
        "five dimensions: (A) HMM regime characterization, (B) gradient-based feature "
        "attribution, (C) LSTM hidden state structure, (D) temporal periodicity patterns, "
        "and (E) regime-conditioned predictive performance."
    )

    # ---- Performance Summary ----
    doc.add_heading("1. Performance Summary", level=1)

    table = doc.add_table(rows=1 + len(horizons), cols=5)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, label in enumerate(["Horizon", "RMSE (ppm)", "MAE (ppm)", "R2", "MAPE (%)"]):
        hdr_cells[i].text = label
    for row_idx, h in enumerate(horizons):
        m = all_metrics[h]
        cells = table.rows[row_idx + 1].cells
        cells[0].text = f"{h}h"
        cells[1].text = f"{m['rmse']:.2f}"
        cells[2].text = f"{m['mae']:.2f}"
        cells[3].text = f"{m['r2']:.4f}"
        cells[4].text = f"{m['mape']:.2f}"

    doc.add_paragraph("")

    for h in horizons:
        doc.add_heading(f"Horizon {h}h Analysis", level=1)

        # ---- Section A: HMM Regime Analysis ----
        doc.add_heading(f"A. HMM Regime Analysis ({h}h)", level=2)
        doc.add_paragraph(
            "The Hidden Markov Model component provides the most interpretable layer of "
            "the HMM-LSTM architecture. It discovers latent environmental regimes from three "
            "observable signals (CO2, Noise, TemperatureExt) and makes these regime beliefs "
            "available to the LSTM as additional input channels."
        )

        hmm_res = all_hmm_results.get(h, {})
        state_labels = hmm_res.get("state_labels", STATE_LABELS_DEFAULT)

        # State table
        doc.add_heading("A.1 State Characterization", level=3)
        means = hmm_res.get("means", [])
        if means:
            table = doc.add_table(rows=1 + len(means), cols=5)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for i, label in enumerate(["State", "Label", "CO2 (ppm)", "Noise (dB)", "Temp (C)"]):
                hdr[i].text = label
            for s_idx, row_means in enumerate(means):
                cells = table.rows[s_idx + 1].cells
                cells[0].text = f"S{s_idx}"
                cells[1].text = state_labels[s_idx] if s_idx < len(state_labels) else f"State {s_idx}"
                for f_idx in range(min(3, len(row_means))):
                    cells[f_idx + 2].text = f"{row_means[f_idx]:.1f}"

        doc.add_paragraph("")

        # State distribution
        for s in range(3):
            pct_key = f"state_{s}_pct"
            dur_key = f"state_{s}_median_duration"
            if pct_key in hmm_res:
                doc.add_paragraph(
                    f"State S{s} ({state_labels[s] if s < len(state_labels) else ''}): "
                    f"{hmm_res[pct_key]:.1f}% of test samples"
                    + (f", median duration {hmm_res[dur_key]:.0f}h" if dur_key in hmm_res else ""),
                    style="List Bullet",
                )

        doc.add_heading("A.2 Transition Matrix", level=3)
        transmat = hmm_res.get("transmat", [])
        if transmat:
            doc.add_paragraph(
                "The transition matrix reveals regime persistence (diagonal dominance) and "
                "transition pathways. High self-transition probabilities indicate stable regimes."
            )
            table = doc.add_table(rows=1 + len(transmat), cols=1 + len(transmat))
            table.style = "Light Grid Accent 1"
            table.rows[0].cells[0].text = "From\\To"
            for j in range(len(transmat)):
                table.rows[0].cells[j + 1].text = f"S{j}"
            for i, row in enumerate(transmat):
                table.rows[i + 1].cells[0].text = f"S{i}"
                for j, val in enumerate(row):
                    table.rows[i + 1].cells[j + 1].text = f"{val:.3f}"

        # Add figures
        fig_dir = output_dir.parent / f"h{h}" / "hmm_regime"
        fig_prefix = f"h{h}"

        figure_specs = [
            (f"A1_transition_matrix_h{h}.png", "Transition matrix heatmap"),
            (f"A2_emission_distributions_h{h}.png", "State emission Gaussian distributions per feature"),
            (f"A3_state_means_table_h{h}.png", "State means summary table"),
            (f"A4_co2_noise_scatter_h{h}.png", "CO2 vs Noise scatter with 95% confidence ellipses"),
            (f"A5_co2_state_timeline_h{h}.png", "Full CO2 time series with state-colored background"),
            (f"A6_zoomed_week_h{h}.png", "One-week zoomed view with Viterbi state sequence"),
            (f"A7_state_duration_hist_h{h}.png", "State duration distribution histogram"),
            (f"A8_posterior_probabilities_h{h}.png", "Posterior probability time series"),
            (f"A9_entropy_error_h{h}.png", "Posterior entropy vs prediction error correlation"),
        ]

        for fig_name, caption in figure_specs:
            fig_path = fig_dir / fig_name
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(6.0))
                last_par = doc.paragraphs[-1]
                last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(f"Figure: {caption}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)

        entropy_corr = hmm_res.get("entropy_error_correlation", None)
        if entropy_corr is not None:
            doc.add_heading("A.3 Entropy-Error Analysis", level=3)
            doc.add_paragraph(
                f"The correlation between posterior entropy and prediction error is "
                f"r = {entropy_corr:.3f}. "
                + ("A positive correlation indicates the model struggles when the HMM "
                   "is uncertain about the regime. "
                   if entropy_corr > 0.1 else
                   "The weak correlation suggests prediction difficulty is not strongly "
                   "related to regime ambiguity. ")
            )

        # ---- Section B: Gradient Attribution ----
        doc.add_heading(f"B. Gradient-Based Feature Attribution ({h}h)", level=2)
        doc.add_paragraph(
            "Input gradient saliency maps reveal which features and timesteps contribute "
            "most to the model's predictions. This analysis uniquely separates the original "
            "sensor features from the 3 HMM posterior probability channels."
        )

        b_figures = [
            (f"B1_gradient_heatmap_h{h}.png", "Full gradient attribution heatmap (lookback x 22 features)"),
            (f"B2_feature_ranking_h{h}.png", "Per-feature gradient importance with HMM channels highlighted"),
            (f"B3_temporal_gradient_split_h{h}.png", "Temporal gradient profile: sensor vs HMM channels"),
        ]
        b_dir = output_dir.parent / f"h{h}" / "gradient_attribution"
        for fig_name, caption in b_figures:
            fig_path = b_dir / fig_name
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(6.0))
                last_par = doc.paragraphs[-1]
                last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(f"Figure: {caption}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)

        # ---- Section C: Hidden State Analysis ----
        doc.add_heading(f"C. Hidden State Structural Analysis ({h}h)", level=2)
        doc.add_paragraph(
            "PCA on the LSTM's final hidden state h_n[-1] reveals how the model organizes "
            "its internal representations. Coloring by HMM regime (unique to this model) "
            "shows whether the LSTM naturally learns to separate states that correspond "
            "to the HMM's regime segmentation."
        )

        c_figures = [
            (f"C1_pca_variance_h{h}.png", "PCA explained variance of LSTM hidden states"),
            (f"C2_pca_co2_level_h{h}.png", "PCA scatter colored by CO2 concentration level"),
            (f"C3_pca_hmm_regime_h{h}.png", "PCA scatter colored by HMM regime (unique to HMM-LSTM)"),
            (f"C4_clustering_crosstab_h{h}.png", "K-Means clustering vs HMM state crosstab"),
        ]
        c_dir = output_dir.parent / f"h{h}" / "hidden_state_analysis"
        for fig_name, caption in c_figures:
            fig_path = c_dir / fig_name
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(6.0))
                last_par = doc.paragraphs[-1]
                last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(f"Figure: {caption}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)

        # ---- Section D: Temporal Patterns ----
        doc.add_heading(f"D. Temporal Pattern Analysis ({h}h)", level=2)
        doc.add_paragraph(
            "Frequency domain analysis reveals periodic structures in the model's internal "
            "representations and prediction errors. The HMM posterior ACF uniquely reveals "
            "the periodicity of regime transitions."
        )

        d_figures = [
            (f"D1_fft_hidden_pcs_h{h}.png", "FFT of hidden state principal components"),
            (f"D2_acf_hidden_pcs_h{h}.png", "Autocorrelation of hidden state PCs"),
            (f"D3_fft_residuals_h{h}.png", "FFT of prediction residuals"),
            (f"D4_rolling_rmse_h{h}.png", "Rolling RMSE over test set"),
            (f"D5_acf_hmm_posteriors_h{h}.png", "Autocorrelation of HMM posterior probabilities"),
        ]
        d_dir = output_dir.parent / f"h{h}" / "temporal_patterns"
        for fig_name, caption in d_figures:
            fig_path = d_dir / fig_name
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(6.0))
                last_par = doc.paragraphs[-1]
                last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(f"Figure: {caption}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)

        # ---- Section E: HMM-LSTM-Specific ----
        doc.add_heading(f"E. HMM-LSTM-Specific Analysis ({h}h)", level=2)

        specific_res = all_specific_results.get(h, {})

        # Regime-conditioned error
        doc.add_heading("E.1 Regime-Conditioned Error", level=3)
        doc.add_paragraph(
            "This analysis decomposes prediction error by HMM regime, revealing which "
            "environmental conditions are hardest to forecast."
        )

        regime_rows = []
        for s in range(3):
            rmse_key = f"rmse_state_{s}"
            count_key = f"count_state_{s}"
            if rmse_key in specific_res:
                regime_rows.append((s, specific_res[count_key], specific_res[rmse_key]))

        if regime_rows:
            table = doc.add_table(rows=1 + len(regime_rows), cols=3)
            table.style = "Light Grid Accent 1"
            for i, label in enumerate(["State", "N Samples", "RMSE (ppm)"]):
                table.rows[0].cells[i].text = label
            for row_idx, (s, cnt, rmse_val) in enumerate(regime_rows):
                cells = table.rows[row_idx + 1].cells
                cells[0].text = f"S{s}: {state_labels[s] if s < len(state_labels) else ''}"
                cells[1].text = str(cnt)
                cells[2].text = f"{rmse_val:.2f}"

        # Permutation importance
        doc.add_heading("E.2 HMM Channel Permutation Importance", level=3)
        perm_imp = specific_res.get("hmm_permutation_importance", {})
        if perm_imp:
            doc.add_paragraph(
                "Permutation importance measures the increase in MSE when HMM channels "
                "are randomly shuffled, quantifying the model's reliance on regime information."
            )
            for ch_name, mse_inc in perm_imp.items():
                doc.add_paragraph(
                    f"{ch_name}: MSE increase = {mse_inc:.4f}",
                    style="List Bullet",
                )

        e_figures = [
            (f"E1_regime_conditioned_error_h{h}.png", "RMSE by HMM regime state"),
            (f"E2_hmm_permutation_importance_h{h}.png", "HMM channel permutation importance"),
            (f"E3_predictions_overlay_h{h}.png", "Prediction overlay: actual vs predicted"),
            (f"E4_scatter_r2_h{h}.png", "Scatter plot with R2 metric"),
            (f"E5_residual_analysis_h{h}.png", "4-panel residual analysis"),
            (f"E6_error_by_co2_level_h{h}.png", "Error breakdown by CO2 concentration level"),
        ]
        e_dir = output_dir.parent / f"h{h}" / "hmm_lstm_specific"
        for fig_name, caption in e_figures:
            fig_path = e_dir / fig_name
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(6.0))
                last_par = doc.paragraphs[-1]
                last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(f"Figure: {caption}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)

    # ---- Avenues of Improvement ----
    doc.add_heading("Avenues of Improvement", level=1)

    improvements = [
        ("Number of HMM States", "The current model uses 3 states. A systematic "
         "search over 2-6 states using BIC/AIC could identify the optimal regime "
         "granularity. Too few states may merge distinct regimes, while too many "
         "may fragment coherent patterns."),
        ("HMM Feature Set", "Currently only CO2, Noise, and TemperatureExt drive "
         "the HMM. Adding humidity (Hrext) or pressure (Pression) could capture "
         "weather-related regime shifts. Time-of-day features could encode "
         "occupancy-driven regime patterns directly."),
        ("Online HMM Adaptation", "The HMM is fit once on training data. An online "
         "adaptation scheme (e.g., incremental EM) could track regime drift as "
         "building usage patterns evolve over seasons."),
        ("Regime-Aware LSTM Architecture", "Instead of simply appending posterior "
         "probabilities, the model could use separate LSTM heads per regime and "
         "blend outputs via the posterior, implementing a mixture-of-experts approach."),
        ("HMM State Duration Modeling", "The geometric duration distribution implicit "
         "in standard HMMs may not match real occupancy patterns. A Hidden Semi-Markov "
         "Model (HSMM) with explicit duration distributions could improve regime "
         "persistence modeling."),
        ("Longer Horizons", "The 24h horizon performance is typically poor for all "
         "models. Regime information could be particularly valuable at long horizons "
         "if the HMM captures diurnal occupancy cycles that persist predictably."),
        ("Multi-Resolution Approach", "Training HMMs at different temporal resolutions "
         "(hourly, daily) and combining their posterior probabilities could capture "
         "both short-term and long-term regime dynamics."),
    ]

    for title_text, desc in improvements:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(desc)

    # ---- Save ----
    report_path = output_dir / "hmm_lstm_interpretability_report.docx"
    doc.save(str(report_path))
    print(f"  DOCX report saved to: {report_path}")


# ======================================================================
#  Main
# ======================================================================

def main() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )

    parser = argparse.ArgumentParser(
        description="HMM-LSTM Interpretability Study on preproc_D (Enhanced 1h)"
    )
    parser.add_argument(
        "--horizons", nargs="+", type=int, default=[1, 24],
        help="Forecast horizons in hours (default: 1 24)",
    )
    parser.add_argument(
        "--epochs", type=int, default=None,
        help="Override max epochs for training",
    )
    args = parser.parse_args()

    print(f"\n{'='*70}")
    print(f"  HMM-LSTM INTERPRETABILITY STUDY")
    print(f"  Variant: preproc_D (Enhanced 1h)")
    print(f"  Horizons: {args.horizons}")
    print(f"  Sections: A(HMM regime) B(gradients) C(hidden) D(temporal) E(HMM-specific)")
    print(f"{'='*70}\n")

    # Load pipeline data once
    pipeline_config = load_interpretability_config(horizon=1, epochs_override=args.epochs)
    seed_everything(pipeline_config["training"]["seed"])

    raw_dir = Path(pipeline_config["data"].get("raw_dir", "data/raw"))
    print("  Loading preprocessing pipeline (preproc_D)...")
    train_df, val_df, test_df = run_preprocessing_pipeline(
        raw_dir=raw_dir, variant_config=pipeline_config,
    )
    print(f"  Pipeline loaded: train={len(train_df)}, val={len(val_df)}, test={len(test_df)}")

    all_metrics: dict[int, dict] = {}
    all_hmm_results: dict[int, dict] = {}
    all_specific_results: dict[int, dict] = {}

    for horizon in args.horizons:
        print(f"\n{'-'*60}")
        print(f"  HORIZON: {horizon}h")
        print(f"{'-'*60}\n")

        config = load_interpretability_config(horizon=horizon, epochs_override=args.epochs)
        seed_everything(config["training"]["seed"])

        output_base = RESULTS_BASE / f"h{horizon}"
        output_base.mkdir(parents=True, exist_ok=True)

        # ---- Train HMM-LSTM ----
        t0 = time.time()
        print(f"  Training HMM-LSTM for {horizon}h horizon...")
        (
            best_model, dm, detector, y_true, y_pred,
            aug_train_df, aug_val_df, aug_test_df,
        ) = train_hmm_lstm(config, train_df.copy(), val_df.copy(), test_df.copy(), horizon)
        elapsed = time.time() - t0
        print(f"  Training completed in {elapsed:.1f}s")

        # ---- Compute metrics (on original ppm scale) ----
        y_true_ppm = inverse_scale_target(
            y_true[:, 0] if y_true.ndim == 2 else y_true.ravel(),
            dm.target_scaler,
        )
        y_pred_ppm = inverse_scale_target(
            y_pred[:, 0] if y_pred.ndim == 2 else y_pred.ravel(),
            dm.target_scaler,
        )
        metrics = compute_metrics(y_true_ppm, y_pred_ppm)
        all_metrics[horizon] = metrics
        print(f"  RMSE={metrics['rmse']:.2f}  MAE={metrics['mae']:.2f}  "
              f"R2={metrics['r2']:.4f}  MAPE={metrics['mape']:.2f}%")

        # ---- Section A: HMM Regime Analysis ----
        hmm_results = run_hmm_regime_analysis(
            detector, test_df, train_df, y_true, y_pred,
            horizon, output_base / "hmm_regime",
        )
        all_hmm_results[horizon] = hmm_results

        # ---- Section B: Gradient Attribution ----
        run_gradient_attribution(
            best_model, dm, config, horizon,
            output_base / "gradient_attribution",
        )

        # ---- Section C: Hidden State Analysis ----
        run_hidden_state_analysis(
            best_model, dm, detector, test_df, config, horizon,
            output_base / "hidden_state_analysis",
        )

        # ---- Section D: Temporal Patterns ----
        run_temporal_patterns(
            best_model, dm, detector, test_df, y_true, y_pred,
            horizon, output_base / "temporal_patterns",
        )

        # ---- Section E: HMM-LSTM-Specific Analysis ----
        specific_results = run_hmm_lstm_specific_analysis(
            best_model, dm, detector, test_df, config,
            y_true, y_pred, metrics, horizon,
            output_base / "hmm_lstm_specific",
        )
        all_specific_results[horizon] = specific_results

        # ---- Save metrics + predictions ----
        save_metrics(
            metrics, f"HMM-LSTM_h{horizon}", output_base / "metrics.json",
            experiment_info={
                "name": "hmm_lstm_interpretability",
                "label": f"HMM-LSTM Deep Analysis h={horizon}",
                "description": "preproc_D Enhanced 1h variant",
            },
        )
        np.savez(output_base / "predictions.npz",
                 y_true=y_true, y_pred=y_pred)

        # ---- Save study results JSON ----
        study_data = {
            "study": "HMM-LSTM Interpretability",
            "horizon": horizon,
            "variant": "preproc_D (Enhanced 1h)",
            "timestamp": datetime.now().isoformat(),
            "metrics": metrics,
            "hmm_results": {k: v for k, v in hmm_results.items()
                           if not isinstance(v, np.ndarray)},
            "specific_results": {k: (v if not isinstance(v, np.ndarray) else str(v))
                                 for k, v in specific_results.items()},
        }
        with open(output_base / "study_results.json", "w", encoding="utf-8") as f:
            json.dump(study_data, f, indent=2, default=str)

        # ---- GPU cleanup ----
        del best_model, dm
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        print(f"  GPU memory freed\n")

    # ---- DOCX Report ----
    print(f"\n{'-'*60}")
    print(f"  GENERATING DOCX REPORT")
    print(f"{'-'*60}\n")

    generate_docx_report(
        all_metrics=all_metrics,
        all_hmm_results=all_hmm_results,
        all_specific_results=all_specific_results,
        horizons=args.horizons,
        output_dir=RESULTS_BASE,
    )

    # ---- Final summary ----
    print(f"\n{'='*70}")
    print(f"  HMM-LSTM INTERPRETABILITY STUDY COMPLETE")
    print(f"  Results saved to: {RESULTS_BASE}")
    print(f"{'='*70}")
    print(f"  Performance Summary:")
    for h in args.horizons:
        m = all_metrics[h]
        print(f"    {h}h: RMSE={m['rmse']:.2f}, MAE={m['mae']:.2f}, "
              f"R2={m['r2']:.4f}, MAPE={m['mape']:.2f}%")
    print(f"  Report: {RESULTS_BASE / 'hmm_lstm_interpretability_report.docx'}")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
