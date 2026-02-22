"""SARIMA Interpretability Study: Deep analysis of statistical time series model.

Performs comprehensive interpretability analysis on preproc_D (Enhanced 1h) data:
  A) Model parameter analysis (coefficients, roots, significance)
  B) Impulse response analysis (IRF, cumulative response)
  C) Residual diagnostics (ACF, PACF, Ljung-Box, heteroscedasticity)
  D) Forecast uncertainty analysis (fan chart, CI coverage, calibration)
  E) Seasonal decomposition comparison (STL, differenced series)
  F) SARIMA-specific analyses (stationarity tests, predictions, metrics)

Generates a DOCX academic report with all figures and quantitative analysis.

Usage:
    python -u scripts/run_sarima_interpretability.py
    python -u scripts/run_sarima_interpretability.py --horizons 1
    python -u scripts/run_sarima_interpretability.py --horizons 1 24
"""

import argparse
import gc
import json
import logging
import sys
import time
import warnings
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.gridspec import GridSpec
from scipy import stats
from scipy.fft import fft, fftfreq
from statsmodels.tsa.statespace.sarimax import SARIMAX
from statsmodels.tsa.stattools import acf, pacf, adfuller, kpss
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
from statsmodels.stats.diagnostic import acorr_ljungbox
from statsmodels.tsa.seasonal import STL

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from src.data.datamodule import CO2DataModule
from src.data.pipeline import run_preprocessing_pipeline
from src.data.preprocessing import inverse_scale_target
from src.evaluation.metrics import compute_metrics, save_metrics
from src.models.sarima import SARIMAForecaster
from src.utils.config import load_config
from src.utils.seed import seed_everything

logger = logging.getLogger(__name__)

plt.style.use("seaborn-v0_8-whitegrid")

RESULTS_BASE = Path("results/sarima_interpretability")

# Color palette (consistent with other studies)
C_PRIMARY = "#2196F3"
C_SECONDARY = "#FF5722"
C_ACCENT = "#4CAF50"
C_WARN = "#FFC107"
C_NEUTRAL = "#607D8B"

# Maximum test windows for batch prediction (SARIMA is slow per window)
MAX_PRED_WINDOWS = 400


# ======================================================================
#  Configuration
# ======================================================================

def load_interpretability_config(
    horizon: int,
) -> dict:
    """Load merged config for SARIMA + preproc_D + specified horizon.

    Args:
        horizon: Forecast horizon in hours (1 or 24).

    Returns:
        Merged configuration dictionary.
    """
    config_files = [
        str(PROJECT_ROOT / "configs" / "training.yaml"),
        str(PROJECT_ROOT / "configs" / "data.yaml"),
        str(PROJECT_ROOT / "configs" / "sarima.yaml"),
        str(PROJECT_ROOT / "configs" / "experiments" / "preproc_D_enhanced_1h.yaml"),
    ]
    config = load_config(config_files)
    config["data"]["forecast_horizon_hours"] = horizon
    # Force correct seasonal period for 1h data: 24 hours/day
    config["model"]["seasonal_order"] = [1, 1, 1, 24]
    config["training"]["results_dir"] = str(
        RESULTS_BASE / f"h{horizon}" / "training_runs"
    )
    return config


# ======================================================================
#  SARIMA Training and Prediction Extraction
# ======================================================================

def train_sarima(
    config: dict,
    train_df: pd.DataFrame,
    val_df: pd.DataFrame,
    test_df: pd.DataFrame,
    horizon: int,
) -> tuple:
    """Fit SARIMA, extract predictions, and compute metrics.

    Args:
        config: Merged configuration dictionary.
        train_df: Training DataFrame from pipeline.
        val_df: Validation DataFrame from pipeline.
        test_df: Test DataFrame from pipeline.
        horizon: Forecast horizon in hours.

    Returns:
        Tuple of (sarima_model, data_module, results_dict).
        results_dict has keys: y_true, y_pred, metrics, train_series.
    """
    # Build DataModule for consistent scaling and windowing
    dm = CO2DataModule(config)
    dm.build_datasets(train_df, val_df, test_df)

    # Extract the scaled training series for SARIMA fitting
    # CO2 is the target column, which is the last column after feature_columns
    train_series = train_df["CO2"].values
    # We need the scaled version for consistency with the test windows
    # The target scaler was fit during build_datasets
    train_series_scaled = dm.target_scaler.transform(
        train_series.reshape(-1, 1)
    ).ravel()

    # Fit SARIMA on the scaled continuous training series
    sarima = SARIMAForecaster(config)
    sarima.fit(train_series_scaled)

    # Get test windows
    test_X = dm.test_dataset.X.numpy()  # (n_samples, lookback, features)
    test_y = dm.test_dataset.y.numpy()  # (n_samples, horizon)

    # Limit number of prediction windows for speed
    n_samples = min(test_X.shape[0], MAX_PRED_WINDOWS)
    # Use evenly spaced indices for representative coverage
    if test_X.shape[0] > MAX_PRED_WINDOWS:
        indices = np.linspace(0, test_X.shape[0] - 1, MAX_PRED_WINDOWS, dtype=int)
        test_X_sub = test_X[indices]
        test_y_sub = test_y[indices]
    else:
        test_X_sub = test_X
        test_y_sub = test_y
        indices = np.arange(test_X.shape[0])

    print(f"  Predicting on {n_samples} test windows (of {test_X.shape[0]} total)...")
    y_pred_scaled = sarima.predict_batch(test_X_sub, target_idx=-1)

    # Inverse scale to original CO2 ppm
    y_pred = inverse_scale_target(y_pred_scaled, dm.target_scaler)
    y_true = inverse_scale_target(test_y_sub, dm.target_scaler)

    # Compute metrics
    metrics = compute_metrics(y_true.ravel(), y_pred.ravel())

    results = {
        "y_true": y_true,
        "y_pred": y_pred,
        "metrics": metrics,
        "train_series_scaled": train_series_scaled,
        "test_indices": indices,
    }

    return sarima, dm, results


# ======================================================================
#  Section A: Model Parameter Analysis
# ======================================================================

def run_parameter_analysis(
    sarima: SARIMAForecaster,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Extract and visualize SARIMA model parameters.

    Analyzes fitted coefficients, p-values, confidence intervals,
    and AR/MA polynomial roots in the complex plane.

    Args:
        sarima: Fitted SARIMAForecaster with result_ attribute.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.

    Returns:
        Dictionary of parameter statistics for the report.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [A] Extracting model parameters...")

    result = sarima.result_
    if result is None:
        print("  [A] WARNING: No fitted result available, skipping.")
        return {}

    param_stats = {}

    # ------------------------------------------------------------------
    # A1: Parameter summary extraction
    # ------------------------------------------------------------------
    try:
        params = result.params
        pvalues = result.pvalues
        bse = result.bse  # standard errors
        conf_int = result.conf_int(alpha=0.05)
        param_names = result.param_names if hasattr(result, "param_names") else [
            f"param_{i}" for i in range(len(params))
        ]

        # Build parameter table data
        param_table = []
        for i, name in enumerate(param_names):
            entry = {
                "name": name,
                "value": float(params[i]),
                "std_error": float(bse[i]) if i < len(bse) else float("nan"),
                "pvalue": float(pvalues[i]) if i < len(pvalues) else float("nan"),
                "ci_lower": float(conf_int[i, 0]) if i < conf_int.shape[0] else float("nan"),
                "ci_upper": float(conf_int[i, 1]) if i < conf_int.shape[0] else float("nan"),
                "significant": bool(pvalues[i] < 0.05) if i < len(pvalues) else False,
            }
            param_table.append(entry)
            param_stats[name] = entry

        # A2: Coefficient bar chart with CIs
        fig, ax = plt.subplots(figsize=(10, 6))
        names = [e["name"] for e in param_table]
        values = [e["value"] for e in param_table]
        ci_low = [e["ci_lower"] for e in param_table]
        ci_high = [e["ci_upper"] for e in param_table]
        errors_low = [v - cl for v, cl in zip(values, ci_low)]
        errors_high = [ch - v for v, ch in zip(values, ci_high)]

        colors = [C_PRIMARY if e["significant"] else C_NEUTRAL for e in param_table]
        y_pos = np.arange(len(names))

        ax.barh(y_pos, values, xerr=[errors_low, errors_high],
                color=colors, capsize=4, edgecolor="black", linewidth=0.5)
        ax.set_yticks(y_pos)
        ax.set_yticklabels(names, fontsize=9)
        ax.axvline(x=0, color="black", linewidth=0.8, linestyle="--")
        ax.set_xlabel("Coefficient Value")
        ax.set_title(f"SARIMA Parameter Estimates with 95% CI (h={horizon}h)")

        # Legend for significance
        from matplotlib.patches import Patch
        legend_elements = [
            Patch(facecolor=C_PRIMARY, label="Significant (p<0.05)"),
            Patch(facecolor=C_NEUTRAL, label="Not significant"),
        ]
        ax.legend(handles=legend_elements, loc="lower right", fontsize=8)

        plt.tight_layout()
        plt.savefig(output_dir / f"coefficient_chart_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

        # A3: AR/MA polynomial root plots
        _plot_roots(result, sarima, horizon, output_dir)

        # A4: Physical interpretation table (as figure)
        _plot_param_summary_table(param_table, horizon, output_dir)

        # Store model info
        param_stats["_aic"] = float(result.aic)
        param_stats["_bic"] = float(result.bic)
        param_stats["_loglik"] = float(result.llf)
        param_stats["_n_obs"] = int(result.nobs)

        print(f"    AIC={result.aic:.2f}  BIC={result.bic:.2f}  LogLik={result.llf:.2f}")

    except Exception as e:
        print(f"  [A] Error in parameter analysis: {e}")

    return param_stats


def _plot_roots(result, sarima, horizon: int, output_dir: Path) -> None:
    """Plot AR and MA polynomial roots in the complex plane."""
    try:
        fig, axes = plt.subplots(1, 2, figsize=(12, 5))

        # Unit circle
        theta = np.linspace(0, 2 * np.pi, 200)
        for ax in axes:
            ax.plot(np.cos(theta), np.sin(theta), "k--", alpha=0.3, linewidth=0.8)
            ax.axhline(y=0, color="k", linewidth=0.5)
            ax.axvline(x=0, color="k", linewidth=0.5)
            ax.set_aspect("equal")

        # AR roots
        try:
            ar_params = result.arparams
            if len(ar_params) > 0:
                ar_poly = np.r_[1, -ar_params]
                ar_roots = np.roots(ar_poly)
                axes[0].scatter(ar_roots.real, ar_roots.imag,
                                c=C_PRIMARY, s=100, zorder=5, marker="o",
                                edgecolors="black", linewidth=0.5)
                for r in ar_roots:
                    axes[0].annotate(f"|r|={abs(r):.3f}",
                                     (r.real, r.imag),
                                     textcoords="offset points",
                                     xytext=(5, 5), fontsize=7)
            else:
                axes[0].text(0, 0, "No AR params", ha="center", va="center")
        except Exception:
            axes[0].text(0, 0, "AR roots unavailable", ha="center", va="center")

        axes[0].set_title(f"AR Polynomial Roots (h={horizon}h)")
        axes[0].set_xlabel("Real")
        axes[0].set_ylabel("Imaginary")

        # MA roots
        try:
            ma_params = result.maparams
            if len(ma_params) > 0:
                ma_poly = np.r_[1, ma_params]
                ma_roots = np.roots(ma_poly)
                axes[1].scatter(ma_roots.real, ma_roots.imag,
                                c=C_SECONDARY, s=100, zorder=5, marker="s",
                                edgecolors="black", linewidth=0.5)
                for r in ma_roots:
                    axes[1].annotate(f"|r|={abs(r):.3f}",
                                     (r.real, r.imag),
                                     textcoords="offset points",
                                     xytext=(5, 5), fontsize=7)
            else:
                axes[1].text(0, 0, "No MA params", ha="center", va="center")
        except Exception:
            axes[1].text(0, 0, "MA roots unavailable", ha="center", va="center")

        axes[1].set_title(f"MA Polynomial Roots (h={horizon}h)")
        axes[1].set_xlabel("Real")
        axes[1].set_ylabel("Imaginary")

        fig.suptitle(
            "Roots should lie outside unit circle for stationarity/invertibility",
            fontsize=9, style="italic",
        )
        plt.tight_layout()
        plt.savefig(output_dir / f"root_plot_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    except Exception as e:
        print(f"    Root plot error: {e}")


def _plot_param_summary_table(param_table: list, horizon: int, output_dir: Path) -> None:
    """Render parameter summary as a matplotlib table figure."""
    try:
        fig, ax = plt.subplots(figsize=(12, max(3, len(param_table) * 0.5 + 1)))
        ax.axis("off")

        col_labels = ["Parameter", "Value", "Std Err", "p-value", "95% CI", "Sig?"]
        cell_data = []
        for e in param_table:
            sig_str = "Yes" if e["significant"] else "No"
            ci_str = f"[{e['ci_lower']:.4f}, {e['ci_upper']:.4f}]"
            cell_data.append([
                e["name"],
                f"{e['value']:.6f}",
                f"{e['std_error']:.6f}",
                f"{e['pvalue']:.4f}",
                ci_str,
                sig_str,
            ])

        tbl = ax.table(
            cellText=cell_data, colLabels=col_labels,
            loc="center", cellLoc="center",
        )
        tbl.auto_set_font_size(False)
        tbl.set_fontsize(9)
        tbl.scale(1.0, 1.6)

        # Color significant rows
        for i, e in enumerate(param_table):
            if e["significant"]:
                for j in range(len(col_labels)):
                    tbl[i + 1, j].set_facecolor("#E3F2FD")

        ax.set_title(f"SARIMA Parameter Summary (h={horizon}h)", fontsize=12, pad=20)
        plt.savefig(output_dir / f"param_summary_table_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    except Exception as e:
        print(f"    Parameter table error: {e}")


# ======================================================================
#  Section B: Impulse Response Analysis
# ======================================================================

def run_impulse_response(
    sarima: SARIMAForecaster,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Compute and visualize impulse response function.

    Args:
        sarima: Fitted SARIMAForecaster.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.

    Returns:
        Dictionary with IRF data.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [B] Computing impulse response function...")

    result = sarima.result_
    if result is None:
        print("  [B] WARNING: No fitted result, skipping.")
        return {}

    irf_data = {}
    irf_steps = 72  # 72 hours = 3 days

    # ------------------------------------------------------------------
    # B1: Impulse Response Function
    # ------------------------------------------------------------------
    try:
        irf = result.impulse_responses(steps=irf_steps)
        irf_values = np.array(irf).ravel()
    except Exception:
        # Fallback: manual computation from MA representation
        try:
            print("    impulse_responses() unavailable, computing manually...")
            # Get the MA infinity representation via simulation
            irf_values = _compute_irf_manual(result, irf_steps)
        except Exception as e2:
            print(f"    Manual IRF also failed: {e2}")
            irf_values = None

    if irf_values is not None:
        irf_data["irf"] = irf_values.tolist()

        # Plot IRF
        fig, axes = plt.subplots(2, 1, figsize=(12, 8))

        # B1: IRF
        axes[0].stem(
            range(len(irf_values)), irf_values,
            linefmt=C_PRIMARY, markerfmt="o", basefmt="k-",
        )
        axes[0].axhline(y=0, color="black", linewidth=0.5)
        axes[0].set_xlabel("Lag (hours)")
        axes[0].set_ylabel("Response")
        axes[0].set_title(f"Impulse Response Function (h={horizon}h)")

        # B2: Cumulative impulse response (step response)
        cumulative_irf = np.cumsum(irf_values)
        irf_data["cumulative_irf"] = cumulative_irf.tolist()

        axes[1].plot(range(len(cumulative_irf)), cumulative_irf,
                     color=C_SECONDARY, linewidth=2)
        axes[1].fill_between(range(len(cumulative_irf)), 0, cumulative_irf,
                             alpha=0.2, color=C_SECONDARY)
        axes[1].axhline(y=0, color="black", linewidth=0.5)
        axes[1].set_xlabel("Lag (hours)")
        axes[1].set_ylabel("Cumulative Response")
        axes[1].set_title(f"Cumulative Impulse Response / Step Response (h={horizon}h)")

        plt.tight_layout()
        plt.savefig(output_dir / f"impulse_response_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

        # B3: Simplified variance decomposition
        _plot_variance_decomposition(irf_values, horizon, output_dir)

        print(f"    IRF computed for {len(irf_values)} steps")
    else:
        print("    IRF computation failed entirely, skipping section B plots.")

    return irf_data


def _compute_irf_manual(result, steps: int) -> np.ndarray:
    """Compute IRF manually by simulating a unit shock."""
    # Create a zero series of sufficient length, apply a unit shock
    # then use the model to forecast the response
    n_warmup = 200
    base = np.zeros(n_warmup)

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        # Baseline forecast from zero history
        model_base = SARIMAX(
            base,
            order=tuple(result.specification["order"]),
            seasonal_order=tuple(result.specification["seasonal_order"]),
            enforce_stationarity=False,
            enforce_invertibility=False,
        )
        res_base = model_base.filter(result.params)
        fc_base = res_base.forecast(steps=steps)

        # Shocked series: unit impulse at the last timestep
        shocked = base.copy()
        shocked[-1] = 1.0
        model_shock = SARIMAX(
            shocked,
            order=tuple(result.specification["order"]),
            seasonal_order=tuple(result.specification["seasonal_order"]),
            enforce_stationarity=False,
            enforce_invertibility=False,
        )
        res_shock = model_shock.filter(result.params)
        fc_shock = res_shock.forecast(steps=steps)

    irf = fc_shock - fc_base
    return np.array(irf).ravel()


def _plot_variance_decomposition(irf_values: np.ndarray, horizon: int, output_dir: Path) -> None:
    """Plot forecast error variance decomposition from IRF."""
    try:
        # FEVD: proportion of forecast variance explained at each horizon
        cumulative_sq = np.cumsum(irf_values ** 2)
        total_var = cumulative_sq[-1] if cumulative_sq[-1] > 0 else 1.0
        fevd_pct = (cumulative_sq / total_var) * 100

        fig, ax = plt.subplots(figsize=(10, 5))
        ax.plot(range(len(fevd_pct)), fevd_pct, color=C_ACCENT, linewidth=2)
        ax.fill_between(range(len(fevd_pct)), 0, fevd_pct,
                        alpha=0.2, color=C_ACCENT)
        ax.axhline(y=50, color=C_WARN, linewidth=1, linestyle="--", label="50%")
        ax.axhline(y=90, color=C_SECONDARY, linewidth=1, linestyle="--", label="90%")
        ax.set_xlabel("Forecast Horizon (hours)")
        ax.set_ylabel("Cumulative Variance Explained (%)")
        ax.set_title(f"Forecast Error Variance Decomposition (h={horizon}h)")
        ax.set_ylim(0, 105)
        ax.legend()

        plt.tight_layout()
        plt.savefig(output_dir / f"variance_decomposition_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    except Exception as e:
        print(f"    Variance decomposition error: {e}")


# ======================================================================
#  Section C: Residual Diagnostics
# ======================================================================

def run_residual_diagnostics(
    sarima: SARIMAForecaster,
    dm: CO2DataModule,
    results: dict,
    test_df: pd.DataFrame,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Comprehensive residual diagnostic analysis.

    Args:
        sarima: Fitted SARIMAForecaster.
        dm: Data module with scaler information.
        results: Results dict with y_true, y_pred.
        test_df: Test DataFrame with datetime column.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.

    Returns:
        Dictionary with residual statistics.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [C] Running residual diagnostics...")

    result = sarima.result_
    resid_stats = {}

    # Use in-sample residuals from the fitted model for diagnostics
    # These are the one-step-ahead prediction errors on the training data
    try:
        in_sample_resids = result.resid.copy() if result is not None else None
    except Exception:
        in_sample_resids = None

    # Also compute out-of-sample residuals from test predictions
    y_true = results["y_true"]
    y_pred = results["y_pred"]
    oos_residuals = (y_true - y_pred).ravel()

    # Use in-sample residuals if available, otherwise use OOS
    if in_sample_resids is not None and len(in_sample_resids) > 100:
        diag_resids = in_sample_resids
        resid_label = "In-Sample"
    else:
        diag_resids = oos_residuals
        resid_label = "Out-of-Sample"

    sigma = np.std(diag_resids)
    resid_stats["mean"] = float(np.mean(diag_resids))
    resid_stats["std"] = float(sigma)
    resid_stats["skewness"] = float(stats.skew(diag_resids))
    resid_stats["kurtosis"] = float(stats.kurtosis(diag_resids))
    resid_stats["n_residuals"] = len(diag_resids)
    resid_stats["source"] = resid_label

    # ------------------------------------------------------------------
    # C1: Residual time series with +/-2*sigma bands
    # ------------------------------------------------------------------
    try:
        fig, ax = plt.subplots(figsize=(14, 4))
        ax.plot(diag_resids, color=C_PRIMARY, linewidth=0.5, alpha=0.7)
        ax.axhline(y=0, color="black", linewidth=0.8)
        ax.axhline(y=2 * sigma, color=C_SECONDARY, linewidth=1, linestyle="--",
                    label=f"+2*sigma ({2*sigma:.4f})")
        ax.axhline(y=-2 * sigma, color=C_SECONDARY, linewidth=1, linestyle="--",
                    label=f"-2*sigma ({-2*sigma:.4f})")
        pct_outside = np.mean(np.abs(diag_resids) > 2 * sigma) * 100
        ax.set_title(
            f"{resid_label} Residuals (h={horizon}h) - "
            f"{pct_outside:.1f}% outside 2-sigma bands"
        )
        ax.set_xlabel("Observation Index")
        ax.set_ylabel("Residual")
        ax.legend(fontsize=8)
        plt.tight_layout()
        plt.savefig(output_dir / f"residual_timeseries_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
        resid_stats["pct_outside_2sigma"] = float(pct_outside)
    except Exception as e:
        print(f"    C1 error: {e}")

    # ------------------------------------------------------------------
    # C2: Residual distribution (histogram + KDE + Q-Q)
    # ------------------------------------------------------------------
    try:
        fig, axes = plt.subplots(1, 2, figsize=(12, 5))

        # Histogram + KDE
        axes[0].hist(diag_resids, bins=80, density=True, alpha=0.6,
                      color=C_PRIMARY, edgecolor="white", linewidth=0.3)
        x_range = np.linspace(diag_resids.min(), diag_resids.max(), 200)
        kde = stats.gaussian_kde(diag_resids)
        axes[0].plot(x_range, kde(x_range), color=C_SECONDARY, linewidth=2,
                      label="KDE")
        # Normal fit
        norm_pdf = stats.norm.pdf(x_range, np.mean(diag_resids), sigma)
        axes[0].plot(x_range, norm_pdf, color=C_ACCENT, linewidth=2,
                      linestyle="--", label="Normal fit")
        axes[0].set_xlabel("Residual")
        axes[0].set_ylabel("Density")
        axes[0].set_title(f"Residual Distribution (h={horizon}h)")
        axes[0].legend(fontsize=8)

        # Q-Q plot
        stats.probplot(diag_resids, dist="norm", plot=axes[1])
        axes[1].set_title(f"Q-Q Plot (h={horizon}h)")
        axes[1].get_lines()[0].set_color(C_PRIMARY)
        axes[1].get_lines()[0].set_markersize(3)
        axes[1].get_lines()[1].set_color(C_SECONDARY)

        # Shapiro-Wilk test (on subset if too many samples)
        sw_sample = diag_resids[:5000] if len(diag_resids) > 5000 else diag_resids
        sw_stat, sw_pval = stats.shapiro(sw_sample)
        resid_stats["shapiro_stat"] = float(sw_stat)
        resid_stats["shapiro_pval"] = float(sw_pval)
        axes[1].text(
            0.05, 0.95,
            f"Shapiro-Wilk: W={sw_stat:.4f}, p={sw_pval:.4f}",
            transform=axes[1].transAxes, fontsize=8,
            verticalalignment="top",
            bbox=dict(boxstyle="round", facecolor="wheat", alpha=0.5),
        )

        plt.tight_layout()
        plt.savefig(output_dir / f"residual_distribution_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    C2 error: {e}")

    # ------------------------------------------------------------------
    # C3/C4: ACF and PACF of residuals
    # ------------------------------------------------------------------
    try:
        fig, axes = plt.subplots(2, 1, figsize=(12, 8))
        max_lags = min(72, len(diag_resids) // 4)
        plot_acf(diag_resids, ax=axes[0], lags=max_lags, alpha=0.05,
                 title=f"ACF of Residuals (h={horizon}h)")
        plot_pacf(diag_resids, ax=axes[1], lags=max_lags, alpha=0.05,
                  method="ywm",
                  title=f"PACF of Residuals (h={horizon}h)")
        plt.tight_layout()
        plt.savefig(output_dir / f"acf_pacf_residuals_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    C3/C4 error: {e}")

    # ------------------------------------------------------------------
    # C5: Ljung-Box test
    # ------------------------------------------------------------------
    try:
        test_lags = [6, 12, 24, 36, 48]
        lb_results = acorr_ljungbox(diag_resids, lags=test_lags, return_df=True)

        fig, ax = plt.subplots(figsize=(10, 5))
        ax.bar(
            range(len(test_lags)),
            lb_results["lb_pvalue"].values,
            color=[C_ACCENT if p > 0.05 else C_SECONDARY
                   for p in lb_results["lb_pvalue"].values],
            edgecolor="black", linewidth=0.5,
        )
        ax.axhline(y=0.05, color="red", linewidth=1.5, linestyle="--",
                    label="alpha=0.05")
        ax.set_xticks(range(len(test_lags)))
        ax.set_xticklabels([str(l) for l in test_lags])
        ax.set_xlabel("Lag")
        ax.set_ylabel("p-value")
        ax.set_title(f"Ljung-Box Test p-values (h={horizon}h)")
        ax.legend()

        # Annotate p-values
        for i, (lag, pv) in enumerate(zip(test_lags, lb_results["lb_pvalue"].values)):
            ax.text(i, pv + 0.02, f"{pv:.4f}", ha="center", fontsize=8)

        plt.tight_layout()
        plt.savefig(output_dir / f"ljungbox_test_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

        resid_stats["ljungbox"] = {
            str(l): float(p)
            for l, p in zip(test_lags, lb_results["lb_pvalue"].values)
        }
    except Exception as e:
        print(f"    C5 error: {e}")

    # ------------------------------------------------------------------
    # C6: Squared residual ACF (heteroscedasticity check)
    # ------------------------------------------------------------------
    try:
        sq_resids = diag_resids ** 2
        fig, ax = plt.subplots(figsize=(12, 4))
        max_lags_sq = min(48, len(sq_resids) // 4)
        plot_acf(sq_resids, ax=ax, lags=max_lags_sq, alpha=0.05,
                 title=f"ACF of Squared Residuals - Heteroscedasticity Check (h={horizon}h)")
        plt.tight_layout()
        plt.savefig(output_dir / f"squared_resid_acf_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

        # Check if significant autocorrelation exists in squared residuals
        sq_acf_vals = acf(sq_resids, nlags=max_lags_sq, alpha=0.05)
        if isinstance(sq_acf_vals, tuple):
            acf_vals_arr = sq_acf_vals[0]
            ci_arr = sq_acf_vals[1]
            # Count lags with significant autocorrelation (outside CI)
            n_sig = 0
            for lag_i in range(1, len(acf_vals_arr)):
                if lag_i < len(ci_arr):
                    if acf_vals_arr[lag_i] < ci_arr[lag_i, 0] or acf_vals_arr[lag_i] > ci_arr[lag_i, 1]:
                        n_sig += 1
            resid_stats["sq_resid_sig_lags"] = n_sig
            resid_stats["heteroscedastic"] = n_sig > 3
    except Exception as e:
        print(f"    C6 error: {e}")

    # ------------------------------------------------------------------
    # C7/C8: Residuals by hour of day and day of week (OOS only)
    # ------------------------------------------------------------------
    try:
        # For OOS residuals with datetime alignment
        test_indices = results.get("test_indices", None)
        datetime_col = "datetime"
        lookback_hours = dm.lookback_steps  # already in steps (=hours for 1h data)

        if datetime_col in test_df.columns and test_indices is not None:
            # Each test window starts at lookback offset
            # The prediction target starts at index (lookback + window_idx)
            test_dates = pd.DatetimeIndex(test_df[datetime_col])
            # Align dates to the prediction window targets
            n_pred = len(oos_residuals)
            # Use stride=1, so window i targets date at (lookback + i)
            pred_dates = []
            for idx in test_indices:
                date_idx = lookback_hours + idx
                if date_idx < len(test_dates):
                    pred_dates.append(test_dates[date_idx])
                else:
                    pred_dates.append(pd.NaT)

            # We need to flatten: each window has horizon steps
            # For simplicity, use mean residual per window
            horizon_steps = y_true.shape[1] if y_true.ndim > 1 else 1
            resid_per_window = (y_true - y_pred).mean(axis=1).ravel() if y_true.ndim > 1 else oos_residuals

            pred_dates_clean = pd.DatetimeIndex([d for d in pred_dates if not pd.isna(d)])
            n_use = min(len(pred_dates_clean), len(resid_per_window))

            if n_use > 50:
                hours = pred_dates_clean[:n_use].hour
                weekdays = pred_dates_clean[:n_use].dayofweek
                resids_use = resid_per_window[:n_use]

                fig, axes = plt.subplots(1, 2, figsize=(14, 5))

                # C7: By hour of day
                hour_data = [resids_use[hours == h] for h in range(24)]
                hour_data = [d for d in hour_data if len(d) > 0]
                bp1 = axes[0].boxplot(
                    hour_data,
                    positions=list(range(len(hour_data))),
                    patch_artist=True,
                    widths=0.6,
                )
                for patch in bp1["boxes"]:
                    patch.set_facecolor(C_PRIMARY)
                    patch.set_alpha(0.5)
                axes[0].axhline(y=0, color="red", linewidth=1, linestyle="--")
                axes[0].set_xlabel("Hour of Day")
                axes[0].set_ylabel("Mean Residual (ppm)")
                axes[0].set_title(f"Residuals by Hour of Day (h={horizon}h)")

                # C8: By day of week
                day_names = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
                day_data = [resids_use[weekdays == d] for d in range(7)]
                day_data = [d for d in day_data if len(d) > 0]
                bp2 = axes[1].boxplot(
                    day_data,
                    positions=list(range(len(day_data))),
                    patch_artist=True,
                    widths=0.6,
                )
                for patch in bp2["boxes"]:
                    patch.set_facecolor(C_ACCENT)
                    patch.set_alpha(0.5)
                axes[1].axhline(y=0, color="red", linewidth=1, linestyle="--")
                axes[1].set_xticks(range(min(7, len(day_data))))
                axes[1].set_xticklabels(day_names[:len(day_data)])
                axes[1].set_xlabel("Day of Week")
                axes[1].set_ylabel("Mean Residual (ppm)")
                axes[1].set_title(f"Residuals by Day of Week (h={horizon}h)")

                plt.tight_layout()
                plt.savefig(output_dir / f"residuals_by_time_h{horizon}.png",
                            dpi=150, bbox_inches="tight")
                plt.close(fig)
    except Exception as e:
        print(f"    C7/C8 error: {e}")

    # ------------------------------------------------------------------
    # C9: Rolling residual variance (24h window)
    # ------------------------------------------------------------------
    try:
        if len(diag_resids) > 48:
            window = 24
            rolling_var = pd.Series(diag_resids).rolling(window=window).var()

            fig, ax = plt.subplots(figsize=(14, 4))
            ax.plot(rolling_var.values, color=C_WARN, linewidth=1)
            ax.set_xlabel("Observation Index")
            ax.set_ylabel("Rolling Variance")
            ax.set_title(
                f"Rolling Residual Variance (24h window, h={horizon}h)"
            )
            ax.axhline(y=np.nanmean(rolling_var), color=C_SECONDARY,
                        linewidth=1, linestyle="--", label="Mean variance")
            ax.legend()
            plt.tight_layout()
            plt.savefig(output_dir / f"rolling_variance_h{horizon}.png",
                        dpi=150, bbox_inches="tight")
            plt.close(fig)
    except Exception as e:
        print(f"    C9 error: {e}")

    # ------------------------------------------------------------------
    # C10: FFT of residuals (power spectrum)
    # ------------------------------------------------------------------
    try:
        n = len(diag_resids)
        if n > 48:
            yf = fft(diag_resids - np.mean(diag_resids))
            xf = fftfreq(n, d=1.0)  # d=1 hour for 1h resolution
            power = np.abs(yf[:n // 2]) ** 2
            freqs = xf[:n // 2]
            periods = np.zeros_like(freqs)
            periods[1:] = 1.0 / freqs[1:]  # avoid div by zero at freq=0

            fig, ax = plt.subplots(figsize=(12, 5))
            ax.semilogy(periods[1:], power[1:], color=C_PRIMARY, linewidth=0.8)
            ax.axvline(x=24, color=C_SECONDARY, linewidth=1.5, linestyle="--",
                        label="24h (daily)")
            ax.axvline(x=12, color=C_WARN, linewidth=1.5, linestyle="--",
                        label="12h (semi-daily)")
            ax.axvline(x=168, color=C_ACCENT, linewidth=1.5, linestyle="--",
                        label="168h (weekly)")
            ax.set_xlim(1, min(500, periods[1]))
            ax.set_xlabel("Period (hours)")
            ax.set_ylabel("Power (log scale)")
            ax.set_title(f"Power Spectrum of Residuals (h={horizon}h)")
            ax.legend()
            plt.tight_layout()
            plt.savefig(output_dir / f"residual_fft_h{horizon}.png",
                        dpi=150, bbox_inches="tight")
            plt.close(fig)
    except Exception as e:
        print(f"    C10 error: {e}")

    print(f"    Residual diagnostics complete: mean={resid_stats['mean']:.4f}, "
          f"std={resid_stats['std']:.4f}")
    return resid_stats


# ======================================================================
#  Section D: Forecast Uncertainty Analysis
# ======================================================================

def run_forecast_uncertainty(
    sarima: SARIMAForecaster,
    dm: CO2DataModule,
    results: dict,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Analyze forecast uncertainty via confidence intervals.

    Args:
        sarima: Fitted SARIMAForecaster.
        dm: Data module with test data and scalers.
        results: Results dict with y_true, y_pred.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.

    Returns:
        Dictionary with uncertainty statistics.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [D] Analyzing forecast uncertainty...")

    uncertainty_stats = {}
    result = sarima.result_

    y_true = results["y_true"]
    y_pred = results["y_pred"]
    test_X = dm.test_dataset.X.numpy()
    test_indices = results.get("test_indices", np.arange(y_true.shape[0]))

    horizon_steps = y_true.shape[1] if y_true.ndim > 1 else 1
    order = sarima.order
    seasonal_order = sarima.seasonal_order

    # ------------------------------------------------------------------
    # D1: Fan chart - forecast with nested CIs for a few example windows
    # ------------------------------------------------------------------
    try:
        n_examples = min(6, len(test_indices))
        example_indices = np.linspace(0, len(test_indices) - 1, n_examples, dtype=int)

        fig, axes = plt.subplots(n_examples, 1, figsize=(14, 3 * n_examples))
        if n_examples == 1:
            axes = [axes]

        ci_widths_all = {50: [], 80: [], 95: []}

        for ax_idx, ex_i in enumerate(example_indices):
            actual_idx = test_indices[ex_i]
            if actual_idx >= test_X.shape[0]:
                continue

            history = test_X[actual_idx, :, -1]  # target is last column

            try:
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore")
                    model_win = SARIMAX(
                        history,
                        order=order,
                        seasonal_order=seasonal_order,
                        enforce_stationarity=False,
                        enforce_invertibility=False,
                    )
                    filtered = model_win.filter(sarima.fitted_params_)
                    forecast_obj = filtered.get_forecast(steps=horizon_steps)

                    mean_fc = forecast_obj.predicted_mean
                    ci_50 = forecast_obj.conf_int(alpha=0.50)
                    ci_80 = forecast_obj.conf_int(alpha=0.20)
                    ci_95 = forecast_obj.conf_int(alpha=0.05)

                # Inverse scale everything
                mean_fc_orig = inverse_scale_target(
                    np.array(mean_fc).reshape(1, -1), dm.target_scaler
                ).ravel()
                ci_50_orig = inverse_scale_target(
                    np.array(ci_50), dm.target_scaler
                )
                ci_80_orig = inverse_scale_target(
                    np.array(ci_80), dm.target_scaler
                )
                ci_95_orig = inverse_scale_target(
                    np.array(ci_95), dm.target_scaler
                )

                actual_vals = y_true[ex_i].ravel() if y_true.ndim > 1 else y_true[ex_i:ex_i+1]
                steps_x = np.arange(horizon_steps)

                # Also show the lookback history in original scale
                hist_orig = inverse_scale_target(
                    history.reshape(1, -1), dm.target_scaler
                ).ravel()
                hist_x = np.arange(-len(hist_orig), 0)

                ax = axes[ax_idx]
                ax.plot(hist_x, hist_orig, color=C_NEUTRAL, linewidth=1,
                        alpha=0.6, label="History")
                ax.fill_between(steps_x, ci_95_orig[:, 0], ci_95_orig[:, 1],
                                alpha=0.15, color=C_PRIMARY, label="95% CI")
                ax.fill_between(steps_x, ci_80_orig[:, 0], ci_80_orig[:, 1],
                                alpha=0.25, color=C_PRIMARY, label="80% CI")
                ax.fill_between(steps_x, ci_50_orig[:, 0], ci_50_orig[:, 1],
                                alpha=0.35, color=C_PRIMARY, label="50% CI")
                ax.plot(steps_x, mean_fc_orig, color=C_PRIMARY, linewidth=2,
                        label="Forecast")
                ax.plot(steps_x[:len(actual_vals)], actual_vals,
                        color=C_SECONDARY, linewidth=2, linestyle="--",
                        label="Actual")
                ax.axvline(x=0, color="black", linewidth=0.5, linestyle=":")
                ax.set_ylabel("CO2 (ppm)")
                if ax_idx == 0:
                    ax.legend(fontsize=7, ncol=3)
                ax.set_title(f"Window {ex_i}", fontsize=9)

                # Track CI widths
                for alpha_key, ci_arr in [(50, ci_50_orig), (80, ci_80_orig), (95, ci_95_orig)]:
                    widths = ci_arr[:, 1] - ci_arr[:, 0]
                    ci_widths_all[alpha_key].append(widths)

            except Exception as e_win:
                axes[ax_idx].text(0.5, 0.5, f"Error: {e_win}",
                                   ha="center", va="center",
                                   transform=axes[ax_idx].transAxes)

        axes[-1].set_xlabel("Forecast Step (hours)")
        fig.suptitle(
            f"SARIMA Fan Chart - Forecast with Confidence Intervals (h={horizon}h)",
            fontsize=13,
        )
        plt.tight_layout()
        plt.savefig(output_dir / f"fan_chart_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    except Exception as e:
        print(f"    D1 error: {e}")
        ci_widths_all = {50: [], 80: [], 95: []}

    # ------------------------------------------------------------------
    # D2: CI width vs horizon step
    # ------------------------------------------------------------------
    try:
        if any(len(v) > 0 for v in ci_widths_all.values()):
            fig, ax = plt.subplots(figsize=(10, 5))
            ci_colors = {50: C_PRIMARY, 80: C_WARN, 95: C_SECONDARY}
            for alpha_key in [50, 80, 95]:
                if len(ci_widths_all[alpha_key]) > 0:
                    widths_arr = np.array(ci_widths_all[alpha_key])
                    mean_widths = widths_arr.mean(axis=0)
                    ax.plot(range(len(mean_widths)), mean_widths,
                            color=ci_colors[alpha_key], linewidth=2,
                            label=f"{alpha_key}% CI")
            ax.set_xlabel("Forecast Step (hours)")
            ax.set_ylabel("CI Width (ppm)")
            ax.set_title(f"Confidence Interval Width vs Horizon (h={horizon}h)")
            ax.legend()
            plt.tight_layout()
            plt.savefig(output_dir / f"ci_width_vs_horizon_h{horizon}.png",
                        dpi=150, bbox_inches="tight")
            plt.close(fig)
    except Exception as e:
        print(f"    D2 error: {e}")

    # ------------------------------------------------------------------
    # D3: Prediction interval coverage
    # ------------------------------------------------------------------
    try:
        # Compute coverage across all prediction windows with CIs
        n_coverage = min(200, len(test_indices))
        coverage_indices = np.linspace(0, len(test_indices) - 1, n_coverage, dtype=int)
        coverages = {50: [], 80: [], 95: []}

        for ci_idx in coverage_indices:
            actual_idx = test_indices[ci_idx]
            if actual_idx >= test_X.shape[0]:
                continue
            history = test_X[actual_idx, :, -1]

            try:
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore")
                    model_cov = SARIMAX(
                        history, order=order, seasonal_order=seasonal_order,
                        enforce_stationarity=False, enforce_invertibility=False,
                    )
                    filtered_cov = model_cov.filter(sarima.fitted_params_)
                    fc_cov = filtered_cov.get_forecast(steps=horizon_steps)

                actual_vals = y_true[ci_idx].ravel() if y_true.ndim > 1 else y_true[ci_idx:ci_idx+1]

                for alpha_key, alpha_val in [(50, 0.50), (80, 0.20), (95, 0.05)]:
                    ci = fc_cov.conf_int(alpha=alpha_val)
                    ci_orig = inverse_scale_target(np.array(ci), dm.target_scaler)
                    in_ci = np.mean(
                        (actual_vals >= ci_orig[:len(actual_vals), 0]) &
                        (actual_vals <= ci_orig[:len(actual_vals), 1])
                    )
                    coverages[alpha_key].append(in_ci)
            except Exception:
                pass

        if any(len(v) > 0 for v in coverages.values()):
            fig, ax = plt.subplots(figsize=(8, 6))
            nominal = [50, 80, 95]
            observed = [np.mean(coverages[k]) * 100 if len(coverages[k]) > 0 else 0
                       for k in nominal]
            x_pos = np.arange(len(nominal))

            bars = ax.bar(x_pos, observed, width=0.5, color=C_PRIMARY,
                          edgecolor="black", linewidth=0.5)
            # Reference line
            for i, nom in enumerate(nominal):
                ax.plot([i - 0.3, i + 0.3], [nom, nom], color=C_SECONDARY,
                        linewidth=2, linestyle="--")

            ax.set_xticks(x_pos)
            ax.set_xticklabels([f"{n}%" for n in nominal])
            ax.set_xlabel("Nominal Coverage Level")
            ax.set_ylabel("Observed Coverage (%)")
            ax.set_title(f"Prediction Interval Coverage (h={horizon}h)")

            for i, (obs, nom) in enumerate(zip(observed, nominal)):
                ax.text(i, obs + 1, f"{obs:.1f}%", ha="center", fontsize=10)

            ax.set_ylim(0, 105)
            plt.tight_layout()
            plt.savefig(output_dir / f"pi_coverage_h{horizon}.png",
                        dpi=150, bbox_inches="tight")
            plt.close(fig)

            uncertainty_stats["coverage"] = {
                str(k): float(np.mean(coverages[k]) * 100) if len(coverages[k]) > 0 else 0.0
                for k in nominal
            }

    except Exception as e:
        print(f"    D3 error: {e}")

    # ------------------------------------------------------------------
    # D4: Calibration plot
    # ------------------------------------------------------------------
    try:
        # Compare predicted quantiles vs observed quantiles
        if len(coverages.get(95, [])) > 10:
            fig, ax = plt.subplots(figsize=(6, 6))
            nominal_levels = np.array([10, 20, 30, 40, 50, 60, 70, 80, 90, 95])
            observed_levels = []

            for nom in nominal_levels:
                alpha_val = 1.0 - nom / 100.0
                cov_list = []
                for ci_idx in coverage_indices[:100]:
                    actual_idx = test_indices[ci_idx]
                    if actual_idx >= test_X.shape[0]:
                        continue
                    history = test_X[actual_idx, :, -1]
                    try:
                        with warnings.catch_warnings():
                            warnings.simplefilter("ignore")
                            m_cal = SARIMAX(
                                history, order=order,
                                seasonal_order=seasonal_order,
                                enforce_stationarity=False,
                                enforce_invertibility=False,
                            )
                            f_cal = m_cal.filter(sarima.fitted_params_)
                            fc_cal = f_cal.get_forecast(steps=horizon_steps)
                            ci_cal = fc_cal.conf_int(alpha=alpha_val)
                            ci_cal_orig = inverse_scale_target(
                                np.array(ci_cal), dm.target_scaler
                            )
                            actual_vals = y_true[ci_idx].ravel() if y_true.ndim > 1 else y_true[ci_idx:ci_idx+1]
                            in_ci = np.mean(
                                (actual_vals >= ci_cal_orig[:len(actual_vals), 0]) &
                                (actual_vals <= ci_cal_orig[:len(actual_vals), 1])
                            )
                            cov_list.append(in_ci)
                    except Exception:
                        pass
                observed_levels.append(np.mean(cov_list) * 100 if cov_list else nom)

            ax.plot([0, 100], [0, 100], "k--", linewidth=1, label="Perfect calibration")
            ax.scatter(nominal_levels, observed_levels, c=C_PRIMARY, s=60,
                      zorder=5, edgecolors="black", linewidth=0.5)
            ax.plot(nominal_levels, observed_levels, color=C_PRIMARY, linewidth=1.5)
            ax.set_xlabel("Nominal Coverage (%)")
            ax.set_ylabel("Observed Coverage (%)")
            ax.set_title(f"Calibration Plot (h={horizon}h)")
            ax.legend()
            ax.set_xlim(0, 100)
            ax.set_ylim(0, 100)
            ax.set_aspect("equal")
            plt.tight_layout()
            plt.savefig(output_dir / f"calibration_plot_h{horizon}.png",
                        dpi=150, bbox_inches="tight")
            plt.close(fig)

    except Exception as e:
        print(f"    D4 error: {e}")

    print(f"    Forecast uncertainty analysis complete")
    return uncertainty_stats


# ======================================================================
#  Section E: Seasonal Decomposition Comparison
# ======================================================================

def run_seasonal_decomposition(
    sarima: SARIMAForecaster,
    train_series_scaled: np.ndarray,
    horizon: int,
    output_dir: Path,
) -> dict:
    """STL decomposition and comparison with SARIMA seasonal component.

    Args:
        sarima: Fitted SARIMAForecaster.
        train_series_scaled: Scaled training series.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.

    Returns:
        Dictionary with decomposition statistics.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [E] Running seasonal decomposition analysis...")

    decomp_stats = {}
    seasonal_period = 24  # 24 hours for 1h data

    # ------------------------------------------------------------------
    # E1: STL decomposition (4-panel: original, trend, seasonal, residual)
    # ------------------------------------------------------------------
    try:
        # Use a subset if series is very long (STL can be slow)
        max_len = min(len(train_series_scaled), 5000)
        series_subset = train_series_scaled[:max_len]

        stl = STL(series_subset, period=seasonal_period, robust=True)
        stl_result = stl.fit()

        fig, axes = plt.subplots(4, 1, figsize=(14, 12), sharex=True)

        axes[0].plot(series_subset, color=C_PRIMARY, linewidth=0.5)
        axes[0].set_ylabel("Original")
        axes[0].set_title(f"STL Decomposition (period={seasonal_period}h, h={horizon}h)")

        axes[1].plot(stl_result.trend, color=C_ACCENT, linewidth=1)
        axes[1].set_ylabel("Trend")

        axes[2].plot(stl_result.seasonal, color=C_WARN, linewidth=0.8)
        axes[2].set_ylabel("Seasonal")

        axes[3].plot(stl_result.resid, color=C_SECONDARY, linewidth=0.5)
        axes[3].set_ylabel("Residual")
        axes[3].set_xlabel("Time (hours)")

        plt.tight_layout()
        plt.savefig(output_dir / f"stl_decomposition_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

        # Compute variance proportions
        total_var = np.var(series_subset)
        if total_var > 0:
            decomp_stats["trend_var_pct"] = float(np.var(stl_result.trend) / total_var * 100)
            decomp_stats["seasonal_var_pct"] = float(np.var(stl_result.seasonal) / total_var * 100)
            decomp_stats["residual_var_pct"] = float(np.var(stl_result.resid) / total_var * 100)

    except Exception as e:
        print(f"    E1 error: {e}")
        stl_result = None

    # ------------------------------------------------------------------
    # E2: Differenced series visualization
    # ------------------------------------------------------------------
    try:
        series_sub = train_series_scaled[:max_len]

        fig, axes = plt.subplots(3, 1, figsize=(14, 9), sharex=False)

        # Original
        axes[0].plot(series_sub, color=C_PRIMARY, linewidth=0.5)
        axes[0].set_title(f"Original Series (h={horizon}h)")
        axes[0].set_ylabel("Value")

        # First difference
        diff1 = np.diff(series_sub, n=1)
        axes[1].plot(diff1, color=C_ACCENT, linewidth=0.5)
        axes[1].set_title("First Difference (d=1)")
        axes[1].set_ylabel("Value")
        axes[1].axhline(y=0, color="black", linewidth=0.5)

        # Seasonal difference
        if len(series_sub) > seasonal_period:
            seasonal_diff = series_sub[seasonal_period:] - series_sub[:-seasonal_period]
            axes[2].plot(seasonal_diff, color=C_SECONDARY, linewidth=0.5)
            axes[2].set_title(f"Seasonal Difference (D=1, s={seasonal_period})")
            axes[2].set_ylabel("Value")
            axes[2].axhline(y=0, color="black", linewidth=0.5)

        axes[2].set_xlabel("Time (hours)")
        plt.tight_layout()
        plt.savefig(output_dir / f"differenced_series_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    except Exception as e:
        print(f"    E2 error: {e}")

    # ------------------------------------------------------------------
    # E3: Compare STL seasonal with SARIMA seasonal extraction
    # ------------------------------------------------------------------
    try:
        if stl_result is not None:
            stl_seasonal = stl_result.seasonal

            # Extract SARIMA's seasonal pattern by looking at the fitted model's
            # one-step predictions minus a non-seasonal version
            # Simpler approach: extract one full period of the STL seasonal
            one_period = stl_seasonal[:seasonal_period]

            fig, ax = plt.subplots(figsize=(10, 5))
            ax.plot(range(seasonal_period), one_period, color=C_PRIMARY,
                    linewidth=2, label="STL Seasonal (1 period)")
            ax.axhline(y=0, color="black", linewidth=0.5, linestyle="--")
            ax.set_xlabel("Hour of Day")
            ax.set_ylabel("Seasonal Component (scaled)")
            ax.set_title(
                f"STL Seasonal Pattern - Daily Cycle (h={horizon}h)"
            )
            ax.legend()
            ax.set_xticks(range(0, 24, 2))
            plt.tight_layout()
            plt.savefig(output_dir / f"seasonal_comparison_h{horizon}.png",
                        dpi=150, bbox_inches="tight")
            plt.close(fig)

    except Exception as e:
        print(f"    E3 error: {e}")

    print(f"    Seasonal decomposition complete")
    return decomp_stats


# ======================================================================
#  Section F: SARIMA-Specific Analyses (Stationarity + Predictions)
# ======================================================================

def run_sarima_specific(
    sarima: SARIMAForecaster,
    dm: CO2DataModule,
    results: dict,
    train_series_scaled: np.ndarray,
    test_df: pd.DataFrame,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Stationarity tests and prediction quality analysis.

    Args:
        sarima: Fitted SARIMAForecaster.
        dm: Data module with scalers.
        results: Results dict with y_true, y_pred, metrics.
        train_series_scaled: Scaled training series.
        test_df: Test DataFrame.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.

    Returns:
        Dictionary with stationarity test results.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [F] Running SARIMA-specific analyses...")

    specific_stats = {}

    # ------------------------------------------------------------------
    # F1: ADF and KPSS stationarity tests
    # ------------------------------------------------------------------
    try:
        series = train_series_scaled
        diff1 = np.diff(series, n=1)
        seasonal_period = 24
        seasonal_diff = series[seasonal_period:] - series[:-seasonal_period]

        test_cases = [
            ("Original", series),
            ("First Diff (d=1)", diff1),
            (f"Seasonal Diff (D=1, s={seasonal_period})", seasonal_diff),
        ]

        adf_results = []
        kpss_results = []

        for label, s in test_cases:
            # ADF test
            adf_stat, adf_pval, adf_usedlag, adf_nobs, adf_crit, _ = adfuller(
                s, autolag="AIC"
            )
            adf_results.append({
                "series": label,
                "statistic": float(adf_stat),
                "pvalue": float(adf_pval),
                "lags_used": int(adf_usedlag),
                "stationary": adf_pval < 0.05,
            })

            # KPSS test
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                kpss_stat, kpss_pval, kpss_lags, kpss_crit = kpss(
                    s, regression="c", nlags="auto"
                )
            kpss_results.append({
                "series": label,
                "statistic": float(kpss_stat),
                "pvalue": float(kpss_pval),
                "stationary": kpss_pval > 0.05,  # KPSS: H0 is stationarity
            })

        specific_stats["adf"] = adf_results
        specific_stats["kpss"] = kpss_results

        # Plot stationarity test results as a table
        fig, ax = plt.subplots(figsize=(12, 4))
        ax.axis("off")

        col_labels = [
            "Series", "ADF Stat", "ADF p-val", "ADF Stationary?",
            "KPSS Stat", "KPSS p-val", "KPSS Stationary?",
        ]
        cell_data = []
        for adf_r, kpss_r in zip(adf_results, kpss_results):
            cell_data.append([
                adf_r["series"],
                f"{adf_r['statistic']:.4f}",
                f"{adf_r['pvalue']:.4f}",
                "Yes" if adf_r["stationary"] else "No",
                f"{kpss_r['statistic']:.4f}",
                f"{kpss_r['pvalue']:.4f}",
                "Yes" if kpss_r["stationary"] else "No",
            ])

        tbl = ax.table(
            cellText=cell_data, colLabels=col_labels,
            loc="center", cellLoc="center",
        )
        tbl.auto_set_font_size(False)
        tbl.set_fontsize(9)
        tbl.scale(1.0, 1.6)

        # Color cells based on stationarity
        for i, (adf_r, kpss_r) in enumerate(zip(adf_results, kpss_results)):
            color_adf = "#C8E6C9" if adf_r["stationary"] else "#FFCDD2"
            color_kpss = "#C8E6C9" if kpss_r["stationary"] else "#FFCDD2"
            tbl[i + 1, 3].set_facecolor(color_adf)
            tbl[i + 1, 6].set_facecolor(color_kpss)

        ax.set_title(f"Stationarity Tests (h={horizon}h)", fontsize=12, pad=20)
        plt.savefig(output_dir / f"stationarity_tests_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)

    except Exception as e:
        print(f"    F1 error: {e}")

    # ------------------------------------------------------------------
    # F2: Prediction overlay, scatter, residual analysis, error by level
    # ------------------------------------------------------------------
    y_true = results["y_true"]
    y_pred = results["y_pred"]
    metrics = results["metrics"]

    # Prediction overlay
    try:
        y_true_flat = y_true.ravel()
        y_pred_flat = y_pred.ravel()

        # Show a representative subset (first 500 points)
        n_show = min(500, len(y_true_flat))

        fig, ax = plt.subplots(figsize=(14, 5))
        ax.plot(range(n_show), y_true_flat[:n_show], color=C_PRIMARY,
                linewidth=1, label="Actual", alpha=0.8)
        ax.plot(range(n_show), y_pred_flat[:n_show], color=C_SECONDARY,
                linewidth=1, label="Predicted", alpha=0.8)
        ax.fill_between(
            range(n_show),
            y_true_flat[:n_show],
            y_pred_flat[:n_show],
            alpha=0.15, color=C_WARN,
        )
        ax.set_xlabel("Prediction Index")
        ax.set_ylabel("CO2 (ppm)")
        ax.set_title(
            f"SARIMA Predictions vs Actual (h={horizon}h) - "
            f"RMSE={metrics['rmse']:.2f}"
        )
        ax.legend()
        plt.tight_layout()
        plt.savefig(output_dir / f"predictions_overlay_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    F2a overlay error: {e}")

    # Scatter plot with R2
    try:
        y_true_flat = y_true.ravel()
        y_pred_flat = y_pred.ravel()

        fig, ax = plt.subplots(figsize=(7, 7))
        ax.scatter(y_true_flat, y_pred_flat, alpha=0.3, s=8,
                   color=C_PRIMARY, edgecolors="none")
        lims = [
            min(y_true_flat.min(), y_pred_flat.min()),
            max(y_true_flat.max(), y_pred_flat.max()),
        ]
        ax.plot(lims, lims, "r--", linewidth=1.5, label="Perfect prediction")
        ax.set_xlabel("Actual CO2 (ppm)")
        ax.set_ylabel("Predicted CO2 (ppm)")
        ax.set_title(
            f"Scatter: Predicted vs Actual (h={horizon}h)\n"
            f"R2={metrics['r2']:.4f}  RMSE={metrics['rmse']:.2f}  "
            f"MAE={metrics['mae']:.2f}"
        )
        ax.legend()
        ax.set_aspect("equal")
        plt.tight_layout()
        plt.savefig(output_dir / f"scatter_r2_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    F2b scatter error: {e}")

    # 4-panel residual analysis
    try:
        y_true_flat = y_true.ravel()
        y_pred_flat = y_pred.ravel()
        resids = y_true_flat - y_pred_flat

        fig, axes = plt.subplots(2, 2, figsize=(14, 10))

        # (a) Residual distribution
        axes[0, 0].hist(resids, bins=60, density=True, alpha=0.6,
                         color=C_PRIMARY, edgecolor="white", linewidth=0.3)
        xr = np.linspace(resids.min(), resids.max(), 200)
        mu, sigma_r = np.mean(resids), np.std(resids)
        axes[0, 0].plot(xr, stats.norm.pdf(xr, mu, sigma_r),
                         color=C_SECONDARY, linewidth=2)
        axes[0, 0].set_title("Residual Distribution")
        axes[0, 0].set_xlabel("Residual (ppm)")

        # (b) Residuals over time
        axes[0, 1].scatter(range(len(resids)), resids, alpha=0.3, s=3,
                            color=C_PRIMARY)
        axes[0, 1].axhline(y=0, color="red", linewidth=1)
        axes[0, 1].set_title("Residuals Over Time")
        axes[0, 1].set_xlabel("Index")
        axes[0, 1].set_ylabel("Residual (ppm)")

        # (c) Residuals vs predicted (check for heteroscedasticity)
        axes[1, 0].scatter(y_pred_flat, resids, alpha=0.3, s=3,
                            color=C_ACCENT)
        axes[1, 0].axhline(y=0, color="red", linewidth=1)
        axes[1, 0].set_title("Residuals vs Predicted")
        axes[1, 0].set_xlabel("Predicted CO2 (ppm)")
        axes[1, 0].set_ylabel("Residual (ppm)")

        # (d) Q-Q plot
        stats.probplot(resids, dist="norm", plot=axes[1, 1])
        axes[1, 1].set_title("Normal Q-Q Plot")

        fig.suptitle(
            f"SARIMA Residual Analysis (h={horizon}h)",
            fontsize=13,
        )
        plt.tight_layout()
        plt.savefig(output_dir / f"residual_analysis_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    F2c residual analysis error: {e}")

    # Error by CO2 level
    try:
        y_true_flat = y_true.ravel()
        y_pred_flat = y_pred.ravel()
        abs_errors = np.abs(y_true_flat - y_pred_flat)

        bins_co2 = [0, 500, 700, 1000, 1500, 5000]
        bin_labels = ["<500", "500-700", "700-1000", "1000-1500", ">1500"]
        bin_indices = np.digitize(y_true_flat, bins_co2) - 1

        fig, ax = plt.subplots(figsize=(10, 6))
        error_data = []
        valid_labels = []
        for bi in range(len(bin_labels)):
            mask = bin_indices == bi
            if mask.sum() > 0:
                error_data.append(abs_errors[mask])
                valid_labels.append(f"{bin_labels[bi]}\n(n={mask.sum()})")

        if error_data:
            bp = ax.boxplot(error_data, patch_artist=True, widths=0.6)
            colors_box = [C_PRIMARY, C_ACCENT, C_WARN, C_SECONDARY, C_NEUTRAL]
            for patch, color in zip(bp["boxes"], colors_box[:len(error_data)]):
                patch.set_facecolor(color)
                patch.set_alpha(0.5)
            ax.set_xticklabels(valid_labels, fontsize=9)
            ax.set_xlabel("CO2 Level (ppm)")
            ax.set_ylabel("Absolute Error (ppm)")
            ax.set_title(f"Prediction Error by CO2 Level (h={horizon}h)")

        plt.tight_layout()
        plt.savefig(output_dir / f"error_by_co2_level_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    F2d error by level: {e}")

    # ------------------------------------------------------------------
    # F3: Metrics summary table figure
    # ------------------------------------------------------------------
    try:
        fig, ax = plt.subplots(figsize=(8, 2.5))
        ax.axis("off")
        col_labels = ["RMSE", "MAE", "R2", "MAPE (%)", "MSE"]
        cell_data = [[
            f"{metrics['rmse']:.2f}",
            f"{metrics['mae']:.2f}",
            f"{metrics['r2']:.4f}",
            f"{metrics['mape']:.2f}",
            f"{metrics['mse']:.2f}",
        ]]
        tbl = ax.table(
            cellText=cell_data, colLabels=col_labels,
            loc="center", cellLoc="center",
        )
        tbl.auto_set_font_size(False)
        tbl.set_fontsize(11)
        tbl.scale(1.2, 1.8)
        ax.set_title(f"SARIMA Performance Metrics (h={horizon}h)", fontsize=12, pad=20)
        plt.savefig(output_dir / f"metrics_table_h{horizon}.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    F3 error: {e}")

    print(f"    SARIMA-specific analyses complete")
    return specific_stats


# ======================================================================
#  Cross-Horizon Comparison
# ======================================================================

def run_cross_horizon_comparison(
    all_results: dict,
    output_dir: Path,
) -> None:
    """Generate cross-horizon comparison plots.

    Args:
        all_results: {horizon: results_dict} with metrics.
        output_dir: Directory to save comparison plots.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  Generating cross-horizon comparison...")

    horizons = sorted(all_results.keys())

    # Metrics comparison bar chart
    try:
        metric_names = ["rmse", "mae", "r2", "mape"]
        metric_labels = ["RMSE (ppm)", "MAE (ppm)", "R2", "MAPE (%)"]

        fig, axes = plt.subplots(2, 2, figsize=(12, 10))
        colors_h = [C_PRIMARY, C_SECONDARY, C_ACCENT, C_WARN]

        for ax_idx, (mname, mlabel) in enumerate(zip(metric_names, metric_labels)):
            ax = axes[ax_idx // 2, ax_idx % 2]
            vals = [all_results[h]["metrics"][mname] for h in horizons]
            x_pos = np.arange(len(horizons))
            bars = ax.bar(x_pos, vals, color=colors_h[:len(horizons)],
                          edgecolor="black", linewidth=0.5, width=0.5)
            ax.set_xticks(x_pos)
            ax.set_xticklabels([f"{h}h" for h in horizons])
            ax.set_ylabel(mlabel)
            ax.set_title(mlabel)

            for i, v in enumerate(vals):
                fmt = f"{v:.4f}" if mname == "r2" else f"{v:.2f}"
                ax.text(i, v * 1.02, fmt, ha="center", fontsize=10)

        fig.suptitle(
            "SARIMA Cross-Horizon Metrics Comparison",
            fontsize=14, fontweight="bold",
        )
        plt.tight_layout()
        plt.savefig(output_dir / "metrics_comparison.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    Metrics comparison error: {e}")

    # Error growth visualization
    try:
        fig, ax = plt.subplots(figsize=(10, 6))
        for h_idx, h in enumerate(horizons):
            y_true = all_results[h]["y_true"]
            y_pred = all_results[h]["y_pred"]
            if y_true.ndim > 1 and y_true.shape[1] > 1:
                # Per-step RMSE
                step_rmse = np.sqrt(np.mean((y_true - y_pred) ** 2, axis=0))
                ax.plot(range(1, len(step_rmse) + 1), step_rmse,
                        color=colors_h[h_idx], linewidth=2,
                        marker="o", markersize=3, label=f"h={h}h")

        ax.set_xlabel("Forecast Step (hours)")
        ax.set_ylabel("RMSE (ppm)")
        ax.set_title("Error Growth: RMSE by Forecast Step")
        ax.legend()
        plt.tight_layout()
        plt.savefig(output_dir / "error_growth.png",
                    dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    Error growth plot error: {e}")

    print(f"  Cross-horizon comparison plots saved")


# ======================================================================
#  Summary Figure and JSON Export
# ======================================================================

def generate_summary_figure(
    all_results: dict,
    all_param_stats: dict,
    all_resid_stats: dict,
    output_dir: Path,
) -> None:
    """Multi-panel summary figure with key SARIMA findings.

    Args:
        all_results: {horizon: results_dict} with metrics.
        all_param_stats: {horizon: param_stats_dict}.
        all_resid_stats: {horizon: resid_stats_dict}.
        output_dir: Base output directory.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    horizons = sorted(all_results.keys())

    fig = plt.figure(figsize=(20, 12))
    gs = GridSpec(2, 3, figure=fig, hspace=0.35, wspace=0.3)

    # [0,0] Metrics table
    ax = fig.add_subplot(gs[0, 0])
    ax.axis("off")
    rows_data = []
    for h in horizons:
        m = all_results[h]["metrics"]
        rows_data.append([
            f"{h}h", f"{m['rmse']:.2f}", f"{m['mae']:.2f}",
            f"{m['r2']:.4f}", f"{m['mape']:.2f}%",
        ])
    tbl = ax.table(
        cellText=rows_data,
        colLabels=["Horizon", "RMSE", "MAE", "R2", "MAPE"],
        loc="center", cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(11)
    tbl.scale(1.0, 1.8)
    ax.set_title("Performance Metrics", fontsize=12, pad=20)

    # [0,1] Model info
    ax = fig.add_subplot(gs[0, 1])
    ax.axis("off")
    info_lines = [
        "SARIMA Model Summary",
        "",
        "Order: (1, 1, 1)",
        "Seasonal: (1, 1, 1, 24)",
        "Resolution: 1 hour",
        "",
    ]
    # Add AIC/BIC if available
    for h in horizons:
        ps = all_param_stats.get(h, {})
        aic = ps.get("_aic", "N/A")
        bic = ps.get("_bic", "N/A")
        if isinstance(aic, float):
            info_lines.append(f"h={h}h: AIC={aic:.1f} BIC={bic:.1f}")
        else:
            info_lines.append(f"h={h}h: AIC={aic} BIC={bic}")

    ax.text(
        0.5, 0.5, "\n".join(info_lines),
        ha="center", va="center", fontsize=10,
        transform=ax.transAxes,
        bbox=dict(boxstyle="round,pad=0.5", facecolor="lightyellow", alpha=0.8),
        family="monospace",
    )

    # [0,2] Residual summary
    ax = fig.add_subplot(gs[0, 2])
    ax.axis("off")
    resid_lines = ["Residual Diagnostics", ""]
    for h in horizons:
        rs = all_resid_stats.get(h, {})
        resid_lines.append(f"h={h}h:")
        resid_lines.append(f"  Mean: {rs.get('mean', 'N/A'):.4f}" if isinstance(rs.get('mean'), float) else f"  Mean: N/A")
        resid_lines.append(f"  Std:  {rs.get('std', 'N/A'):.4f}" if isinstance(rs.get('std'), float) else f"  Std: N/A")
        resid_lines.append(f"  Skew: {rs.get('skewness', 'N/A'):.4f}" if isinstance(rs.get('skewness'), float) else f"  Skew: N/A")

    ax.text(
        0.5, 0.5, "\n".join(resid_lines),
        ha="center", va="center", fontsize=10,
        transform=ax.transAxes,
        bbox=dict(boxstyle="round,pad=0.5", facecolor="lightyellow", alpha=0.8),
        family="monospace",
    )

    # Bottom row: Section summaries
    section_names = [
        "Parameters (A) + Impulse (B)\n- Coefficient significance\n- AR/MA roots\n- IRF analysis",
        "Residual Diagnostics (C)\n- ACF/PACF analysis\n- Ljung-Box test\n- Heteroscedasticity",
        "Forecasting (D-F)\n- Uncertainty analysis\n- Seasonal decomposition\n- Stationarity tests",
    ]
    for col in range(3):
        ax = fig.add_subplot(gs[1, col])
        ax.text(
            0.5, 0.5, section_names[col],
            ha="center", va="center", fontsize=10,
            transform=ax.transAxes,
            bbox=dict(boxstyle="round,pad=0.5", facecolor="lightyellow", alpha=0.8),
        )
        ax.set_frame_on(False)
        ax.set_xticks([])
        ax.set_yticks([])

    fig.suptitle(
        "SARIMA Interpretability Study Summary - preproc_D (Enhanced 1h)",
        fontsize=15, fontweight="bold",
    )
    plt.savefig(output_dir / "summary.png", dpi=150, bbox_inches="tight")
    plt.close(fig)


def save_study_results(
    all_results: dict,
    all_param_stats: dict,
    all_resid_stats: dict,
    all_decomp_stats: dict,
    all_uncertainty_stats: dict,
    all_specific_stats: dict,
    output_dir: Path,
) -> None:
    """Save all study data to a single JSON file.

    Args:
        all_results: {horizon: results_dict}.
        all_param_stats: {horizon: param_stats}.
        all_resid_stats: {horizon: resid_stats}.
        all_decomp_stats: {horizon: decomp_stats}.
        all_uncertainty_stats: {horizon: uncertainty_stats}.
        all_specific_stats: {horizon: specific_stats}.
        output_dir: Base output directory.
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    def _safe_serialize(obj):
        """Convert numpy types to Python types for JSON."""
        if isinstance(obj, (np.integer,)):
            return int(obj)
        if isinstance(obj, (np.floating,)):
            return float(obj)
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        if isinstance(obj, dict):
            return {k: _safe_serialize(v) for k, v in obj.items()}
        if isinstance(obj, (list, tuple)):
            return [_safe_serialize(v) for v in obj]
        return obj

    results_json = {
        "study": "SARIMA Interpretability",
        "variant": "preproc_D (Enhanced 1h)",
        "timestamp": datetime.now().isoformat(),
        "horizons": {},
    }

    for h in sorted(all_results.keys()):
        horizon_data = {
            "metrics": all_results[h]["metrics"],
        }
        if h in all_param_stats:
            # Filter out non-serializable items
            ps = {k: v for k, v in all_param_stats[h].items()
                  if isinstance(v, (dict, float, int, str, bool))}
            horizon_data["parameters"] = _safe_serialize(ps)
        if h in all_resid_stats:
            horizon_data["residual_diagnostics"] = _safe_serialize(all_resid_stats[h])
        if h in all_decomp_stats:
            horizon_data["decomposition"] = _safe_serialize(all_decomp_stats[h])
        if h in all_uncertainty_stats:
            horizon_data["uncertainty"] = _safe_serialize(all_uncertainty_stats[h])
        if h in all_specific_stats:
            horizon_data["stationarity"] = _safe_serialize(all_specific_stats[h])

        results_json["horizons"][str(h)] = horizon_data

    with open(output_dir / "study_results.json", "w", encoding="utf-8") as f:
        json.dump(results_json, f, indent=2, default=str)

    print(f"  Results saved to: {output_dir / 'study_results.json'}")


# ======================================================================
#  DOCX Report Generation
# ======================================================================

def generate_docx_report(
    all_results: dict,
    all_param_stats: dict,
    all_resid_stats: dict,
    all_irf_data: dict,
    all_decomp_stats: dict,
    all_uncertainty_stats: dict,
    all_specific_stats: dict,
    horizons: list,
    output_dir: Path,
) -> Path:
    """Generate comprehensive DOCX academic report.

    Args:
        all_results: {horizon: results_dict}.
        all_param_stats: {horizon: param_stats}.
        all_resid_stats: {horizon: resid_stats}.
        all_irf_data: {horizon: irf_data}.
        all_decomp_stats: {horizon: decomp_stats}.
        all_uncertainty_stats: {horizon: uncertainty_stats}.
        all_specific_stats: {horizon: specific_stats}.
        horizons: List of analyzed horizons.
        output_dir: Base output directory.

    Returns:
        Path to the generated DOCX file.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print("  Generating DOCX report...")
    doc = Document()

    # --- Title ---
    title = doc.add_heading(
        "SARIMA Interpretability Study: Statistical Time Series Analysis "
        "for Indoor CO2 Forecasting",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Preprocessing Variant: preproc_D (Enhanced 1h)\n"
        f"Model: SARIMA(1,1,1)x(1,1,1,24)\n"
        f"Horizons: {', '.join(str(h) + 'h' for h in horizons)}\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Section A: Model Parameter Analysis",
        "3. Section B: Impulse Response Analysis",
        "4. Section C: Residual Diagnostics",
        "5. Section D: Forecast Uncertainty Analysis",
        "6. Section E: Seasonal Decomposition",
        "7. Section F: SARIMA-Specific Analyses",
        "8. Cross-Horizon Comparison",
        "9. Avenues of Improvement",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # Helper to safely add figures
    def add_figure(fig_path: Path, caption: str, width: float = 6.0) -> None:
        if fig_path.exists():
            doc.add_picture(str(fig_path), width=Inches(width))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.font.size = Pt(9)
            r.font.italic = True
        else:
            doc.add_paragraph(f"[Figure not found: {fig_path.name}]")

    # ============================================================
    #  1. Executive Summary
    # ============================================================
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report presents a comprehensive interpretability study of a SARIMA "
        "(Seasonal AutoRegressive Integrated Moving Average) model trained for indoor "
        "CO2 concentration forecasting. The model uses the preproc_D preprocessing variant "
        "(Enhanced, 1-hour resolution) and is configured as SARIMA(1,1,1)x(1,1,1,24), "
        "where the seasonal period of 24 captures the daily occupancy cycle. Unlike the "
        "neural network models in this project, SARIMA is a univariate statistical model "
        "that uses only the CO2 target variable, making it a pure baseline that captures "
        "temporal structure (trend, seasonality, autocorrelation) without feature engineering. "
        "The study investigates six aspects: (A) model parameter significance and polynomial "
        "root stability, (B) impulse response characterization, (C) residual diagnostics "
        "including autocorrelation and heteroscedasticity tests, (D) forecast uncertainty "
        "quantification, (E) seasonal decomposition comparison, and (F) stationarity tests "
        "and prediction quality assessment."
    )

    # Summary metrics table
    doc.add_heading("Performance Summary", level=2)
    table = doc.add_table(rows=1 + len(horizons), cols=5)
    table.style = "Light Grid Accent 1"
    headers = ["Horizon", "RMSE (ppm)", "MAE (ppm)", "R2", "MAPE (%)"]
    for i, hdr in enumerate(headers):
        table.rows[0].cells[i].text = hdr
    for row_idx, h in enumerate(horizons):
        m = all_results[h]["metrics"]
        table.rows[row_idx + 1].cells[0].text = f"{h}h"
        table.rows[row_idx + 1].cells[1].text = f"{m['rmse']:.2f}"
        table.rows[row_idx + 1].cells[2].text = f"{m['mae']:.2f}"
        table.rows[row_idx + 1].cells[3].text = f"{m['r2']:.4f}"
        table.rows[row_idx + 1].cells[4].text = f"{m['mape']:.2f}"

    doc.add_page_break()

    # ============================================================
    #  2. Section A: Model Parameter Analysis
    # ============================================================
    doc.add_heading("2. Section A: Model Parameter Analysis", level=1)

    doc.add_paragraph(
        "SARIMA model parameters are estimated via maximum likelihood. Each coefficient "
        "has an associated standard error and p-value indicating its statistical significance. "
        "The AR and MA polynomial roots determine model stability: AR roots outside the unit "
        "circle confirm stationarity, while MA roots outside the unit circle confirm "
        "invertibility. The SARIMA(1,1,1)x(1,1,1,24) specification includes non-seasonal "
        "AR(1), MA(1) and seasonal AR(1), MA(1) parameters, plus variance and differencing "
        "components."
    )

    for h in horizons:
        doc.add_heading(f"2.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "parameters"

        add_figure(
            h_dir / f"coefficient_chart_h{h}.png",
            f"Figure: SARIMA coefficient estimates with 95% confidence intervals. "
            f"Blue bars indicate statistically significant parameters (p<0.05), "
            f"gray bars indicate non-significant parameters.",
        )

        add_figure(
            h_dir / f"root_plot_h{h}.png",
            f"Figure: AR and MA polynomial roots in the complex plane. "
            f"Roots outside the unit circle (dashed) confirm stationarity "
            f"(AR) and invertibility (MA) of the fitted model.",
        )

        add_figure(
            h_dir / f"param_summary_table_h{h}.png",
            f"Figure: Complete parameter summary table with coefficient values, "
            f"standard errors, p-values, and confidence intervals.",
        )

        # Parameter table in DOCX
        ps = all_param_stats.get(h, {})
        param_entries = {k: v for k, v in ps.items()
                        if isinstance(v, dict) and "value" in v}
        if param_entries:
            doc.add_heading("Parameter Table", level=3)
            tbl = doc.add_table(rows=1 + len(param_entries), cols=5)
            tbl.style = "Light Grid Accent 1"
            for ci, ch in enumerate(["Parameter", "Value", "Std Err", "p-value", "Sig?"]):
                tbl.rows[0].cells[ci].text = ch
            for ri, (pname, pdata) in enumerate(param_entries.items()):
                tbl.rows[ri + 1].cells[0].text = pname
                tbl.rows[ri + 1].cells[1].text = f"{pdata['value']:.6f}"
                tbl.rows[ri + 1].cells[2].text = f"{pdata['std_error']:.6f}"
                tbl.rows[ri + 1].cells[3].text = f"{pdata['pvalue']:.4f}"
                tbl.rows[ri + 1].cells[4].text = "Yes" if pdata["significant"] else "No"

            # Model fit statistics
            aic_val = ps.get("_aic", "N/A")
            bic_val = ps.get("_bic", "N/A")
            ll_val = ps.get("_loglik", "N/A")
            if isinstance(aic_val, float):
                doc.add_paragraph(
                    f"Model fit: AIC = {aic_val:.2f}, BIC = {bic_val:.2f}, "
                    f"Log-Likelihood = {ll_val:.2f}"
                )

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The parameter analysis reveals the relative importance of each SARIMA component. "
            "A significant AR(1) coefficient indicates that recent CO2 levels are predictive "
            "of future values. A significant seasonal AR(1) at lag 24 means that the CO2 "
            "level 24 hours ago contains useful information, reflecting daily occupancy "
            "patterns. The MA terms capture short-lived shocks (ventilation events, sudden "
            "occupancy changes). Roots close to the unit circle indicate near-unit-root "
            "behavior, which is common for environmental time series with strong persistence."
        )

    doc.add_page_break()

    # ============================================================
    #  3. Section B: Impulse Response Analysis
    # ============================================================
    doc.add_heading("3. Section B: Impulse Response Analysis", level=1)

    doc.add_paragraph(
        "The impulse response function (IRF) shows how the system responds to a unit shock "
        "over time. For SARIMA, a unit increase in the innovation (white noise) propagates "
        "through the AR and MA polynomials, with the response decaying as the effect is "
        "absorbed. The cumulative impulse response (step response) shows the total accumulated "
        "effect, revealing the long-run multiplier. The forecast error variance decomposition "
        "shows how quickly forecast uncertainty accumulates over the prediction horizon."
    )

    for h in horizons:
        doc.add_heading(f"3.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "impulse_response"

        add_figure(
            h_dir / f"impulse_response_h{h}.png",
            f"Figure: Top - Impulse Response Function showing the dynamic response "
            f"to a unit shock. Bottom - Cumulative impulse response (step response) "
            f"showing the accumulated effect over time.",
        )

        add_figure(
            h_dir / f"variance_decomposition_h{h}.png",
            f"Figure: Forecast Error Variance Decomposition showing how forecast "
            f"uncertainty accumulates with horizon length.",
        )

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The IRF characterizes the dynamic behavior of the fitted SARIMA model. "
            "A rapidly decaying IRF indicates that shocks are quickly absorbed, while "
            "a slowly decaying or oscillating IRF suggests persistent effects. The seasonal "
            "component typically produces a periodic pattern in the IRF with peaks at "
            "multiples of 24 hours. The cumulative response converging to a finite value "
            "confirms model stability, while a diverging response would indicate instability."
        )

    doc.add_page_break()

    # ============================================================
    #  4. Section C: Residual Diagnostics
    # ============================================================
    doc.add_heading("4. Section C: Residual Diagnostics", level=1)

    doc.add_paragraph(
        "Residual diagnostics are critical for assessing SARIMA model adequacy. A well-"
        "specified model should produce residuals that resemble white noise: zero mean, "
        "constant variance, no autocorrelation, and approximately normal distribution. "
        "Systematic patterns in the residuals indicate model misspecification and suggest "
        "directions for improvement."
    )

    for h in horizons:
        doc.add_heading(f"4.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "residual_diagnostics"

        add_figure(
            h_dir / f"residual_timeseries_h{h}.png",
            f"Figure: Residual time series with +/-2-sigma bands. Points outside "
            f"the bands indicate unusual observations or model inadequacy.",
        )

        add_figure(
            h_dir / f"residual_distribution_h{h}.png",
            f"Figure: Left - Residual histogram with KDE and normal fit. "
            f"Right - Q-Q plot for normality assessment. Departures from the "
            f"diagonal indicate non-normality.",
        )

        add_figure(
            h_dir / f"acf_pacf_residuals_h{h}.png",
            f"Figure: ACF and PACF of residuals. Significant spikes outside the "
            f"confidence bounds indicate remaining autocorrelation structure not "
            f"captured by the model.",
        )

        add_figure(
            h_dir / f"ljungbox_test_h{h}.png",
            f"Figure: Ljung-Box test p-values at various lags. Values below 0.05 "
            f"(red line) indicate significant autocorrelation, suggesting the model "
            f"has not fully captured the temporal dependence structure.",
        )

        add_figure(
            h_dir / f"squared_resid_acf_h{h}.png",
            f"Figure: ACF of squared residuals. Significant autocorrelation indicates "
            f"conditional heteroscedasticity (time-varying variance), which could be "
            f"addressed by GARCH-type extensions.",
        )

        add_figure(
            h_dir / f"residuals_by_time_h{h}.png",
            f"Figure: Left - Residuals by hour of day. Right - Residuals by day of "
            f"week. Systematic patterns indicate the model fails to capture certain "
            f"temporal variations.",
        )

        add_figure(
            h_dir / f"rolling_variance_h{h}.png",
            f"Figure: Rolling residual variance with a 24-hour window. Non-constant "
            f"variance over time indicates heteroscedasticity.",
        )

        add_figure(
            h_dir / f"residual_fft_h{h}.png",
            f"Figure: Power spectrum of residuals. Peaks at specific periods "
            f"indicate periodic patterns not captured by the SARIMA model.",
        )

        # Residual statistics table
        rs = all_resid_stats.get(h, {})
        if rs:
            doc.add_heading("Residual Statistics", level=3)
            tbl = doc.add_table(rows=2, cols=6)
            tbl.style = "Light Grid Accent 1"
            for ci, ch in enumerate(["Mean", "Std", "Skewness", "Kurtosis", "Shapiro W", "Shapiro p"]):
                tbl.rows[0].cells[ci].text = ch
            tbl.rows[1].cells[0].text = f"{rs.get('mean', 0):.4f}"
            tbl.rows[1].cells[1].text = f"{rs.get('std', 0):.4f}"
            tbl.rows[1].cells[2].text = f"{rs.get('skewness', 0):.4f}"
            tbl.rows[1].cells[3].text = f"{rs.get('kurtosis', 0):.4f}"
            tbl.rows[1].cells[4].text = f"{rs.get('shapiro_stat', 0):.4f}"
            tbl.rows[1].cells[5].text = f"{rs.get('shapiro_pval', 0):.4f}"

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The residual diagnostics reveal the extent to which the SARIMA model captures "
            "the temporal structure of indoor CO2 dynamics. Significant ACF spikes at low "
            "lags suggest the model order may be insufficient, while spikes at seasonal lags "
            "indicate incomplete seasonal modeling. The Ljung-Box test provides a formal "
            "omnibus test for residual autocorrelation. Heavy tails in the Q-Q plot (common "
            "for CO2 data) reflect occasional extreme events like sudden ventilation changes "
            "or occupancy spikes. Significant squared residual ACF indicates conditional "
            "heteroscedasticity, suggesting GARCH-type extensions could improve the model."
        )

    doc.add_page_break()

    # ============================================================
    #  5. Section D: Forecast Uncertainty Analysis
    # ============================================================
    doc.add_heading("5. Section D: Forecast Uncertainty Analysis", level=1)

    doc.add_paragraph(
        "One of SARIMA's key advantages over neural network models is its natural ability "
        "to produce calibrated forecast uncertainty intervals. The model's error variance "
        "and parameter estimates yield analytically-derived confidence intervals that widen "
        "with the forecast horizon, reflecting increasing uncertainty for longer-range "
        "predictions."
    )

    for h in horizons:
        doc.add_heading(f"5.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "forecast_uncertainty"

        add_figure(
            h_dir / f"fan_chart_h{h}.png",
            f"Figure: Fan chart showing SARIMA forecasts with nested confidence "
            f"intervals (50%, 80%, 95%) for representative test windows. The "
            f"widening intervals reflect increasing forecast uncertainty.",
        )

        add_figure(
            h_dir / f"ci_width_vs_horizon_h{h}.png",
            f"Figure: Confidence interval width as a function of forecast horizon. "
            f"The rate of widening reflects the model's uncertainty growth rate.",
        )

        add_figure(
            h_dir / f"pi_coverage_h{h}.png",
            f"Figure: Prediction interval coverage comparison. The dashed red lines "
            f"show nominal coverage levels. Under-coverage indicates overconfident "
            f"intervals; over-coverage indicates conservative intervals.",
        )

        add_figure(
            h_dir / f"calibration_plot_h{h}.png",
            f"Figure: Calibration plot comparing nominal vs observed coverage. "
            f"Points on the diagonal indicate perfect calibration.",
        )

        # Coverage table
        us = all_uncertainty_stats.get(h, {})
        cov = us.get("coverage", {})
        if cov:
            doc.add_heading("Coverage Statistics", level=3)
            tbl = doc.add_table(rows=2, cols=len(cov))
            tbl.style = "Light Grid Accent 1"
            for ci, (nom, obs) in enumerate(sorted(cov.items())):
                tbl.rows[0].cells[ci].text = f"Nominal {nom}%"
                tbl.rows[1].cells[ci].text = f"{obs:.1f}%"

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The forecast uncertainty analysis evaluates the reliability of SARIMA's "
            "prediction intervals. Well-calibrated intervals (observed coverage matching "
            "nominal levels) indicate the model correctly quantifies its own uncertainty. "
            "Under-coverage at longer horizons is common and suggests the Gaussian assumption "
            "breaks down for multi-step predictions. The fan chart visualizations illustrate "
            "how uncertainty grows with the forecast horizon, providing actionable information "
            "for building management decision-making."
        )

    doc.add_page_break()

    # ============================================================
    #  6. Section E: Seasonal Decomposition
    # ============================================================
    doc.add_heading("6. Section E: Seasonal Decomposition", level=1)

    doc.add_paragraph(
        "STL (Seasonal and Trend decomposition using Loess) provides a model-free "
        "decomposition of the CO2 time series into trend, seasonal, and residual "
        "components. Comparing this decomposition with the SARIMA model's implicit "
        "seasonal handling reveals how well the parametric model captures the true "
        "seasonal structure."
    )

    for h in horizons:
        doc.add_heading(f"6.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "seasonal_decomposition"

        add_figure(
            h_dir / f"stl_decomposition_h{h}.png",
            f"Figure: STL decomposition of the training series into trend, "
            f"seasonal, and residual components with period=24 hours.",
        )

        add_figure(
            h_dir / f"differenced_series_h{h}.png",
            f"Figure: Original, first-differenced, and seasonally-differenced "
            f"series. Differencing removes trend and seasonal components, "
            f"which is what SARIMA's 'd' and 'D' parameters accomplish.",
        )

        add_figure(
            h_dir / f"seasonal_comparison_h{h}.png",
            f"Figure: STL-extracted seasonal pattern showing the typical daily "
            f"CO2 cycle. This represents the average diurnal pattern driven by "
            f"occupancy and ventilation schedules.",
        )

        # Decomposition stats
        ds = all_decomp_stats.get(h, {})
        if ds:
            doc.add_heading("Variance Decomposition", level=3)
            doc.add_paragraph(
                f"Trend explains {ds.get('trend_var_pct', 0):.1f}% of total variance, "
                f"seasonal component explains {ds.get('seasonal_var_pct', 0):.1f}%, "
                f"and residual accounts for {ds.get('residual_var_pct', 0):.1f}%."
            )

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The STL decomposition reveals the relative importance of trend, seasonal, "
            "and irregular components in indoor CO2 dynamics. A dominant seasonal component "
            "validates the choice of seasonal ARIMA over non-seasonal alternatives. The "
            "differenced series plots show the effect of the d=1 (first differencing) and "
            "D=1 (seasonal differencing) operations that SARIMA applies internally. If the "
            "differenced series appears stationary (constant mean and variance), the "
            "differencing orders are appropriate."
        )

    doc.add_page_break()

    # ============================================================
    #  7. Section F: SARIMA-Specific Analyses
    # ============================================================
    doc.add_heading("7. Section F: SARIMA-Specific Analyses", level=1)

    doc.add_paragraph(
        "This section presents stationarity testing, prediction quality assessment, "
        "and error stratification by CO2 concentration level. The ADF (Augmented "
        "Dickey-Fuller) and KPSS (Kwiatkowski-Phillips-Schmidt-Shin) tests provide "
        "complementary evidence about the stationarity of the original and differenced "
        "series."
    )

    for h in horizons:
        doc.add_heading(f"7.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "sarima_specific"

        add_figure(
            h_dir / f"stationarity_tests_h{h}.png",
            f"Figure: ADF and KPSS stationarity test results for the original, "
            f"first-differenced, and seasonally-differenced series. Green indicates "
            f"stationarity, red indicates non-stationarity.",
        )

        h_pred_dir = output_dir / f"h{h}" / "predictions"

        add_figure(
            h_pred_dir / f"predictions_overlay_h{h}.png",
            f"Figure: Time series overlay of actual vs predicted CO2 concentrations.",
        )

        add_figure(
            h_pred_dir / f"scatter_r2_h{h}.png",
            f"Figure: Scatter plot of predicted vs actual CO2 with R-squared, "
            f"RMSE, and MAE statistics.",
        )

        add_figure(
            h_pred_dir / f"residual_analysis_h{h}.png",
            f"Figure: Four-panel residual analysis: distribution, time trend, "
            f"residuals vs predicted, and Q-Q plot.",
        )

        add_figure(
            h_pred_dir / f"error_by_co2_level_h{h}.png",
            f"Figure: Prediction error stratified by CO2 concentration level.",
        )

        # Metrics table
        doc.add_heading("Performance Metrics", level=3)
        m = all_results[h]["metrics"]
        tbl = doc.add_table(rows=2, cols=5)
        tbl.style = "Light Grid Accent 1"
        for ci, ch in enumerate(["RMSE", "MAE", "R2", "MAPE", "MSE"]):
            tbl.rows[0].cells[ci].text = ch
        tbl.rows[1].cells[0].text = f"{m['rmse']:.2f}"
        tbl.rows[1].cells[1].text = f"{m['mae']:.2f}"
        tbl.rows[1].cells[2].text = f"{m['r2']:.4f}"
        tbl.rows[1].cells[3].text = f"{m['mape']:.2f}%"
        tbl.rows[1].cells[4].text = f"{m['mse']:.2f}"

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The stationarity tests confirm whether the differencing operations applied "
            "by SARIMA are appropriate. The ADF test (null: unit root present) and KPSS "
            "test (null: series is stationary) provide complementary evidence. Both tests "
            "agreeing on stationarity after differencing validates the model specification. "
            "The prediction analysis shows the model's accuracy across different CO2 "
            "concentration regimes. Higher errors at extreme concentrations are expected "
            "due to nonlinear dynamics that the linear SARIMA model cannot fully capture."
        )

    doc.add_page_break()

    # ============================================================
    #  8. Cross-Horizon Comparison
    # ============================================================
    if len(horizons) > 1:
        doc.add_heading("8. Cross-Horizon Comparison", level=1)

        doc.add_paragraph(
            "Comparing SARIMA performance across forecast horizons reveals how prediction "
            "quality degrades with increasing lead time. SARIMA generates multi-step "
            "forecasts by iterating forward one step at a time, conditioning on its own "
            "predictions. This recursive structure causes errors to compound over longer "
            "horizons."
        )

        comp_dir = output_dir / "comparison"
        add_figure(
            comp_dir / "metrics_comparison.png",
            "Figure: Side-by-side performance metrics across horizons.",
        )
        add_figure(
            comp_dir / "error_growth.png",
            "Figure: RMSE growth as a function of forecast step, showing "
            "how prediction errors compound over the horizon.",
        )

        # Comparison table
        doc.add_heading("Metrics Comparison Table", level=2)
        tbl = doc.add_table(rows=1 + len(horizons), cols=5)
        tbl.style = "Light Grid Accent 1"
        for ci, ch in enumerate(["Horizon", "RMSE", "MAE", "R2", "MAPE"]):
            tbl.rows[0].cells[ci].text = ch
        for ri, h in enumerate(horizons):
            m = all_results[h]["metrics"]
            tbl.rows[ri + 1].cells[0].text = f"{h}h"
            tbl.rows[ri + 1].cells[1].text = f"{m['rmse']:.2f}"
            tbl.rows[ri + 1].cells[2].text = f"{m['mae']:.2f}"
            tbl.rows[ri + 1].cells[3].text = f"{m['r2']:.4f}"
            tbl.rows[ri + 1].cells[4].text = f"{m['mape']:.2f}%"

        doc.add_heading("Discussion", level=2)
        doc.add_paragraph(
            "The cross-horizon comparison typically shows significant performance degradation "
            "from 1h to 24h forecasting. This is expected for SARIMA because: (1) recursive "
            "multi-step forecasting compounds one-step errors; (2) the linear model cannot "
            "adapt to nonlinear regime changes over longer horizons; (3) the univariate "
            "specification ignores exogenous drivers (temperature, humidity, occupancy) that "
            "become increasingly important for longer-term predictions. The error growth curve "
            "reveals the rate at which predictive skill decays."
        )

    doc.add_page_break()

    # ============================================================
    #  9. Avenues of Improvement
    # ============================================================
    doc.add_heading("9. Avenues of Improvement", level=1)

    improvements = [
        (
            "SARIMAX with Exogenous Variables",
            "The current SARIMA model is univariate, ignoring available features like "
            "temperature, humidity, pressure, and noise level. SARIMAX extends SARIMA "
            "with exogenous regressors, allowing the model to incorporate these physical "
            "drivers of CO2 dynamics. This is likely the single largest improvement "
            "opportunity, as environmental conditions directly affect ventilation rates "
            "and occupancy patterns that drive CO2 variations."
        ),
        (
            "GARCH Extensions for Conditional Heteroscedasticity",
            "If the squared residual ACF analysis reveals significant autocorrelation, "
            "this indicates conditional heteroscedasticity (time-varying variance). "
            "SARIMA-GARCH models explicitly model the conditional variance process, "
            "producing more accurate prediction intervals. This is particularly relevant "
            "for CO2 data where variance may be higher during occupied hours than "
            "during unoccupied periods."
        ),
        (
            "Higher-Order Seasonal Components",
            "The current model captures only the daily (s=24) seasonal cycle. Indoor "
            "CO2 also exhibits weekly patterns (weekday vs weekend occupancy). Adding "
            "a second seasonal component (s=168 for weekly) or using Fourier terms "
            "via SARIMAX's exogenous variable interface could capture this structure. "
            "Fourier terms are computationally cheaper than full seasonal ARIMA with "
            "large seasonal periods."
        ),
        (
            "Regime-Switching Models (Markov-Switching ARIMA)",
            "Indoor CO2 dynamics exhibit distinct regimes: occupied (rising CO2, high "
            "variance) and unoccupied (falling CO2, low variance). Markov-Switching "
            "ARIMA models explicitly model these regime transitions, allowing different "
            "AR/MA parameters in each regime. This can substantially improve forecasting "
            "during regime transitions."
        ),
        (
            "Hybrid SARIMA + Neural Network Correction",
            "The SARIMA residuals contain information about nonlinear patterns the "
            "linear model cannot capture. Training a neural network (LSTM, MLP) on "
            "the SARIMA residuals with access to the full feature set creates a "
            "complementary hybrid: SARIMA handles the linear temporal structure while "
            "the neural network captures nonlinear residual patterns. This approach "
            "often outperforms either model alone."
        ),
        (
            "Automated Order Selection (auto_arima)",
            "The current (1,1,1)x(1,1,1,24) specification was chosen manually. "
            "Automated order selection using AIC/BIC grid search (e.g., pmdarima's "
            "auto_arima) could identify a better-fitting specification. Higher-order "
            "AR or MA terms might capture dynamics the current model misses."
        ),
        (
            "Longer Lookback Windows for Prediction Conditioning",
            "The current sliding-window approach conditions each forecast on a fixed "
            "lookback window. For SARIMA, longer conditioning histories provide more "
            "context for the Kalman filter state initialization, potentially improving "
            "forecast accuracy especially for seasonal patterns."
        ),
        (
            "Robust Estimation Methods",
            "CO2 time series contain occasional spikes from extreme events (e.g., "
            "gatherings, cooking, sensor glitches). M-estimation or other robust "
            "estimation methods reduce the influence of these outliers on parameter "
            "estimates, producing a model that better reflects the typical CO2 dynamics "
            "rather than being distorted by rare extreme observations."
        ),
    ]

    for imp_title, imp_desc in improvements:
        doc.add_heading(imp_title, level=2)
        doc.add_paragraph(imp_desc)

    # Save report
    report_path = output_dir / "sarima_interpretability_report.docx"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(report_path))
    print(f"  DOCX report saved to: {report_path}")

    return report_path


# ======================================================================
#  Main
# ======================================================================

def main() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )

    parser = argparse.ArgumentParser(
        description="SARIMA Interpretability Study on preproc_D (Enhanced 1h)"
    )
    parser.add_argument(
        "--horizons", nargs="+", type=int, default=[1, 24],
        help="Forecast horizons in hours (default: 1 24)",
    )
    args = parser.parse_args()

    print(f"\n{'='*70}")
    print(f"  SARIMA INTERPRETABILITY STUDY (DEEP ANALYSIS)")
    print(f"  Variant: preproc_D (Enhanced 1h)")
    print(f"  Model: SARIMA(1,1,1)x(1,1,1,24)")
    print(f"  Horizons: {args.horizons}")
    print(f"  Sections: A(params) B(impulse) C(residuals) D(uncertainty)")
    print(f"            E(seasonal) F(stationarity+predictions)")
    print(f"{'='*70}\n")

    # Load pipeline data once (shared across horizons)
    pipeline_config = load_interpretability_config(horizon=1)
    seed_everything(pipeline_config["training"]["seed"])

    raw_dir = Path(pipeline_config["data"].get("raw_dir", "data/raw"))
    print("  Loading preprocessing pipeline (preproc_D)...")
    train_df, val_df, test_df = run_preprocessing_pipeline(
        raw_dir=raw_dir, variant_config=pipeline_config,
    )
    print(f"  Pipeline loaded: train={len(train_df)}, val={len(val_df)}, test={len(test_df)}")

    all_results: dict = {}
    all_param_stats: dict = {}
    all_irf_data: dict = {}
    all_resid_stats: dict = {}
    all_uncertainty_stats: dict = {}
    all_decomp_stats: dict = {}
    all_specific_stats: dict = {}

    for horizon in args.horizons:
        print(f"\n{'-'*60}")
        print(f"  HORIZON: {horizon}h")
        print(f"{'-'*60}\n")

        config = load_interpretability_config(horizon=horizon)
        seed_everything(config["training"]["seed"])

        output_dir = RESULTS_BASE / f"h{horizon}"
        output_dir.mkdir(parents=True, exist_ok=True)

        # ---- Fit SARIMA and get predictions ----
        t0 = time.time()
        print(f"  Fitting SARIMA for {horizon}h horizon...")
        sarima, dm, results = train_sarima(
            config, train_df.copy(), val_df.copy(), test_df.copy(), horizon,
        )
        elapsed = time.time() - t0
        print(f"  SARIMA fitting + prediction completed in {elapsed:.1f}s")

        metrics = results["metrics"]
        all_results[horizon] = results
        print(f"  RMSE={metrics['rmse']:.2f}  MAE={metrics['mae']:.2f}  "
              f"R2={metrics['r2']:.4f}  MAPE={metrics['mape']:.2f}%")

        train_series_scaled = results["train_series_scaled"]

        # ---- Section A: Parameter Analysis ----
        try:
            param_stats = run_parameter_analysis(
                sarima, horizon, output_dir / "parameters"
            )
            all_param_stats[horizon] = param_stats
        except Exception as e:
            print(f"  Section A failed: {e}")
            all_param_stats[horizon] = {}

        # ---- Section B: Impulse Response ----
        try:
            irf_data = run_impulse_response(
                sarima, horizon, output_dir / "impulse_response"
            )
            all_irf_data[horizon] = irf_data
        except Exception as e:
            print(f"  Section B failed: {e}")
            all_irf_data[horizon] = {}

        # ---- Section C: Residual Diagnostics ----
        try:
            resid_stats = run_residual_diagnostics(
                sarima, dm, results, test_df, horizon,
                output_dir / "residual_diagnostics",
            )
            all_resid_stats[horizon] = resid_stats
        except Exception as e:
            print(f"  Section C failed: {e}")
            all_resid_stats[horizon] = {}

        # ---- Section D: Forecast Uncertainty ----
        try:
            uncertainty_stats = run_forecast_uncertainty(
                sarima, dm, results, horizon,
                output_dir / "forecast_uncertainty",
            )
            all_uncertainty_stats[horizon] = uncertainty_stats
        except Exception as e:
            print(f"  Section D failed: {e}")
            all_uncertainty_stats[horizon] = {}

        # ---- Section E: Seasonal Decomposition ----
        try:
            decomp_stats = run_seasonal_decomposition(
                sarima, train_series_scaled, horizon,
                output_dir / "seasonal_decomposition",
            )
            all_decomp_stats[horizon] = decomp_stats
        except Exception as e:
            print(f"  Section E failed: {e}")
            all_decomp_stats[horizon] = {}

        # ---- Section F: SARIMA-Specific + Predictions ----
        try:
            # Stationarity tests go in sarima_specific subfolder
            specific_stats = run_sarima_specific(
                sarima, dm, results, train_series_scaled, test_df, horizon,
                output_dir / "sarima_specific",
            )
            all_specific_stats[horizon] = specific_stats

            # Prediction plots go in predictions subfolder
            # (already generated inside run_sarima_specific, but the overlay/scatter
            # plots are saved in the sarima_specific dir -- let's also save in predictions)
            pred_dir = output_dir / "predictions"
            pred_dir.mkdir(parents=True, exist_ok=True)
            # Copy prediction-related plots to predictions/ subfolder
            import shutil
            for fname in [
                f"predictions_overlay_h{horizon}.png",
                f"scatter_r2_h{horizon}.png",
                f"residual_analysis_h{horizon}.png",
                f"error_by_co2_level_h{horizon}.png",
                f"metrics_table_h{horizon}.png",
            ]:
                src = output_dir / "sarima_specific" / fname
                dst = pred_dir / fname
                if src.exists():
                    shutil.copy2(str(src), str(dst))
        except Exception as e:
            print(f"  Section F failed: {e}")
            all_specific_stats[horizon] = {}

        # ---- Save metrics + predictions ----
        save_metrics(
            metrics, f"SARIMA_h{horizon}", output_dir / "metrics.json",
            experiment_info={
                "name": "sarima_interpretability",
                "label": f"SARIMA Deep Analysis h={horizon}",
                "description": "preproc_D Enhanced 1h variant",
            },
        )
        np.savez(
            output_dir / "predictions.npz",
            y_true=results["y_true"],
            y_pred=results["y_pred"],
        )

        # ---- Cleanup ----
        del sarima, dm
        gc.collect()
        print(f"  Memory freed\n")

    # ---- Cross-horizon comparison ----
    if len(args.horizons) > 1:
        print(f"\n{'-'*60}")
        print(f"  CROSS-HORIZON COMPARISON")
        print(f"{'-'*60}\n")
        run_cross_horizon_comparison(
            all_results, RESULTS_BASE / "comparison"
        )

    # ---- Summary ----
    generate_summary_figure(
        all_results, all_param_stats, all_resid_stats, RESULTS_BASE
    )
    save_study_results(
        all_results, all_param_stats, all_resid_stats,
        all_decomp_stats, all_uncertainty_stats, all_specific_stats,
        RESULTS_BASE,
    )

    # ---- DOCX Report ----
    print(f"\n{'-'*60}")
    print(f"  GENERATING DOCX REPORT")
    print(f"{'-'*60}\n")
    report_path = generate_docx_report(
        all_results, all_param_stats, all_resid_stats,
        all_irf_data, all_decomp_stats, all_uncertainty_stats,
        all_specific_stats, args.horizons, RESULTS_BASE,
    )

    print(f"\n{'='*70}")
    print(f"  SARIMA INTERPRETABILITY STUDY COMPLETE")
    print(f"  Results saved to: {RESULTS_BASE}")
    print(f"  Report: {report_path}")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
