"""Generate an academic DOCX report on occupancy inference experiments.

Produces a formatted Word document covering:
  - Rule-based detectors and consensus evaluation
  - HMM with temporal features (Gaussian sin/cos)
  - HMM with circular distributions (von Mises)
  - Physics-informed Switching AR-HMM
  - Comparative analysis and discussion
"""

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
RESULTS = PROJECT_ROOT / "results"
REPORT_DIR = PROJECT_ROOT / "reports"
REPORT_DIR.mkdir(parents=True, exist_ok=True)

OUTPUT_PATH = REPORT_DIR / "Occupancy_Inference_Report.docx"

# ── Result paths ─────────────────────────────────────────────────
OCC_ACTUAL = RESULTS / "occupancy" / "actual" / "consensus.json"
OCC_PREDICTED = RESULTS / "occupancy" / "predicted" / "consensus.json"
OCC_COMPARISON = RESULTS / "occupancy" / "comparison" / "actual_vs_predicted.json"
HMM_TEMPORAL = RESULTS / "hmm_temporal_occupancy" / "experiment_results.json"
HMM_CIRCULAR = RESULTS / "hmm_circular_vs_gaussian" / "experiment_results.json"
SWITCHING_AR = RESULTS / "switching_ar_occupancy" / "experiment_results.json"

# ── Figure paths ─────────────────────────────────────────────────
OCC_FIGS = RESULTS / "occupancy" / "figures"
HMM_TEMP_FIGS = RESULTS / "hmm_temporal_occupancy"
HMM_CIRC_FIGS = RESULTS / "hmm_circular_vs_gaussian"
SWAR_FIGS = RESULTS / "switching_ar_occupancy"


def load_json(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def set_cell_shading(cell, color_hex: str):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:fill"): color_hex},
    )
    shading_elm.append(shading)


def add_styled_table(doc, headers, rows, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
        if highlight_rows and i in highlight_rows:
            for j in range(len(row)):
                set_cell_shading(table.rows[i + 1].cells[j], "E2EFDA")

    return table


def add_figure(doc, path, caption, width=Inches(5.8)):
    if Path(path).exists():
        doc.add_picture(str(path), width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
    else:
        doc.add_paragraph(f"[Figure not found: {path}]")


def add_bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def add_equation(doc, text):
    """Add a centered, italic equation-like paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    return p


def build_report():
    # ── Load all results ─────────────────────────────────────────
    actual_consensus = load_json(OCC_ACTUAL)
    predicted_consensus = load_json(OCC_PREDICTED)
    comparison = load_json(OCC_COMPARISON)
    hmm_temporal = load_json(HMM_TEMPORAL)
    hmm_circular = load_json(HMM_CIRCULAR)
    switching_ar = load_json(SWITCHING_AR)

    doc = Document()

    # ── Styles ───────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ══════════════════════════════════════════════════════════════
    #  TITLE PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Occupancy Inference from Indoor CO\u2082 Measurements:\n"
        "A Comparative Study of Rule-Based, HMM,\n"
        "and Physics-Informed Approaches"
    )
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "SalonBM Dataset \u2014 Hourly CO\u2082, Noise, and Derived Features"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Abstract",
        "2. Introduction",
        "3. Dataset and Features",
        "4. Rule-Based Occupancy Detectors",
        "   4.1 Detector Descriptions",
        "   4.2 Consensus Evaluation",
        "   4.3 Actual vs. Predicted CO\u2082",
        "5. HMM-Based Occupancy Detection",
        "   5.1 Gaussian HMM Baseline",
        "   5.2 HMM with Temporal Features (sin/cos)",
        "   5.3 Mixed Gaussian\u2013von Mises HMM",
        "6. Physics-Informed Switching AR-HMM",
        "   6.1 Model Formulation",
        "   6.2 Occupancy Mapping Strategy",
        "   6.3 Experimental Results",
        "7. Comparative Analysis",
        "8. Discussion",
        "9. Limitations",
        "10. Future Work",
        "11. Conclusions",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  1. ABSTRACT
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. Abstract", level=1)

    doc.add_paragraph(
        "Occupancy detection in buildings is critical for energy-efficient HVAC control, "
        "demand-driven ventilation, and smart building management. This report presents a "
        "comprehensive evaluation of occupancy inference methods using indoor CO\u2082 concentration "
        "measurements from the SalonBM dataset, a residential living room monitored at hourly "
        "intervals. We compare six rule-based detectors, Gaussian Hidden Markov Models (HMMs) "
        "with various feature sets, a custom Mixed Gaussian\u2013von Mises HMM for circular temporal "
        "features, and a physics-informed Switching Autoregressive HMM grounded in the CO\u2082 "
        "mass-balance equation."
    )
    doc.add_paragraph(
        "In the absence of ground-truth occupancy labels, we adopt a majority-vote consensus "
        "among the six rule-based detectors as a pseudo-ground-truth reference. The best "
        "performing approach is the Gaussian HMM using CO\u2082, dCO\u2082/dt, and noise level "
        "(\u03ba = 0.53 vs. consensus), while adding temporal features consistently degrades "
        "performance. The physics-informed Switching AR model offers superior interpretability "
        "through learned ventilation time constants and CO\u2082 generation rates, and achieves "
        "the highest agreement with the rate-of-change detector (\u03ba = 0.70), but scores "
        "lower against consensus (\u03ba = 0.32). These findings highlight a fundamental tension "
        "between dynamics-based and level-based occupancy definitions."
    )

    # ══════════════════════════════════════════════════════════════
    #  2. INTRODUCTION
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. Introduction", level=1)

    doc.add_paragraph(
        "Indoor air quality monitoring has become increasingly important for occupant health, "
        "energy management, and building automation. CO\u2082 concentration serves as a reliable "
        "proxy for human presence because metabolic respiration directly generates CO\u2082, "
        "causing measurable increases relative to the outdoor baseline (~420 ppm). When "
        "occupants leave, ventilation dilutes the indoor CO\u2082 back toward ambient levels."
    )
    doc.add_paragraph(
        "The key challenge is that CO\u2082 is a lagging indicator: concentrations remain elevated "
        "after occupants depart (due to the ventilation time constant \u03c4) and may not rise "
        "immediately upon arrival (due to the slow accumulation rate relative to sampling "
        "frequency). This makes simple thresholding unreliable, motivating more sophisticated "
        "approaches that consider dynamics (rate of change), temporal patterns (time of day), "
        "and physics-based models (ventilation equations)."
    )
    doc.add_paragraph(
        "This report evaluates four families of approaches across a common dataset and "
        "evaluation protocol:"
    )
    doc.add_paragraph(
        "Rule-based detectors: Six heuristic methods ranging from absolute CO\u2082 thresholds "
        "to state machines and the literature-based Diarra method.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Gaussian HMM: Standard multivariate Gaussian emission HMM with exhaustive "
        "binary mapping search over state-to-occupancy assignments.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Mixed Gaussian\u2013von Mises HMM: A custom HMM combining Gaussian emissions for "
        "continuous features with von Mises distributions for circular temporal features.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Physics-informed Switching AR-HMM: An autoregressive model derived from the CO\u2082 "
        "mass-balance ODE, with learned ventilation and generation parameters.",
        style="List Bullet",
    )

    # ══════════════════════════════════════════════════════════════
    #  3. DATASET AND FEATURES
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. Dataset and Features", level=1)

    doc.add_paragraph(
        "The SalonBM dataset contains continuous sensor measurements from a residential "
        "living room in France, resampled to 1-hour intervals. The test set comprises "
        "2,260 hourly observations (~94 days). No ground-truth occupancy labels are available."
    )

    doc.add_heading("Sensor Variables", level=2)
    features = [
        ("CO\u2082 (ppm)", "Indoor carbon dioxide concentration; primary occupancy indicator"),
        ("dCO\u2082/dt (ppm/h)", "First derivative of CO\u2082; captures arrival/departure dynamics"),
        ("Noise (dB)", "Ambient sound level; secondary human activity indicator"),
    ]
    for name, desc in features:
        p = doc.add_paragraph(style="List Bullet")
        add_bold_run(p, f"{name}: ")
        p.add_run(desc)

    doc.add_heading("Derived Temporal Features", level=2)
    temporal = [
        ("Day_sin, Day_cos", "24-hour cycle encoded as sin/cos of 2\u03c0\u00b7hour/24"),
        ("Weekday_sin, Weekday_cos", "7-day cycle encoded as sin/cos of 2\u03c0\u00b7weekday/7"),
        ("hour_weekday_sin/cos", "168-hour combined weekly cycle"),
        ("hour (angle)", "Hour of day as an angle in [0, 2\u03c0) for von Mises distributions"),
        ("day-of-week (angle)", "Day of week as an angle in [0, 2\u03c0) for von Mises distributions"),
    ]
    for name, desc in temporal:
        p = doc.add_paragraph(style="List Bullet")
        add_bold_run(p, f"{name}: ")
        p.add_run(desc)

    # ══════════════════════════════════════════════════════════════
    #  4. RULE-BASED OCCUPANCY DETECTORS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. Rule-Based Occupancy Detectors", level=1)

    # 4.1 Detector Descriptions
    doc.add_heading("4.1 Detector Descriptions", level=2)

    detectors_desc = [
        (
            "Absolute Threshold",
            "The simplest approach: occupied = CO\u2082 > 500 ppm. Effective when ventilation "
            "is stable but unable to distinguish sustained elevated CO\u2082 after departure from "
            "active presence.",
        ),
        (
            "Rate of Change",
            "Occupied when the smoothed dCO\u2082/dt > 5.0 ppm/h (3-step rolling window). "
            "Captures arrival events through rising CO\u2082 but misses stable high-occupancy "
            "periods where the rate flattens.",
        ),
        (
            "Adaptive Threshold",
            "Learns per-hour CO\u2082 thresholds from training data as the 75th percentile + 50 ppm "
            "margin. Most conservative detector, yielding only 10.4% occupancy.",
        ),
        (
            "Hybrid",
            "OR-gate fusion: occupied = (absolute threshold) OR (rate of change). Reduces "
            "false negatives by combining level and dynamics signals.",
        ),
        (
            "State Machine",
            "Four-state finite automaton (UNOCCUPIED \u2192 ONSET \u2192 OCCUPIED \u2192 DECAY) with "
            "configurable thresholds (onset dCO\u2082 > 5 ppm/h, sustain CO\u2082 > 500 ppm, "
            "6-step timeout, 2-step minimum onset). The stickiest detector at 91.8% occupancy.",
        ),
        (
            "Diarra (Diarra et al., Sensors 2023)",
            "Four-state classification using CO\u2082, dCO\u2082, and noise simultaneously: "
            "prolonged absence (CO\u2082 low, dCO\u2082 \u2264 0, noise low), presence (CO\u2082 > 600 ppm "
            "OR dCO\u2082 > 5 OR noise > 55 dB), absence (dCO\u2082 < 0, noise low), and inactive "
            "presence (CO\u2082 high, dCO\u2082 \u2264 0, noise low). Thresholds tuned for SalonBM.",
        ),
    ]
    for name, desc in detectors_desc:
        p = doc.add_paragraph()
        add_bold_run(p, f"{name}. ")
        p.add_run(desc)

    # 4.2 Consensus Evaluation
    doc.add_heading("4.2 Consensus Evaluation", level=2)

    doc.add_paragraph(
        "Without ground-truth labels, we define the consensus as a majority vote across all "
        "six detectors: a timestep is labeled as occupied if at least four of six detectors "
        "agree. This provides a pseudo-ground-truth that balances diverse detection philosophies."
    )

    # Occupancy rates table
    occ_rates = actual_consensus["occupancy_rates"]
    det_names = ["absolute_threshold", "rate_of_change", "adaptive_threshold",
                 "hybrid", "state_machine", "diarra"]
    det_labels = ["Absolute Threshold", "Rate of Change", "Adaptive Threshold",
                  "Hybrid", "State Machine", "Diarra"]

    rows_occ = []
    for name, label in zip(det_names, det_labels):
        rows_occ.append([label, f"{occ_rates[name]*100:.1f}%"])
    rows_occ.append(["Consensus (majority vote)", f"{actual_consensus['majority_vote_occupancy_rate']*100:.1f}%"])

    doc.add_paragraph()
    p = doc.add_paragraph()
    add_bold_run(p, "Table 1. ")
    p.add_run("Occupancy rates by detector (actual CO\u2082).")
    add_styled_table(doc, ["Detector", "Occupancy Rate"], rows_occ)

    doc.add_paragraph()
    doc.add_paragraph(
        f"The inter-detector agreement measured by Fleiss' kappa is "
        f"{actual_consensus['fleiss_kappa']:.4f}, indicating weak overall agreement. "
        f"This is expected given the fundamentally different detection philosophies: "
        f"the rate-of-change detector fires only during dynamic events (36.9% occupancy), "
        f"while the state machine is sticky by design (91.8%). The mean consistency "
        f"(fraction of timesteps where a majority agrees) is "
        f"{actual_consensus['mean_consistency']*100:.1f}%."
    )

    # Cohen's kappa matrix
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_bold_run(p, "Table 2. ")
    p.add_run("Pairwise Cohen's kappa between detectors (actual CO\u2082).")

    kappa_matrix = actual_consensus["cohens_kappa_matrix"]
    kappa_headers = [""] + [l[:8] for l in det_labels]
    kappa_rows = []
    for i, label in enumerate(det_labels):
        row = [label[:12]]
        for j in range(len(det_labels)):
            if i == j:
                row.append("1.00")
            else:
                row.append(f"{kappa_matrix[i][j]:.2f}")
        kappa_rows.append(row)
    add_styled_table(doc, kappa_headers, kappa_rows)

    doc.add_paragraph()
    doc.add_paragraph(
        "The strongest pairwise agreement is between the Absolute Threshold and Hybrid "
        "detectors (\u03ba = 0.81), which is expected since the Hybrid includes the absolute "
        "threshold as one of its two components. The weakest agreement is between the Rate "
        "of Change and State Machine detectors (\u03ba \u2248 0.00), reflecting their fundamentally "
        "different operating principles."
    )

    # Figures
    add_figure(
        doc,
        OCC_FIGS / "fig1_detector_timeline_actual.png",
        "Figure 1. Timeline of all six detector outputs on actual CO\u2082 test data.",
    )
    add_figure(
        doc,
        OCC_FIGS / "fig3_occupancy_rates.png",
        "Figure 2. Comparison of occupancy rates across detectors.",
        width=Inches(4.5),
    )
    add_figure(
        doc,
        OCC_FIGS / "fig4_agreement_heatmap_kappa.png",
        "Figure 3. Heatmap of pairwise Cohen's kappa between detectors.",
        width=Inches(4.5),
    )

    # 4.3 Actual vs Predicted
    doc.add_heading("4.3 Actual vs. Predicted CO\u2082", level=2)

    doc.add_paragraph(
        "To assess whether a forecasting model's predicted CO\u2082 preserves occupancy-relevant "
        "information, we run all six detectors on predicted CO\u2082 and compare the resulting "
        "occupancy labels to those from actual CO\u2082."
    )

    comp_rows = []
    for name, label in zip(det_names, det_labels):
        k = comparison[name]["kappa"]
        a = comparison[name]["agreement"]
        comp_rows.append([label, f"{k:.4f}", f"{a*100:.1f}%"])

    p = doc.add_paragraph()
    add_bold_run(p, "Table 3. ")
    p.add_run("Agreement between actual and predicted CO\u2082 occupancy detections.")
    add_styled_table(
        doc,
        ["Detector", "Cohen's \u03ba", "Raw Agreement"],
        comp_rows,
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "All detectors show strong agreement (\u03ba > 0.73) between actual and predicted CO\u2082, "
        "indicating that the forecasting model successfully preserves occupancy-relevant "
        "signal characteristics. The adaptive threshold achieves the highest raw agreement "
        "(97.5%) due to its conservative nature."
    )

    add_figure(
        doc,
        OCC_FIGS / "fig6_actual_vs_predicted.png",
        "Figure 4. Actual vs. predicted CO\u2082 occupancy detection comparison.",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  5. HMM-BASED OCCUPANCY DETECTION
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. HMM-Based Occupancy Detection", level=1)

    doc.add_paragraph(
        "Hidden Markov Models offer a principled probabilistic framework for occupancy "
        "detection. Each HMM state implicitly captures a distinct CO\u2082 regime (e.g., "
        "unoccupied baseline, occupied with rising CO\u2082, decay after departure). The "
        "state-to-occupancy mapping is determined post-hoc via exhaustive search over "
        "all 2\u1d37 \u2212 2 possible binary assignments of K states to {occupied, unoccupied}, "
        "selecting the mapping that maximizes Cohen's kappa against the consensus."
    )

    # 5.1 Gaussian HMM Baseline
    doc.add_heading("5.1 Gaussian HMM Baseline", level=2)

    baseline = hmm_temporal["Exp0: CO2+dCO2+Noise (baseline, 4 states)"]
    doc.add_paragraph(
        f"The Gaussian HMM baseline uses three features (CO\u2082, dCO\u2082/dt, Noise) with "
        f"K=4 states. After training on the training set and decoding the test set via "
        f"the Viterbi algorithm, exhaustive bitmask search yields the best consensus "
        f"agreement at \u03ba = {baseline['best_binary_mapping']['kappa_vs_consensus']:.4f} "
        f"with an occupancy rate of "
        f"{baseline['best_binary_mapping']['occupancy_rate']}%."
    )

    # State profiles table
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_bold_run(p, "Table 4. ")
    p.add_run("Gaussian HMM baseline state profiles (4 states, CO\u2082+dCO\u2082+Noise).")

    state_rows = []
    for sp in baseline["state_profiles"]:
        state_rows.append([
            str(sp["state"]),
            f"{sp['pct']}%",
            f"{sp['means']['CO2']:.0f}",
            f"{sp['means']['dCO2']:+.1f}",
            f"{sp['means']['Noise']:.1f}",
            f"{sp['consensus_occupied_pct']:.0f}%",
        ])
    add_styled_table(
        doc,
        ["State", "Time %", "CO\u2082 (ppm)", "dCO\u2082/dt", "Noise (dB)", "Consensus Occ."],
        state_rows,
        highlight_rows=[1, 2, 3],  # occupied states
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "The four states have clear physical interpretations: State 0 represents unoccupied "
        "baseline (CO\u2082 \u2248 485 ppm, 47.7% consensus), State 1 captures active occupancy with "
        "rising CO\u2082 (657 ppm, +9.8 ppm/h, 95.4% consensus), State 2 is a rare high-CO\u2082 "
        "state (2.7% of time, 100% consensus), and State 3 represents post-occupancy decay "
        "(608 ppm, \u221212.4 ppm/h, 94.6% consensus)."
    )

    # 5.2 HMM with Temporal Features
    doc.add_heading("5.2 HMM with Temporal Features (sin/cos)", level=2)

    doc.add_paragraph(
        "We hypothesized that encoding time-of-day and day-of-week as sin/cos features "
        "would help the HMM discover occupancy-aligned states by learning diurnal and "
        "weekly patterns. Seven configurations were tested with 4\u20136 states."
    )

    # Results table
    temporal_exps = [
        ("Exp0: CO2+dCO2+Noise (baseline, 4 states)", "CO\u2082+dCO\u2082+Noise", 4),
        ("Exp1: +Day_sin/cos (24h cycle)", "+Day sin/cos", 4),
        ("Exp2: +Day+Weekday (24h+7d cycles)", "+Day+Weekday", 4),
        ("Exp3: +hour_weekday (168h cycle)", "+hour_weekday", 4),
        ("Exp4: All temporal features", "+All temporal", 4),
        ("Exp5: +Day_sin/cos (24h), 5 states", "+Day sin/cos", 5),
        ("Exp6: +Day+Weekday, 6 states", "+Day+Weekday", 6),
    ]

    temp_rows = []
    for key, label, k in temporal_exps:
        exp = hmm_temporal[key]
        bm = exp["best_binary_mapping"]
        temp_rows.append([
            label,
            str(k),
            f"{bm['kappa_vs_consensus']:.4f}",
            f"{bm['occupancy_rate']}%",
            str(bm['occupied_states']),
        ])

    p = doc.add_paragraph()
    add_bold_run(p, "Table 5. ")
    p.add_run("Gaussian HMM occupancy results with temporal features.")
    add_styled_table(
        doc,
        ["Features", "K", "\u03ba vs. Consensus", "Occ. Rate", "Occupied States"],
        temp_rows,
        highlight_rows=[0],  # baseline is best
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "All temporal feature configurations degrade performance relative to the baseline. "
        "The worst degradation occurs with the 168-hour weekly cycle (Exp3, \u03ba = 0.04), "
        "where the HMM learns to segment by time-of-day rather than by CO\u2082 dynamics. "
        "Even with additional states (Exp5: 5 states, Exp6: 6 states), temporal features "
        "fail to recover baseline performance."
    )
    doc.add_paragraph(
        "The root cause is that Gaussian HMM treats sin/cos features as independent "
        "Gaussian dimensions, ignoring their circular topology. The HMM can learn that "
        "State A occurs when Day_sin \u2248 \u22120.87 (morning) and State B when Day_sin \u2248 0.87 "
        "(evening), but this temporal segmentation overrides the CO\u2082-based occupancy "
        "signal rather than complementing it."
    )

    add_figure(
        doc,
        HMM_TEMP_FIGS / "hmm_temporal_vs_occupancy.png",
        "Figure 5. HMM kappa vs. consensus across temporal feature configurations.",
    )

    # 5.3 Mixed Gaussian-von Mises HMM
    doc.add_heading("5.3 Mixed Gaussian\u2013von Mises HMM", level=2)

    doc.add_paragraph(
        "To properly handle circular temporal features, we implemented a custom HMM that "
        "combines Gaussian emissions for continuous features (CO\u2082, dCO\u2082, Noise) with "
        "von Mises emissions for circular features (hour of day, day of week). The von Mises "
        "distribution is the circular analogue of the Gaussian, parameterized by a mean "
        "direction \u03bc \u2208 [0, 2\u03c0) and concentration \u03ba \u2265 0."
    )

    add_equation(doc, "f(\u03b8 | \u03bc, \u03ba) = exp(\u03ba \u00b7 cos(\u03b8 \u2212 \u03bc)) / (2\u03c0 \u00b7 I\u2080(\u03ba))")

    doc.add_paragraph(
        "where I\u2080 is the modified Bessel function of the first kind of order zero."
    )

    # Results table
    circular_exps = [
        ("A1: Gaussian CO2+dCO2+Noise (4st)", "Gaussian: CO\u2082+dCO\u2082+Noise", "Gaussian"),
        ("A2: Gaussian +Day_sin/cos (4st)", "Gaussian: +Day sin/cos", "Gaussian"),
        ("A3: Gaussian +Day+Weekday sin/cos (4st)", "Gaussian: +Day+Weekday", "Gaussian"),
        ("B1: Mixed +hour_angle (4st)", "Mixed: +hour(VM) 4st", "Mixed"),
        ("B1: Mixed +hour_angle (5st)", "Mixed: +hour(VM) 5st", "Mixed"),
        ("B1: Mixed +hour_angle (6st)", "Mixed: +hour(VM) 6st", "Mixed"),
        ("B2: Mixed +hour+dow (4st)", "Mixed: +hour+dow(VM) 4st", "Mixed"),
        ("B2: Mixed +hour+dow (5st)", "Mixed: +hour+dow(VM) 5st", "Mixed"),
        ("B2: Mixed +hour+dow (6st)", "Mixed: +hour+dow(VM) 6st", "Mixed"),
        ("B3: Mixed +hour (8st, more granular)", "Mixed: +hour(VM) 8st", "Mixed"),
    ]

    circ_rows = []
    best_mixed_kappa = -1
    best_mixed_idx = -1
    for i, (key, label, model_type) in enumerate(circular_exps):
        exp = hmm_circular[key]
        bm = exp["best_binary_mapping"]
        kappa = bm["kappa_vs_consensus"]
        circ_rows.append([
            label,
            str(exp["n_states"]),
            f"{kappa:.4f}",
            f"{bm['occupancy_rate']}%",
        ])
        if model_type == "Mixed" and kappa > best_mixed_kappa:
            best_mixed_kappa = kappa
            best_mixed_idx = i

    p = doc.add_paragraph()
    add_bold_run(p, "Table 6. ")
    p.add_run("Gaussian vs. Mixed Gaussian\u2013von Mises HMM comparison.")
    add_styled_table(
        doc,
        ["Configuration", "K", "\u03ba vs. Consensus", "Occ. Rate"],
        circ_rows,
        highlight_rows=[0],  # Gaussian baseline A1 is best
    )

    doc.add_paragraph()
    doc.add_paragraph(
        f"The Mixed Gaussian\u2013von Mises HMM partially recovers the performance lost by "
        f"treating temporal features as Gaussian. The best Mixed configuration achieves "
        f"\u03ba = {best_mixed_kappa:.4f}, approximately 3\u20134\u00d7 better than the Gaussian "
        f"equivalent with sin/cos features (\u03ba \u2248 0.07), but still below the baseline "
        f"Gaussian HMM without temporal features (\u03ba = 0.53)."
    )
    doc.add_paragraph(
        "This indicates that while the von Mises distribution correctly models the circular "
        "topology, temporal features are fundamentally uninformative for occupancy detection "
        "in this dataset. The CO\u2082 dynamics alone (level and rate of change) carry the "
        "essential occupancy signal."
    )

    add_figure(
        doc,
        HMM_CIRC_FIGS / "gaussian_vs_mixed_hmm.png",
        "Figure 6. Gaussian vs. Mixed HMM kappa comparison across configurations.",
    )
    add_figure(
        doc,
        HMM_CIRC_FIGS / "mixed_hmm_state_interpretation.png",
        "Figure 7. Von Mises circular state interpretation (polar plot).",
        width=Inches(5.0),
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  6. PHYSICS-INFORMED SWITCHING AR-HMM
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. Physics-Informed Switching AR-HMM", level=1)

    # 6.1 Model Formulation
    doc.add_heading("6.1 Model Formulation", level=2)

    doc.add_paragraph(
        "The Switching Autoregressive HMM is grounded in the indoor CO\u2082 mass-balance "
        "ordinary differential equation. The underlying continuous-time model describes "
        "the CO\u2082 concentration y(t) as:"
    )

    add_equation(doc, "dy/dt = \u2212(1/\u03c4) \u00b7 y(t) + n(t) \u00b7 r")

    doc.add_paragraph(
        "where \u03c4 is the ventilation time constant (hours), n(t) is the number of occupants, "
        "and r is the per-person CO\u2082 generation rate (ppm/h). Discretizing with time step "
        "\u0394t gives a first-order autoregressive model:"
    )

    add_equation(doc, "y\u209c = c\u2096 \u00b7 y\u209c\u208b\u2081 + \u03bc\u2096 + \u03b5\u209c")

    doc.add_paragraph(
        "where the regime-specific parameters are:"
    )
    p = doc.add_paragraph(style="List Bullet")
    add_bold_run(p, "c\u2096 = exp(\u2212\u0394t/\u03c4\u2096)")
    p.add_run(
        ": the autoregressive coefficient, encoding the ventilation time constant \u03c4\u2096. "
        "Physically constrained to (0, 1)."
    )
    p = doc.add_paragraph(style="List Bullet")
    add_bold_run(p, "\u03bc\u2096 = (1 \u2212 c\u2096) \u00b7 r \u00b7 n\u2096")
    p.add_run(
        ": the drift term, encoding occupancy-driven CO\u2082 generation. Positive \u03bc implies "
        "active CO\u2082 sources (people present); negative or zero \u03bc implies decay toward "
        "ambient (unoccupied)."
    )
    p = doc.add_paragraph(style="List Bullet")
    add_bold_run(p, "\u03b5\u209c ~ N(0, \u03c3\u00b2\u2096)")
    p.add_run(": regime-specific observation noise.")

    doc.add_paragraph(
        "The model is trained via the Expectation-Maximization algorithm with the "
        "forward\u2013backward algorithm computed in log-space for numerical stability. "
        "The M-step uses weighted least squares with Cramer's rule to update (c\u2096, \u03bc\u2096) "
        "from the 2\u00d72 normal equations. Global initialization uses Yule\u2013Walker "
        "estimation of the AR(1) coefficient from the sample autocorrelation at lag 1."
    )

    # 6.2 Occupancy Mapping Strategy
    doc.add_heading("6.2 Occupancy Mapping Strategy", level=2)

    doc.add_paragraph(
        "Two mapping strategies are evaluated:"
    )
    p = doc.add_paragraph()
    add_bold_run(p, "Physics-based (Otsu on \u03bc). ")
    p.add_run(
        "Apply Otsu's thresholding to the K drift parameters \u03bc\u2096 to find the binary "
        "split that minimizes within-class variance. A physical floor enforces \u03bc\u2096 > 0 "
        "for occupied states (negative drift means CO\u2082 is decaying, implying no occupants). "
        "This approach is fully unsupervised and physically interpretable."
    )
    p = doc.add_paragraph()
    add_bold_run(p, "Exhaustive bitmask search. ")
    p.add_run(
        "Try all 2\u1d37 \u2212 2 possible state-to-occupancy mappings and select the one "
        "maximizing \u03ba vs. consensus. This provides an upper bound on the model's state "
        "segmentation quality, independent of the mapping strategy."
    )

    doc.add_paragraph(
        "An initial implementation using Otsu thresholding on the generation rate "
        "\u03bc/(1\u2212c) was found to be numerically unstable when c \u2192 1, as (1\u2212c) \u2192 0 "
        "causes the generation rate to explode. The \u03bc-based approach avoids this issue."
    )

    # 6.3 Experimental Results
    doc.add_heading("6.3 Experimental Results", level=2)

    ar_rows = []
    best_ar_bitmask = -1
    best_ar_idx = -1
    for i, (key, exp) in enumerate(switching_ar.items()):
        k = exp["n_states"]
        ar_rows.append([
            f"K={k}",
            f"{exp['kappa_physics']:.4f}",
            f"{exp['occupancy_rate_physics']}%",
            f"{exp['kappa_bitmask']:.4f}",
            f"{exp['occupancy_rate_bitmask']}%",
            str(exp["bitmask_occ_states"]),
            str(exp["n_iter"]),
        ])
        if exp["kappa_bitmask"] > best_ar_bitmask:
            best_ar_bitmask = exp["kappa_bitmask"]
            best_ar_idx = i

    p = doc.add_paragraph()
    add_bold_run(p, "Table 7. ")
    p.add_run("Switching AR-HMM results: physics-based vs. exhaustive bitmask mapping.")
    add_styled_table(
        doc,
        ["Model", "\u03ba (physics)", "Occ% (phys)", "\u03ba (bitmask)", "Occ% (bitmask)",
         "Bitmask States", "Iter."],
        ar_rows,
        highlight_rows=[best_ar_idx],
    )

    doc.add_paragraph()

    # Per-detector kappa for best model
    best_key = list(switching_ar.keys())[best_ar_idx]
    best_exp = switching_ar[best_key]
    doc.add_paragraph(
        f"The best configuration (K={best_exp['n_states']}) achieves \u03ba = "
        f"{best_exp['kappa_bitmask']:.4f} against consensus via exhaustive search, "
        f"and \u03ba = {best_exp['kappa_physics']:.4f} via physics-based mapping."
    )

    # Per-detector kappa table for K=8 (best rate_of_change)
    k8_exp = switching_ar.get("SwitchingAR K=8", {})
    if k8_exp:
        k8_kappas = k8_exp.get("kappas", {})
        det_kappa_rows = []
        for name, label in zip(det_names, det_labels):
            det_kappa_rows.append([label, f"{k8_kappas.get(name, 0):.4f}"])
        det_kappa_rows.append(["Consensus", f"{k8_kappas.get('consensus', 0):.4f}"])

        doc.add_paragraph()
        p = doc.add_paragraph()
        add_bold_run(p, "Table 8. ")
        p.add_run("Switching AR K=8: per-detector Cohen's kappa (physics-based mapping).")
        add_styled_table(
            doc,
            ["Detector", "\u03ba"],
            det_kappa_rows,
            highlight_rows=[1],  # rate_of_change
        )

        doc.add_paragraph()
        doc.add_paragraph(
            "The switching AR model achieves its highest agreement with the Rate of Change "
            "detector (\u03ba = 0.70), which is the single highest model-detector agreement "
            "across all experiments. This reflects the AR model's strength in capturing CO\u2082 "
            "dynamics rather than absolute levels."
        )

    # Physical parameter profiles for best K
    k5_exp = switching_ar.get("SwitchingAR K=5", {})
    if k5_exp:
        doc.add_paragraph()
        p = doc.add_paragraph()
        add_bold_run(p, "Table 9. ")
        p.add_run("Switching AR K=5 learned physical parameters.")

        param_rows = []
        for sp in k5_exp["state_profiles"]:
            param_rows.append([
                str(sp["state"]),
                f"{sp['pct']}%",
                f"{sp['co2_mean']:.0f}",
                f"{sp['c']:.4f}",
                f"{sp['mu']:+.2f}",
                f"{sp['sigma']:.1f}",
                f"{sp['generation']:+.0f}",
                sp["label"],
            ])
        add_styled_table(
            doc,
            ["State", "Time%", "CO\u2082", "c", "\u03bc (ppm)", "\u03c3", "Gen. rate", "Label"],
            param_rows,
            highlight_rows=[3, 4],  # occupied states
        )

        doc.add_paragraph()
        doc.add_paragraph(
            "The learned parameters are physically interpretable: states with high c "
            "(slow ventilation) and positive \u03bc (CO\u2082 generation) correspond to occupied "
            "periods, while states with low c (fast ventilation) and negative \u03bc correspond "
            "to unoccupied periods with active air exchange."
        )

    add_figure(
        doc,
        SWAR_FIGS / "switching_ar_physical_params.png",
        "Figure 8. Switching AR-HMM learned physical parameters across K values.",
    )
    add_figure(
        doc,
        SWAR_FIGS / "switching_ar_vs_baseline.png",
        "Figure 9. Switching AR-HMM vs. Gaussian HMM baseline comparison.",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  7. COMPARATIVE ANALYSIS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("7. Comparative Analysis", level=1)

    doc.add_paragraph(
        "The following table summarizes the best result from each method family."
    )

    summary_rows = [
        ["Gaussian HMM (CO\u2082+dCO\u2082+Noise, K=4)", "0.5324", "71.3%", "Exhaustive bitmask"],
        [f"Mixed Gauss-VM HMM (best)", f"{best_mixed_kappa:.4f}",
         f"{hmm_circular[circular_exps[best_mixed_idx][0]]['best_binary_mapping']['occupancy_rate']}%",
         "Exhaustive bitmask"],
        ["Gaussian HMM + sin/cos (best)", "0.0855", "41.7%", "Exhaustive bitmask"],
        [f"Switching AR-HMM (best bitmask)", f"{best_ar_bitmask:.4f}",
         f"{switching_ar[best_key]['occupancy_rate_bitmask']}%",
         "Exhaustive bitmask"],
        [f"Switching AR-HMM (best physics)",
         f"{max(e['kappa_physics'] for e in switching_ar.values()):.4f}",
         f"{[e for e in switching_ar.values() if e['kappa_physics'] == max(e2['kappa_physics'] for e2 in switching_ar.values())][0]['occupancy_rate_physics']}%",
         "Otsu on \u03bc"],
    ]

    p = doc.add_paragraph()
    add_bold_run(p, "Table 10. ")
    p.add_run("Summary of best results across all method families.")
    add_styled_table(
        doc,
        ["Method", "\u03ba vs. Consensus", "Occ. Rate", "Mapping"],
        summary_rows,
        highlight_rows=[0],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "The ranking is clear: the Gaussian HMM with sensor features only "
        "(CO\u2082, dCO\u2082/dt, Noise) outperforms all other approaches on consensus agreement. "
        "Adding temporal features degrades all HMM variants. The Mixed von Mises HMM "
        "partially recovers the damage caused by Gaussian sin/cos encoding but cannot "
        "exceed the non-temporal baseline. The Switching AR model brings unique advantages "
        "in physical interpretability but scores lower on the consensus metric."
    )

    # ══════════════════════════════════════════════════════════════
    #  8. DISCUSSION
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("8. Discussion", level=1)

    doc.add_heading("Why temporal features hurt", level=3)
    doc.add_paragraph(
        "Occupancy events in the SalonBM dataset do not follow strict temporal patterns. "
        "Being a residential environment, occupant behavior is irregular\u2014people may be "
        "home at any hour. When temporal features are added, the HMM preferentially "
        "segments by time-of-day (morning/afternoon/evening/night) rather than by CO\u2082 "
        "dynamics, producing states that are temporally coherent but occupancy-irrelevant."
    )

    doc.add_heading("The consensus paradox", level=3)
    doc.add_paragraph(
        "The consensus is dominated by 'sticky' detectors (absolute threshold at 75.9%, "
        "state machine at 91.8%, hybrid at 82.3%) that maintain occupied status as long "
        "as CO\u2082 remains elevated. The rate-of-change detector (36.9%) only fires during "
        "dynamic events. The Switching AR model aligns well with rate-of-change "
        "(\u03ba = 0.70) because both capture CO\u2082 dynamics, but this means the AR model "
        "disagrees with the sticky majority that defines the consensus."
    )

    doc.add_heading("Physical interpretability", level=3)
    doc.add_paragraph(
        "Despite lower consensus \u03ba, the Switching AR model provides uniquely valuable "
        "insights. The learned ventilation time constants \u03c4 range from ~5 hours (fast "
        "ventilation states) to ~27 hours (poor ventilation/high occupancy), and the "
        "drift parameters \u03bc cleanly separate active CO\u2082 generation from passive decay. "
        "These parameters are directly actionable for HVAC control\u2014for example, "
        "estimating the number of occupants from the generation rate."
    )

    doc.add_heading("Von Mises: correct distribution, wrong information", level=3)
    doc.add_paragraph(
        "The Mixed Gaussian\u2013von Mises HMM demonstrates that the circular distribution "
        "is mathematically appropriate for temporal features (3\u20134\u00d7 better than Gaussian "
        "sin/cos). However, the fundamental issue is that temporal information adds no "
        "occupancy-relevant signal beyond what CO\u2082 dynamics already capture."
    )

    # ══════════════════════════════════════════════════════════════
    #  9. LIMITATIONS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("9. Limitations", level=1)

    limitations = [
        "No ground-truth occupancy labels: All results are measured against a "
        "pseudo-consensus that may not reflect true occupancy. The consensus "
        "inherits the biases of its constituent detectors.",
        "Single-room residential dataset: The SalonBM dataset represents one "
        "specific environment. Results may not generalize to offices, classrooms, "
        "or multi-zone buildings with different ventilation characteristics.",
        "Hourly resolution: The 1-hour sampling interval limits the ability to "
        "detect brief occupancy events (< 1 hour) and blurs transition dynamics.",
        "Static ventilation assumption: The Switching AR model assumes that "
        "ventilation rates change only between regimes. In reality, window opening "
        "and HVAC activation can change ventilation within a regime.",
        "Otsu threshold sensitivity: The physics-based mapping depends on the "
        "distribution of \u03bc values across K states, which varies with K. No single "
        "mapping strategy is optimal across all K values.",
        "No multi-occupant modeling: The current framework detects binary "
        "occupied/unoccupied without estimating the number of occupants.",
        "Consensus weighting: All detectors contribute equally to the consensus. "
        "Weighting by detector reliability could improve the reference standard.",
        "EM convergence: The Switching AR model's EM algorithm may converge to "
        "local optima. Multiple restarts with different initializations are not "
        "currently implemented.",
    ]
    for i, lim in enumerate(limitations, 1):
        doc.add_paragraph(f"{i}. {lim}")

    # ══════════════════════════════════════════════════════════════
    #  10. FUTURE WORK
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("10. Future Work", level=1)

    future = [
        "Multi-feature Switching AR: Extend the AR model to incorporate noise "
        "level as an additional observation, combining physics-based CO\u2082 dynamics "
        "with acoustic occupancy indicators.",
        "Supervised calibration: Use a small labeled dataset (even a few days) to "
        "calibrate the state-to-occupancy mapping, bridging the gap between "
        "physics-based and consensus-based evaluation.",
        "Multi-occupant estimation: Extend the Switching AR framework to estimate "
        "the number of occupants n(t) from the generation rate \u03bc/(1\u2212c), given "
        "known per-person CO\u2082 generation rates.",
        "Transfer learning across rooms: Test whether learned ventilation time "
        "constants generalize across rooms or require per-room calibration.",
        "Hierarchical HMM: Model occupancy at multiple timescales (hourly events "
        "nested within daily patterns) using a hierarchical HMM structure.",
        "Real-time deployment: Implement online versions of the detectors and "
        "HMM for real-time occupancy inference in smart building controllers.",
    ]
    for item in future:
        doc.add_paragraph(item, style="List Bullet")

    # ══════════════════════════════════════════════════════════════
    #  11. CONCLUSIONS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("11. Conclusions", level=1)

    doc.add_paragraph(
        "This study presents a comprehensive evaluation of occupancy inference methods "
        "using indoor CO\u2082 measurements. The key findings are:"
    )
    conclusions = [
        "The Gaussian HMM with sensor features only (CO\u2082, dCO\u2082/dt, Noise) "
        "achieves the highest consensus agreement (\u03ba = 0.53), outperforming all "
        "temporal-enriched and physics-based variants.",
        "Adding temporal features (sin/cos or von Mises) consistently degrades "
        "HMM performance. The von Mises distribution is mathematically correct "
        "for circular features but the temporal signal is uninformative for "
        "occupancy in this residential dataset.",
        "The physics-informed Switching AR-HMM provides uniquely interpretable "
        "parameters (ventilation time constant \u03c4, CO\u2082 generation rate \u03bc) and "
        "excels at capturing CO\u2082 dynamics (\u03ba = 0.70 vs. rate-of-change detector), "
        "but scores lower against the level-dominated consensus (\u03ba = 0.34).",
        "The consensus metric is not neutral: it favors detectors that track "
        "CO\u2082 levels over those that track dynamics. This highlights the need "
        "for ground-truth occupancy labels to fairly evaluate all approaches.",
        "Predicted CO\u2082 from forecasting models preserves occupancy-relevant "
        "information well (\u03ba > 0.73 for all detectors), enabling downstream "
        "occupancy inference from model outputs.",
    ]
    for i, c in enumerate(conclusions, 1):
        doc.add_paragraph(f"{i}. {c}")

    # ── Save ─────────────────────────────────────────────────────
    doc.save(str(OUTPUT_PATH))
    print(f"Report saved to: {OUTPUT_PATH}")
    print(f"File size: {OUTPUT_PATH.stat().st_size / 1024:.0f} KB")
    return 0


if __name__ == "__main__":
    sys.exit(build_report())
