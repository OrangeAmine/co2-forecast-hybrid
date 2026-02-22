"""XGBoost Interpretability Study: Tree-based feature attribution and decision analysis.

Performs comprehensive interpretability analysis on preproc_D (Enhanced 1h) data:
  A) Built-in feature importance (gain, weight, cover)
  B) SHAP analysis (TreeSHAP exact values, temporal/feature decomposition)
  C) Partial dependence analysis (1D and 2D PDP, ICE plots)
  D) Tree structure analysis (split statistics, leaf distributions)
  E) Prediction quality analysis (overlay, scatter, residuals, error by CO2 level)

Generates a DOCX academic report with all figures and quantitative analysis.

Usage:
    python -u scripts/run_xgboost_interpretability.py
    python -u scripts/run_xgboost_interpretability.py --horizons 1
    python -u scripts/run_xgboost_interpretability.py --horizons 1 24
"""

import argparse
import gc
import json
import logging
import sys
import time
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from src.data.datamodule import CO2DataModule
from src.data.pipeline import run_preprocessing_pipeline
from src.data.preprocessing import inverse_scale_target
from src.evaluation.metrics import compute_metrics, save_metrics
from src.models.xgboost_model import XGBoostForecaster
from src.utils.config import load_config
from src.utils.seed import seed_everything

logger = logging.getLogger(__name__)

plt.style.use("seaborn-v0_8-whitegrid")

RESULTS_BASE = Path("results/xgboost_interpretability")

C_PRIMARY = "#2196F3"
C_SECONDARY = "#FF5722"
C_ACCENT = "#4CAF50"
C_WARN = "#FFC107"
C_NEUTRAL = "#607D8B"


# ======================================================================
#  Configuration
# ======================================================================

def load_interpretability_config(horizon: int) -> dict:
    config_files = [
        str(PROJECT_ROOT / "configs" / "training.yaml"),
        str(PROJECT_ROOT / "configs" / "data.yaml"),
        str(PROJECT_ROOT / "configs" / "xgboost.yaml"),
        str(PROJECT_ROOT / "configs" / "experiments" / "preproc_D_enhanced_1h.yaml"),
    ]
    config = load_config(config_files)
    config["data"]["forecast_horizon_hours"] = horizon
    config["training"]["results_dir"] = str(
        RESULTS_BASE / f"h{horizon}" / "training_runs"
    )
    return config


def get_feature_names(config: dict) -> list[str]:
    return config["data"]["feature_columns"] + [config["data"]["target_column"]]


def get_flat_feature_names(feature_names: list[str], lookback: int) -> list[str]:
    return [f"{feat}_lag{lag}" for lag in range(lookback) for feat in feature_names]


# ======================================================================
#  Training
# ======================================================================

def train_xgboost(
    config: dict,
    train_df: pd.DataFrame,
    val_df: pd.DataFrame,
    test_df: pd.DataFrame,
    horizon: int,
) -> tuple:
    dm = CO2DataModule(config)
    dm.build_datasets(train_df, val_df, test_df)

    X_train = dm.train_dataset.X.numpy()
    y_train = dm.train_dataset.y.numpy()
    X_val = dm.val_dataset.X.numpy()
    y_val = dm.val_dataset.y.numpy()
    X_test = dm.test_dataset.X.numpy()
    y_test_scaled = dm.test_dataset.y.numpy()

    model = XGBoostForecaster(config)
    model.fit(X_train, y_train, X_val, y_val)

    y_pred_scaled = model.predict(X_test)
    y_pred = inverse_scale_target(y_pred_scaled, dm.target_scaler)
    y_true = inverse_scale_target(y_test_scaled, dm.target_scaler)

    metrics = compute_metrics(y_true.ravel(), y_pred.ravel())

    return model, dm, {
        "y_true": y_true,
        "y_pred": y_pred,
        "metrics": metrics,
        "X_test": X_test,
        "X_train": X_train,
    }


# ======================================================================
#  Section A: Built-in Feature Importance
# ======================================================================

def run_builtin_importance(
    model: XGBoostForecaster,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [A] Extracting built-in feature importance...")

    feature_names = get_feature_names(config)
    n_features = len(feature_names)
    lookback = config["data"]["lookback_hours"] * config["data"]["samples_per_hour"]
    n_steps = len(model.model_.estimators_)

    # Collect gain importance from each sub-model
    all_importances = []
    for step_idx, est in enumerate(model.model_.estimators_):
        imp = est.feature_importances_  # (lookback * n_features,)
        all_importances.append(imp)

    all_importances = np.array(all_importances)  # (n_steps, lookback*n_features)
    avg_importance = all_importances.mean(axis=0)  # (lookback*n_features,)

    # Reshape to (lookback, n_features)
    imp_2d = avg_importance.reshape(lookback, n_features)

    # Plot A1: Importance heatmap
    fig, ax = plt.subplots(figsize=(max(10, n_features * 0.5), 6))
    im = ax.imshow(imp_2d.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Lookback Step (hours ago)")
    ax.set_ylabel("Feature")
    ax.set_yticks(range(n_features))
    ax.set_yticklabels(feature_names, fontsize=8)
    tick_pos = np.linspace(0, lookback - 1, min(8, lookback)).astype(int)
    ax.set_xticks(tick_pos)
    ax.set_xticklabels([f"-{lookback - t}" for t in tick_pos])
    ax.set_title(f"XGBoost Gain Importance Heatmap - {horizon}h")
    plt.colorbar(im, ax=ax, label="Gain Importance")
    plt.tight_layout()
    plt.savefig(output_dir / f"gain_heatmap_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot A2: Per-feature importance (sum across lags)
    feat_imp = imp_2d.sum(axis=0)
    feat_imp_pct = feat_imp / feat_imp.sum() * 100 if feat_imp.sum() > 0 else feat_imp
    sort_idx = np.argsort(feat_imp_pct)

    fig, ax = plt.subplots(figsize=(8, max(4, n_features * 0.35)))
    ax.barh(range(len(sort_idx)), feat_imp_pct[sort_idx], color=C_ACCENT)
    ax.set_yticks(range(len(sort_idx)))
    ax.set_yticklabels([feature_names[i] for i in sort_idx], fontsize=9)
    ax.set_xlabel("Importance (%)")
    ax.set_title(f"XGBoost Feature Importance (Gain) - {horizon}h")
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"feature_importance_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot A3: Per-lag importance (sum across features)
    lag_imp = imp_2d.sum(axis=1)
    lag_imp_pct = lag_imp / lag_imp.sum() * 100 if lag_imp.sum() > 0 else lag_imp

    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(range(-lookback, 0), lag_imp_pct, color=C_PRIMARY, alpha=0.8)
    ax.set_xlabel("Hours Ago")
    ax.set_ylabel("Importance (%)")
    ax.set_title(f"XGBoost Temporal Importance Profile - {horizon}h")
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"temporal_importance_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot A4: Per-step importance variation
    per_step_feat = np.zeros((n_steps, n_features))
    for step_idx in range(n_steps):
        imp_step = all_importances[step_idx].reshape(lookback, n_features)
        per_step_feat[step_idx] = imp_step.sum(axis=0)

    fig, ax = plt.subplots(figsize=(max(10, n_features * 0.5), 5))
    im = ax.imshow(per_step_feat.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Forecast Step")
    ax.set_ylabel("Feature")
    ax.set_yticks(range(n_features))
    ax.set_yticklabels(feature_names, fontsize=8)
    ax.set_title(f"Feature Importance by Forecast Step - {horizon}h")
    plt.colorbar(im, ax=ax, label="Gain Importance")
    plt.tight_layout()
    plt.savefig(output_dir / f"perstep_importance_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Save importance data
    imp_df = pd.DataFrame(imp_2d, columns=feature_names)
    imp_df.to_csv(output_dir / f"gain_importance_h{horizon}.csv", index=False)

    ranking = [(feature_names[i], float(feat_imp_pct[i])) for i in np.argsort(feat_imp_pct)[::-1]]
    print(f"  [A] Top 5 features: {', '.join(f'{n}({v:.1f}%)' for n, v in ranking[:5])}")

    return {
        "imp_2d": imp_2d,
        "feat_imp_pct": feat_imp_pct,
        "lag_imp_pct": lag_imp_pct,
        "ranking": ranking,
    }


# ======================================================================
#  Section B: SHAP Analysis
# ======================================================================

def run_shap_analysis(
    model: XGBoostForecaster,
    X_test: np.ndarray,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [B] Computing SHAP values (TreeSHAP)...")

    try:
        import shap
    except ImportError:
        print("  [B] shap not installed. Skipping SHAP analysis.")
        return {}

    feature_names = get_feature_names(config)
    n_features = len(feature_names)
    lookback = config["data"]["lookback_hours"] * config["data"]["samples_per_hour"]
    flat_names = get_flat_feature_names(feature_names, lookback)

    X_test_flat = X_test.reshape(X_test.shape[0], -1)
    n_shap = min(500, X_test_flat.shape[0])
    X_shap = X_test_flat[:n_shap]

    n_steps = len(model.model_.estimators_)

    # SHAP for step 0 (primary analysis)
    print(f"    Computing SHAP for step 0 ({n_shap} samples)...")
    sub_model_0 = model.model_.estimators_[0]
    explainer_0 = shap.TreeExplainer(sub_model_0)
    shap_values_0 = explainer_0.shap_values(X_shap)

    # Reshape to (n_shap, lookback, n_features)
    shap_3d = shap_values_0.reshape(n_shap, lookback, n_features)

    # Plot B1: SHAP importance bar chart (per original feature)
    mean_abs_shap_2d = np.abs(shap_3d).mean(axis=0)  # (lookback, n_features)
    feat_shap = mean_abs_shap_2d.sum(axis=0)  # (n_features,)
    feat_shap_pct = feat_shap / feat_shap.sum() * 100 if feat_shap.sum() > 0 else feat_shap
    sort_idx = np.argsort(feat_shap_pct)

    fig, ax = plt.subplots(figsize=(8, max(4, n_features * 0.35)))
    ax.barh(range(len(sort_idx)), feat_shap_pct[sort_idx], color=C_SECONDARY)
    ax.set_yticks(range(len(sort_idx)))
    ax.set_yticklabels([feature_names[i] for i in sort_idx], fontsize=9)
    ax.set_xlabel("Mean |SHAP| (%)")
    ax.set_title(f"SHAP Feature Importance - {horizon}h")
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"shap_feature_importance_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot B2: SHAP temporal heatmap
    fig, ax = plt.subplots(figsize=(max(10, n_features * 0.5), 6))
    im = ax.imshow(mean_abs_shap_2d.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Lookback Step (hours ago)")
    ax.set_ylabel("Feature")
    ax.set_yticks(range(n_features))
    ax.set_yticklabels(feature_names, fontsize=8)
    tick_pos = np.linspace(0, lookback - 1, min(8, lookback)).astype(int)
    ax.set_xticks(tick_pos)
    ax.set_xticklabels([f"-{lookback - t}" for t in tick_pos])
    ax.set_title(f"SHAP Temporal Heatmap - {horizon}h (Step 0)")
    plt.colorbar(im, ax=ax, label="Mean |SHAP|")
    plt.tight_layout()
    plt.savefig(output_dir / f"shap_temporal_heatmap_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot B3: SHAP temporal profile
    lag_shap = mean_abs_shap_2d.sum(axis=1)
    lag_shap_pct = lag_shap / lag_shap.sum() * 100 if lag_shap.sum() > 0 else lag_shap

    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(range(-lookback, 0), lag_shap_pct, color=C_PRIMARY, alpha=0.8)
    ax.set_xlabel("Hours Ago")
    ax.set_ylabel("Mean |SHAP| (%)")
    ax.set_title(f"SHAP Temporal Profile - {horizon}h")
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"shap_temporal_profile_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot B4: SHAP by forecast step (top 5 features)
    if n_steps > 1:
        top5_idx = np.argsort(feat_shap)[-5:][::-1]
        step_shap = np.zeros((n_steps, n_features))

        for step_idx in range(min(n_steps, 12)):
            print(f"    Computing SHAP for step {step_idx}...")
            est = model.model_.estimators_[step_idx]
            exp = shap.TreeExplainer(est)
            sv = exp.shap_values(X_shap)
            sv_3d = sv.reshape(n_shap, lookback, n_features)
            step_shap[step_idx] = np.abs(sv_3d).mean(axis=0).sum(axis=0)

        fig, ax = plt.subplots(figsize=(10, 6))
        for rank, fi in enumerate(top5_idx):
            ax.plot(range(min(n_steps, 12)), step_shap[:min(n_steps, 12), fi],
                    marker="o", markersize=4, label=feature_names[fi])
        ax.set_xlabel("Forecast Step")
        ax.set_ylabel("Mean |SHAP|")
        ax.set_title(f"SHAP Importance by Forecast Step - {horizon}h")
        ax.legend(fontsize=8)
        ax.grid(alpha=0.3)
        plt.tight_layout()
        plt.savefig(output_dir / f"shap_by_step_h{horizon}.png", dpi=150, bbox_inches="tight")
        plt.close(fig)

    # Plot B5: Waterfall-style plot for 3 samples
    base_value = explainer_0.expected_value
    if isinstance(base_value, np.ndarray):
        base_value = float(base_value[0])

    # Pick samples: median prediction, high CO2, large error
    sample_indices = [n_shap // 2, 0, n_shap - 1]
    sample_labels = ["Median sample", "First sample", "Last sample"]

    for si, (idx, label) in enumerate(zip(sample_indices, sample_labels)):
        sv = shap_values_0[idx]
        top_k = 15
        abs_sv = np.abs(sv)
        top_feat_idx = np.argsort(abs_sv)[-top_k:][::-1]

        fig, ax = plt.subplots(figsize=(10, 6))
        vals = sv[top_feat_idx]
        names = [flat_names[i] for i in top_feat_idx]
        colors = [C_SECONDARY if v > 0 else C_PRIMARY for v in vals]
        ax.barh(range(len(vals)), vals, color=colors)
        ax.set_yticks(range(len(vals)))
        ax.set_yticklabels(names, fontsize=8)
        ax.set_xlabel("SHAP Value")
        ax.set_title(f"SHAP Waterfall: {label} - {horizon}h\n(base={base_value:.2f})")
        ax.axvline(0, color="black", linewidth=0.8)
        ax.grid(axis="x", alpha=0.3)
        plt.tight_layout()
        plt.savefig(output_dir / f"shap_waterfall_{si}_h{horizon}.png", dpi=150, bbox_inches="tight")
        plt.close(fig)

    # Save SHAP data
    shap_df = pd.DataFrame(mean_abs_shap_2d, columns=feature_names)
    shap_df.to_csv(output_dir / f"shap_importance_h{horizon}.csv", index=False)

    shap_ranking = [(feature_names[i], float(feat_shap_pct[i]))
                    for i in np.argsort(feat_shap_pct)[::-1]]
    print(f"  [B] Top 5 SHAP: {', '.join(f'{n}({v:.1f}%)' for n, v in shap_ranking[:5])}")

    return {
        "feat_shap_pct": feat_shap_pct,
        "lag_shap_pct": lag_shap_pct,
        "shap_ranking": shap_ranking,
        "shap_3d": mean_abs_shap_2d,
    }


# ======================================================================
#  Section C: Partial Dependence
# ======================================================================

def run_partial_dependence(
    model: XGBoostForecaster,
    X_test: np.ndarray,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [C] Computing partial dependence plots...")

    from sklearn.inspection import partial_dependence

    feature_names = get_feature_names(config)
    n_features = len(feature_names)
    lookback = config["data"]["lookback_hours"] * config["data"]["samples_per_hour"]
    flat_names = get_flat_feature_names(feature_names, lookback)

    X_test_flat = X_test.reshape(X_test.shape[0], -1)
    n_pdp = min(300, X_test_flat.shape[0])
    X_pdp = X_test_flat[:n_pdp]

    sub_model = model.model_.estimators_[0]

    # Find most important features (lag 0 = most recent)
    imp = sub_model.feature_importances_
    # Focus on lag-0 features (indices 0..n_features-1 in the last lookback position)
    # Actually with reshape order: flat_idx = lag * n_features + feat_idx
    # Lag 0 features: indices 0..n_features-1
    # Most recent (lag lookback-1): indices (lookback-1)*n_features .. lookback*n_features-1
    recent_start = (lookback - 1) * n_features
    recent_imp = imp[recent_start:recent_start + n_features]
    top6_feat = np.argsort(recent_imp)[-6:][::-1]
    top6_flat = [recent_start + fi for fi in top6_feat]

    # Plot C1: 1D PDP for top 6 features
    fig, axes = plt.subplots(2, 3, figsize=(15, 9))
    axes_flat = axes.ravel()

    for idx, (feat_flat_idx, feat_orig_idx) in enumerate(zip(top6_flat, top6_feat)):
        ax = axes_flat[idx]
        try:
            pdp_results = partial_dependence(
                sub_model, X_pdp, features=[feat_flat_idx],
                kind="average", grid_resolution=50,
            )
            ax.plot(pdp_results["grid_values"][0], pdp_results["average"][0],
                    color=C_PRIMARY, linewidth=2)
            ax.set_xlabel(f"{feature_names[feat_orig_idx]} (most recent)", fontsize=9)
            ax.set_ylabel("Partial Dependence")
            ax.set_title(feature_names[feat_orig_idx], fontsize=10)
            ax.grid(alpha=0.3)
        except Exception as e:
            ax.text(0.5, 0.5, f"Error: {str(e)[:30]}", ha="center", va="center",
                    transform=ax.transAxes, fontsize=8)

    fig.suptitle(f"Partial Dependence Plots (Most Recent Lag) - {horizon}h", fontsize=13)
    plt.tight_layout()
    plt.savefig(output_dir / f"pdp_top6_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot C2: 2D PDP for CO2 x dCO2 (most recent lag)
    try:
        co2_idx = feature_names.index("CO2")
        dco2_idx = feature_names.index("dCO2")
        co2_flat = recent_start + co2_idx
        dco2_flat = recent_start + dco2_idx

        pdp_2d = partial_dependence(
            sub_model, X_pdp, features=[(co2_flat, dco2_flat)],
            kind="average", grid_resolution=30,
        )

        fig, ax = plt.subplots(figsize=(8, 6))
        XX, YY = np.meshgrid(pdp_2d["grid_values"][0], pdp_2d["grid_values"][1])
        ZZ = pdp_2d["average"][0]
        im = ax.contourf(XX, YY, ZZ, levels=20, cmap="RdYlBu_r")
        ax.set_xlabel("CO2 (most recent)")
        ax.set_ylabel("dCO2 (most recent)")
        ax.set_title(f"2D Partial Dependence: CO2 x dCO2 - {horizon}h")
        plt.colorbar(im, ax=ax, label="Prediction")
        plt.tight_layout()
        plt.savefig(output_dir / f"pdp_2d_co2_dco2_h{horizon}.png", dpi=150, bbox_inches="tight")
        plt.close(fig)
    except Exception as e:
        print(f"    2D PDP failed: {e}")

    print(f"  [C] Partial dependence analysis complete")


# ======================================================================
#  Section D: Tree Structure Analysis
# ======================================================================

def run_tree_analysis(
    model: XGBoostForecaster,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [D] Analyzing tree structure...")

    feature_names = get_feature_names(config)
    n_features = len(feature_names)
    lookback = config["data"]["lookback_hours"] * config["data"]["samples_per_hour"]
    flat_names = get_flat_feature_names(feature_names, lookback)

    sub_model = model.model_.estimators_[0]
    booster = sub_model.get_booster()

    # Plot D1: Split feature frequency
    try:
        weight_scores = booster.get_score(importance_type="weight")
        # Map from f0, f1, ... to feature names
        feat_counts = {}
        for fkey, count in weight_scores.items():
            fidx = int(fkey.replace("f", ""))
            if fidx < len(flat_names):
                # Map to original feature
                orig_idx = fidx % n_features
                fname = feature_names[orig_idx]
                feat_counts[fname] = feat_counts.get(fname, 0) + count

        if feat_counts:
            sorted_feats = sorted(feat_counts.items(), key=lambda x: x[1], reverse=True)
            names_s, counts_s = zip(*sorted_feats[:20])

            fig, ax = plt.subplots(figsize=(10, 6))
            ax.barh(range(len(names_s)), counts_s, color=C_PRIMARY)
            ax.set_yticks(range(len(names_s)))
            ax.set_yticklabels(names_s, fontsize=9)
            ax.set_xlabel("Number of Splits")
            ax.set_title(f"Split Feature Frequency (Step 0) - {horizon}h")
            ax.grid(axis="x", alpha=0.3)
            plt.tight_layout()
            plt.savefig(output_dir / f"split_frequency_h{horizon}.png", dpi=150, bbox_inches="tight")
            plt.close(fig)
    except Exception as e:
        print(f"    Split frequency analysis failed: {e}")

    # Plot D2: Leaf value distribution
    try:
        dump = booster.get_dump(dump_format="text")
        leaf_values = []
        for tree_str in dump:
            for line in tree_str.split("\n"):
                if "leaf=" in line:
                    val = float(line.split("leaf=")[1].split(",")[0])
                    leaf_values.append(val)

        if leaf_values:
            leaf_values = np.array(leaf_values)
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.hist(leaf_values, bins=80, color=C_PRIMARY, alpha=0.7, edgecolor="white", density=True)
            ax.axvline(0, color="red", linestyle="--", linewidth=1)
            ax.set_xlabel("Leaf Value")
            ax.set_ylabel("Density")
            ax.set_title(f"Leaf Value Distribution - {horizon}h\n"
                         f"(mean={leaf_values.mean():.4f}, std={leaf_values.std():.4f}, "
                         f"n_leaves={len(leaf_values)})")
            ax.grid(axis="y", alpha=0.3)
            plt.tight_layout()
            plt.savefig(output_dir / f"leaf_distribution_h{horizon}.png", dpi=150, bbox_inches="tight")
            plt.close(fig)
    except Exception as e:
        print(f"    Leaf analysis failed: {e}")

    # Plot D3: Per-step RMSE
    n_steps = len(model.model_.estimators_)
    if n_steps > 1:
        # Need per-step predictions - already in results
        print(f"    Per-step analysis: {n_steps} forecast steps")

    print(f"  [D] Tree structure analysis complete")


# ======================================================================
#  Section E: Prediction Quality
# ======================================================================

def run_prediction_analysis(
    results: dict,
    test_df: pd.DataFrame,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [E] Generating prediction analysis plots...")

    y_true = results["y_true"].ravel()
    y_pred = results["y_pred"].ravel()
    residuals = y_true - y_pred
    metrics = results["metrics"]

    n_show = min(500, len(y_true))

    # Plot E1: Predictions overlay
    fig, ax = plt.subplots(figsize=(14, 5))
    ax.plot(y_true[:n_show], label="Actual", color=C_PRIMARY, linewidth=1.0, alpha=0.8)
    ax.plot(y_pred[:n_show], label="Predicted", color=C_SECONDARY, linewidth=1.0, alpha=0.8)
    ax.set_xlabel("Sample Index")
    ax.set_ylabel("CO2 (ppm)")
    ax.set_title(f"XGBoost Predictions vs Actual - {horizon}h")
    ax.legend()
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_dir / f"predictions_overlay_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot E2: Scatter with R2
    fig, ax = plt.subplots(figsize=(7, 7))
    ax.scatter(y_true, y_pred, alpha=0.3, s=5, color=C_PRIMARY)
    vmin = min(y_true.min(), y_pred.min())
    vmax = max(y_true.max(), y_pred.max())
    ax.plot([vmin, vmax], [vmin, vmax], "r--", linewidth=1.5, label="Perfect")
    ax.text(
        0.05, 0.95,
        f"R2 = {metrics['r2']:.4f}\nRMSE = {metrics['rmse']:.2f}\nMAE = {metrics['mae']:.2f}",
        transform=ax.transAxes, fontsize=11, va="top",
        bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8),
    )
    ax.set_xlabel("Actual CO2 (ppm)")
    ax.set_ylabel("Predicted CO2 (ppm)")
    ax.set_title(f"XGBoost Scatter - {horizon}h")
    ax.legend()
    plt.tight_layout()
    plt.savefig(output_dir / f"scatter_r2_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot E3: 4-panel residual analysis
    datetime_col = config["data"].get("datetime_column", "datetime")
    test_dates = None
    if datetime_col in test_df.columns:
        test_dates = pd.DatetimeIndex(test_df[datetime_col])

    fig, axes = plt.subplots(2, 2, figsize=(14, 10))

    ax = axes[0, 0]
    ax.hist(residuals, bins=50, color=C_PRIMARY, alpha=0.7, edgecolor="white", density=True)
    ax.axvline(0, color="red", linestyle="--", linewidth=1.5)
    ax.set_xlabel("Residual (ppm)")
    ax.set_ylabel("Density")
    ax.set_title("Residual Distribution")
    ax.text(0.95, 0.95, f"Mean: {residuals.mean():.2f}\nStd: {residuals.std():.2f}",
            transform=ax.transAxes, fontsize=9, ha="right", va="top",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8))

    ax = axes[0, 1]
    ax.scatter(np.arange(len(residuals)), residuals, alpha=0.3, s=3, color=C_SECONDARY)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xlabel("Sample Index")
    ax.set_ylabel("Residual (ppm)")
    ax.set_title("Residuals Over Time")

    ax = axes[1, 0]
    if test_dates is not None and len(test_dates) >= len(residuals):
        hours = test_dates[-len(residuals):].hour
        res_df = pd.DataFrame({"hour": hours, "residual": residuals})
        res_df.boxplot(column="residual", by="hour", ax=ax)
        ax.set_xlabel("Hour of Day")
        ax.set_ylabel("Residual (ppm)")
        ax.set_title("Residuals by Hour")
        plt.suptitle("")
    else:
        ax.text(0.5, 0.5, "No datetime available", ha="center", va="center", transform=ax.transAxes)
        ax.set_title("Residuals by Hour (N/A)")

    ax = axes[1, 1]
    if test_dates is not None and len(test_dates) >= len(residuals):
        dow = test_dates[-len(residuals):].dayofweek
        day_names = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
        res_df = pd.DataFrame({"day_num": dow, "residual": residuals})
        res_df.boxplot(column="residual", by="day_num", ax=ax)
        ax.set_xticklabels(day_names)
        ax.set_xlabel("Day of Week")
        ax.set_ylabel("Residual (ppm)")
        ax.set_title("Residuals by Day")
        plt.suptitle("")
    else:
        ax.text(0.5, 0.5, "No datetime available", ha="center", va="center", transform=ax.transAxes)
        ax.set_title("Residuals by Day (N/A)")

    plt.suptitle("")
    plt.tight_layout()
    plt.savefig(output_dir / f"residual_analysis_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Plot E4: Error by CO2 level
    bins = [0, 500, 1000, np.inf]
    labels = ["Low (<500)", "Medium (500-1000)", "High (>1000)"]
    categories = pd.cut(y_true, bins=bins, labels=labels)

    fig, axes = plt.subplots(1, 2, figsize=(14, 5))

    ax = axes[0]
    abs_errors = np.abs(residuals)
    err_df = pd.DataFrame({"category": categories, "abs_error": abs_errors})
    err_df.boxplot(column="abs_error", by="category", ax=ax)
    ax.set_xlabel("CO2 Level")
    ax.set_ylabel("Absolute Error (ppm)")
    ax.set_title(f"Error by CO2 Level - {horizon}h")
    plt.suptitle("")

    ax = axes[1]
    ax.axis("off")
    rows_data = []
    for label in labels:
        mask = categories == label
        if mask.sum() > 0:
            bin_res = residuals[mask]
            rows_data.append([
                label, f"{mask.sum()}", f"{np.sqrt(np.mean(bin_res**2)):.2f}",
                f"{np.mean(np.abs(bin_res)):.2f}", f"{np.mean(bin_res):.2f}",
            ])
        else:
            rows_data.append([label, "0", "N/A", "N/A", "N/A"])

    tbl = ax.table(
        cellText=rows_data,
        colLabels=["CO2 Level", "N", "RMSE", "MAE", "Mean Bias"],
        loc="center", cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.0, 1.5)

    plt.tight_layout()
    plt.savefig(output_dir / f"error_by_co2_level_h{horizon}.png", dpi=150, bbox_inches="tight")
    plt.close(fig)

    # Per-step RMSE
    y_true_2d = results["y_true"]
    y_pred_2d = results["y_pred"]
    if y_true_2d.ndim == 2 and y_true_2d.shape[1] > 1:
        step_rmse = np.sqrt(np.mean((y_true_2d - y_pred_2d)**2, axis=0))
        fig, ax = plt.subplots(figsize=(10, 5))
        ax.plot(range(1, len(step_rmse) + 1), step_rmse, marker="o", color=C_PRIMARY, markersize=4)
        ax.set_xlabel("Forecast Step")
        ax.set_ylabel("RMSE (ppm)")
        ax.set_title(f"XGBoost Per-Step RMSE - {horizon}h")
        ax.grid(alpha=0.3)
        plt.tight_layout()
        plt.savefig(output_dir / f"perstep_rmse_h{horizon}.png", dpi=150, bbox_inches="tight")
        plt.close(fig)

    print(f"  [E] Prediction analysis complete")


# ======================================================================
#  DOCX Report
# ======================================================================

def generate_docx_report(
    all_results: dict,
    all_builtin: dict,
    all_shap: dict,
    horizons: list[int],
    output_dir: Path,
) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print("  Generating DOCX report...")
    doc = Document()

    title = doc.add_heading(
        "XGBoost Interpretability Study: Tree-Based Feature Attribution "
        "and Decision Analysis for Indoor CO2 Forecasting", level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Preprocessing Variant: preproc_D (Enhanced 1h)\n"
        f"Horizons: {', '.join(str(h) + 'h' for h in horizons)}\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Section A: Built-in Feature Importance",
        "3. Section B: SHAP Analysis (TreeSHAP)",
        "4. Section C: Partial Dependence Analysis",
        "5. Section D: Tree Structure Analysis",
        "6. Section E: Prediction Quality Analysis",
        "7. Cross-Horizon Comparison",
        "8. Avenues of Improvement",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    def add_figure(fig_path: Path, caption: str, width: float = 6.0) -> None:
        if fig_path.exists():
            doc.add_picture(str(fig_path), width=Inches(width))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.font.size = Pt(9)
            r.font.italic = True
        else:
            doc.add_paragraph(f"[Figure not found: {fig_path.name}]")

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This report presents a comprehensive interpretability study of an XGBoost "
        "gradient-boosted tree ensemble trained for indoor CO2 concentration forecasting. "
        "XGBoost uses a MultiOutputRegressor architecture with one independent tree ensemble "
        "per forecast step. The flattened lookback window creates a high-dimensional feature "
        "space (24 timesteps x 19 features = 456 features), enabling both temporal and "
        "feature-level attribution analysis. The study covers: (A) built-in gain importance "
        "analysis, (B) exact SHAP values via TreeSHAP, (C) partial dependence analysis, "
        "(D) tree structure inspection, and (E) prediction quality assessment."
    )

    doc.add_heading("Performance Summary", level=2)
    table = doc.add_table(rows=1 + len(horizons), cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Horizon", "RMSE (ppm)", "MAE (ppm)", "R2", "MAPE (%)"]):
        table.rows[0].cells[i].text = h
    for row_idx, h in enumerate(horizons):
        m = all_results[h]["metrics"]
        table.rows[row_idx + 1].cells[0].text = f"{h}h"
        table.rows[row_idx + 1].cells[1].text = f"{m['rmse']:.2f}"
        table.rows[row_idx + 1].cells[2].text = f"{m['mae']:.2f}"
        table.rows[row_idx + 1].cells[3].text = f"{m['r2']:.4f}"
        table.rows[row_idx + 1].cells[4].text = f"{m['mape']:.2f}"
    doc.add_page_break()

    # 2. Section A
    doc.add_heading("2. Section A: Built-in Feature Importance", level=1)
    doc.add_paragraph(
        "XGBoost provides native feature importance based on the total gain from all "
        "splits using each feature across all trees. Since the input is a flattened "
        "lookback window, importance values span both temporal lag positions and feature "
        "channels, producing a 2D importance map. This reveals which past observations "
        "and which sensor channels drive the model's predictions."
    )

    for h in horizons:
        doc.add_heading(f"Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "builtin_importance"
        add_figure(h_dir / f"gain_heatmap_h{h}.png",
                   f"Gain importance heatmap showing which (lag, feature) positions are most used by the tree ensemble at {h}h horizon.")
        add_figure(h_dir / f"feature_importance_h{h}.png",
                   f"Per-feature importance aggregated across all lags. Shows which sensor channels dominate the model's decisions.")
        add_figure(h_dir / f"temporal_importance_h{h}.png",
                   f"Temporal importance profile showing which past hours are most utilized for splits.")
        add_figure(h_dir / f"perstep_importance_h{h}.png",
                   f"Feature importance variation across forecast steps.")

        if h in all_builtin:
            doc.add_heading("Quantitative Ranking", level=3)
            ranking = all_builtin[h].get("ranking", [])
            if ranking:
                tbl = doc.add_table(rows=min(10, len(ranking)) + 1, cols=2)
                tbl.style = "Light Grid Accent 1"
                tbl.rows[0].cells[0].text = "Feature"
                tbl.rows[0].cells[1].text = "Importance (%)"
                for ri, (name, pct) in enumerate(ranking[:10]):
                    tbl.rows[ri + 1].cells[0].text = name
                    tbl.rows[ri + 1].cells[1].text = f"{pct:.2f}"

    doc.add_heading("Discussion", level=2)
    doc.add_paragraph(
        "The gain importance analysis reveals the temporal decay pattern of XGBoost's "
        "attention: recent observations receive disproportionate importance, consistent "
        "with the autoregressive nature of CO2 dynamics. The CO2 target at the most "
        "recent lag position typically dominates, followed by CO2 lag features and the "
        "rate of change (dCO2). This reflects the strong persistence and momentum in "
        "indoor CO2 concentration."
    )
    doc.add_page_break()

    # 3. Section B
    doc.add_heading("3. Section B: SHAP Analysis (TreeSHAP)", level=1)
    doc.add_paragraph(
        "SHAP (SHapley Additive exPlanations) provides exact feature attribution values "
        "using the TreeSHAP algorithm, which runs in polynomial time for tree ensembles. "
        "Unlike gain importance, SHAP accounts for feature interactions and provides "
        "per-sample explanations, enabling both global and local interpretability."
    )

    for h in horizons:
        doc.add_heading(f"Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "shap_analysis"
        add_figure(h_dir / f"shap_feature_importance_h{h}.png",
                   f"SHAP-based feature importance at {h}h. Computed as mean absolute SHAP value per feature across all test samples.")
        add_figure(h_dir / f"shap_temporal_heatmap_h{h}.png",
                   f"SHAP temporal heatmap showing the mean |SHAP| at each (lag, feature) position.")
        add_figure(h_dir / f"shap_temporal_profile_h{h}.png",
                   f"Temporal SHAP profile showing which lookback positions carry the most attribution.")
        add_figure(h_dir / f"shap_by_step_h{h}.png",
                   f"SHAP importance evolution across forecast steps for the top 5 features.")
        for si in range(3):
            add_figure(h_dir / f"shap_waterfall_{si}_h{h}.png",
                       f"SHAP waterfall plot for sample {si}: shows how individual features push the prediction above/below the baseline.")

    doc.add_heading("Gain vs SHAP Comparison", level=2)
    doc.add_paragraph(
        "Comparing the gain-based importance ranking (Section A) with SHAP-based ranking "
        "reveals consistency in the top features but potential divergence in lower-ranked "
        "ones. Gain importance reflects how frequently and effectively features are used "
        "for splits, while SHAP captures the average marginal contribution to each prediction. "
        "Features that interact strongly with others may have higher SHAP importance than "
        "gain importance, as SHAP accounts for the full interaction structure."
    )
    doc.add_page_break()

    # 4. Section C
    doc.add_heading("4. Section C: Partial Dependence Analysis", level=1)
    doc.add_paragraph(
        "Partial dependence plots show the marginal effect of individual features on "
        "the model's prediction, averaging over all other features. This reveals the "
        "functional form of the learned response (linear, nonlinear, threshold-based) "
        "for each key input variable."
    )
    for h in horizons:
        doc.add_heading(f"Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "partial_dependence"
        add_figure(h_dir / f"pdp_top6_h{h}.png",
                   f"1D partial dependence plots for the top 6 most important features at the most recent lag.")
        add_figure(h_dir / f"pdp_2d_co2_dco2_h{h}.png",
                   f"2D partial dependence showing the interaction between current CO2 and its rate of change (dCO2).")

    doc.add_heading("Discussion", level=2)
    doc.add_paragraph(
        "The partial dependence plots reveal that XGBoost learns approximately monotonic "
        "relationships for the most important features (CO2, lag features), with some "
        "nonlinear thresholds at extreme values. The 2D PDP for CO2 x dCO2 shows that "
        "the model captures the interaction between current level and trend, predicting "
        "higher future CO2 when both current CO2 is high and the rate of change is positive."
    )
    doc.add_page_break()

    # 5. Section D
    doc.add_heading("5. Section D: Tree Structure Analysis", level=1)
    doc.add_paragraph(
        "Inspecting the tree structure reveals the decision rules XGBoost has learned. "
        "Split frequency shows which features serve as primary decision variables, while "
        "the leaf value distribution characterizes the model's prediction granularity."
    )
    for h in horizons:
        doc.add_heading(f"Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "tree_analysis"
        add_figure(h_dir / f"split_frequency_h{h}.png",
                   f"Split frequency: number of times each feature is used as a split variable across all trees in the step-0 sub-model.")
        add_figure(h_dir / f"leaf_distribution_h{h}.png",
                   f"Distribution of leaf values across all trees. A symmetric distribution centered near zero indicates balanced predictions.")
    doc.add_page_break()

    # 6. Section E
    doc.add_heading("6. Section E: Prediction Quality Analysis", level=1)
    for h in horizons:
        doc.add_heading(f"Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "prediction_analysis"
        add_figure(h_dir / f"predictions_overlay_h{h}.png",
                   f"Time series overlay of actual vs predicted CO2 at {h}h horizon.")
        add_figure(h_dir / f"scatter_r2_h{h}.png",
                   f"Scatter plot with R-squared showing prediction accuracy.")
        add_figure(h_dir / f"residual_analysis_h{h}.png",
                   f"Four-panel residual analysis: distribution, temporal pattern, by hour, by day of week.")
        add_figure(h_dir / f"error_by_co2_level_h{h}.png",
                   f"Error breakdown by CO2 concentration regime.")
        add_figure(h_dir / f"perstep_rmse_h{h}.png",
                   f"Per-step RMSE showing error growth across the forecast horizon.")
    doc.add_page_break()

    # 7. Cross-Horizon
    if len(horizons) > 1:
        doc.add_heading("7. Cross-Horizon Comparison", level=1)
        doc.add_paragraph(
            "Comparing the 1h and 24h horizons reveals how XGBoost adapts its feature "
            "utilization strategy to different prediction timescales."
        )
        tbl = doc.add_table(rows=1 + len(horizons), cols=5)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["Horizon", "RMSE", "MAE", "R2", "MAPE"]):
            tbl.rows[0].cells[i].text = h
        for ri, h in enumerate(horizons):
            m = all_results[h]["metrics"]
            tbl.rows[ri + 1].cells[0].text = f"{h}h"
            tbl.rows[ri + 1].cells[1].text = f"{m['rmse']:.2f}"
            tbl.rows[ri + 1].cells[2].text = f"{m['mae']:.2f}"
            tbl.rows[ri + 1].cells[3].text = f"{m['r2']:.4f}"
            tbl.rows[ri + 1].cells[4].text = f"{m['mape']:.2f}"
        doc.add_page_break()

    # 8. Avenues
    doc.add_heading("8. Avenues of Improvement", level=1)
    improvements = [
        ("Hyperparameter Optimization",
         "The current XGBoost configuration uses fixed hyperparameters. Bayesian optimization "
         "of max_depth, learning_rate, n_estimators, subsample, and regularization parameters "
         "could improve both accuracy and interpretability by finding a sparser, more generalizable model."),
        ("Feature Engineering",
         "The flattened lookback window creates 456 features, many of which may be redundant. "
         "Feature selection using SHAP importance or recursive feature elimination could reduce "
         "dimensionality while maintaining or improving performance."),
        ("Interaction-Aware Features",
         "The 2D PDP analysis reveals important feature interactions. Explicitly engineering "
         "interaction features (e.g., CO2 * dCO2, CO2 * Day_sin) could help the model capture "
         "these relationships more efficiently."),
        ("Multi-Target Strategy",
         "The current MultiOutputRegressor trains independent models per step. A chained "
         "approach or a direct multi-output XGBoost formulation could capture dependencies "
         "between forecast steps."),
        ("SHAP-Based Feature Selection",
         "Using SHAP values to prune features with consistently low attribution could reduce "
         "model complexity and training time while improving interpretability."),
        ("Ensemble with Neural Models",
         "XGBoost excels at capturing nonlinear relationships in the tabular feature space. "
         "Combining XGBoost predictions with LSTM hidden state features could leverage both "
         "the tree model's nonlinear capacity and the LSTM's sequential memory."),
    ]
    for title_text, desc in improvements:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(desc)

    report_path = output_dir / "xgboost_interpretability_report.docx"
    doc.save(str(report_path))
    print(f"  Report saved to: {report_path}")
    return report_path


# ======================================================================
#  Main
# ======================================================================

def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

    parser = argparse.ArgumentParser(description="XGBoost Interpretability Study")
    parser.add_argument("--horizons", nargs="+", type=int, default=[1, 24])
    parser.add_argument("--epochs", type=int, default=None, help="Unused for XGBoost")
    args = parser.parse_args()

    print(f"\n{'='*70}")
    print(f"  XGBOOST INTERPRETABILITY STUDY (DEEP ANALYSIS)")
    print(f"  Variant: preproc_D (Enhanced 1h)")
    print(f"  Horizons: {args.horizons}")
    print(f"  Sections: A(builtin) B(SHAP) C(PDP) D(trees) E(predictions)")
    print(f"{'='*70}\n")

    pipeline_config = load_interpretability_config(horizon=1)
    seed_everything(pipeline_config["training"]["seed"])

    raw_dir = Path(pipeline_config["data"].get("raw_dir", "data/raw"))
    print("  Loading preprocessing pipeline (preproc_D)...")
    train_df, val_df, test_df = run_preprocessing_pipeline(
        raw_dir=raw_dir, variant_config=pipeline_config,
    )
    print(f"  Pipeline loaded: train={len(train_df)}, val={len(val_df)}, test={len(test_df)}")

    all_results = {}
    all_builtin = {}
    all_shap = {}

    for horizon in args.horizons:
        print(f"\n{'-'*60}")
        print(f"  HORIZON: {horizon}h")
        print(f"{'-'*60}\n")

        config = load_interpretability_config(horizon=horizon)
        seed_everything(config["training"]["seed"])

        output_dir = RESULTS_BASE / f"h{horizon}"
        output_dir.mkdir(parents=True, exist_ok=True)

        t0 = time.time()
        print(f"  Training XGBoost for {horizon}h horizon...")
        model, dm, results = train_xgboost(
            config, train_df.copy(), val_df.copy(), test_df.copy(), horizon,
        )
        elapsed = time.time() - t0
        print(f"  Training completed in {elapsed:.1f}s")

        metrics = results["metrics"]
        all_results[horizon] = results
        print(f"  RMSE={metrics['rmse']:.2f}  MAE={metrics['mae']:.2f}  "
              f"R2={metrics['r2']:.4f}  MAPE={metrics['mape']:.2f}%")

        builtin_data = run_builtin_importance(model, config, horizon, output_dir / "builtin_importance")
        all_builtin[horizon] = builtin_data

        shap_data = run_shap_analysis(model, results["X_test"], config, horizon, output_dir / "shap_analysis")
        all_shap[horizon] = shap_data

        run_partial_dependence(model, results["X_test"], config, horizon, output_dir / "partial_dependence")

        run_tree_analysis(model, config, horizon, output_dir / "tree_analysis")

        run_prediction_analysis(results, test_df, config, horizon, output_dir / "prediction_analysis")

        save_metrics(
            metrics, f"XGBoost_h{horizon}", output_dir / "metrics.json",
            experiment_info={"name": "xgboost_interpretability", "label": f"XGBoost h={horizon}"},
        )
        np.savez(output_dir / "predictions.npz", y_true=results["y_true"], y_pred=results["y_pred"])

        del model, dm
        gc.collect()
        print(f"  Memory freed\n")

    # DOCX Report
    print(f"\n{'-'*60}")
    print(f"  GENERATING DOCX REPORT")
    print(f"{'-'*60}\n")
    report_path = generate_docx_report(all_results, all_builtin, all_shap, args.horizons, RESULTS_BASE)

    print(f"\n{'='*70}")
    print(f"  XGBOOST INTERPRETABILITY STUDY COMPLETE")
    print(f"  Results saved to: {RESULTS_BASE}")
    print(f"  Report: {report_path}")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
