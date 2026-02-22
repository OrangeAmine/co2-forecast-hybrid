"""LSTM Interpretability Study: Deep analysis of internal representations.

Performs comprehensive interpretability analysis on preproc_D (Enhanced 1h) data:
  A) Gate dynamics analysis (forget, input, candidate, output gates via manual unrolling)
  B) Gradient-based feature attribution (input saliency maps + SmoothGrad)
  C) Hidden state structural analysis (PCA, clustering on LSTM hidden states)
  D) Temporal pattern analysis (FFT, autocorrelation on hidden state PCs and residuals)
  E) LSTM-specific + prediction analysis (weight matrices, predictions, residuals)

Generates a DOCX academic report with all figures and quantitative analysis.

Usage:
    python -u scripts/run_lstm_interpretability.py
    python -u scripts/run_lstm_interpretability.py --horizons 1
    python -u scripts/run_lstm_interpretability.py --horizons 1 24 --epochs 30
"""

import argparse
import gc
import json
import logging
import sys
import time
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import torch
import torch.nn as nn
from matplotlib.gridspec import GridSpec
from scipy.fft import fft, fftfreq
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from src.data.datamodule import CO2DataModule
from src.data.pipeline import run_preprocessing_pipeline
from src.data.preprocessing import inverse_scale_target
from src.evaluation.metrics import compute_metrics, save_metrics
from src.models.lstm import LSTMForecaster
from src.training.trainer import create_trainer
from src.utils.config import load_config
from src.utils.seed import seed_everything

logger = logging.getLogger(__name__)

plt.style.use("seaborn-v0_8-whitegrid")

RESULTS_BASE = Path("results/lstm_interpretability")

# Color palette (consistent with TFT study)
C_PRIMARY = "#2196F3"
C_SECONDARY = "#FF5722"
C_ACCENT = "#4CAF50"
C_WARN = "#FFC107"
C_NEUTRAL = "#607D8B"


# ======================================================================
#  Configuration
# ======================================================================

def load_interpretability_config(
    horizon: int,
    epochs_override: int | None = None,
) -> dict:
    """Load merged config for LSTM + preproc_D + specified horizon.

    Args:
        horizon: Forecast horizon in hours (1 or 24).
        epochs_override: If set, overrides max_epochs from config.

    Returns:
        Merged configuration dictionary.
    """
    config_files = [
        str(PROJECT_ROOT / "configs" / "training.yaml"),
        str(PROJECT_ROOT / "configs" / "data.yaml"),
        str(PROJECT_ROOT / "configs" / "lstm.yaml"),
        str(PROJECT_ROOT / "configs" / "experiments" / "preproc_D_enhanced_1h.yaml"),
    ]
    config = load_config(config_files)
    config["data"]["forecast_horizon_hours"] = horizon
    if epochs_override is not None:
        config["training"]["max_epochs"] = epochs_override
    config["training"]["results_dir"] = str(
        RESULTS_BASE / f"h{horizon}" / "training_runs"
    )
    return config


def get_feature_names(config: dict) -> list[str]:
    """Return full list of input feature names including target.

    Args:
        config: Merged configuration dictionary.

    Returns:
        List of feature names: feature_columns + [target_column].
    """
    feature_cols = list(config["data"]["feature_columns"])
    target_col = config["data"]["target_column"]
    return feature_cols + [target_col]


# ======================================================================
#  LSTM Training and Prediction Extraction
# ======================================================================

def train_lstm(
    config: dict,
    train_df: pd.DataFrame,
    val_df: pd.DataFrame,
    test_df: pd.DataFrame,
    horizon: int,
) -> tuple[LSTMForecaster, CO2DataModule, dict]:
    """Train LSTM, extract best model, predictions, and metrics.

    Args:
        config: Merged configuration dictionary.
        train_df: Training DataFrame from pipeline.
        val_df: Validation DataFrame from pipeline.
        test_df: Test DataFrame from pipeline.
        horizon: Forecast horizon in hours.

    Returns:
        Tuple of (best_model, data_module, results_dict).
        results_dict has keys: y_true, y_pred, metrics.
    """
    # Build DataModule
    dm = CO2DataModule(config)
    dm.build_datasets(train_df, val_df, test_df)

    # Create model
    model = LSTMForecaster(config)

    # Create trainer
    trainer, run_dir = create_trainer(config, "lstm_interp")
    trainer.fit(model, datamodule=dm)

    # Load best checkpoint
    best_ckpt = trainer.checkpoint_callback.best_model_path
    best_model = LSTMForecaster.load_from_checkpoint(best_ckpt, config=config)

    # Predictions on original scale
    predictions = trainer.predict(best_model, dm.test_dataloader())
    y_pred_scaled = torch.cat(predictions).numpy()
    y_pred = inverse_scale_target(y_pred_scaled, dm.target_scaler)

    y_true_scaled = dm.test_dataset.y.numpy()
    y_true = inverse_scale_target(y_true_scaled, dm.target_scaler)

    # Compute metrics
    metrics = compute_metrics(y_true.ravel(), y_pred.ravel())

    results = {
        "y_true": y_true,
        "y_pred": y_pred,
        "metrics": metrics,
    }

    return best_model, dm, results


# ======================================================================
#  Section A: Gate Dynamics Analysis
# ======================================================================

def manual_lstm_unroll(model: LSTMForecaster, x: torch.Tensor) -> dict:
    """Manually unroll LSTM to capture gate activations at each timestep.

    PyTorch's nn.LSTM fuses gate computations for speed, so we manually
    replicate the gate equations using the weight matrices.

    PyTorch gate order in weight matrices: (input_gate, forget_gate, cell_gate, output_gate)
    Each gate has hidden_size rows, so for hidden_size=H, the weight matrix
    has shape (4*H, input_size) for W_ih and (4*H, H) for W_hh.

    Gate equations at timestep t:
        gates = W_ih @ x_t + b_ih + W_hh @ h_{t-1} + b_hh
        i_t = sigmoid(gates[0:H])         -- input gate
        f_t = sigmoid(gates[H:2H])        -- forget gate
        g_t = tanh(gates[2H:3H])          -- candidate cell
        o_t = sigmoid(gates[3H:4H])       -- output gate
        c_t = f_t * c_{t-1} + i_t * g_t   -- cell state update
        h_t = o_t * tanh(c_t)              -- hidden state

    Args:
        model: Trained LSTMForecaster model (eval mode, on correct device).
        x: Input tensor of shape (batch, seq_len, input_size).

    Returns:
        Dictionary with gate activations per layer:
        {
            'layer_0': {
                'input_gate': (batch, seq_len, hidden_size),
                'forget_gate': ..., 'cell_gate': ..., 'output_gate': ...,
                'hidden': ..., 'cell_state': ...
            },
            'layer_1': { ... },
        }
    """
    device = x.device
    batch_size, seq_len, _ = x.shape
    num_layers = model.config["model"]["num_layers"]
    hidden_size = model.config["model"]["hidden_size"]

    result = {}
    # Current input to this layer
    layer_input = x

    for layer_idx in range(num_layers):
        # Extract weight matrices for this layer
        W_ih = getattr(model.lstm, f"weight_ih_l{layer_idx}").detach()
        W_hh = getattr(model.lstm, f"weight_hh_l{layer_idx}").detach()
        b_ih = getattr(model.lstm, f"bias_ih_l{layer_idx}").detach()
        b_hh = getattr(model.lstm, f"bias_hh_l{layer_idx}").detach()

        # Initialize hidden and cell states
        h_t = torch.zeros(batch_size, hidden_size, device=device)
        c_t = torch.zeros(batch_size, hidden_size, device=device)

        # Storage for gate activations at each timestep
        all_input_gates = []
        all_forget_gates = []
        all_cell_gates = []
        all_output_gates = []
        all_hiddens = []
        all_cells = []

        for t in range(seq_len):
            x_t = layer_input[:, t, :]  # (batch, input_size_for_layer)

            # Compute all four gates at once
            # gates shape: (batch, 4 * hidden_size)
            gates = x_t @ W_ih.T + b_ih + h_t @ W_hh.T + b_hh

            # Split into four gates - PyTorch order: i, f, g, o
            i_t = torch.sigmoid(gates[:, 0:hidden_size])
            f_t = torch.sigmoid(gates[:, hidden_size:2*hidden_size])
            g_t = torch.tanh(gates[:, 2*hidden_size:3*hidden_size])
            o_t = torch.sigmoid(gates[:, 3*hidden_size:4*hidden_size])

            # Cell state update: c_t = f_t * c_{t-1} + i_t * g_t
            c_t = f_t * c_t + i_t * g_t
            # Hidden state: h_t = o_t * tanh(c_t)
            h_t = o_t * torch.tanh(c_t)

            all_input_gates.append(i_t.detach().cpu())
            all_forget_gates.append(f_t.detach().cpu())
            all_cell_gates.append(g_t.detach().cpu())
            all_output_gates.append(o_t.detach().cpu())
            all_hiddens.append(h_t.detach().cpu())
            all_cells.append(c_t.detach().cpu())

        # Stack: (batch, seq_len, hidden_size)
        result[f"layer_{layer_idx}"] = {
            "input_gate": torch.stack(all_input_gates, dim=1).numpy(),
            "forget_gate": torch.stack(all_forget_gates, dim=1).numpy(),
            "cell_gate": torch.stack(all_cell_gates, dim=1).numpy(),
            "output_gate": torch.stack(all_output_gates, dim=1).numpy(),
            "hidden": torch.stack(all_hiddens, dim=1).numpy(),
            "cell_state": torch.stack(all_cells, dim=1).numpy(),
        }

        # Next layer input is this layer's hidden states (on device)
        layer_input = torch.stack(all_hiddens, dim=1).to(device)

    return result


def run_gate_dynamics(
    best_model: LSTMForecaster,
    dm: CO2DataModule,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Extract and visualize LSTM gate activations via manual unrolling.

    Args:
        best_model: Trained LSTM model.
        dm: Data module with test data.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.

    Returns:
        Dictionary of aggregated gate statistics for the report.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [A] Extracting gate dynamics via manual LSTM unrolling...")

    best_model.eval()
    device = next(best_model.parameters()).device

    # Collect gate activations from test batches
    all_gates = {}
    max_batches = 30

    with torch.no_grad():
        for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
            if i >= max_batches:
                break
            x_batch = x_batch.to(device)
            batch_gates = manual_lstm_unroll(best_model, x_batch)

            for layer_key, layer_data in batch_gates.items():
                if layer_key not in all_gates:
                    all_gates[layer_key] = {k: [] for k in layer_data}
                for gate_name, gate_val in layer_data.items():
                    all_gates[layer_key][gate_name].append(gate_val)

    # Concatenate across batches
    gate_data = {}
    for layer_key, layer_dict in all_gates.items():
        gate_data[layer_key] = {}
        for gate_name, arrays in layer_dict.items():
            gate_data[layer_key][gate_name] = np.concatenate(arrays, axis=0)

    num_layers = len(gate_data)
    n_samples = gate_data["layer_0"]["forget_gate"].shape[0]
    seq_len = gate_data["layer_0"]["forget_gate"].shape[1]
    hidden_size = gate_data["layer_0"]["forget_gate"].shape[2]
    print(f"  [A] Captured {n_samples} samples, {num_layers} layers, "
          f"seq_len={seq_len}, hidden={hidden_size}")

    gate_stats = {}

    # --- Plot 1: Gate activation distributions (4 histograms per layer) ---
    gate_names_ordered = ["input_gate", "forget_gate", "cell_gate", "output_gate"]
    gate_labels = ["Input Gate (i)", "Forget Gate (f)", "Candidate (g)", "Output Gate (o)"]
    gate_colors = [C_PRIMARY, C_SECONDARY, C_ACCENT, C_WARN]

    for layer_key in sorted(gate_data.keys()):
        layer_idx = int(layer_key.split("_")[1])
        fig, axes = plt.subplots(1, 4, figsize=(20, 4))

        for g_idx, (gname, glabel, gcolor) in enumerate(
            zip(gate_names_ordered, gate_labels, gate_colors)
        ):
            ax = axes[g_idx]
            vals = gate_data[layer_key][gname].ravel()
            # Subsample for performance
            if len(vals) > 200000:
                rng = np.random.default_rng(42)
                vals = rng.choice(vals, 200000, replace=False)
            ax.hist(vals, bins=80, color=gcolor, alpha=0.7, edgecolor="white", density=True)
            ax.axvline(np.mean(vals), color="black", linestyle="--", linewidth=1,
                       label=f"Mean={np.mean(vals):.3f}")
            ax.set_title(glabel, fontsize=11)
            ax.set_xlabel("Activation")
            ax.set_ylabel("Density")
            ax.legend(fontsize=8)

            # Collect stats
            stat_key = f"L{layer_idx}_{gname}"
            gate_stats[stat_key] = {
                "mean": float(np.mean(vals)),
                "std": float(np.std(vals)),
                "min": float(np.min(vals)),
                "max": float(np.max(vals)),
                "near_zero_pct": float((np.abs(vals) < 0.05).mean() * 100),
                "near_one_pct": float((vals > 0.95).mean() * 100),
            }

        fig.suptitle(
            f"LSTM Gate Activation Distributions - Layer {layer_idx} - {horizon}h",
            fontsize=13,
        )
        plt.tight_layout()
        plt.savefig(
            output_dir / f"gate_distributions_L{layer_idx}_h{horizon}.png",
            dpi=150, bbox_inches="tight",
        )
        plt.close(fig)

    # --- Plot 2: Gate evolution over lookback window (mean across samples & hidden dims) ---
    for layer_key in sorted(gate_data.keys()):
        layer_idx = int(layer_key.split("_")[1])
        fig, ax = plt.subplots(figsize=(12, 5))
        time_axis = np.arange(-seq_len, 0)

        for gname, glabel, gcolor in zip(gate_names_ordered, gate_labels, gate_colors):
            # Mean across samples and hidden dims -> (seq_len,)
            temporal = gate_data[layer_key][gname].mean(axis=(0, 2))
            ax.plot(time_axis, temporal, label=glabel, color=gcolor, linewidth=2)

        ax.set_xlabel("Timestep (hours ago)")
        ax.set_ylabel("Mean Gate Activation")
        ax.set_title(f"Gate Evolution Over Lookback - Layer {layer_idx} - {horizon}h")
        ax.legend()
        ax.grid(alpha=0.3)
        plt.tight_layout()
        plt.savefig(
            output_dir / f"gate_evolution_L{layer_idx}_h{horizon}.png",
            dpi=150, bbox_inches="tight",
        )
        plt.close(fig)

    # --- Plot 3: Forget gate heatmap (first layer, mean across samples) ---
    # Shows which hidden units retain vs forget information at each timestep
    fg_mean = gate_data["layer_0"]["forget_gate"].mean(axis=0)  # (seq_len, hidden_size)
    n_show_dims = min(64, hidden_size)

    fig, ax = plt.subplots(figsize=(14, 6))
    im = ax.imshow(
        fg_mean[:, :n_show_dims].T, aspect="auto", cmap="RdYlGn",
        interpolation="nearest", vmin=0, vmax=1,
    )
    ax.set_xlabel("Timestep (lookback)")
    ax.set_ylabel("Hidden Dimension")
    ax.set_title(f"Forget Gate Heatmap (Layer 0, mean across samples) - {horizon}h")
    tick_positions = np.linspace(0, seq_len - 1, min(8, seq_len)).astype(int)
    ax.set_xticks(tick_positions)
    ax.set_xticklabels([f"-{seq_len - t}" for t in tick_positions])
    plt.colorbar(im, ax=ax, label="Forget Gate Value (0=forget, 1=retain)")
    plt.tight_layout()
    plt.savefig(
        output_dir / f"forget_gate_heatmap_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 4: Gate statistics table ---
    fig, ax = plt.subplots(figsize=(14, max(3, 0.4 * len(gate_stats))))
    ax.axis("off")
    rows = []
    for stat_key in sorted(gate_stats.keys()):
        s = gate_stats[stat_key]
        rows.append([
            stat_key,
            f"{s['mean']:.4f}",
            f"{s['std']:.4f}",
            f"{s['min']:.4f}",
            f"{s['max']:.4f}",
            f"{s['near_zero_pct']:.1f}%",
            f"{s['near_one_pct']:.1f}%",
        ])

    table = ax.table(
        cellText=rows,
        colLabels=["Gate", "Mean", "Std", "Min", "Max", "<0.05", ">0.95"],
        loc="center",
        cellLoc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1.0, 1.5)
    ax.set_title(f"LSTM Gate Statistics - {horizon}h", fontsize=12, pad=20)
    plt.tight_layout()
    plt.savefig(
        output_dir / f"gate_statistics_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    print(f"  [A] Gate dynamics: {num_layers} layers analyzed, plots saved")
    return gate_stats


# ======================================================================
#  Section B: Gradient-Based Feature Attribution
# ======================================================================

def run_gradient_attribution(
    best_model: LSTMForecaster,
    dm: CO2DataModule,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> dict:
    """Compute input gradients and SmoothGrad for feature attribution.

    Standard gradient saliency:
        Enable gradients on input x, forward pass, backward on prediction mean.
        Collect |grad| of shape (batch, lookback, n_features), average across samples.

    SmoothGrad (Smilkov et al., 2017):
        Adds Gaussian noise to the input N times and averages the gradients,
        producing smoother and more interpretable saliency maps.

    Args:
        best_model: Trained LSTM model.
        dm: Data module with test data.
        config: Merged configuration dictionary.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.

    Returns:
        Dictionary with feature importance data for the report.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [B] Computing gradient-based feature attribution...")

    best_model.eval()
    device = next(best_model.parameters()).device
    feature_names = get_feature_names(config)

    # --- Standard gradient saliency ---
    all_grads = []
    max_batches = 50

    for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
        if i >= max_batches:
            break
        x_input = x_batch.clone().detach().to(device).requires_grad_(True)
        pred = best_model(x_input)
        target = pred.mean()
        target.backward()

        if x_input.grad is not None:
            grad_abs = x_input.grad.abs().detach().cpu().numpy()
            all_grads.append(grad_abs)

        best_model.zero_grad()

    if not all_grads:
        print("  [B] No gradients collected. Skipping.")
        return {}

    # Mean across all samples: (lookback, n_features)
    grads = np.concatenate(all_grads, axis=0)
    avg_grad = grads.mean(axis=0)

    # --- SmoothGrad ---
    print("  [B] Computing SmoothGrad (N=30 perturbations)...")
    n_smooth = 30
    smooth_grads = []
    noise_std = 0.1  # Standard deviation of Gaussian noise

    for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
        if i >= max_batches:
            break
        x_base = x_batch.to(device)
        batch_smooth = torch.zeros_like(x_base)

        for _ in range(n_smooth):
            noise = torch.randn_like(x_base) * noise_std
            x_noisy = (x_base + noise).requires_grad_(True)
            pred = best_model(x_noisy)
            pred.mean().backward()
            if x_noisy.grad is not None:
                batch_smooth += x_noisy.grad.abs().detach()
            best_model.zero_grad()

        batch_smooth /= n_smooth
        smooth_grads.append(batch_smooth.cpu().numpy())

    smooth_grads_arr = np.concatenate(smooth_grads, axis=0)
    avg_smooth_grad = smooth_grads_arr.mean(axis=0)

    lookback = avg_grad.shape[0]
    n_features = avg_grad.shape[1]

    # Truncate or pad feature names
    if len(feature_names) != n_features:
        feature_names = [f"feat_{i}" for i in range(n_features)]

    # --- Plot 1: Gradient heatmap (lookback x features) ---
    fig, ax = plt.subplots(figsize=(max(10, n_features * 0.6), 6))
    im = ax.imshow(avg_grad.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Lookback Timestep (hours ago)")
    ax.set_ylabel("Feature")
    ax.set_yticks(range(n_features))
    ax.set_yticklabels(feature_names, fontsize=8)
    tick_positions = np.linspace(0, lookback - 1, min(8, lookback)).astype(int)
    ax.set_xticks(tick_positions)
    ax.set_xticklabels([f"-{lookback - t}" for t in tick_positions])
    ax.set_title(f"Gradient Attribution Heatmap - {horizon}h Horizon")
    plt.colorbar(im, ax=ax, label="|Gradient|")
    plt.tight_layout()
    plt.savefig(
        output_dir / f"gradient_heatmap_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 2: Per-feature gradient magnitude (bar chart) ---
    feat_importance = avg_grad.mean(axis=0)
    feat_importance_pct = feat_importance / feat_importance.sum() * 100
    sort_idx = np.argsort(feat_importance_pct)

    fig, ax = plt.subplots(figsize=(8, max(4, n_features * 0.35)))
    ax.barh(range(len(sort_idx)), feat_importance_pct[sort_idx], color=C_ACCENT)
    ax.set_yticks(range(len(sort_idx)))
    ax.set_yticklabels([feature_names[i] for i in sort_idx], fontsize=9)
    ax.set_xlabel("Gradient Importance (%)")
    ax.set_title(f"Feature Attribution (Gradient) - {horizon}h")
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    plt.savefig(
        output_dir / f"gradient_feature_ranking_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # --- Plot 3: Temporal gradient profile ---
    temporal_importance = avg_grad.sum(axis=1)
    temporal_importance_norm = temporal_importance / temporal_importance.sum()
    time_idx = np.arange(-lookback, 0)

    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(time_idx, temporal_importance_norm, color=C_PRIMARY, alpha=0.8)
    ax.set_xlabel("Hours Ago")
    ax.set_ylabel("Normalized Gradient Magnitude")
    ax.set_title(f"Temporal Gradient Profile - {horizon}h Horizon")
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(
        output_dir / f"gradient_temporal_profile_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # --- Plot 4: SmoothGrad comparison (side by side with vanilla) ---
    fig, axes = plt.subplots(1, 2, figsize=(20, 6))

    ax = axes[0]
    im = ax.imshow(avg_grad.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Lookback Timestep")
    ax.set_ylabel("Feature")
    ax.set_yticks(range(n_features))
    ax.set_yticklabels(feature_names, fontsize=7)
    ax.set_title("Vanilla Gradient")
    plt.colorbar(im, ax=ax, fraction=0.046)

    ax = axes[1]
    im = ax.imshow(avg_smooth_grad.T, aspect="auto", cmap="YlOrRd", interpolation="nearest")
    ax.set_xlabel("Lookback Timestep")
    ax.set_ylabel("Feature")
    ax.set_yticks(range(n_features))
    ax.set_yticklabels(feature_names, fontsize=7)
    ax.set_title("SmoothGrad (N=30)")
    plt.colorbar(im, ax=ax, fraction=0.046)

    fig.suptitle(f"Gradient vs SmoothGrad Attribution - {horizon}h", fontsize=13)
    plt.tight_layout()
    plt.savefig(
        output_dir / f"gradient_vs_smoothgrad_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # --- Plot 5: SmoothGrad feature ranking ---
    smooth_feat_importance = avg_smooth_grad.mean(axis=0)
    smooth_feat_pct = smooth_feat_importance / smooth_feat_importance.sum() * 100
    smooth_sort_idx = np.argsort(smooth_feat_pct)

    fig, ax = plt.subplots(figsize=(8, max(4, n_features * 0.35)))
    ax.barh(range(len(smooth_sort_idx)), smooth_feat_pct[smooth_sort_idx], color=C_SECONDARY)
    ax.set_yticks(range(len(smooth_sort_idx)))
    ax.set_yticklabels([feature_names[i] for i in smooth_sort_idx], fontsize=9)
    ax.set_xlabel("SmoothGrad Importance (%)")
    ax.set_title(f"Feature Attribution (SmoothGrad) - {horizon}h")
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    plt.savefig(
        output_dir / f"smoothgrad_feature_ranking_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # Save gradient data
    grad_df = pd.DataFrame(avg_grad, columns=feature_names)
    grad_df.to_csv(output_dir / f"gradient_attribution_h{horizon}.csv", index=False)
    smooth_df = pd.DataFrame(avg_smooth_grad, columns=feature_names)
    smooth_df.to_csv(output_dir / f"smoothgrad_attribution_h{horizon}.csv", index=False)

    # Build importance dictionary
    importance_data = {
        "feature_names": feature_names,
        "gradient_pct": feat_importance_pct.tolist(),
        "smoothgrad_pct": smooth_feat_pct.tolist(),
        "gradient_ranking": [feature_names[i] for i in np.argsort(-feat_importance_pct)],
        "smoothgrad_ranking": [feature_names[i] for i in np.argsort(-smooth_feat_pct)],
    }

    print(f"  [B] Gradient attribution: {grads.shape[0]} samples, {n_features} features")
    print("  Top 5 features (vanilla gradient):")
    for rank, fname in enumerate(importance_data["gradient_ranking"][:5]):
        idx = feature_names.index(fname)
        print(f"    {rank+1}. {fname}: {feat_importance_pct[idx]:.1f}%")

    return importance_data


# ======================================================================
#  Section C: Hidden State Structural Analysis
# ======================================================================

def run_hidden_state_analysis(
    best_model: LSTMForecaster,
    dm: CO2DataModule,
    test_df: pd.DataFrame,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> None:
    """PCA and clustering on LSTM hidden states.

    Hooks into model.lstm to capture h_n (final hidden state per layer)
    and full lstm_out (hidden state at every timestep). Performs PCA
    dimensionality reduction and K-Means clustering to identify learned
    operational regimes.

    Args:
        best_model: Trained LSTM model.
        dm: Data module with test data.
        test_df: Test DataFrame with datetime column.
        config: Merged configuration dictionary.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [C] Extracting LSTM hidden states via forward hooks...")

    best_model.eval()
    device = next(best_model.parameters()).device

    # Hook to capture LSTM outputs
    lstm_outputs = []
    lstm_h_n = []

    def lstm_hook(module, inp, output):
        # output = (lstm_out, (h_n, c_n))
        lstm_outputs.append(output[0].detach().cpu())  # (batch, seq, hidden)
        lstm_h_n.append(output[1][0].detach().cpu())   # (n_layers, batch, hidden)

    hook_handle = best_model.lstm.register_forward_hook(lstm_hook)

    max_batches = 30
    all_targets = []

    with torch.no_grad():
        for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
            if i >= max_batches:
                break
            x_batch = x_batch.to(device)
            best_model(x_batch)
            all_targets.append(y_batch.numpy()[:, 0] if y_batch.ndim > 1 else y_batch.numpy())

    hook_handle.remove()

    if not lstm_h_n:
        print("  [C] No hidden states captured. Skipping.")
        return

    # h_n shape per batch: (n_layers, batch, hidden_size)
    # Take last layer's hidden state
    hidden_states = np.concatenate(
        [h[-1].numpy() for h in lstm_h_n], axis=0
    )  # (total_samples, hidden_size)
    target_values = np.concatenate(all_targets, axis=0).ravel()

    # Trim to match
    n = min(len(hidden_states), len(target_values))
    hidden_states = hidden_states[:n]
    target_values = target_values[:n]

    # Inverse scale targets for CO2 level categorization
    target_values_orig = inverse_scale_target(target_values, dm.target_scaler)

    print(f"  [C] Hidden states shape: {hidden_states.shape}")

    # PCA
    n_components = min(20, hidden_states.shape[1])
    pca = PCA(n_components=n_components)
    pca_result = pca.fit_transform(hidden_states)

    # --- Plot 1: PCA explained variance ---
    fig, ax = plt.subplots(figsize=(8, 5))
    cumvar = np.cumsum(pca.explained_variance_ratio_) * 100
    ax.bar(range(1, n_components + 1), pca.explained_variance_ratio_ * 100,
           color=C_PRIMARY, alpha=0.7, label="Individual")
    ax.plot(range(1, n_components + 1), cumvar, "r-o", markersize=4, label="Cumulative")
    ax.set_xlabel("Principal Component")
    ax.set_ylabel("Explained Variance (%)")
    ax.set_title(f"PCA on LSTM Hidden States (Layer -1) - {horizon}h")
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(
        output_dir / f"pca_variance_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # CO2 level categories
    co2_bins = [0, 500, 1000, np.inf]
    co2_labels = ["Low (<500)", "Medium (500-1000)", "High (>1000)"]
    co2_cats = pd.cut(target_values_orig, bins=co2_bins, labels=co2_labels)
    cat_colors = {
        "Low (<500)": C_ACCENT,
        "Medium (500-1000)": C_PRIMARY,
        "High (>1000)": C_SECONDARY,
    }

    # --- Plot 2: PCA scatter colored by CO2 level ---
    fig, ax = plt.subplots(figsize=(8, 7))
    for label in co2_labels:
        mask = co2_cats == label
        if mask.sum() > 0:
            ax.scatter(
                pca_result[mask, 0], pca_result[mask, 1],
                c=cat_colors[label], label=f"{label} (n={mask.sum()})",
                alpha=0.5, s=10,
            )
    ax.set_xlabel(f"PC1 ({pca.explained_variance_ratio_[0]*100:.1f}%)")
    ax.set_ylabel(f"PC2 ({pca.explained_variance_ratio_[1]*100:.1f}%)")
    ax.set_title(f"LSTM Hidden States by CO2 Level - {horizon}h")
    ax.legend()
    plt.tight_layout()
    plt.savefig(
        output_dir / f"pca_co2_level_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 3: PCA scatter colored by hour of day ---
    datetime_col = config["data"]["datetime_column"]
    if datetime_col in test_df.columns:
        test_dates = pd.DatetimeIndex(test_df[datetime_col])
        if len(test_dates) >= n:
            hours = test_dates[-n:].hour
        else:
            hours = np.zeros(n)
    else:
        hours = np.zeros(n)

    fig, ax = plt.subplots(figsize=(8, 7))
    scatter = ax.scatter(
        pca_result[:, 0], pca_result[:, 1],
        c=hours, cmap="twilight", alpha=0.5, s=10,
    )
    ax.set_xlabel(f"PC1 ({pca.explained_variance_ratio_[0]*100:.1f}%)")
    ax.set_ylabel(f"PC2 ({pca.explained_variance_ratio_[1]*100:.1f}%)")
    ax.set_title(f"LSTM Hidden States by Hour of Day - {horizon}h")
    plt.colorbar(scatter, ax=ax, label="Hour of Day")
    plt.tight_layout()
    plt.savefig(
        output_dir / f"pca_hour_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 4: K-Means clustering (K=4) with regime crosstab ---
    n_clusters = 4
    kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
    cluster_labels = kmeans.fit_predict(hidden_states)

    fig, axes = plt.subplots(1, 2, figsize=(15, 6))

    # Scatter colored by cluster
    ax = axes[0]
    cluster_colors = [C_PRIMARY, C_SECONDARY, C_ACCENT, C_WARN]
    for c in range(n_clusters):
        mask = cluster_labels == c
        ax.scatter(
            pca_result[mask, 0], pca_result[mask, 1],
            c=cluster_colors[c], label=f"Cluster {c} (n={mask.sum()})",
            alpha=0.5, s=10,
        )
    ax.set_xlabel("PC1")
    ax.set_ylabel("PC2")
    ax.set_title("K-Means Clusters in PCA Space")
    ax.legend(fontsize=8)

    # Crosstab table
    ax = axes[1]
    ax.axis("off")
    cross = pd.crosstab(
        pd.Series(cluster_labels, name="Cluster"),
        pd.Series(co2_cats, name="CO2 Level"),
    )
    rows_data = []
    for c_idx in range(n_clusters):
        row = [f"Cluster {c_idx}"]
        for label in co2_labels:
            val = cross.loc[c_idx, label] if c_idx in cross.index and label in cross.columns else 0
            row.append(str(val))
        rows_data.append(row)

    tbl = ax.table(
        cellText=rows_data,
        colLabels=["Cluster"] + co2_labels,
        loc="center",
        cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.0, 1.5)
    ax.set_title("Cluster vs CO2 Regime Crosstab", fontsize=11, pad=20)

    fig.suptitle(f"Hidden State Clustering - {horizon}h", fontsize=13)
    plt.tight_layout()
    plt.savefig(
        output_dir / f"clustering_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 5: Hidden state trajectories for 5 example sequences ---
    # Use lstm_outputs (full sequence hidden states) to visualize trajectories in PCA space
    if lstm_outputs:
        # Concatenate all full sequence outputs
        # Each is (batch, seq_len, hidden_size)
        full_seqs = torch.cat(lstm_outputs, dim=0).numpy()  # (total, seq_len, hidden)
        n_seqs = min(5, len(full_seqs))
        # Select 5 evenly spaced sequences
        seq_indices = np.linspace(0, len(full_seqs) - 1, n_seqs).astype(int)

        fig, ax = plt.subplots(figsize=(10, 8))
        trajectory_cmap = plt.colormaps.get_cmap("tab10")

        for plot_idx, seq_idx in enumerate(seq_indices):
            seq_hidden = full_seqs[seq_idx]  # (seq_len, hidden_size)
            # Project into PCA space
            seq_pca = pca.transform(seq_hidden)
            ax.plot(
                seq_pca[:, 0], seq_pca[:, 1],
                color=trajectory_cmap(plot_idx), alpha=0.7,
                linewidth=1.5, label=f"Seq {seq_idx}",
            )
            # Mark start and end
            ax.scatter(seq_pca[0, 0], seq_pca[0, 1], marker="o",
                       c=[trajectory_cmap(plot_idx)], s=60, zorder=5)
            ax.scatter(seq_pca[-1, 0], seq_pca[-1, 1], marker="X",
                       c=[trajectory_cmap(plot_idx)], s=80, zorder=5)

        ax.set_xlabel("PC1")
        ax.set_ylabel("PC2")
        ax.set_title(f"Hidden State Trajectories in PCA Space - {horizon}h\n(o=start, X=end)")
        ax.legend(fontsize=9)
        ax.grid(alpha=0.3)
        plt.tight_layout()
        plt.savefig(
            output_dir / f"hidden_trajectories_h{horizon}.png",
            dpi=150, bbox_inches="tight",
        )
        plt.close(fig)

    print(f"  [C] Hidden state analysis complete: {n} samples, {n_components} PCs")


# ======================================================================
#  Section D: Temporal Pattern Analysis
# ======================================================================

def run_temporal_patterns(
    best_model: LSTMForecaster,
    dm: CO2DataModule,
    results: dict,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> None:
    """FFT and autocorrelation on hidden state PCs, FFT of residuals, rolling RMSE.

    Args:
        best_model: Trained LSTM model.
        dm: Data module with test data.
        results: Dictionary with y_true and y_pred arrays.
        config: Merged configuration dictionary.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [D] Analyzing temporal patterns...")

    best_model.eval()
    device = next(best_model.parameters()).device

    # Extract hidden states for temporal analysis
    lstm_outputs_list = []

    def lstm_hook(module, inp, output):
        lstm_outputs_list.append(output[1][0].detach().cpu())  # h_n: (n_layers, batch, hidden)

    hook_handle = best_model.lstm.register_forward_hook(lstm_hook)
    max_batches = 50

    with torch.no_grad():
        for i, (x_batch, y_batch) in enumerate(dm.test_dataloader()):
            if i >= max_batches:
                break
            x_batch = x_batch.to(device)
            best_model(x_batch)

    hook_handle.remove()

    if not lstm_outputs_list:
        print("  [D] No hidden states captured. Skipping temporal analysis.")
        return

    # Last layer hidden states: concatenate across batches
    hidden_concat = np.concatenate(
        [h[-1].numpy() for h in lstm_outputs_list], axis=0
    )  # (total_samples, hidden_size)
    n_samples_h = hidden_concat.shape[0]

    # PCA for temporal analysis
    pca_temp = PCA(n_components=min(5, hidden_concat.shape[1]))
    pcs = pca_temp.fit_transform(hidden_concat)

    # --- Plot 1: FFT of hidden state PCs ---
    if n_samples_h > 50:
        fig, axes = plt.subplots(3, 1, figsize=(12, 10))
        for pc_idx in range(min(3, pcs.shape[1])):
            ax = axes[pc_idx]
            pc_series = pcs[:, pc_idx]
            pc_series_centered = pc_series - pc_series.mean()
            n_fft = len(pc_series_centered)

            yf = np.abs(fft(pc_series_centered))
            xf = fftfreq(n_fft, d=1.0)  # d=1 hour
            pos_mask = xf > 0
            periods = 1.0 / xf[pos_mask]
            magnitudes = yf[pos_mask]

            valid = periods < min(200, n_fft // 2)
            ax.plot(periods[valid], magnitudes[valid], color=C_PRIMARY, linewidth=1)
            ax.set_ylabel(f"PC{pc_idx+1} FFT Mag")
            ax.set_title(
                f"PC{pc_idx+1} ({pca_temp.explained_variance_ratio_[pc_idx]*100:.1f}% var)"
            )
            ax.grid(alpha=0.3)

        axes[-1].set_xlabel("Period (hours)")
        fig.suptitle(f"FFT of Hidden State PCs - {horizon}h", fontsize=13)
        plt.tight_layout()
        plt.savefig(
            output_dir / f"fft_hidden_pcs_h{horizon}.png", dpi=150, bbox_inches="tight"
        )
        plt.close(fig)

    # --- Plot 2: ACF of hidden state PCs ---
    if n_samples_h > 50:
        max_lag = min(72, n_samples_h // 3)

        fig, axes = plt.subplots(3, 1, figsize=(12, 9))
        for pc_idx in range(min(3, pcs.shape[1])):
            ax = axes[pc_idx]
            pc_series = pcs[:, pc_idx]
            pc_series = pc_series - pc_series.mean()
            norm = np.sum(pc_series**2)

            if norm > 0:
                autocorr = np.correlate(pc_series, pc_series, mode="full")
                autocorr = autocorr[len(autocorr) // 2:]
                autocorr = autocorr[:max_lag] / norm
            else:
                autocorr = np.zeros(max_lag)

            ax.bar(range(max_lag), autocorr, color=C_PRIMARY, alpha=0.7, width=0.8)
            ax.axhline(0, color="black", linewidth=0.5)
            ax.set_ylabel(f"PC{pc_idx+1} ACF")
            ax.set_title(
                f"PC{pc_idx+1} ({pca_temp.explained_variance_ratio_[pc_idx]*100:.1f}% var)"
            )
            ax.grid(axis="y", alpha=0.3)

        axes[-1].set_xlabel("Lag (hours)")
        fig.suptitle(f"Autocorrelation of Hidden State PCs - {horizon}h", fontsize=13)
        plt.tight_layout()
        plt.savefig(
            output_dir / f"autocorrelation_hidden_pcs_h{horizon}.png",
            dpi=150, bbox_inches="tight",
        )
        plt.close(fig)

    # --- Plot 3: FFT of residuals ---
    y_true = results["y_true"].ravel()
    y_pred = results["y_pred"].ravel()
    residuals = y_true - y_pred

    if len(residuals) > 50:
        fig, ax = plt.subplots(figsize=(12, 5))
        res_centered = residuals - residuals.mean()
        n_res = len(res_centered)
        yf = np.abs(fft(res_centered))
        xf = fftfreq(n_res, d=1.0)
        pos_mask = xf > 0
        periods = 1.0 / xf[pos_mask]
        magnitudes = yf[pos_mask]
        valid = periods < min(200, n_res // 2)

        ax.plot(periods[valid], magnitudes[valid], color=C_SECONDARY, linewidth=1)
        ax.set_xlabel("Period (hours)")
        ax.set_ylabel("FFT Magnitude")
        ax.set_title(f"FFT of Prediction Residuals - {horizon}h")
        ax.grid(alpha=0.3)

        # Mark key periods (24h daily cycle)
        for p_mark in [24, 12, 6]:
            if p_mark < periods[valid].max():
                ax.axvline(p_mark, color=C_WARN, linestyle="--", alpha=0.6,
                           label=f"{p_mark}h")
        ax.legend(fontsize=9)
        plt.tight_layout()
        plt.savefig(
            output_dir / f"fft_residuals_h{horizon}.png", dpi=150, bbox_inches="tight"
        )
        plt.close(fig)

    # --- Plot 4: Rolling RMSE (24h window) ---
    if len(residuals) > 48:
        window = 24
        rolling_sq_err = pd.Series(residuals**2).rolling(window=window, min_periods=1).mean()
        rolling_rmse = np.sqrt(rolling_sq_err.values)

        fig, ax = plt.subplots(figsize=(14, 5))
        ax.plot(rolling_rmse, color=C_PRIMARY, linewidth=1, alpha=0.8)
        ax.axhline(np.sqrt(np.mean(residuals**2)), color=C_SECONDARY, linestyle="--",
                    linewidth=1.5, label=f"Overall RMSE = {np.sqrt(np.mean(residuals**2)):.2f}")
        ax.set_xlabel("Sample Index")
        ax.set_ylabel("Rolling RMSE (ppm)")
        ax.set_title(f"Rolling RMSE (24h window) - {horizon}h Horizon")
        ax.legend()
        ax.grid(alpha=0.3)
        plt.tight_layout()
        plt.savefig(
            output_dir / f"rolling_rmse_h{horizon}.png", dpi=150, bbox_inches="tight"
        )
        plt.close(fig)

    print(f"  [D] Temporal pattern analysis complete")


# ======================================================================
#  Section E: LSTM-Specific + Prediction Analysis
# ======================================================================

def run_lstm_specific_analysis(
    best_model: LSTMForecaster,
    dm: CO2DataModule,
    results: dict,
    test_df: pd.DataFrame,
    config: dict,
    horizon: int,
    output_dir: Path,
) -> None:
    """Weight matrix visualization, prediction overlay, scatter, residuals, error by CO2 level.

    Args:
        best_model: Trained LSTM model.
        dm: Data module with test data.
        results: Dictionary with y_true, y_pred, metrics.
        test_df: Test DataFrame with datetime column.
        config: Merged configuration dictionary.
        horizon: Forecast horizon in hours.
        output_dir: Directory to save plots.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    print("  [E] Generating LSTM-specific + prediction analysis...")

    feature_names = get_feature_names(config)
    metrics = results["metrics"]
    y_true = results["y_true"].ravel()
    y_pred = results["y_pred"].ravel()
    residuals = y_true - y_pred

    # --- Plot 1: Weight matrix visualization (W_ih for layer 0) ---
    W_ih = best_model.lstm.weight_ih_l0.detach().cpu().numpy()
    hidden_size = W_ih.shape[0] // 4
    input_size = W_ih.shape[1]

    gate_slices = {
        "Input Gate (i)": W_ih[0:hidden_size, :],
        "Forget Gate (f)": W_ih[hidden_size:2*hidden_size, :],
        "Candidate (g)": W_ih[2*hidden_size:3*hidden_size, :],
        "Output Gate (o)": W_ih[3*hidden_size:4*hidden_size, :],
    }

    fig, axes = plt.subplots(2, 2, figsize=(16, 12))
    for idx, (gate_name, weights) in enumerate(gate_slices.items()):
        ax = axes[idx // 2, idx % 2]
        # Show first 32 hidden dims for readability
        n_show = min(32, hidden_size)
        im = ax.imshow(
            weights[:n_show, :].T, aspect="auto", cmap="RdBu_r",
            interpolation="nearest",
            vmin=-np.abs(weights[:n_show, :]).max(),
            vmax=np.abs(weights[:n_show, :]).max(),
        )
        ax.set_xlabel("Hidden Dimension")
        ax.set_ylabel("Input Feature")
        ax.set_yticks(range(input_size))
        yticklabels = feature_names[:input_size] if len(feature_names) >= input_size else [
            f"feat_{i}" for i in range(input_size)
        ]
        ax.set_yticklabels(yticklabels, fontsize=7)
        ax.set_title(gate_name)
        plt.colorbar(im, ax=ax, fraction=0.046)

    fig.suptitle(f"LSTM Layer 0 W_ih Weight Matrices - {horizon}h", fontsize=13)
    plt.tight_layout()
    plt.savefig(
        output_dir / f"weight_matrix_W_ih_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 2: Prediction overlay (actual vs predicted) ---
    n_show = min(500, len(y_true))
    fig, ax = plt.subplots(figsize=(14, 5))
    ax.plot(y_true[:n_show], label="Actual", color=C_PRIMARY, linewidth=1.0, alpha=0.8)
    ax.plot(y_pred[:n_show], label="Predicted", color=C_SECONDARY, linewidth=1.0, alpha=0.8)
    ax.fill_between(
        range(n_show),
        y_true[:n_show], y_pred[:n_show],
        alpha=0.15, color=C_WARN,
    )
    ax.set_xlabel("Sample Index")
    ax.set_ylabel("CO2 (ppm)")
    ax.set_title(f"LSTM Predictions vs Actual - {horizon}h Horizon")
    ax.legend()
    ax.grid(alpha=0.3)
    plt.tight_layout()
    plt.savefig(
        output_dir / f"predictions_overlay_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 3: Scatter plot with R2 ---
    fig, ax = plt.subplots(figsize=(7, 7))
    ax.scatter(y_true, y_pred, alpha=0.3, s=5, color=C_PRIMARY)
    vmin = min(y_true.min(), y_pred.min())
    vmax = max(y_true.max(), y_pred.max())
    ax.plot([vmin, vmax], [vmin, vmax], "r--", linewidth=1.5, label="Perfect Prediction")
    ax.text(
        0.05, 0.95,
        f"R2 = {metrics['r2']:.4f}\nRMSE = {metrics['rmse']:.2f}\nMAE = {metrics['mae']:.2f}",
        transform=ax.transAxes, fontsize=11, va="top",
        bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8),
    )
    ax.set_xlabel("Actual CO2 (ppm)")
    ax.set_ylabel("Predicted CO2 (ppm)")
    ax.set_title(f"LSTM Scatter Plot - {horizon}h")
    ax.legend()
    plt.tight_layout()
    plt.savefig(
        output_dir / f"scatter_r2_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 4: 4-panel residual analysis ---
    datetime_col = config["data"]["datetime_column"]
    test_dates = None
    if datetime_col in test_df.columns:
        test_dates = pd.DatetimeIndex(test_df[datetime_col])

    fig, axes = plt.subplots(2, 2, figsize=(14, 10))

    # (a) Distribution
    ax = axes[0, 0]
    ax.hist(residuals, bins=50, color=C_PRIMARY, alpha=0.7, edgecolor="white", density=True)
    ax.axvline(0, color="red", linestyle="--", linewidth=1.5)
    ax.set_xlabel("Residual (ppm)")
    ax.set_ylabel("Density")
    ax.set_title("(a) Residual Distribution")
    ax.text(
        0.95, 0.95,
        f"Mean: {residuals.mean():.2f}\nStd: {residuals.std():.2f}\n"
        f"Skew: {float(pd.Series(residuals).skew()):.2f}",
        transform=ax.transAxes, fontsize=9, ha="right", va="top",
        bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8),
    )

    # (b) Over time
    ax = axes[0, 1]
    ax.scatter(np.arange(len(residuals)), residuals, alpha=0.3, s=3, color=C_SECONDARY)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xlabel("Sample Index")
    ax.set_ylabel("Residual (ppm)")
    ax.set_title("(b) Residuals Over Time")

    # (c) By hour of day
    ax = axes[1, 0]
    if test_dates is not None and len(test_dates) >= len(residuals):
        hours = test_dates[-len(residuals):].hour
        res_df = pd.DataFrame({"hour": hours, "residual": residuals})
        res_df.boxplot(column="residual", by="hour", ax=ax)
        ax.set_xlabel("Hour of Day")
        ax.set_ylabel("Residual (ppm)")
        ax.set_title("(c) Residuals by Hour")
        plt.suptitle("")  # Remove auto-title from boxplot
    else:
        ax.text(0.5, 0.5, "No datetime available", ha="center", va="center",
                transform=ax.transAxes)
        ax.set_title("(c) Residuals by Hour (N/A)")

    # (d) By day of week
    ax = axes[1, 1]
    if test_dates is not None and len(test_dates) >= len(residuals):
        dow = test_dates[-len(residuals):].dayofweek
        day_names = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
        res_df = pd.DataFrame({"day_num": dow, "residual": residuals})
        res_df.boxplot(column="residual", by="day_num", ax=ax)
        ax.set_xticklabels(day_names)
        ax.set_xlabel("Day of Week")
        ax.set_ylabel("Residual (ppm)")
        ax.set_title("(d) Residuals by Day")
        plt.suptitle("")
    else:
        ax.text(0.5, 0.5, "No datetime available", ha="center", va="center",
                transform=ax.transAxes)
        ax.set_title("(d) Residuals by Day (N/A)")

    plt.suptitle("")
    plt.tight_layout()
    plt.savefig(
        output_dir / f"residual_analysis_h{horizon}.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 5: Error by CO2 level ---
    bins = [0, 500, 1000, np.inf]
    labels = ["Low (<500)", "Medium (500-1000)", "High (>1000)"]
    categories = pd.cut(y_true, bins=bins, labels=labels)

    fig, axes = plt.subplots(1, 2, figsize=(14, 5))

    ax = axes[0]
    abs_errors = np.abs(residuals)
    err_df = pd.DataFrame({"category": categories, "abs_error": abs_errors})
    err_df.boxplot(column="abs_error", by="category", ax=ax)
    ax.set_xlabel("CO2 Level")
    ax.set_ylabel("Absolute Error (ppm)")
    ax.set_title(f"Error by CO2 Level - {horizon}h")
    plt.suptitle("")

    ax = axes[1]
    ax.axis("off")
    rows_data = []
    for label in labels:
        mask = categories == label
        if mask.sum() > 0:
            bin_res = residuals[mask]
            rows_data.append([
                label,
                f"{mask.sum()}",
                f"{np.sqrt(np.mean(bin_res**2)):.2f}",
                f"{np.mean(np.abs(bin_res)):.2f}",
                f"{np.mean(bin_res):.2f}",
            ])
        else:
            rows_data.append([label, "0", "N/A", "N/A", "N/A"])

    tbl = ax.table(
        cellText=rows_data,
        colLabels=["CO2 Level", "N", "RMSE", "MAE", "Mean Bias"],
        loc="center",
        cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    tbl.scale(1.0, 1.5)

    plt.tight_layout()
    plt.savefig(
        output_dir / f"error_by_co2_level_h{horizon}.png",
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    print(f"  [E] LSTM-specific + prediction analysis complete")


# ======================================================================
#  Cross-Horizon Comparison
# ======================================================================

def run_cross_horizon_comparison(
    all_results: dict[int, dict],
    all_importance: dict[int, dict],
    output_dir: Path,
) -> None:
    """Compare metrics and feature attribution across horizons.

    Args:
        all_results: {horizon: results_dict} with y_true, y_pred, metrics.
        all_importance: {horizon: importance_data} with gradient rankings.
        output_dir: Directory to save plots.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    horizons = sorted(all_results.keys())
    print(f"  Generating cross-horizon comparison for horizons: {horizons}")

    if len(horizons) < 2:
        print("  Need >= 2 horizons for comparison. Skipping.")
        return

    # --- Plot 1: Side-by-side metrics table ---
    fig, ax = plt.subplots(figsize=(10, 3))
    ax.axis("off")
    rows_data = []
    for h in horizons:
        m = all_results[h]["metrics"]
        rows_data.append([
            f"{h}h",
            f"{m['rmse']:.2f}",
            f"{m['mae']:.2f}",
            f"{m['r2']:.4f}",
            f"{m['mape']:.2f}%",
        ])

    tbl = ax.table(
        cellText=rows_data,
        colLabels=["Horizon", "RMSE", "MAE", "R2", "MAPE"],
        loc="center",
        cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(12)
    tbl.scale(1.2, 2.0)
    ax.set_title("LSTM Performance Summary (preproc_D Enhanced 1h)", fontsize=13, pad=30)
    plt.tight_layout()
    plt.savefig(
        output_dir / "metrics_comparison.png", dpi=150, bbox_inches="tight"
    )
    plt.close(fig)

    # --- Plot 2: Feature attribution comparison ---
    if all_importance:
        n_horizons = len(horizons)
        fig, axes = plt.subplots(1, n_horizons, figsize=(8 * n_horizons, 8))
        if n_horizons == 1:
            axes = [axes]
        colors = [C_PRIMARY, C_SECONDARY, C_ACCENT, C_WARN]

        for idx, h in enumerate(horizons):
            ax = axes[idx]
            if h in all_importance and "feature_names" in all_importance[h]:
                imp_data = all_importance[h]
                fnames = imp_data["feature_names"]
                pcts = np.array(imp_data["gradient_pct"])
                sort_order = np.argsort(pcts)
                ax.barh(range(len(fnames)), pcts[sort_order],
                        color=colors[idx % len(colors)])
                ax.set_yticks(range(len(fnames)))
                ax.set_yticklabels([fnames[i] for i in sort_order], fontsize=8)
                ax.set_xlabel("Gradient Importance (%)")
                ax.set_title(f"{h}h Horizon")
                ax.grid(axis="x", alpha=0.3)
            else:
                ax.text(0.5, 0.5, f"No importance data for {h}h",
                        ha="center", va="center", transform=ax.transAxes)

        fig.suptitle("Feature Attribution Comparison Across Horizons", fontsize=14)
        plt.tight_layout()
        plt.savefig(
            output_dir / "feature_attribution_comparison.png",
            dpi=150, bbox_inches="tight",
        )
        plt.close(fig)

    print(f"  Cross-horizon comparison plots saved")


# ======================================================================
#  DOCX Report Generation
# ======================================================================

def generate_docx_report(
    all_results: dict[int, dict],
    all_gate_stats: dict[int, dict],
    all_importance: dict[int, dict],
    horizons: list[int],
    output_dir: Path,
) -> Path:
    """Generate comprehensive DOCX academic report with all figures and analysis.

    Args:
        all_results: {horizon: results_dict} with y_true, y_pred, metrics.
        all_gate_stats: {horizon: gate_stats_dict} from section A.
        all_importance: {horizon: importance_data} from section B.
        horizons: List of analyzed horizons.
        output_dir: Base output directory.

    Returns:
        Path to the generated DOCX file.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print("  Generating DOCX report...")
    doc = Document()

    # --- Title ---
    title = doc.add_heading(
        "LSTM Interpretability Study: Deep Analysis of Internal "
        "Representations for Indoor CO2 Forecasting",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle / metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Preprocessing Variant: preproc_D (Enhanced 1h)\n"
        f"Horizons: {', '.join(str(h) + 'h' for h in horizons)}\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_page_break()

    # --- Table of Contents (manual) ---
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Section A: Gate Dynamics Analysis",
        "3. Section B: Gradient-Based Feature Attribution",
        "4. Section C: Hidden State Structural Analysis",
        "5. Section D: Temporal Pattern Analysis",
        "6. Section E: LSTM-Specific and Prediction Analysis",
        "7. Cross-Horizon Comparison",
        "8. Avenues of Improvement",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # ============================================================
    #  1. Executive Summary
    # ============================================================
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report presents a comprehensive interpretability study of a standard "
        "multi-layer LSTM model trained for indoor CO2 concentration forecasting. "
        "The model uses the preproc_D preprocessing variant (Enhanced, 1-hour resolution) "
        "with 18 engineered features plus the CO2 target, a 24-hour lookback window, "
        "and two forecast horizons (1h and 24h). The study investigates five aspects: "
        "(A) internal gate dynamics through manual LSTM unrolling, "
        "(B) gradient-based feature attribution including SmoothGrad, "
        "(C) hidden state structure via PCA and K-Means clustering, "
        "(D) temporal patterns via FFT and autocorrelation analysis, and "
        "(E) weight matrix visualization and prediction quality assessment."
    )

    # Summary metrics table
    doc.add_heading("Performance Summary", level=2)
    table = doc.add_table(rows=1 + len(horizons), cols=5)
    table.style = "Light Grid Accent 1"
    headers = ["Horizon", "RMSE (ppm)", "MAE (ppm)", "R2", "MAPE (%)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_idx, h in enumerate(horizons):
        m = all_results[h]["metrics"]
        table.rows[row_idx + 1].cells[0].text = f"{h}h"
        table.rows[row_idx + 1].cells[1].text = f"{m['rmse']:.2f}"
        table.rows[row_idx + 1].cells[2].text = f"{m['mae']:.2f}"
        table.rows[row_idx + 1].cells[3].text = f"{m['r2']:.4f}"
        table.rows[row_idx + 1].cells[4].text = f"{m['mape']:.2f}"

    doc.add_page_break()

    # Helper to safely add figures
    def add_figure(fig_path: Path, caption: str, width: float = 6.0) -> None:
        """Add a figure with caption to the document."""
        if fig_path.exists():
            doc.add_picture(str(fig_path), width=Inches(width))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
        else:
            doc.add_paragraph(f"[Figure not found: {fig_path.name}]")

    # ============================================================
    #  2. Section A: Gate Dynamics Analysis
    # ============================================================
    doc.add_heading("2. Section A: Gate Dynamics Analysis", level=1)

    doc.add_paragraph(
        "LSTM networks use gating mechanisms to control information flow through the "
        "recurrent cell. Each timestep involves four gates: the input gate (i) controls "
        "how much new information enters the cell state; the forget gate (f) controls how "
        "much old information is retained; the candidate gate (g) proposes new cell content; "
        "and the output gate (o) controls how much of the cell state is exposed as the "
        "hidden state. By manually unrolling the LSTM using its weight matrices "
        "(W_ih, W_hh, b_ih, b_hh), we capture each gate's activation at every timestep "
        "and hidden dimension, providing insight into how the model processes temporal "
        "information."
    )

    for h in horizons:
        doc.add_heading(f"2.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "gate_dynamics"

        # Gate distributions
        for layer_idx in range(2):  # 2 layers
            fig_path = h_dir / f"gate_distributions_L{layer_idx}_h{h}.png"
            add_figure(
                fig_path,
                f"Figure: Gate activation distributions for Layer {layer_idx} at {h}h horizon. "
                f"The input and output gates use sigmoid activation (range [0,1]), "
                f"while the candidate gate uses tanh (range [-1,1]). "
                f"Distributions centered near 0.5 indicate active gating, while "
                f"distributions near 0 or 1 indicate gates that are predominantly closed or open.",
            )

        # Gate evolution
        for layer_idx in range(2):
            fig_path = h_dir / f"gate_evolution_L{layer_idx}_h{h}.png"
            add_figure(
                fig_path,
                f"Figure: Mean gate activation evolution over the lookback window for Layer {layer_idx}. "
                f"This reveals how the LSTM dynamically adjusts its gates as it processes "
                f"the input sequence from oldest to most recent timestep. Increasing forget gate "
                f"values toward recent steps indicate the model prioritizes recent information.",
            )

        # Forget gate heatmap
        fig_path = h_dir / f"forget_gate_heatmap_h{h}.png"
        add_figure(
            fig_path,
            f"Figure: Forget gate heatmap (Layer 0) showing per-dimension gate values across "
            f"the lookback window. Green (near 1) indicates information retention, red (near 0) "
            f"indicates forgetting. Patterns reveal which hidden dimensions specialize in "
            f"long-term vs short-term memory.",
        )

        # Gate statistics table
        fig_path = h_dir / f"gate_statistics_h{h}.png"
        add_figure(
            fig_path,
            f"Figure: Quantitative gate statistics including mean, std, min, max, and "
            f"sparsity metrics. Near-zero percentages indicate how often gates are effectively "
            f"closed, while near-one percentages show how often they are fully open.",
        )

        # Gate statistics DOCX table
        if h in all_gate_stats and all_gate_stats[h]:
            doc.add_heading("Gate Statistics (Quantitative)", level=3)
            gs = all_gate_stats[h]
            tbl = doc.add_table(rows=1 + len(gs), cols=7)
            tbl.style = "Light Grid Accent 1"
            tbl_headers = ["Gate", "Mean", "Std", "Min", "Max", "<0.05", ">0.95"]
            for ci, ch in enumerate(tbl_headers):
                tbl.rows[0].cells[ci].text = ch
            for ri, (stat_key, s) in enumerate(sorted(gs.items())):
                tbl.rows[ri + 1].cells[0].text = stat_key
                tbl.rows[ri + 1].cells[1].text = f"{s['mean']:.4f}"
                tbl.rows[ri + 1].cells[2].text = f"{s['std']:.4f}"
                tbl.rows[ri + 1].cells[3].text = f"{s['min']:.4f}"
                tbl.rows[ri + 1].cells[4].text = f"{s['max']:.4f}"
                tbl.rows[ri + 1].cells[5].text = f"{s['near_zero_pct']:.1f}%"
                tbl.rows[ri + 1].cells[6].text = f"{s['near_one_pct']:.1f}%"

        # Discussion
        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The gate dynamics reveal how the LSTM processes temporal CO2 data. "
            "The forget gate distribution and evolution across the lookback window "
            "indicate which timescales the model considers most important. A forget gate "
            "that stays high throughout the sequence suggests the model maintains a long-term "
            "memory of CO2 trends, while a gate that drops at earlier timesteps indicates "
            "the model focuses primarily on recent observations. The heatmap shows "
            "specialization across hidden dimensions - some dimensions act as long-term "
            "memory banks while others respond primarily to recent changes."
        )

    doc.add_page_break()

    # ============================================================
    #  3. Section B: Gradient-Based Feature Attribution
    # ============================================================
    doc.add_heading("3. Section B: Gradient-Based Feature Attribution", level=1)

    doc.add_paragraph(
        "Gradient-based attribution identifies which input features and timesteps most "
        "influence the model's predictions. We compute the gradient of the mean prediction "
        "with respect to the input tensor, yielding a saliency map of shape "
        "(lookback, n_features). Two methods are compared: vanilla gradient (raw absolute "
        "gradients) and SmoothGrad, which averages gradients over 30 noise-perturbed versions "
        "of the input to produce smoother, more interpretable saliency maps."
    )

    for h in horizons:
        doc.add_heading(f"3.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "gradient_attribution"

        # Gradient heatmap
        add_figure(
            h_dir / f"gradient_heatmap_h{h}.png",
            f"Figure: Gradient attribution heatmap showing |grad| magnitude across lookback "
            f"timesteps (x-axis) and features (y-axis). Brighter colors indicate stronger "
            f"influence on the prediction. Recent timesteps and CO2-related features "
            f"typically show highest attribution.",
        )

        # Feature ranking
        add_figure(
            h_dir / f"gradient_feature_ranking_h{h}.png",
            f"Figure: Per-feature gradient importance ranking (averaged across all timesteps "
            f"and test samples). This reveals the overall importance hierarchy of input features.",
        )

        # Temporal profile
        add_figure(
            h_dir / f"gradient_temporal_profile_h{h}.png",
            f"Figure: Temporal gradient profile showing how gradient magnitude varies across "
            f"the lookback window. A peak near t=0 (most recent) indicates the model relies "
            f"heavily on recent observations, while significant mass at earlier timesteps "
            f"suggests longer-term dependencies.",
        )

        # Vanilla vs SmoothGrad comparison
        add_figure(
            h_dir / f"gradient_vs_smoothgrad_h{h}.png",
            f"Figure: Side-by-side comparison of vanilla gradient and SmoothGrad (N=30) "
            f"attribution heatmaps. SmoothGrad produces smoother, less noisy saliency maps "
            f"by averaging over multiple noise-perturbed inputs.",
        )

        # SmoothGrad ranking
        add_figure(
            h_dir / f"smoothgrad_feature_ranking_h{h}.png",
            f"Figure: Feature importance ranking based on SmoothGrad. This provides a "
            f"more robust importance estimate compared to vanilla gradients.",
        )

        # Importance table
        if h in all_importance and all_importance[h]:
            doc.add_heading("Feature Importance (Quantitative)", level=3)
            imp = all_importance[h]
            fnames = imp["feature_names"]
            grad_pcts = imp["gradient_pct"]
            smooth_pcts = imp["smoothgrad_pct"]

            sort_idx = np.argsort(-np.array(grad_pcts))
            tbl = doc.add_table(rows=1 + len(fnames), cols=3)
            tbl.style = "Light Grid Accent 1"
            for ci, ch in enumerate(["Feature", "Gradient (%)", "SmoothGrad (%)"]):
                tbl.rows[0].cells[ci].text = ch
            for ri, si in enumerate(sort_idx):
                tbl.rows[ri + 1].cells[0].text = fnames[si]
                tbl.rows[ri + 1].cells[1].text = f"{grad_pcts[si]:.2f}"
                tbl.rows[ri + 1].cells[2].text = f"{smooth_pcts[si]:.2f}"

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The gradient analysis reveals which features the LSTM most relies upon for "
            "forecasting. CO2-related lag features (CO2_lag_1, CO2_lag_6, CO2_lag_24) and "
            "the CO2 target itself typically dominate, confirming the autoregressive nature "
            "of the forecasting task. Environmental features like temperature, humidity, and "
            "noise contribute secondary information. The temporal gradient profile shows "
            "exponential recency bias for the 1h horizon, while the 24h horizon should "
            "show broader temporal attention. SmoothGrad reduces gradient noise and may "
            "reveal subtle contributions from features that appear noisy under vanilla "
            "gradient analysis."
        )

    doc.add_page_break()

    # ============================================================
    #  4. Section C: Hidden State Structural Analysis
    # ============================================================
    doc.add_heading("4. Section C: Hidden State Structural Analysis", level=1)

    doc.add_paragraph(
        "The LSTM hidden state at the final timestep encodes the model's compressed "
        "representation of the input sequence. By applying PCA dimensionality reduction "
        "and K-Means clustering to these hidden state vectors across all test samples, "
        "we can visualize the geometric structure of the learned representation space "
        "and identify whether the model has learned to distinguish different operational "
        "regimes (e.g., low vs high CO2, daytime vs nighttime)."
    )

    for h in horizons:
        doc.add_heading(f"4.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "hidden_state_analysis"

        add_figure(
            h_dir / f"pca_variance_h{h}.png",
            f"Figure: PCA explained variance for LSTM hidden states. The cumulative curve "
            f"shows how many principal components are needed to capture the majority of "
            f"variance in the hidden state space.",
        )

        add_figure(
            h_dir / f"pca_co2_level_h{h}.png",
            f"Figure: PCA projection of hidden states colored by CO2 concentration level "
            f"(Low < 500 ppm, Medium 500-1000 ppm, High > 1000 ppm). Clear separation "
            f"indicates the model has learned CO2-level-dependent representations.",
        )

        add_figure(
            h_dir / f"pca_hour_h{h}.png",
            f"Figure: PCA projection colored by hour of day. Diurnal patterns in the "
            f"hidden state space indicate the model encodes time-of-day information, "
            f"which is relevant for occupancy-driven CO2 patterns.",
        )

        add_figure(
            h_dir / f"clustering_h{h}.png",
            f"Figure: K-Means clustering (K=4) in hidden state space with regime "
            f"crosstab. The crosstab shows how clusters correspond to CO2 concentration "
            f"regimes, revealing whether the LSTM has learned distinct operational modes.",
        )

        add_figure(
            h_dir / f"hidden_trajectories_h{h}.png",
            f"Figure: Hidden state trajectories for 5 example sequences projected into "
            f"PCA space. Each trajectory shows how the hidden state evolves as the model "
            f"processes the input sequence from start (circle) to end (X). Similar "
            f"trajectories across different sequences indicate consistent internal "
            f"dynamics.",
        )

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The hidden state analysis reveals the geometric structure of the LSTM's "
            "learned representations. Clear clustering by CO2 level indicates the model "
            "has effectively learned to differentiate between low-occupancy (low CO2) and "
            "high-occupancy (high CO2) scenarios. The PCA explained variance curve shows "
            "the effective dimensionality of the learned representations - if a small number "
            "of PCs capture most variance, the model has learned a compact representation. "
            "The hidden state trajectories show how the representation evolves through the "
            "lookback window, with convergent trajectories suggesting the model stabilizes "
            "its internal state as more context is provided."
        )

    doc.add_page_break()

    # ============================================================
    #  5. Section D: Temporal Pattern Analysis
    # ============================================================
    doc.add_heading("5. Section D: Temporal Pattern Analysis", level=1)

    doc.add_paragraph(
        "Temporal pattern analysis examines periodic structures in the LSTM's internal "
        "representations and prediction errors. FFT (Fast Fourier Transform) decomposes "
        "signals into frequency components, revealing periodicities like daily (24h) and "
        "half-daily (12h) cycles. Autocorrelation analysis measures temporal self-similarity "
        "at different lags, while rolling RMSE tracks how prediction quality varies over time."
    )

    for h in horizons:
        doc.add_heading(f"5.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "temporal_patterns"

        add_figure(
            h_dir / f"fft_hidden_pcs_h{h}.png",
            f"Figure: FFT spectrum of the top 3 principal components of the hidden states. "
            f"Peaks at specific periods indicate the model has learned to encode periodic "
            f"patterns in its hidden representations.",
        )

        add_figure(
            h_dir / f"autocorrelation_hidden_pcs_h{h}.png",
            f"Figure: Autocorrelation function (ACF) of hidden state PCs. Significant "
            f"autocorrelation at lag 24 would confirm the model captures daily periodicity. "
            f"Slowly decaying ACF indicates long-range temporal dependencies.",
        )

        add_figure(
            h_dir / f"fft_residuals_h{h}.png",
            f"Figure: FFT of prediction residuals. Peaks indicate periodic patterns that "
            f"the model fails to capture. A peak at 24h suggests the model does not fully "
            f"account for the daily CO2 cycle, while a flat spectrum indicates well-captured "
            f"temporal dynamics.",
        )

        add_figure(
            h_dir / f"rolling_rmse_h{h}.png",
            f"Figure: Rolling RMSE with a 24-hour window showing how prediction quality "
            f"varies over time. Spikes indicate periods where the model performs poorly, "
            f"potentially due to unusual events or regime changes.",
        )

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The temporal analysis reveals how well the LSTM captures periodic patterns "
            "in CO2 dynamics. Residual FFT peaks at 24h and 12h periods suggest the model "
            "partially captures but does not fully account for the daily occupancy cycle. "
            "Hidden state PC autocorrelation shows the temporal memory span of the learned "
            "representations. The rolling RMSE highlights temporal non-stationarity in "
            "model performance, which may correlate with seasonal changes, unusual weather "
            "events, or changes in building occupancy patterns."
        )

    doc.add_page_break()

    # ============================================================
    #  6. Section E: LSTM-Specific + Prediction Analysis
    # ============================================================
    doc.add_heading("6. Section E: LSTM-Specific and Prediction Analysis", level=1)

    doc.add_paragraph(
        "This section examines the LSTM's learned weight matrices and prediction quality. "
        "The input-to-hidden weight matrix (W_ih) reveals how each input feature connects "
        "to each gate across all hidden dimensions. The prediction analysis includes "
        "time series overlays, scatter plots, and multi-faceted residual diagnostics."
    )

    for h in horizons:
        doc.add_heading(f"6.{horizons.index(h)+1} Horizon: {h}h", level=2)
        h_dir = output_dir / f"h{h}" / "prediction_analysis"

        add_figure(
            h_dir / f"weight_matrix_W_ih_h{h}.png",
            f"Figure: Input-to-hidden weight matrices (W_ih) for Layer 0, separated by "
            f"gate type. The heatmaps show the learned weight values connecting each input "
            f"feature (y-axis) to each hidden unit (x-axis). Strong weights indicate "
            f"features that directly influence specific gate activations.",
        )

        add_figure(
            h_dir / f"predictions_overlay_h{h}.png",
            f"Figure: Time series overlay of actual vs predicted CO2 concentrations. "
            f"The shaded area represents prediction error magnitude. Close tracking "
            f"indicates good model performance.",
        )

        add_figure(
            h_dir / f"scatter_r2_h{h}.png",
            f"Figure: Scatter plot of predicted vs actual CO2 with R-squared, RMSE, "
            f"and MAE statistics. Points near the diagonal (red dashed line) indicate "
            f"accurate predictions.",
        )

        add_figure(
            h_dir / f"residual_analysis_h{h}.png",
            f"Figure: Four-panel residual analysis: (a) residual distribution with "
            f"normality assessment, (b) residuals over time for trend detection, "
            f"(c) residuals by hour of day for diurnal patterns, "
            f"(d) residuals by day of week for weekly patterns.",
        )

        add_figure(
            h_dir / f"error_by_co2_level_h{h}.png",
            f"Figure: Prediction error stratified by CO2 concentration level, with "
            f"quantitative metrics per regime. Higher errors at extreme CO2 levels "
            f"are typical due to fewer training samples in those ranges.",
        )

        # Metrics table
        doc.add_heading("Performance Metrics", level=3)
        m = all_results[h]["metrics"]
        tbl = doc.add_table(rows=2, cols=5)
        tbl.style = "Light Grid Accent 1"
        for ci, ch in enumerate(["RMSE", "MAE", "R2", "MAPE", "MSE"]):
            tbl.rows[0].cells[ci].text = ch
        tbl.rows[1].cells[0].text = f"{m['rmse']:.2f}"
        tbl.rows[1].cells[1].text = f"{m['mae']:.2f}"
        tbl.rows[1].cells[2].text = f"{m['r2']:.4f}"
        tbl.rows[1].cells[3].text = f"{m['mape']:.2f}%"
        tbl.rows[1].cells[4].text = f"{m['mse']:.2f}"

        doc.add_heading("Discussion", level=3)
        doc.add_paragraph(
            "The weight matrix visualization reveals the LSTM's learned feature "
            "connectivity patterns. Features with large-magnitude weights across many "
            "hidden dimensions are strongly utilized by the model, while sparse connectivity "
            "suggests selective feature use. The residual analysis provides diagnostic "
            "information: a mean-zero, symmetric residual distribution suggests unbiased "
            "predictions; diurnal patterns in residuals by hour indicate systematic "
            "time-of-day errors; and higher errors at extreme CO2 levels reflect the "
            "challenge of predicting rare events with limited training examples."
        )

    doc.add_page_break()

    # ============================================================
    #  7. Cross-Horizon Comparison
    # ============================================================
    if len(horizons) > 1:
        doc.add_heading("7. Cross-Horizon Comparison", level=1)

        doc.add_paragraph(
            "Comparing the LSTM's behavior across forecast horizons reveals how the model "
            "adapts its internal representations and feature utilization for short-term vs "
            "long-term prediction. As the forecast horizon increases, the task becomes "
            "fundamentally harder, requiring the model to capture longer-range dependencies "
            "and being less able to rely on simple autoregressive patterns."
        )

        comp_dir = output_dir / "comparison"
        add_figure(
            comp_dir / "metrics_comparison.png",
            "Figure: Side-by-side performance metrics across horizons. Degradation from "
            "1h to 24h indicates the inherent difficulty of longer-range forecasting.",
        )
        add_figure(
            comp_dir / "feature_attribution_comparison.png",
            "Figure: Feature importance comparison across horizons. Changes in the "
            "importance hierarchy reveal how the model shifts its reliance from "
            "autoregressive features (short horizon) to contextual features (long horizon).",
        )

        # Comparison table
        doc.add_heading("Metrics Comparison Table", level=2)
        tbl = doc.add_table(rows=1 + len(horizons), cols=5)
        tbl.style = "Light Grid Accent 1"
        for ci, ch in enumerate(["Horizon", "RMSE", "MAE", "R2", "MAPE"]):
            tbl.rows[0].cells[ci].text = ch
        for ri, h in enumerate(horizons):
            m = all_results[h]["metrics"]
            tbl.rows[ri + 1].cells[0].text = f"{h}h"
            tbl.rows[ri + 1].cells[1].text = f"{m['rmse']:.2f}"
            tbl.rows[ri + 1].cells[2].text = f"{m['mae']:.2f}"
            tbl.rows[ri + 1].cells[3].text = f"{m['r2']:.4f}"
            tbl.rows[ri + 1].cells[4].text = f"{m['mape']:.2f}%"

        doc.add_heading("Discussion", level=2)
        doc.add_paragraph(
            "The cross-horizon comparison typically shows significant performance degradation "
            "from 1h to 24h forecasting. This is expected because: (1) short-term predictions "
            "can leverage strong autoregressive patterns in CO2 concentration, "
            "(2) 24-hour ahead prediction requires capturing longer-range dependencies that "
            "may be beyond the effective memory capacity of the LSTM, and (3) accumulation "
            "of small errors over the longer horizon amplifies prediction uncertainty. "
            "The feature importance shift from lag features (dominant at 1h) to environmental "
            "features (more important at 24h) reflects the model learning to use contextual "
            "information when direct autoregressive signals become less reliable."
        )

    doc.add_page_break()

    # ============================================================
    #  8. Avenues of Improvement
    # ============================================================
    doc.add_heading("8. Avenues of Improvement", level=1)

    improvements = [
        (
            "Attention Mechanisms",
            "Incorporating attention mechanisms (e.g., Temporal Attention, as in TFT) "
            "would allow the LSTM to dynamically weight different timesteps in the lookback "
            "window rather than relying solely on the hidden state bottleneck. The temporal "
            "gradient analysis shows exponential recency bias, which attention could help "
            "balance by providing direct access to earlier timesteps."
        ),
        (
            "Bidirectional Processing",
            "While the current model is unidirectional, a bidirectional LSTM could capture "
            "both forward and backward temporal patterns during training. However, for "
            "real-time forecasting this is only applicable during feature extraction, not "
            "during actual prediction."
        ),
        (
            "Variational / Bayesian LSTM",
            "Adding dropout-based uncertainty estimation (MC Dropout) or variational "
            "inference would provide prediction confidence intervals, which are critical "
            "for decision-making in building management systems."
        ),
        (
            "Curriculum Learning for Long Horizons",
            "Training on progressively longer horizons (starting with 1h, gradually "
            "extending to 24h) could improve long-range prediction by helping the model "
            "first learn short-range dynamics before attempting longer-range forecasting."
        ),
        (
            "Feature Engineering Refinements",
            "The gradient analysis identifies which features contribute most. Features "
            "with consistently low gradient attribution could be removed to reduce model "
            "complexity, while new features (e.g., CO2 rate of change over longer windows, "
            "weather forecasts) could provide additional predictive signal."
        ),
        (
            "Ensemble Methods",
            "Combining the LSTM with tree-based models (XGBoost, CatBoost) in an ensemble "
            "could capture both sequential patterns (LSTM strength) and nonlinear feature "
            "interactions (tree model strength), potentially improving overall performance."
        ),
        (
            "Cell State Regularization",
            "Applying regularization directly to the cell state dynamics (e.g., encouraging "
            "smoother cell state trajectories or penalizing cell state magnitude) could "
            "improve generalization and prevent the forget gate from either fully opening "
            "or closing, maintaining more dynamic memory management."
        ),
        (
            "Transfer Learning Across Horizons",
            "Pre-training on the 1h horizon (where the model performs well) and fine-tuning "
            "for the 24h horizon could leverage the learned feature representations while "
            "adapting the output layer for longer-range prediction."
        ),
    ]

    for title, desc in improvements:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    # Save report
    report_path = output_dir / "lstm_interpretability_report.docx"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(report_path))
    print(f"  DOCX report saved to: {report_path}")

    return report_path


# ======================================================================
#  Summary Figure and JSON Export
# ======================================================================

def generate_summary_figure(
    all_results: dict[int, dict],
    all_importance: dict[int, dict],
    output_dir: Path,
) -> None:
    """Multi-panel summary figure with key findings.

    Args:
        all_results: {horizon: results_dict} with metrics.
        all_importance: {horizon: importance_data} with rankings.
        output_dir: Directory to save the summary figure.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    horizons = sorted(all_results.keys())

    fig = plt.figure(figsize=(20, 12))
    gs = GridSpec(2, 3, figure=fig, hspace=0.35, wspace=0.3)

    # [0,0] Metrics table
    ax = fig.add_subplot(gs[0, 0])
    ax.axis("off")
    rows_data = []
    for h in horizons:
        m = all_results[h]["metrics"]
        rows_data.append([
            f"{h}h", f"{m['rmse']:.2f}", f"{m['mae']:.2f}",
            f"{m['r2']:.4f}", f"{m['mape']:.2f}%",
        ])
    tbl = ax.table(
        cellText=rows_data,
        colLabels=["Horizon", "RMSE", "MAE", "R2", "MAPE"],
        loc="center", cellLoc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(11)
    tbl.scale(1.0, 1.8)
    ax.set_title("Performance Metrics", fontsize=12, pad=20)

    # [0,1] and [0,2]: Feature importance per horizon (top 10)
    colors_list = [C_PRIMARY, C_SECONDARY]
    for col_idx, h in enumerate(horizons[:2]):
        ax = fig.add_subplot(gs[0, col_idx + 1])
        if h in all_importance and "feature_names" in all_importance[h]:
            imp = all_importance[h]
            fnames = imp["feature_names"]
            pcts = np.array(imp["gradient_pct"])
            # Top 10
            top_idx = np.argsort(-pcts)[:10]
            top_idx = top_idx[::-1]  # Reverse for horizontal bar
            ax.barh(
                range(len(top_idx)),
                pcts[top_idx],
                color=colors_list[col_idx],
            )
            ax.set_yticks(range(len(top_idx)))
            ax.set_yticklabels([fnames[i] for i in top_idx], fontsize=8)
            ax.set_xlabel("Importance (%)")
            ax.set_title(f"Top Features - {h}h", fontsize=10)
        else:
            ax.text(0.5, 0.5, f"No data for {h}h", ha="center", va="center",
                    transform=ax.transAxes)

    # Bottom row: Analysis section summaries
    section_names = [
        "Gate Dynamics (A)\n- 4 gate types analyzed\n- Manual LSTM unrolling\n- Temporal evolution",
        "Gradient Attribution (B)\n- Vanilla + SmoothGrad\n- Per-feature ranking\n- Temporal profile",
        "Hidden States (C)\n- PCA decomposition\n- K-Means clustering\n- Trajectory analysis",
    ]
    for col in range(3):
        ax = fig.add_subplot(gs[1, col])
        ax.text(
            0.5, 0.5, section_names[col],
            ha="center", va="center", fontsize=10,
            transform=ax.transAxes,
            bbox=dict(boxstyle="round,pad=0.5", facecolor="lightyellow", alpha=0.8),
        )
        ax.set_frame_on(False)
        ax.set_xticks([])
        ax.set_yticks([])

    fig.suptitle(
        "LSTM Interpretability Study Summary - preproc_D (Enhanced 1h)",
        fontsize=15, fontweight="bold",
    )
    plt.savefig(output_dir / "summary.png", dpi=150, bbox_inches="tight")
    plt.close(fig)


def save_study_results(
    all_results: dict[int, dict],
    all_importance: dict[int, dict],
    all_gate_stats: dict[int, dict],
    output_dir: Path,
) -> None:
    """Save all metrics and importance data to JSON.

    Args:
        all_results: {horizon: results_dict}.
        all_importance: {horizon: importance_data}.
        all_gate_stats: {horizon: gate_stats_dict}.
        output_dir: Base output directory.
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    results_json: dict = {
        "study": "LSTM Interpretability",
        "variant": "preproc_D (Enhanced 1h)",
        "timestamp": datetime.now().isoformat(),
        "horizons": {},
    }

    for h in sorted(all_results.keys()):
        horizon_data: dict = {
            "metrics": all_results[h]["metrics"],
        }
        if h in all_importance:
            horizon_data["feature_importance"] = {
                "gradient_ranking": all_importance[h].get("gradient_ranking", []),
                "smoothgrad_ranking": all_importance[h].get("smoothgrad_ranking", []),
            }
        if h in all_gate_stats:
            horizon_data["gate_statistics"] = all_gate_stats[h]
        results_json["horizons"][str(h)] = horizon_data

    with open(output_dir / "study_results.json", "w", encoding="utf-8") as f:
        json.dump(results_json, f, indent=2)

    print(f"  Results saved to: {output_dir / 'study_results.json'}")


# ======================================================================
#  Main
# ======================================================================

def main() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )

    parser = argparse.ArgumentParser(
        description="LSTM Interpretability Study on preproc_D (Enhanced 1h)"
    )
    parser.add_argument(
        "--horizons", nargs="+", type=int, default=[1, 24],
        help="Forecast horizons in hours (default: 1 24)",
    )
    parser.add_argument(
        "--epochs", type=int, default=None,
        help="Override max epochs for LSTM training",
    )
    args = parser.parse_args()

    print(f"\n{'='*70}")
    print(f"  LSTM INTERPRETABILITY STUDY (DEEP ANALYSIS)")
    print(f"  Variant: preproc_D (Enhanced 1h)")
    print(f"  Horizons: {args.horizons}")
    print(f"  Sections: A(gates) B(gradients) C(hidden) D(temporal) E(LSTM-specific)")
    print(f"{'='*70}\n")

    # Load pipeline data once
    pipeline_config = load_interpretability_config(
        horizon=1, epochs_override=args.epochs
    )
    seed_everything(pipeline_config["training"]["seed"])

    raw_dir = Path(pipeline_config["data"].get("raw_dir", "data/raw"))
    print("  Loading preprocessing pipeline (preproc_D)...")
    train_df, val_df, test_df = run_preprocessing_pipeline(
        raw_dir=raw_dir, variant_config=pipeline_config,
    )
    print(f"  Pipeline loaded: train={len(train_df)}, val={len(val_df)}, test={len(test_df)}")

    all_results: dict[int, dict] = {}
    all_gate_stats: dict[int, dict] = {}
    all_importance: dict[int, dict] = {}

    for horizon in args.horizons:
        print(f"\n{'-'*60}")
        print(f"  HORIZON: {horizon}h")
        print(f"{'-'*60}\n")

        config = load_interpretability_config(
            horizon=horizon, epochs_override=args.epochs
        )
        seed_everything(config["training"]["seed"])

        output_dir = RESULTS_BASE / f"h{horizon}"
        output_dir.mkdir(parents=True, exist_ok=True)

        # ---- Train LSTM ----
        t0 = time.time()
        print(f"  Training LSTM for {horizon}h horizon...")
        best_model, dm, results = train_lstm(
            config, train_df.copy(), val_df.copy(), test_df.copy(), horizon,
        )
        elapsed = time.time() - t0
        print(f"  Training completed in {elapsed:.1f}s")

        metrics = results["metrics"]
        all_results[horizon] = results
        print(f"  RMSE={metrics['rmse']:.2f}  MAE={metrics['mae']:.2f}  "
              f"R2={metrics['r2']:.4f}  MAPE={metrics['mape']:.2f}%")

        # ---- Section A: Gate Dynamics ----
        gate_stats = run_gate_dynamics(
            best_model, dm, horizon, output_dir / "gate_dynamics"
        )
        all_gate_stats[horizon] = gate_stats

        # ---- Section B: Gradient Attribution ----
        importance_data = run_gradient_attribution(
            best_model, dm, config, horizon, output_dir / "gradient_attribution"
        )
        all_importance[horizon] = importance_data

        # ---- Section C: Hidden State Analysis ----
        run_hidden_state_analysis(
            best_model, dm, test_df, config, horizon,
            output_dir / "hidden_state_analysis",
        )

        # ---- Section D: Temporal Patterns ----
        run_temporal_patterns(
            best_model, dm, results, config, horizon,
            output_dir / "temporal_patterns",
        )

        # ---- Section E: LSTM-Specific + Predictions ----
        run_lstm_specific_analysis(
            best_model, dm, results, test_df, config, horizon,
            output_dir / "prediction_analysis",
        )

        # ---- Save metrics + predictions ----
        save_metrics(
            metrics, f"LSTM_h{horizon}", output_dir / "metrics.json",
            experiment_info={
                "name": "lstm_interpretability",
                "label": f"LSTM Deep Analysis h={horizon}",
                "description": "preproc_D Enhanced 1h variant",
            },
        )
        np.savez(
            output_dir / "predictions.npz",
            y_true=results["y_true"],
            y_pred=results["y_pred"],
        )

        # ---- GPU cleanup ----
        del best_model, dm
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        print(f"  GPU memory freed\n")

    # ---- Cross-horizon comparison ----
    if len(args.horizons) > 1:
        print(f"\n{'-'*60}")
        print(f"  CROSS-HORIZON COMPARISON")
        print(f"{'-'*60}\n")
        run_cross_horizon_comparison(
            all_results, all_importance, RESULTS_BASE / "comparison"
        )

    # ---- Summary ----
    generate_summary_figure(all_results, all_importance, RESULTS_BASE)
    save_study_results(all_results, all_importance, all_gate_stats, RESULTS_BASE)

    # ---- DOCX Report ----
    print(f"\n{'-'*60}")
    print(f"  GENERATING DOCX REPORT")
    print(f"{'-'*60}\n")
    report_path = generate_docx_report(
        all_results, all_gate_stats, all_importance, args.horizons, RESULTS_BASE,
    )

    print(f"\n{'='*70}")
    print(f"  LSTM INTERPRETABILITY STUDY COMPLETE")
    print(f"  Results saved to: {RESULTS_BASE}")
    print(f"  Report: {report_path}")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
